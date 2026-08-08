#!/usr/bin/env python3
"""
Banksia OS — HMO Operations API Blueprint.
Provides all HMO operations endpoints for daily team use.
Mounts at /api/banksia-os/

Architecture: Route definitions only — business logic lives in services/ modules.
"""
import ipaddress
import json, os, sys, re
from datetime import datetime, timezone, timedelta

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from flask import Blueprint, jsonify, request, session, current_app, g, has_request_context
from functools import wraps

from banksia_os_db import get_db, get_dict_db, count, dict_from_row, raw_query

# ── Service-layer imports ──
from services.db_service import (
    bool_fields, paginate, json_success, json_error, clean_none,
    int_param, float_param, build_search_clause, build_order_by, record_change,
    safe_error
)
from services.activity_service import (
    create_activity_log, log_activity, _format_value, _redact_if_sensitive,
    _get_entity_label, _derive_timeline_type, _redact_sensitive_fields, _enhance_timeline_item
)
from services.property_service import sync_unit_vacancy, get_monday_token, _monday_graphql, ensure_landlord_link
from services.maintenance_service import safe_status, safe_priority, parse_monday_cols, parse_photo_paths, parse_invoice_paths
from services.auth_service import (
    _hash_password, _verify_password, _validate_password_strength, _load_users, _save_users,
    VALID_ROLES, check_role_access, check_rate_limit, record_login_attempt, log_auth_event,
    ensure_audit_table
)
from services.finance_service import calculate_arrears, get_tenancy_summary
from services.notification_service import create_notification, get_user_notifications, get_my_updates

banksia_os_bp = Blueprint("banksia_os", __name__, url_prefix="/api/banksia-os")

# Versioned alias — /api/v1/banksia-os routes to the same blueprint
# DEPRECATED: After frontend migration, change url_prefix to /api/v1/banksia-os

# ── Change log table init ──
try:
    _cl_db = get_dict_db()
    _cl_db.execute("""
        CREATE TABLE IF NOT EXISTS change_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_name TEXT NOT NULL,
            action TEXT NOT NULL,
            entity_type TEXT NOT NULL,
            entity_id TEXT,
            summary TEXT,
            details TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    _cl_db.commit()
    _cl_db.close()
except Exception as _e:
    current_app.logger.error(f"Error in line ~54: {_e}")
    pass  # Expected before change_log table exists on first run


# ── Global auth for the entire blueprint ──
# ── Role → blocked route-family policy ──
# Route families are matched against the path relative to /api/banksia-os.
# super_admin and admin are never scoped here (handled by the early return in
# the guard). Every other role is denied the families listed in "block".
_FAM_FINANCE = ("/transactions", "/invoices", "/rent", "/deposits", "/finance")
_FAM_PII     = ("/tenants", "/tenancies", "/guarantors")
_FAM_APPS    = ("/applicants", "/submissions")
_FAM_MAINT   = ("/maintenance", "/contractors", "/orders")
_FAM_DOCS    = ("/documents", "/entity-documents")

_ROLE_POLICY = {
    # Finance: money + read tenant context; no applications, no maintenance.
    "finance":     {"block": _FAM_APPS + _FAM_MAINT,                         "read_only": False},
    # HMO / STR managers: full ops incl. tenant PII + maintenance; no finance.
    "hmo_manager": {"block": _FAM_FINANCE,                                   "read_only": False},
    "str_manager": {"block": _FAM_FINANCE,                                   "read_only": False},
    # Maintenance: jobs + contractors + read tenant contact; no money, no apps.
    "maintenance": {"block": _FAM_FINANCE + _FAM_APPS,                       "read_only": False},
    # Lettings & viewings: applications only; no live-tenant PII, no money.
    "lettings":    {"block": _FAM_FINANCE + _FAM_PII + _FAM_MAINT + _FAM_DOCS, "read_only": False},
    # Projects (dev/PM): properties/units + applications; no PII, no money.
    "projects":    {"block": _FAM_FINANCE + _FAM_PII + _FAM_MAINT + _FAM_DOCS, "read_only": False},
    # Viewer: read-only dashboards/reports/properties; nothing sensitive.
    "viewer":      {"block": _FAM_FINANCE + _FAM_PII + _FAM_APPS + _FAM_MAINT + _FAM_DOCS, "read_only": True},
    # Unknown roles get the most restrictive treatment.
    "_default":    {"block": _FAM_FINANCE + _FAM_PII + _FAM_APPS + _FAM_MAINT + _FAM_DOCS, "read_only": True},
}


def _client_ip():
    """Address Traefik actually saw.

    Traefik appends the peer to X-Forwarded-For, so the last entry is the one
    it observed and the only one a caller cannot spoof by sending its own header.
    """
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        return xff.split(",")[-1].strip()
    return request.remote_addr or ""


def _api_key_ip_allowed(entry):
    """Whether this key may be used from the calling address.

    Loopback and private ranges always pass so anything running on the box keeps
    working. Off-box, a key is refused unless the address is in its allowed_ips.
    """
    try:
        ip = ipaddress.ip_address(_client_ip())
    except ValueError:
        return False
    if ip.is_loopback or ip.is_private or ip.is_link_local:
        return True
    for cidr in entry.get("allowed_ips") or []:
        try:
            if ip in ipaddress.ip_network(cidr, strict=False):
                return True
        except ValueError:
            continue
    return False


@banksia_os_bp.before_request
def _require_banksia_auth():
    """All routes in this blueprint require a logged-in session or valid API key."""
    # Public routes that don't need auth
    # Missive posts here with an HMAC signature instead of a session, so this one
    # route verifies itself (see api_missive_hook). Matched on the full path
    # because request.path carries the blueprint prefix.
    if request.path == _MISSIVE_HOOK_PATH:
        return None
    public_prefixes = ("/submissions/public", "/applicants/public", "/tenancies/public")
    if request.path.startswith(public_prefixes):
        return None
    # Check API key first (for programmatic access)
    # Header only. A key in the query string ends up in access logs, browser
    # history and Referer headers, which is not a place for a super_admin key.
    api_key = request.headers.get("X-API-Key")
    if api_key:
        _ak_path = os.path.join(os.path.dirname(__file__), "api_keys.json")
        if os.path.exists(_ak_path):
            try:
                _ak_data = json.load(open(_ak_path))
                _entry = _ak_data.get(api_key)
                if _entry:
                    if not _api_key_ip_allowed(_entry):
                        return jsonify({"success": False, "error": "API key not permitted from this address"}), 403
                    request.current_user = {"username": _entry.get("name", "API"), "role": _entry.get("role", "admin")}
                    return None
            except Exception:
                pass
        return jsonify({"success": False, "error": "Invalid API key"}), 401
    user = session.get("user")
    if not user:
        return jsonify({"success": False, "error": "Not logged in"}), 401
    request.current_user = user

    # ── Role-based data scoping (server-side source of truth) ──
    # Backend mirror of packages/permissions/roles.ts. The UI hides what a role
    # cannot use; this makes the API refuse it even if the UI is bypassed.
    # Each role is blocked from whole route families it has no business in.
    role = (user.get("role") or "").lower()
    if role not in ("super_admin", "admin"):
        rel = request.path[len("/api/banksia-os"):] or "/"
        policy = _ROLE_POLICY.get(role, _ROLE_POLICY["_default"])
        # Read-only roles may only issue GET/HEAD/OPTIONS.
        if policy["read_only"] and request.method not in ("GET", "HEAD", "OPTIONS"):
            return jsonify({
                "success": False,
                "error": "Your role has read-only access",
            }), 403
        # Whole-family blocks (any method).
        if rel.startswith(policy["block"]):
            return jsonify({
                "success": False,
                "error": "You do not have permission to access this data",
            }), 403

    # ── Destructive-action guard ──
    # Only super_admin may permanently delete a CORE ENTITY (property, owner,
    # unit, tenancy, tenant, applicant). This mirrors the frontend permission
    # map (roles.ts), where 'delete' is granted to super_admin only, and makes
    # the backend the source of truth so the API can't be bypassed when the UI
    # hides the button. Operational sub-resource deletes (documents, images,
    # tags, messages, invoices, referencing) are intentionally NOT covered here
    # so admins can still manage day-to-day records.
    if request.method == "DELETE" and role != "super_admin":
        rel = request.path[len("/api/banksia-os"):] or "/"
        if re.match(r"^/(properties|property-owners|units|tenancies|tenants|applicants)/\d+/?$", rel):
            return jsonify({
                "success": False,
                "error": "Only super admins can permanently delete this record",
            }), 403


# ── Helpers imported from services/db_service.py ──
#   bool_fields, paginate, json_success, json_error, clean_none,
#   int_param, float_param, build_search_clause, build_order_by, record_change


@banksia_os_bp.route("/sync/fingerprint", methods=["GET"])
def api_sync_fingerprint():
    """Lightweight endpoint returning data fingerprints for change detection.
    Returns the latest change_log id so the client can detect changes."""
    _db = get_dict_db()
    try:
        row = _db.execute("SELECT MAX(id) AS max_id, MAX(created_at) AS latest_ts FROM change_log").fetchone()
        return json_success({
            "fingerprint": row["max_id"] or 0,
            "latest_ts": row["latest_ts"] or "",
        })
    except Exception:
        return json_success({"fingerprint": 0, "latest_ts": ""})
    finally:
        _db.close()


@banksia_os_bp.route("/sync/activity", methods=["GET"])
def api_sync_activity():
    """Return recent activity log entries."""
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 50, max_val=MAX_PAGE_SIZE)

    since_id = request.args.get("since_id")
    if since_id:
        try:
            since_id = int(since_id)
            where = "WHERE id > ?"
            params = [since_id]
        except ValueError:
            where = ""
            params = []
    else:
        where = ""
        params = []

    rows, total = paginate(
        f"SELECT * FROM change_log {where} ORDER BY created_at DESC, id DESC",
        f"SELECT COUNT(*) AS cnt FROM change_log {where}",
        params, page, per_page
    )

    return json_success({"items": rows, "total": total})


@banksia_os_bp.route("/activity-feed", methods=["GET"])
def api_activity_feed():
    """Unified audit feed from activity_log with before/after + page context.
    Super admins/admins see EVERYONE's activity; every other role sees only their own."""
    username, role = _get_current_user()
    is_admin = (role or "").lower() in ("super_admin", "admin")

    page = int_param(request.args.get("page"), 1)
    per_page = int_param(request.args.get("per_page"), 50, max_val=MAX_PAGE_SIZE)
    search = (request.args.get("search") or "").strip()

    where, params = [], []
    if not is_admin:
        where.append("user_name = ?")
        params.append(username)
    if search:
        like = f"%{search}%"
        where.append("(user_name LIKE ? OR action LIKE ? OR entity_type LIKE ? OR "
                     "COALESCE(field_changed,'') LIKE ? OR COALESCE(notes,'') LIKE ? OR "
                     "COALESCE(old_value,'') LIKE ? OR COALESCE(new_value,'') LIKE ?)")
        params.extend([like] * 7)
    where_sql = ("WHERE " + " AND ".join(where)) if where else ""

    PAGE_BY_ENTITY = {
        "property": ("Properties", "/properties"),
        "unit": ("Units", "/properties"),
        "tenancy": ("Tenancies", "/tenancies"),
        "tenant": ("Tenants", "/tenants"),
        "applicant": ("Applicants", "/applicants"),
        "referencing_form": ("Referencing", "/referencing"),
        "guarantor": ("Guarantors", "/applicants"),
        "deposit": ("Deposits", "/deposits"),
        "property_owner": ("Property Owners", "/owners"),
        "maintenance_job": ("Maintenance", "/maintenance"),
        "invoice": ("Invoices", "/invoices"),
        "message_thread": ("Messages", "/messages"),
        "tag": ("Settings", "/settings"),
    }

    db = get_dict_db()
    try:
        total = db.execute(f"SELECT COUNT(*) AS c FROM activity_log {where_sql}", params).fetchone()["c"]
        offset = (page - 1) * per_page
        rows = db.execute(
            f"SELECT * FROM activity_log {where_sql} ORDER BY created DESC, id DESC LIMIT ? OFFSET ?",
            params + [per_page, offset]
        ).fetchall()

        items = []
        for r in rows:
            d = dict(r)
            ent = d.get("entity_type") or ""
            page_label, page_path = PAGE_BY_ENTITY.get(ent, (ent.replace("_", " ").title() or "System", ""))
            label = _get_entity_label(db, ent, d.get("entity_id"))
            fld = d.get("field_changed")
            action = d.get("action") or "update"
            if fld:
                summary = f"{str(fld).replace('_', ' ')} changed on {ent.replace('_', ' ')} {label}"
            elif d.get("notes"):
                summary = d.get("notes")
            else:
                summary = f"{action.replace('_', ' ')} · {ent.replace('_', ' ')} {label}"
            items.append({
                "id": d.get("id"),
                "user_name": d.get("user_name") or "system",
                "action": action,
                "entity_type": ent,
                "entity_id": d.get("entity_id"),
                "entity_label": label,
                "page_label": page_label,
                "page_path": page_path,
                "field_changed": fld,
                "old_value": d.get("old_value"),
                "new_value": d.get("new_value"),
                "summary": summary,
                "notes": d.get("notes"),
                "created_at": d.get("created"),
            })
        return json_success({
            "items": items, "total": total, "page": page, "per_page": per_page,
            "scope": "all" if is_admin else "own",
            "viewer": username, "viewer_role": role,
        })
    finally:
        db.close()


@banksia_os_bp.route("/activity-feed/fingerprint", methods=["GET"])
def api_activity_feed_fingerprint():
    """Lightweight change signal for the activity feed's realtime refresh."""
    db = get_dict_db()
    try:
        row = db.execute("SELECT MAX(id) AS max_id, MAX(created) AS latest FROM activity_log").fetchone()
        return json_success({"fingerprint": row["max_id"] or 0, "latest_ts": row["latest"] or ""})
    finally:
        db.close()


# ── User helpers imported from services/auth_service.py ──
#   _hash_password, _validate_password_strength, _load_users, _save_users
# ── Utility helpers imported from services/db_service.py ──
#   int_param, float_param, build_search_clause, build_order_by

# Hard ceiling on any page-size / limit param to prevent a single request
# loading an entire table into memory (memory-exhaustion protection under load).
MAX_PAGE_SIZE = 200


def api_update_resource(table, item_id):
    """Generic PATCH handler — updates any field on any table by item ID."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")
    # Build SET clause from provided fields
    set_parts = []
    params = []
    valid_tables = {"properties", "units", "tenancies", "tenants", "applicants", "property_owners", "message_threads"}
    if table not in valid_tables:
        return json_error(f"Invalid table: {table}", 400)
    # Tables mirrored from Arthur carry dirty-tracking columns. Any local edit
    # must flag the row so (a) the inbound pull sync won't overwrite it and
    # (b) the push-back sync knows to send the change to Arthur.
    SYNCED_TABLES = {"properties", "units", "tenancies", "tenants", "applicants"}
    protected_keys = {"sync_dirty", "local_modified", "sync_origin", "pushed_at", "arthur_id", "id"}
    # Introspect the real columns so an unknown field from the client is ignored
    # rather than crashing the UPDATE with a 500 "no such column" error.
    _col_db = get_dict_db()
    try:
        real_cols = {r["name"] for r in _col_db.execute(f"PRAGMA table_info({table})").fetchall()}
    finally:
        _col_db.close()
    ignored = []
    for key, val in data.items():
        if key in protected_keys:
            continue  # never let the client set tracking/identity fields directly
        if key not in real_cols:
            ignored.append(key)
            continue  # skip fields that don't exist on this table
        set_parts.append(f"{key} = ?")
        params.append(val)
    if not set_parts:
        return json_error(f"No valid fields to update (ignored: {', '.join(ignored) or 'none'})")
    if table in SYNCED_TABLES:
        _now = datetime.now(timezone.utc).isoformat()
        set_parts.append("sync_dirty = ?");    params.append(1)
        set_parts.append("local_modified = ?"); params.append(_now)
        set_parts.append("sync_origin = ?");    params.append("banksia_os")
    params.append(item_id)
    ENTITY_TYPE_BY_TABLE = {
        "properties": "property", "units": "unit", "tenancies": "tenancy",
        "tenants": "tenant", "applicants": "applicant",
        "property_owners": "property_owner", "message_threads": "message_thread",
        "tags": "tag",
    }
    _ent = ENTITY_TYPE_BY_TABLE.get(table, table)
    db = get_dict_db()
    try:
        _old_row = db.execute(f"SELECT * FROM {table} WHERE id = ?", [item_id]).fetchone()
        _old_row = dict(_old_row) if _old_row else {}
        db.execute(f"UPDATE {table} SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()
        updated_fields = [k for k in data.keys() if k in real_cols and k not in protected_keys]
        # -- Audit: log before/after for every field that actually changed --
        try:
            _logged = 0
            for _k in updated_fields:
                _o = _old_row.get(_k)
                _n = data.get(_k)
                if str(_o if _o is not None else "") != str(_n if _n is not None else ""):
                    _log_activity(_ent, item_id, "update", _k, _o, _n, db=db)
                    _logged += 1
            if _logged:
                db.commit()
                _actor = getattr(request, "current_user", {}).get("username", "system")
                _lbl = (_old_row.get("name") or _old_row.get("ref") or _old_row.get("unit_ref")
                        or _old_row.get("main_tenant_name") or str(item_id))
                record_change(_actor, "updated", _ent, str(item_id), str(_lbl))
        except Exception:
            pass
        return json_success({"updated": True, "id": item_id, "fields": updated_fields, "ignored": ignored})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

# ═══════════════════════════════════════════════
# 0. UNIT VACANCY SYNC — derives unit_vacant from active tenancies
# ═══════════════════════════════════════════════

def sync_unit_vacancy(db=None):
    """
    Derive unit_vacant from active tenancies.
    A unit is vacant (unit_vacant=1) if it has NO tenancies in
    Current/Periodic/Active status. Otherwise it's occupied (unit_vacant=0).
    Idempotent — safe to run any time.
    Returns a dict with counts of updated rows.
    """
    if db is None:
        db = get_dict_db()
        close_after = True
    else:
        close_after = False
    try:
        active_statuses = ("'Current','current','Periodic','periodic','Active','active'")
        # Units that should be occupied (have active tenancy) but are flagged vacant
        now_occupied = db.execute(
            f"UPDATE units SET unit_vacant = 0, modified = datetime('now') "
            f"WHERE unit_vacant = 1 AND id IN ("
            f"  SELECT DISTINCT t.unit_id FROM tenancies t "
            f"  WHERE t.unit_id IS NOT NULL AND t.status IN ({active_statuses})"
            f")"
        ).rowcount

        # Units that should be vacant (no active tenancy) but are flagged occupied
        now_vacant = db.execute(
            f"UPDATE units SET unit_vacant = 1, modified = datetime('now') "
            f"WHERE unit_vacant = 0 AND id NOT IN ("
            f"  SELECT DISTINCT t.unit_id FROM tenancies t "
            f"  WHERE t.unit_id IS NOT NULL AND t.status IN ({active_statuses})"
            f") AND id NOT IN (SELECT unit_id FROM tenancies WHERE unit_id IS NOT NULL AND status='Past')"
        ).rowcount

        db.commit()
        return {"occupied_fixed": now_occupied, "vacant_fixed": now_vacant, "total_changed": now_occupied + now_vacant}
    except Exception as e:
        db.rollback()
        raise e
    finally:
        if close_after:
            db.close()

@banksia_os_bp.route("/units/sync-vacancy", methods=["POST"])
def api_sync_unit_vacancy():
    """Manually trigger a unit vacancy sync from tenancies."""
    try:
        result = sync_unit_vacancy()
        return json_success(result)
    except Exception as e:
        return json_error(safe_error(e), 500)

# ═══════════════════════════════════════════════
# 1. DASHBOARD SUMMARY
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/dashboard")
def api_dashboard():
    db = get_dict_db()
    try:
        # Total properties
        total_properties = db.execute("SELECT COUNT(*) AS cnt FROM properties WHERE status IS NULL OR status <> 'archived'").fetchone()["cnt"]

        # Total units
        total_units = db.execute("SELECT COUNT(*) AS cnt FROM units WHERE is_active IS NULL OR is_active = 1").fetchone()["cnt"]

        # Occupied / vacant units
        occupied_units = db.execute(
            "SELECT COUNT(*) AS cnt FROM units WHERE unit_vacant = 0 AND (is_active IS NULL OR is_active = 1)"
        ).fetchone()["cnt"]
        vacant_units = db.execute(
            "SELECT COUNT(*) AS cnt FROM units WHERE unit_vacant = 1 AND (is_active IS NULL OR is_active = 1)"
        ).fetchone()["cnt"]

        # Total tenancies & tenants
        total_tenancies = db.execute("SELECT COUNT(*) AS cnt FROM tenancies WHERE status <> 'Archived'").fetchone()["cnt"]
        total_tenants = db.execute("SELECT COUNT(*) AS cnt FROM tenants").fetchone()["cnt"]
        total_applicants = db.execute("SELECT COUNT(*) AS cnt FROM applicants").fetchone()["cnt"]

        # Active tenancies — Arthur statuses: Current, Periodic, Active
        active_statuses = ("'Current', 'current', 'Periodic', 'periodic', 'Active', 'active'")
        active_tenancies = db.execute(
            f"SELECT COUNT(*) AS cnt FROM tenancies WHERE status IN ({active_statuses})"
        ).fetchone()["cnt"]

        # Monthly rent roll — active tenancies only
        monthly_rent_roll = db.execute(
            f"SELECT COALESCE(SUM(rent_amount), 0) AS total FROM tenancies "
            f"WHERE status IN ({active_statuses})"
        ).fetchone()["total"]

        # Total arrears
        total_arrears = db.execute(
            "SELECT COALESCE(SUM(amount_outstanding), 0) AS total FROM transactions "
            "WHERE is_outstanding = 1"
        ).fetchone()["total"]

        # Pending applicants
        pending_applicants = db.execute(
            "SELECT COUNT(*) AS cnt FROM applicants WHERE status IN ('Active', 'active', 'Pending', 'pending', 'New', 'new', 'Viewing', 'viewing', 'Application', 'application', 'Referencing', 'referencing')"
        ).fetchone()["cnt"]

        # Deposits — currently held (from deposits table)
        currently_held = db.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total FROM deposits WHERE current_status = 'held'"
        ).fetchone()["total"]

        all_time_deposits = db.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total FROM deposits"
        ).fetchone()["total"]

        deposits_unregistered = db.execute(
            "SELECT COUNT(*) AS cnt FROM deposits WHERE protection_status != 'protected' AND current_status = 'held'"
        ).fetchone()["cnt"]

        now = datetime.now(timezone.utc)
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
        if now.month == 12:
            next_month = now.replace(year=now.year + 1, month=1, day=1)
        else:
            next_month = now.replace(month=now.month + 1, day=1)
        month_end = next_month.isoformat()

        tenants_moving_in_this_month = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies "
            "WHERE move_in_date >= ? AND move_in_date < ? "
            "AND status IN ('Active', 'active', 'Periodic', 'periodic')",
            (month_start, month_end)
        ).fetchone()["cnt"]

        tenants_moving_out_this_month = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies "
            "WHERE move_out_date >= ? AND move_out_date < ? "
            "AND status IN ('Active', 'active', 'Periodic', 'periodic')",
            (month_start, month_end)
        ).fetchone()["cnt"]

        # Unit occupancy rate
        unit_occupancy_rate = round((occupied_units / total_units * 100) if total_units > 0 else 0, 1)

        # ── Portal / referencing submissions awaiting the team ──
        # Referencing forms the applicant has actually submitted but nobody has reviewed
        pending_referencing_submissions = db.execute(
            "SELECT COUNT(*) AS cnt FROM referencing_forms "
            "WHERE submitted_at IS NOT NULL AND reviewed_at IS NULL "
            "AND status IN ('submitted', 'Submitted')"
        ).fetchone()["cnt"]
        # Tenant-portal maintenance requests still open
        open_maintenance_requests = db.execute(
            "SELECT COUNT(*) AS cnt FROM maintenance_requests "
            "WHERE LOWER(COALESCE(status, 'open')) IN ('open', 'new', '')"
        ).fetchone()["cnt"]
        # Portal message threads still open
        open_message_threads = db.execute(
            "SELECT COUNT(*) AS cnt FROM message_threads "
            "WHERE LOWER(COALESCE(status, 'open')) IN ('open', 'new', '')"
        ).fetchone()["cnt"]
        # Applicant-uploaded documents awaiting the team to verify
        pending_document_uploads = db.execute(
            "SELECT COUNT(*) AS cnt FROM referencing_documents "
            "WHERE LOWER(COALESCE(uploaded_by, '')) = 'applicant' "
            "AND COALESCE(is_verified, 0) = 0"
        ).fetchone()["cnt"]
        new_submissions_total = (
            pending_referencing_submissions + open_maintenance_requests
            + open_message_threads + pending_document_uploads
        )

        # Leading property (highest total rent)
        leading = db.execute(
            "SELECT p.id, COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS name, "
            "SUM(t.rent_amount) AS total_rent FROM tenancies t "
            "JOIN properties p ON t.property_id = p.id "
            "WHERE t.status IN ('Current', 'Active', 'Periodic', 'current', 'active', 'periodic') "
            "GROUP BY p.id ORDER BY total_rent DESC LIMIT 1"
        ).fetchone()

        # ── Phase 2 additions ──

        # Vacant units list with property names
        vacant_units_list = db.execute(
            "SELECT u.id, u.unit_ref, u.market_rent, u.property_id, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name "
            "FROM units u "
            "JOIN properties p ON u.property_id = p.id "
            "WHERE u.unit_vacant = 1 "
            "ORDER BY p.name ASC, u.sort_order ASC, u.unit_ref ASC"
        ).fetchall()

        # Upcoming move-ins with tenant/property/unit details (this calendar month only)
        month_end_exclusive = next_month.isoformat()
        upcoming_move_ins = db.execute(
            "SELECT t.id AS tenancy_id, t.move_in_date, t.main_tenant_name, "
            "t.property_id, t.unit_id, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "u.unit_ref "
            "FROM tenancies t "
            "JOIN properties p ON t.property_id = p.id "
            "JOIN units u ON t.unit_id = u.id "
            "WHERE t.move_in_date >= ? AND t.move_in_date < ? "
            "AND t.status IN ('Active', 'active', 'Periodic', 'periodic') "
            "ORDER BY t.move_in_date ASC",
            (month_start, month_end_exclusive)
        ).fetchall()

        # Upcoming move-outs with tenant/property/unit details (this calendar month only)
        upcoming_move_outs = db.execute(
            "SELECT t.id AS tenancy_id, t.move_out_date, t.main_tenant_name, "
            "t.property_id, t.unit_id, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "u.unit_ref "
            "FROM tenancies t "
            "JOIN properties p ON t.property_id = p.id "
            "JOIN units u ON t.unit_id = u.id "
            "WHERE t.move_out_date >= ? AND t.move_out_date < ? "
            "AND t.status IN ('Active', 'active', 'Periodic', 'periodic') "
            "ORDER BY t.move_out_date ASC",
            (month_start, month_end_exclusive)
        ).fetchall()

        # Referencing pipeline breakdown
        referencing_pipeline_raw = db.execute(
            "SELECT status, COUNT(*) AS count FROM referencing_forms GROUP BY status"
        ).fetchall()
        pipeline_map = {}
        for r in referencing_pipeline_raw:
            st = (r["status"] or "unknown").lower()
            pipeline_map[st] = r["count"]
        referencing_pipeline = {
            "new": pipeline_map.get("draft", 0) + pipeline_map.get("sent", 0),
            "submitted": pipeline_map.get("submitted", 0),
            "under_review": pipeline_map.get("under_review", 0),
            "approved": pipeline_map.get("approved", 0),
            "rejected": pipeline_map.get("rejected", 0),
            "declined": pipeline_map.get("rejected", 0) + pipeline_map.get("declined", 0),
            "tenancy_created": pipeline_map.get("tenancy_created", 0),
            "total": sum(pipeline_map.values()),
        }

        # Arrears by tenancy — count of affected tenancies and top arrears list
        tenancies_in_arrears_count = db.execute(
            "SELECT COUNT(DISTINCT tenancy_id) AS cnt FROM transactions "
            "WHERE is_outstanding = 1 AND tenancy_id IS NOT NULL "
            "AND amount_outstanding > 0"
        ).fetchone()["cnt"]

        arrears_by_tenancy = db.execute(
            "SELECT txn.tenancy_id, t.id AS local_tenancy_id, t.ref AS tenancy_ref, "
            "t.main_tenant_name, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "SUM(COALESCE(txn.amount_outstanding, 0)) AS arrears_total "
            "FROM transactions txn "
            "LEFT JOIN tenancies t ON t.arthur_id = txn.tenancy_id "
            "LEFT JOIN properties p ON t.property_id = p.id "
            "WHERE txn.is_outstanding = 1 AND txn.amount_outstanding > 0 "
            "AND txn.tenancy_id IS NOT NULL "
            "GROUP BY txn.tenancy_id "
            "ORDER BY arrears_total DESC "
            "LIMIT 20"
        ).fetchall()
        arrears_by_tenancy_list = [
            {
                "tenancy_id": r["tenancy_id"],
                "tenancy_ref": r["tenancy_ref"],
                "tenant_name": r["main_tenant_name"],
                "property_name": r["property_name"],
                "arrears_total": round(r["arrears_total"], 2),
            }
            for r in arrears_by_tenancy
        ]

        return json_success({
            "total_properties": total_properties,
            "total_units": total_units,
            "occupied_units": occupied_units,
            "vacant_units": vacant_units,
            "total_tenancies": total_tenancies,
            "active_tenancies": active_tenancies,
            "total_tenants": total_tenants,
            "total_applicants": total_applicants,
            "monthly_rent_roll": round(monthly_rent_roll, 2),
            "monthly_rent_income": round(monthly_rent_roll, 2),
            "total_arrears": round(total_arrears, 2),
            "total_deposits_held": round(currently_held, 2),
            "total_deposits": round(currently_held, 2),
            "total_deposits_all_time": round(all_time_deposits, 2),
            "pending_applicants": pending_applicants,
            "total_pending_applicants": pending_applicants,
            "unit_occupancy_rate": unit_occupancy_rate,
            "recent_arrivals_count": tenants_moving_in_this_month,
            "upcoming_move_outs_count": tenants_moving_out_this_month,
            "tenants_moving_in_this_month": tenants_moving_in_this_month,
            "tenants_moving_out_this_month": tenants_moving_out_this_month,
            "deposits_unregistered": deposits_unregistered,
            "leading_property": ({"id": leading["id"], "name": leading["name"], "total_rent": round(leading["total_rent"] or 0, 2)} if leading and leading["name"] else None),
            "pending_referencing_submissions": pending_referencing_submissions,
            "open_maintenance_requests": open_maintenance_requests,
            "open_message_threads": open_message_threads,
            "pending_document_uploads": pending_document_uploads,
            "new_submissions_total": new_submissions_total,
            # Phase 2 additions
            "vacant_units_list": [{"id": r["id"], "unit_ref": r["unit_ref"], "market_rent": r["market_rent"], "property_id": r["property_id"], "property_name": r["property_name"]} for r in vacant_units_list],
            "upcoming_move_ins": [{"tenancy_id": r["tenancy_id"], "move_in_date": r["move_in_date"], "tenant_name": r["main_tenant_name"], "property_name": r["property_name"], "property_id": r["property_id"], "unit_ref": r["unit_ref"], "unit_id": r["unit_id"]} for r in upcoming_move_ins],
            "upcoming_move_outs": [{"tenancy_id": r["tenancy_id"], "move_out_date": r["move_out_date"], "tenant_name": r["main_tenant_name"], "property_name": r["property_name"], "property_id": r["property_id"], "unit_ref": r["unit_ref"], "unit_id": r["unit_id"]} for r in upcoming_move_outs],
            "referencing_pipeline": referencing_pipeline,
            "tenancies_in_arrears_count": tenancies_in_arrears_count,
            "arrears_by_tenancy": arrears_by_tenancy_list,
            # E-signature stats
            "esign_pending": db.execute(
                "SELECT COUNT(*) AS cnt FROM esignature_requests WHERE status IN ('draft', 'sent', 'viewed', 'applicant_signed')"
            ).fetchone()["cnt"],
            "esign_completed_this_month": db.execute(
                "SELECT COUNT(*) AS cnt FROM esignature_requests WHERE status = 'completed' AND completed_at >= datetime('now', 'start of month')"
            ).fetchone()["cnt"],
            "esign_needing_action": db.execute(
                "SELECT COUNT(*) AS cnt FROM esignature_requests WHERE status = 'applicant_signed'"
            ).fetchone()["cnt"],
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 1a. RECENT ACTIVITY FEED
#     Synthetic union of recent events across submissions, referencing,
#     maintenance requests, and tenancy changes. No new table required.
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/dashboard/activity")
def api_dashboard_activity():
    db = get_dict_db()
    try:
        limit = int_param(request.args.get("limit", 30), 30, max_val=MAX_PAGE_SIZE)
        since = request.args.get("since")
        activity = []
        has_since = bool(since)

        # Build the WHERE/NULLIF based on whether 'since' param is provided
        def build_activity_query(base_select, base_from, date_col, extra_where="", extra_join=""):
            if has_since:
                where_clause = f"WHERE {date_col} IS NOT NULL AND {date_col} >= ?{extra_where}"
                params = [since]
            else:
                where_clause = f"WHERE {date_col} IS NOT NULL{extra_where}"
                params = []
            return f"{base_select} FROM {base_from} {extra_join} {where_clause}", params

        # 1. Referencing form submissions
        sql, params = build_activity_query(
            "SELECT id, 'referencing_submitted' AS event_type, submitted_at AS ts, "
            "COALESCE(NULLIF(first_name, ''), 'Applicant') || ' ' || COALESCE(NULLIF(last_name, ''), '') AS title, "
            "'Referencing form submitted' AS description, "
            "'referencing' AS category, 'referencing_form' AS link_type, id AS link_id, "
            "applicant_id AS related_id",
            "referencing_forms", "submitted_at"
        )
        rows = db.execute(sql, params).fetchall()
        for r in rows:
            activity.append(dict(r))

        # 2. Referencing reviews
        sql, params = build_activity_query(
            "SELECT id, 'referencing_reviewed' AS event_type, reviewed_at AS ts, "
            "COALESCE(NULLIF(first_name, ''), 'Applicant') || ' ' || COALESCE(NULLIF(last_name, ''), '') AS title, "
            "'Referencing reviewed by ' || COALESCE(reviewed_by, 'team') AS description, "
            "'referencing' AS category, 'referencing_form' AS link_type, id AS link_id, "
            "applicant_id AS related_id",
            "referencing_forms", "reviewed_at"
        )
        rows = db.execute(sql, params).fetchall()
        for r in rows:
            activity.append(dict(r))

        # 3. Maintenance requests
        sql, params = build_activity_query(
            "SELECT id, 'maintenance_created' AS event_type, created AS ts, "
            "COALESCE(title, 'Maintenance request') AS title, "
            "COALESCE(category, 'General') || ' - ' || COALESCE(reporter_name, 'Tenant') AS description, "
            "'maintenance' AS category, 'maintenance_request' AS link_type, id AS link_id, "
            "property_id AS related_id",
            "maintenance_requests", "created"
        )
        rows = db.execute(sql, params).fetchall()
        for r in rows:
            activity.append(dict(r))

        # 4. Tenancy changes (new tenancies created)
        sql, params = build_activity_query(
            "SELECT t.id, 'tenancy_created' AS event_type, t.created AS ts, "
            "COALESCE(t.main_tenant_name, 'Tenant') || ' - ' || "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS title, "
            "'New tenancy created' AS description, "
            "'tenancy' AS category, 'tenancy' AS link_type, t.id AS link_id, "
            "t.property_id AS related_id",
            "tenancies t", "t.created",
            extra_join="JOIN properties p ON t.property_id = p.id"
        )
        rows = db.execute(sql, params).fetchall()
        for r in rows:
            activity.append(dict(r))

        # 5. Message threads
        sql, params = build_activity_query(
            "SELECT id, 'message_created' AS event_type, created AS ts, "
            "COALESCE(title, 'Message thread') AS title, "
            "'New message thread opened' AS description, "
            "'message' AS category, 'message_thread' AS link_type, id AS link_id, "
            "property_id AS related_id",
            "message_threads", "created"
        )
        rows = db.execute(sql, params).fetchall()
        for r in rows:
            activity.append(dict(r))

        # Sort by timestamp descending and limit.
        # Timestamps arrive in two shapes across source tables: tz-aware ISO
        # ('2026-07-12T17:06:22+00:00') and naive ('2026-07-12 17:28:05').
        # Normalise both to 'YYYY-MM-DD HH:MM:SS' so same-day events from
        # different sources order truly chronologically, not by raw byte value.
        def _norm_ts(item):
            ts = item.get("ts") or ""
            ts = ts.replace("T", " ")
            if len(ts) >= 20 and ts[19] in "+-":
                ts = ts[:19]
            elif ts.endswith("Z"):
                ts = ts[:-1]
            return ts[:19]

        activity.sort(key=_norm_ts, reverse=True)
        activity = activity[:limit]

        return json_success(activity, total=len(activity), page=1, per_page=limit)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 1b. SUBMISSIONS INBOX
#     Unified feed of everything submitted via the tenant portal and the
#     referencing portal, so the team can find it all from one screen.
# ═══════════════════════════════════════════════
@banksia_os_bp.route("/submissions")
def api_submissions():
    db = get_dict_db()
    try:
        limit = int_param(request.args.get("limit", 60), 60, max_val=MAX_PAGE_SIZE)
        stype = (request.args.get("type") or "all").lower()      # all|referencing|maintenance|message
        only_new = str(request.args.get("new", "")).lower() in ("1", "true", "yes")
        items = []

        # 1. Referencing form submissions — only those the applicant actually submitted
        if stype in ("all", "referencing"):
            rows = db.execute(
                "SELECT id, applicant_id, status, submitted_at, reviewed_at, reviewed_by, "
                "first_name, last_name, email, mobile_phone, preferred_move_in_date "
                "FROM referencing_forms WHERE submitted_at IS NOT NULL "
                "ORDER BY submitted_at DESC"
            ).fetchall()
            for r in rows:
                st = (r["status"] or "").lower()
                needs = (r["reviewed_at"] is None) and st == "submitted"
                name = f"{(r['first_name'] or '').strip()} {(r['last_name'] or '').strip()}".strip()
                items.append({
                    "kind": "referencing",
                    "kind_label": "Referencing",
                    "id": r["id"],
                    "ref": "REF-%05d" % r["id"],
                    "title": name or "Applicant referencing",
                    "subtitle": r["email"] or r["mobile_phone"] or "",
                    "status": r["status"] or "submitted",
                    "needs_attention": needs,
                    "timestamp": r["submitted_at"],
                    "link_type": "referencing_form",
                    "link_id": r["id"],
                    "applicant_id": r["applicant_id"],
                })

        # 2. Maintenance requests raised from the tenant portal
        if stype in ("all", "maintenance"):
            rows = db.execute(
                "SELECT id, reference, tenancy_id, property_id, reporter_name, reporter_email, "
                "category, title, priority, status, created "
                "FROM maintenance_requests ORDER BY created DESC"
            ).fetchall()
            for r in rows:
                needs = (r["status"] or "open").lower() in ("open", "new", "")
                items.append({
                    "kind": "maintenance",
                    "kind_label": "Maintenance",
                    "id": r["id"],
                    "ref": r["reference"] or ("MR-%05d" % r["id"]),
                    "title": r["title"] or "Maintenance request",
                    "subtitle": "%s · %s" % ((r["reporter_name"] or "Tenant"), (r["category"] or "General")),
                    "status": r["status"] or "open",
                    "priority": r["priority"],
                    "needs_attention": needs,
                    "timestamp": r["created"],
                    "link_type": "maintenance_request",
                    "link_id": r["id"],
                    "property_id": r["property_id"],
                    "tenancy_id": r["tenancy_id"],
                })

        # 3. Portal message threads
        if stype in ("all", "message"):
            rows = db.execute(
                "SELECT t.id, t.title, t.status, t.property_id, t.tenancy_id, t.created, "
                "COUNT(m.id) AS msg_count, MAX(m.created) AS last_message "
                "FROM message_threads t "
                "LEFT JOIN messages m ON m.thread_id = t.id "
                "AND (m.is_deleted IS NULL OR m.is_deleted = 0) "
                "GROUP BY t.id ORDER BY COALESCE(MAX(m.created), t.created) DESC"
            ).fetchall()
            for r in rows:
                needs = (r["status"] or "open").lower() in ("open", "new", "")
                cnt = r["msg_count"] or 0
                items.append({
                    "kind": "message",
                    "kind_label": "Message",
                    "id": r["id"],
                    "ref": "MSG-%05d" % r["id"],
                    "title": r["title"] or "Message thread",
                    "subtitle": "%d message%s" % (cnt, "" if cnt == 1 else "s"),
                    "status": r["status"] or "open",
                    "needs_attention": needs,
                    "timestamp": r["last_message"] or r["created"],
                    "link_type": "message_thread",
                    "link_id": r["id"],
                    "property_id": r["property_id"],
                    "tenancy_id": r["tenancy_id"],
                })

        # 4. Documents uploaded by the applicant via the tenant / referencing portal
        if stype in ("all", "document"):
            rows = db.execute(
                "SELECT d.id, d.form_id, d.category, d.original_filename, d.file_size, "
                "d.mime_type, d.uploaded_at, d.is_verified, "
                "f.first_name, f.last_name, f.email "
                "FROM referencing_documents d "
                "LEFT JOIN referencing_forms f ON f.id = d.form_id "
                "WHERE LOWER(COALESCE(d.uploaded_by, '')) = 'applicant' "
                "ORDER BY d.uploaded_at DESC"
            ).fetchall()
            for r in rows:
                needs = not bool(r["is_verified"])
                name = f"{(r['first_name'] or '').strip()} {(r['last_name'] or '').strip()}".strip()
                cat = (r["category"] or "document").replace("_", " ")
                items.append({
                    "kind": "document",
                    "kind_label": "Document",
                    "id": r["id"],
                    "ref": "DOC-%05d" % r["id"],
                    "title": r["original_filename"] or "Uploaded document",
                    "subtitle": "%s · %s" % ((name or r["email"] or "Applicant"), cat),
                    "status": "verified" if r["is_verified"] else "awaiting review",
                    "needs_attention": needs,
                    "timestamp": r["uploaded_at"],
                    "link_type": "referencing_document",
                    "link_id": r["id"],
                    "form_id": r["form_id"],
                })

        # Unified feed — newest first (rows with no timestamp fall to the bottom)
        items.sort(key=lambda x: (x["timestamp"] or ""), reverse=True)
        if only_new:
            items = [i for i in items if i["needs_attention"]]
        counts = {
            "referencing": sum(1 for i in items if i["kind"] == "referencing"),
            "maintenance": sum(1 for i in items if i["kind"] == "maintenance"),
            "message": sum(1 for i in items if i["kind"] == "message"),
            "document": sum(1 for i in items if i["kind"] == "document"),
            "needs_attention": sum(1 for i in items if i["needs_attention"]),
        }
        items = items[:limit]
        return json_success({"items": items, "counts": counts}, total=len(items))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 3. MAINTENANCE OPERATIONS PORTAL
# ═══════════════════════════════════════════════

# The board's four groups (Norbert, 2026-08-07). A job's group IS its status, so
# moving it between groups is the only status change there is. The older values
# below are kept so anything still writing them is tolerated rather than rejected;
# the board files anything it does not recognise under "TO BE ARRANGED".
MAINT_STATUSES = [
    "NEW REPORT",
    "URGENT", "TO BE ARRANGED", "LIVE", "COMPLETED",
    "PENDING", "IN PROGRESS", "ON HOLD", "CANCELLED",
    "ACKNOWLEDGED", "WAITING INVOICE", "No Invoice Found", "Invoice Uploaded"
]

# Also the categories a contractor can be put in (Norbert, 2026-08-07): the same
# vocabulary on purpose, so "who does plumbing" and "this is a plumbing job" are
# the same word and a job can find its trade.
MAINT_TYPES = [
    "Plumbing", "Gas", "Heating", "Electrical", "Utilities", "Furniture", "NA", "Cleaning",
    "Structural", "Builder", "General Maintenance", "Appliances", "Refurbishment",
    "Certificate", "Orders", "Wall Repairs", "Painting", "Removal", "Locksmith",
    "Pest Control", "Small Repair", "Licenses", "Inspection", "Gardening"
]

# What a contractor can be categorised as. "NA" and "Orders" are job bookkeeping
# rather than trades, so nobody can be filed under them.
MAINT_TRADES = [t for t in MAINT_TYPES if t not in ("NA", "Orders")]

MAINT_PRIORITIES = ["Emergency", "Critical", "High", "Medium", "Low"]


@banksia_os_bp.route("/maintenance/jobs", methods=["GET", "POST"])
def api_maintenance_jobs():
    if request.method == "POST":
        return api_create_maintenance_job()
    db = get_dict_db()
    try:
        page = int_param(request.args.get("page"))
        per_page = int_param(request.args.get("per_page"), 50, max_val=MAX_PAGE_SIZE)
        search = (request.args.get("search") or "").strip()
        status_filter = request.args.get("status", "")
        type_filter = request.args.get("type", "")
        priority_filter = request.args.get("priority", "")
        contractor_filter = request.args.get("contractor", "")
        bill_ll_only = request.args.get("bill_ll", "") == "1"
        ll_not_informed = request.args.get("ll_uninformed", "") == "1"

        where = ["1=1"]
        params = []

        if status_filter:
            where.append("mj.status = ?")
            params.append(status_filter)
        if type_filter:
            where.append("mj.type = ?")
            params.append(type_filter)
        if priority_filter:
            where.append("mj.priority = ?")
            params.append(priority_filter)
        if contractor_filter:
            where.append("mj.contractor = ?")
            params.append(contractor_filter)
        if bill_ll_only:
            where.append("mj.bill_ll = 1")
        if ll_not_informed:
            where.append("mj.bill_ll = 1 AND mj.ll_informed = 0")
        if search:
            where.append("(mj.title LIKE ? OR mj.description LIKE ? OR mj.address LIKE ? OR mj.reference LIKE ? OR mj.contractor LIKE ? OR mj.type LIKE ? OR mj.reporter_name LIKE ? OR mj.team_notes LIKE ?)")
            s = f"%{search}%"
            params.extend([s, s, s, s, s, s, s, s])

        where_clause = " AND ".join(where)

        total = db.execute(
            f"SELECT COUNT(*) AS cnt FROM maintenance_jobs mj WHERE {where_clause}",
            params
        ).fetchone()["cnt"]

        offset = (page - 1) * per_page
        rows = db.execute(
            f"""SELECT mj.*, COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN '' ELSE p.name END, ''), p.address_line_1, p.ref, p.name) AS property_name,
                       p.management_type AS property_management_type
                FROM maintenance_jobs mj
                LEFT JOIN properties p ON mj.property_id = p.id
                WHERE {where_clause}
                ORDER BY
                    CASE mj.priority
                        WHEN 'Emergency' THEN 0 WHEN 'Critical' THEN 1
                        WHEN 'High' THEN 2 WHEN 'Medium' THEN 3 WHEN 'Low' THEN 4
                    END,
                    mj.created DESC
                LIMIT ? OFFSET ?""",
            params + [per_page, offset]
        ).fetchall()

        for r in rows:
            r["bill_ll"] = bool(r["bill_ll"])
            r["emergency"] = bool(r["emergency"])
            r["ll_informed"] = bool(r["ll_informed"])
            # Fetch order count
            o = db.execute(
                "SELECT COUNT(*) AS cnt FROM maintenance_orders WHERE job_id = ?",
                [r["id"]]
            ).fetchone()
            r["order_count"] = o["cnt"] if o else 0

        counts = {}
        for s in MAINT_STATUSES:
            c = db.execute("SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE status = ?", [s]).fetchone()
            counts[s] = c["cnt"] if c else 0

        return json_success(rows, total=total, page=page, per_page=per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


MAINT_REF_PREFIX = "REF"

# A job can only be Live once somebody could actually turn up and be paid
# (Norbert, 2026-08-07). Materials of zero is a real answer, so the test is
# "has a figure been entered", not "is it more than nothing" -- which is why new
# jobs leave the cost columns NULL rather than defaulting them to 0.
# "type" is in the list because it decides whether evidence is required at all,
# so a request that changes the type has to be judged on the new one.
MAINT_LIVE_REQUIRED = ("contractor", "labour_cost", "materials_cost", "photo_paths", "type")

# A certificate job produces the certificate, and that is filed against the
# property on the compliance board rather than photographed here (Norbert,
# 2026-08-08). Asking for a photo as well only produces a token one taken to
# clear the block, which is worse than not asking.
# "compliance" is here because the compliance board has been writing that type
# since long before this rule; the job is the same thing by another name.
MAINT_NO_EVIDENCE_TYPES = {"certificate", "compliance"}


def _next_maintenance_reference(db):
    """Next free REF-#### reference.

    Was COUNT(*) + 1, which is wrong the moment a job is deleted or the numbers
    arrive from an import with gaps -- the Monday import left references up to
    0401 against 197 rows, so counting would have handed out numbers already in
    use. Take the highest number actually present and step past it, then check.

    Old MJ- references are read alongside REF- ones so renaming the prefix cannot
    reissue a number somebody already has on an invoice.
    """
    rows = db.execute(
        "SELECT reference FROM maintenance_jobs "
        "WHERE reference LIKE 'REF-%' OR reference LIKE 'MJ-%'"
    ).fetchall()
    highest = 0
    for r in rows:
        ref = str(dict(r).get("reference") or "")
        tail = ref.split("-", 1)[1] if "-" in ref else ""
        if tail.isdigit():
            highest = max(highest, int(tail))
    n = highest + 1
    # Belt and braces: the column has no unique constraint, so confirm rather
    # than assume the arithmetic is enough.
    while db.execute("SELECT 1 AS hit FROM maintenance_jobs WHERE reference = ?",
                     ["%s-%s" % (MAINT_REF_PREFIX, str(n).zfill(4))]).fetchone():
        n += 1
    return "%s-%s" % (MAINT_REF_PREFIX, str(n).zfill(4))


def _is_management_fee(db, property_id):
    """Management-fee properties are billed for their own repairs.

    On a fixed rent we carry the cost; on a management fee the landlord does, so
    the tick is a property of the agreement rather than a decision somebody makes
    per job. It is still a normal checkbox afterwards -- this only sets the
    starting position, it does not lock it.
    """
    if not property_id:
        return False
    row = db.execute("SELECT management_type FROM properties WHERE id = ?", [property_id]).fetchone()
    return "management fee" in str(dict(row or {}).get("management_type") or "").strip().lower()


def _needs_evidence(job):
    """Whether this job has to show its work before it can go Live."""
    return str(job.get("type") or "").strip().lower() not in MAINT_NO_EVIDENCE_TYPES


# Somebody did the work and somebody has to be paid for it. A job closed with
# nobody named cannot be matched to an invoice later (Norbert, 2026-08-08).
MAINT_COMPLETED_REQUIRED = ("contractor",)


def _completed_blockers(job):
    """What is stopping this job being closed."""
    if not str(job.get("contractor") or "").strip():
        return ["a contractor"]
    return []


def _live_blockers(job):
    """What is stopping this job going Live, in words a person can act on."""
    missing = []
    if not str(job.get("contractor") or "").strip():
        missing.append("a contractor")
    if job.get("labour_cost") is None:
        missing.append("the labour")
    if job.get("materials_cost") is None:
        missing.append("the materials")
    if _needs_evidence(job) and not [p for p in str(job.get("photo_paths") or "").split(",") if p.strip()]:
        missing.append("evidence")
    return missing


def _london_today():
    """Today where the business is, not where the server thinks it is.

    SQLite's date('now') is UTC, so a job moved to Live at half past midnight in
    summer would record yesterday.
    """
    from datetime import datetime
    try:
        from zoneinfo import ZoneInfo
        return datetime.now(ZoneInfo("Europe/London")).date().isoformat()
    except Exception:
        return datetime.now().date().isoformat()


def _ensure_maintenance_cert_key(db):
    """Which certificate a certificate job is for.

    The type only says "Certificate"; the price depends on whether it is an EICR
    or a gas safety check, so the board has to hold the distinction.
    """
    try:
        db.execute("ALTER TABLE maintenance_jobs ADD COLUMN cert_key TEXT DEFAULT ''")
        db.commit()
    except Exception:
        pass  # already present


def _standard_quote(db, cert_key):
    """The agreed price for a certificate, or None where there is no agreed price.

    None matters: an FRA has no standard figure ("depends on the location"), and
    billing a landlord zero for one is worse than falling back to the labour rule.
    """
    key = str(cert_key or "").strip()
    if not key:
        return None
    try:
        row = db.execute("SELECT amount FROM compliance_quotes WHERE cert_key = ?",
                         (key,)).fetchone()
    except Exception:
        return None
    raw = str(dict(row or {}).get("amount") or "").strip()
    if not raw:
        return None
    try:
        return round(float(raw), 2)
    except (TypeError, ValueError):
        return None


def _ensure_maintenance_cost_ll(db):
    """What we charge the landlord, kept separate from what the contractor charges
    us. Added on demand so the board works on a database that predates it."""
    try:
        db.execute("ALTER TABLE maintenance_jobs ADD COLUMN cost_ll REAL DEFAULT 0")
        db.commit()
    except Exception:
        pass  # already present


def _ensure_compliance_jobs(db):
    """A certificate job: the quote we asked a contractor for, and what became of it.

    Separate from maintenance_jobs because it exists BEFORE there is a work order --
    from the moment we ask for a price to the moment it is booked. The work order is
    the outcome, not the record of the conversation.
    """
    db.execute("""
        CREATE TABLE IF NOT EXISTS compliance_jobs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            compliance_id INTEGER NOT NULL,
            cert_key TEXT NOT NULL,
            contractor_name TEXT DEFAULT '',
            contractor_group TEXT DEFAULT '',
            -- quote_requested -> quoted -> booked -> cancelled
            status TEXT NOT NULL DEFAULT 'quote_requested',
            contractor_quote REAL,
            scheduled_date TEXT DEFAULT '',
            cost_ll REAL,
            reference TEXT DEFAULT '',
            maintenance_job_id INTEGER,
            expiry_date TEXT DEFAULT '',
            note TEXT DEFAULT '',
            created_by TEXT DEFAULT '',
            requested_at TEXT,
            quoted_at TEXT,
            booked_at TEXT,
            created TEXT DEFAULT (datetime('now')),
            modified TEXT DEFAULT (datetime('now'))
        )
    """)
    db.execute("CREATE INDEX IF NOT EXISTS idx_compliance_jobs_row "
               "ON compliance_jobs (compliance_id, cert_key)")
    db.commit()


def api_create_maintenance_job():
    data = request.get_json(force=True, silent=True) or {}
    required = ["title"]
    for f in required:
        if not data.get(f):
            return json_error(f"'{f}' is required")
    if data.get("contractor"):
        known = _known_contractor(data.get("contractor"))
        if known is None:
            return json_error(
                "%s is not on the Contractors page — add them there first, then pick them here."
                % data.get("contractor"), 422)
        data = dict(data)
        data["contractor"] = known

    db = get_dict_db()
    try:
        _ensure_maintenance_cost_ll(db)
        _ensure_maintenance_cert_key(db)
        reference = _next_maintenance_reference(db)

        cur = db.execute(
            """INSERT INTO maintenance_jobs
               (reference, title, description, type, cert_key, priority, status, location,
                property_id, address, contractor, labour_cost, materials_cost,
                bill_ll, emergency, reporter_name, reporter_email, team_notes, source,
                cost_ll)
               VALUES (?, ?, ?, ?, ?, ?, 'TO BE ARRANGED', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            [
                reference,
                data.get("title"),
                data.get("description", ""),
                data.get("type"),
                str(data.get("cert_key") or ""),
                data.get("priority", "Medium"),
                data.get("location"),
                data.get("property_id"),
                data.get("address"),
                data.get("contractor"),
                None if data.get("labour_cost") is None else float(data.get("labour_cost")),
                None if data.get("materials_cost") is None else float(data.get("materials_cost")),
                1 if data.get("bill_ll") else 0,
                1 if data.get("emergency") else 0,
                data.get("reporter_name", ""),
                data.get("reporter_email", ""),
                data.get("team_notes", ""),
                data.get("source", "board"),
                float(data.get("cost_ll") or 0),
            ]
        )
        db.commit()
        job_id = cur.lastrowid
        if data.get("bill_ll") is None and _is_management_fee(db, data.get("property_id")):
            db.execute("UPDATE maintenance_jobs SET bill_ll = 1 WHERE id = ?", [job_id])
            db.commit()

        wanted = str(data.get("status") or "").strip()
        if wanted.upper() == "COMPLETED":
            missing = _completed_blockers(dict(data))
            if missing:
                db.execute("DELETE FROM maintenance_jobs WHERE id = ?", [job_id])
                db.commit()
                return json_error(
                    "A job cannot be raised straight into Completed — it still needs %s."
                    % _join_words(missing), 422)
        if wanted.upper() == "LIVE":
            missing = _live_blockers(dict(data))
            if missing:
                db.execute("DELETE FROM maintenance_jobs WHERE id = ?", [job_id])
                db.commit()
                return json_error(
                    "A job cannot start Live — it still needs %s. Raise it in another group and "
                    "move it across once it has them." % _join_words(missing), 422)
        if wanted and wanted in MAINT_STATUSES:
            db.execute("UPDATE maintenance_jobs SET status = ? WHERE id = ?", [wanted, job_id])
            if wanted.upper() == "COMPLETED":
                db.execute("UPDATE maintenance_jobs SET completed_date = datetime('now') "
                           "WHERE id = ? AND completed_date IS NULL", [job_id])
            db.commit()

        # A Cost LL supplied at creation is a decision (the compliance booking
        # sends the standard quote), so it must not be recalculated to labour+15%.
        _ensure_cost_ll_override(db)
        if data.get("cost_ll") is not None:
            db.execute("UPDATE maintenance_jobs SET cost_ll_override = 1 WHERE id = ?", [job_id])
            db.commit()
        else:
            _sync_cost_ll(db, job_id, {})
        job = db.execute(
            """SELECT mj.*, COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN '' ELSE p.name END, ''), p.address_line_1, p.ref, p.name) AS property_name
               FROM maintenance_jobs mj
               LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.id = ?""",
            [job_id]
        ).fetchone()
        return json_success(dict(job)), 201
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>", methods=["GET", "PATCH"])
def api_maintenance_job(job_id):
    db = get_dict_db()
    try:
        if request.method == "GET":
            job = db.execute(
                """SELECT mj.*, COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN '' ELSE p.name END, ''), p.address_line_1, p.ref, p.name) AS property_name
                   FROM maintenance_jobs mj
                   LEFT JOIN properties p ON mj.property_id = p.id
                   WHERE mj.id = ?""",
                [job_id]
            ).fetchone()
            if not job:
                return json_error("Job not found", 404)
            # Get orders for this job
            orders = db.execute(
                "SELECT * FROM maintenance_orders WHERE job_id = ? ORDER BY created DESC",
                [job_id]
            ).fetchall()
            # Get LL communications
            ll_comms = db.execute(
                "SELECT * FROM ll_communications WHERE job_id = ? ORDER BY sent_at DESC",
                [job_id]
            ).fetchall()
            result = dict(job)
            result["orders"] = [dict(o) for o in orders]
            result["ll_comms"] = [dict(c) for c in ll_comms]
            result["bill_ll"] = bool(result["bill_ll"])
            result["emergency"] = bool(result["emergency"])
            result["ll_informed"] = bool(result["ll_informed"])
            return json_success(result)

        # PATCH
        # The column has to exist before the UPDATE below can name it. Ensuring it
        # inside _sync_cost_ll was too late -- that runs after the write.
        _ensure_maintenance_cert_key(db)
        data = request.get_json(force=True, silent=True) or {}
        allowed = [
            "title", "description", "type", "priority", "status", "location",
            "address", "contractor", "labour_cost", "materials_cost",
            "bill_ll", "ll_informed", "ll_informed_via", "ll_notes",
            "emergency", "reporter_name", "reporter_email", "photo_paths",
            "invoice_paths", "team_notes", "start_date", "completed_date",
            "property_id", "unit", "cost_ll", "cost_ll_override", "cert_key"
        ]
        if "contractor" in data and str(data.get("contractor") or "").strip():
            known = _known_contractor(data.get("contractor"))
            if known is None:
                return json_error(
                    "%s is not on the Contractors page — add them there first, then pick them here."
                    % data.get("contractor"), 422)
            data = dict(data)
            data["contractor"] = known
            # The lookup closes the shared connection on its way past; take a
            # fresh one rather than carrying on with a dead handle.
            db = get_dict_db()

        # Guard before writing, not after: a rejected move must leave the job
        # exactly as it was, not half-applied.
        if str(data.get("status") or "").strip().upper() == "COMPLETED":
            current = db.execute(
                "SELECT contractor FROM maintenance_jobs WHERE id = ?", [job_id]
            ).fetchone()
            merged = dict(current or {})
            for f in MAINT_COMPLETED_REQUIRED:
                if f in data:
                    merged[f] = data[f]
            missing = _completed_blockers(merged)
            if missing:
                return json_error(
                    "This job cannot be marked Completed — it still needs %s."
                    % _join_words(missing), 422)

        if str(data.get("status") or "").strip().upper() == "LIVE":
            current = db.execute(
                "SELECT contractor, labour_cost, materials_cost, photo_paths, type "
                "FROM maintenance_jobs WHERE id = ?",
                [job_id]
            ).fetchone()
            merged = dict(current or {})
            for f in MAINT_LIVE_REQUIRED:
                if f in data:
                    merged[f] = data[f]
            missing = _live_blockers(merged)
            if missing:
                return json_error(
                    "This job cannot go Live yet — it still needs %s." % _join_words(missing), 422)

        # Pointing a job at a management-fee property ticks Bill LL, unless the
        # same request is explicitly saying otherwise.
        if "property_id" in data and "bill_ll" not in data and _is_management_fee(db, data.get("property_id")):
            data = dict(data)
            data["bill_ll"] = True

        updates = []
        params = []
        for field in allowed:
            if field in data:
                val = data[field]
                if field in ("bill_ll", "emergency", "ll_informed", "cost_ll_override"):
                    val = 1 if val else 0
                updates.append(f"{field} = ?")
                params.append(val)
        if not updates:
            return json_error("No valid fields to update")
        updates.append("modified = datetime('now')")
        params.append(job_id)
        db.execute(
            f"UPDATE maintenance_jobs SET {', '.join(updates)} WHERE id = ?",
            params
        )
        db.commit()

        # Cost LL is derived from the labour, so it has to be recalculated after
        # anything that could move it -- not only when it is sent explicitly.
        _sync_cost_ll(db, job_id, data)

        # Going Live is the day the work starts, so the board should not have to be
        # told twice (Norbert, 2026-08-08). Only filled when empty: a date already
        # agreed with a contractor beats today's date.
        if str(data.get("status") or "").strip().upper() == "LIVE":
            db.execute(
                "UPDATE maintenance_jobs SET start_date = ? WHERE id = ? "
                "AND (start_date IS NULL OR TRIM(start_date) = '')",
                [_london_today(), job_id]
            )
            db.commit()

        # If status changed to COMPLETED, set completed_date
        if data.get("status") == "COMPLETED":
            db.execute(
                "UPDATE maintenance_jobs SET completed_date = datetime('now') WHERE id = ? AND completed_date IS NULL",
                [job_id]
            )
            db.commit()

        job = db.execute("SELECT * FROM maintenance_jobs WHERE id = ?", [job_id]).fetchone()
        # Mark for push-back to Monday (async sync will pick it up)
        try:
            db.execute(
                "UPDATE maintenance_jobs SET sync_pending = 1 WHERE id = ?",
                [job_id]
            )
            db.commit()
        except Exception:
            db.rollback()
        return json_success(dict(job))
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Promote portal maintenance request to tracked job ──
@banksia_os_bp.route("/maintenance/promote-from-portal", methods=["POST"])
def api_promote_portal_request():
    """Copy a maintenance_requests row into maintenance_jobs so it becomes
    visible on the team dashboard and can receive orders / LL comms / contractors."""
    data = request.get_json(force=True, silent=True) or {}
    req_id = data.get("request_id")
    if not req_id:
        return json_error("request_id is required")

    db = get_dict_db()
    try:
        req = db.execute(
            "SELECT * FROM maintenance_requests WHERE id = ?", [req_id]
        ).fetchone()
        if not req:
            return json_error("Portal request not found", 404)

        # Build reference
        ref_prefix = "MJ"
        count = db.execute("SELECT COUNT(*) AS cnt FROM maintenance_jobs").fetchone()["cnt"]
        reference = f"{ref_prefix}-{str(count + 1).zfill(4)}"

        cur = db.execute(
            """INSERT INTO maintenance_jobs
               (reference, title, description, type, priority, status, location,
                property_id, address, reporter_name, reporter_email, source, team_notes)
               VALUES (?, ?, ?, ?, ?, 'PENDING', ?, ?, ?, ?, ?, 'portal', ?)""",
            [
                reference,
                req.get("title") or req.get("category", "Maintenance request"),
                req.get("description", ""),
                req.get("category"),
                req.get("priority", "Medium"),
                req.get("location"),
                data.get("property_id") or req.get("property_id"),
                data.get("address", ""),
                req.get("reporter_name", ""),
                req.get("reporter_email", ""),
                data.get("notes", ""),
            ]
        )
        db.commit()
        job_id = cur.lastrowid

        # Update original request status to 'promoted'
        db.execute(
            "UPDATE maintenance_requests SET status = 'promoted' WHERE id = ?",
            [req_id]
        )
        db.commit()

        job = db.execute(
            "SELECT * FROM maintenance_jobs WHERE id = ?", [job_id]
        ).fetchone()
        return json_success(dict(job)), 201
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Maintenance Orders ──

@banksia_os_bp.route("/maintenance/orders", methods=["GET", "POST"])
def api_maintenance_orders():
    db = get_dict_db()
    try:
        if request.method == "POST":
            data = request.get_json(force=True, silent=True) or {}
            if not data.get("job_id"):
                return json_error("job_id is required")
            cur = db.execute(
                """INSERT INTO maintenance_orders
                   (job_id, item_name, supplier, order_ref, cost, status,
                    tracking_url, estimated_delivery, notes)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                [
                    data["job_id"], data.get("item_name"),
                    data.get("supplier"), data.get("order_ref"),
                    float(data.get("cost", 0)),
                    data.get("status", "ordered"),
                    data.get("tracking_url"),
                    data.get("estimated_delivery"),
                    data.get("notes", ""),
                ]
            )
            db.commit()
            return json_success({"id": cur.lastrowid}), 201

        # GET — list orders, optionally filtered by job_id
        job_id = request.args.get("job_id")
        if job_id:
            orders = db.execute(
                "SELECT * FROM maintenance_orders WHERE job_id = ? ORDER BY created DESC",
                [job_id]
            ).fetchall()
        else:
            orders = db.execute(
                """SELECT mo.*, mj.title AS job_title
                   FROM maintenance_orders mo
                   JOIN maintenance_jobs mj ON mo.job_id = mj.id
                   ORDER BY mo.created DESC LIMIT 100"""
            ).fetchall()
        return json_success([dict(o) for o in orders])
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/orders/<int:order_id>", methods=["PATCH"])
def api_maintenance_order(order_id):
    data = request.get_json(force=True, silent=True) or {}
    allowed = ["item_name", "supplier", "order_ref", "cost", "status",
               "tracking_url", "estimated_delivery", "delivered_at",
               "received_by", "notes"]
    updates = []
    params = []
    for field in allowed:
        if field in data:
            val = data[field]
            if field == "cost":
                val = float(val)
            updates.append(f"{field} = ?")
            params.append(val)
    if not updates:
        return json_error("No valid fields")
    updates.append("modified = datetime('now')")
    params.append(order_id)
    db = get_dict_db()
    try:
        db.execute(
            f"UPDATE maintenance_orders SET {', '.join(updates)} WHERE id = ?",
            params
        )
        db.commit()
        order = db.execute("SELECT * FROM maintenance_orders WHERE id = ?", [order_id]).fetchone()
        return json_success(dict(order))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── LL Communications ──

@banksia_os_bp.route("/maintenance/ll-comms", methods=["GET", "POST"])
def api_ll_comms():
    db = get_dict_db()
    try:
        if request.method == "POST":
            data = request.get_json(force=True, silent=True) or {}
            if not data.get("job_id"):
                return json_error("job_id is required")
            cur = db.execute(
                """INSERT INTO ll_communications
                   (job_id, contact_method, contact_ref, summary, ll_response, sent_at)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                [
                    data["job_id"], data.get("contact_method"),
                    data.get("contact_ref"), data.get("summary", ""),
                    data.get("ll_response", ""), data.get("sent_at"),
                ]
            )
            # Mark job as ll_informed
            db.execute(
                "UPDATE maintenance_jobs SET ll_informed = 1, ll_informed_via = ? WHERE id = ?",
                [data.get("contact_method"), data["job_id"]]
            )
            db.commit()
            return json_success({"id": cur.lastrowid}), 201

        job_id = request.args.get("job_id")
        if not job_id:
            return json_error("job_id is required")
        comms = db.execute(
            "SELECT * FROM ll_communications WHERE job_id = ? ORDER BY sent_at DESC",
            [job_id]
        ).fetchall()
        return json_success([dict(c) for c in comms])
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Maintenance lookup data ──

@banksia_os_bp.route("/maintenance/lookup")
def api_maintenance_lookup():
    return json_success({
        "statuses": MAINT_STATUSES,
        "types": MAINT_TYPES,
        "priorities": MAINT_PRIORITIES,
    })


# ── Monday.com sync endpoint ──

def _monday_graphql(mtok, query):
    """Execute a Monday.com GraphQL query and return the parsed result."""
    import urllib.request
    req = urllib.request.Request(
        "https://api.monday.com/v2",
        data=json.dumps({"query": query}).encode(),
        headers={"Authorization": mtok, "Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, timeout=30) as r:
        return json.loads(r.read())


def _parse_monday_cols(column_values):
    """Build a flat dict of {column_id: text} from Monday column_values list."""
    cols = {}
    for cv in column_values:
        cols[cv["id"]] = cv.get("text") or ""
    return cols


def _safe_status(val):
    if val not in MAINT_STATUSES:
        return "PENDING"
    return val


def _safe_priority(val):
    if val not in MAINT_PRIORITIES:
        return "Medium"
    return val


def _parse_photo_paths(cols):
    """Extract photo evidence URLs (comma-separated)."""
    val = cols.get("file_mm0v10xk", "")
    if not val:
        return ""
    # Multiple URLs are comma-separated in the text field
    return val


def _parse_invoice_paths(cols):
    """Extract contractor invoice URLs."""
    val = cols.get("file_mm0pryh", "")
    return val


@banksia_os_bp.route("/maintenance/sync-from-monday", methods=["POST"])
def api_sync_from_monday():
    """Pull jobs from Monday.com Maintenance Reports board into local DB.

    Performs a full re-sync:
      - Inserts new items (monday_id not seen before)
      - Updates existing items whose data has changed on Monday
      - Handles pagination (cursor-based) for boards with 200+ items
      - Maps the full set of columns to DB fields
    """
    mtok = None
    try:
        mtok = open("/root/.hermes/secrets/monday_token.txt").read().strip()
    except Exception:
        pass
    if not mtok:
        return json_error("Monday token not found")

    db = get_dict_db()
    try:
        # ── Fetch ALL items with cursor-based pagination ──
        all_items = []
        cursor = None
        page = 0

        while True:
            page_ql = f"items_page(limit:200" + (f',cursor:"{cursor}"' if cursor else "") + ")"
            q = (
                "{ boards(ids: [18401159622]) { id name "
                + page_ql
                + """ { cursor items {
                        id name column_values { id text value }
                    } } } }"""
            )
            data = _monday_graphql(mtok, q)
            page_data = (
                data.get("data", {})
                .get("boards", [{}])[0]
                .get("items_page", {})
            )
            items = page_data.get("items", [])
            cursor = page_data.get("cursor")
            all_items.extend(items)
            page += 1

            if not cursor or len(items) < 200:
                break

        # ── Process every item (INSERT or UPDATE) ──
        inserted = 0
        updated = 0
        unchanged = 0
        pending = 0  # writes accumulated since last commit — batched to keep
                     # the write lock from being held for the whole loop

        for item in all_items:
            cols = _parse_monday_cols(item.get("column_values", []))
            monday_id = item["id"]
            title = item.get("name", "")

            # Map Monday column IDs → DB fields
            status = _safe_status(cols.get("status", "PENDING"))
            priority = _safe_priority(cols.get("color_mm0p8qna", "Medium"))
            maint_type = cols.get("color_mm0vfxmq", "")
            address = (
                cols.get("short_text041ydfbp", "")
                or cols.get("long_text_mm50g0j6", "")
                or cols.get("board_relation_mm0p7cv6", "")
            )
            contractor = cols.get("color_mm0p4947", "")
            location = cols.get("dropdown_mm0p6nzm", "")

            # Labour & materials costs
            labour_raw = cols.get("numeric_mm0pndmj", "") or "0"
            materials_raw = cols.get("numeric_mm0p7jdn", "") or "0"
            try:
                labour_cost = float(labour_raw.replace("£", "").replace(",", "").strip())
            except (ValueError, AttributeError):
                labour_cost = 0.0
            try:
                materials_cost = float(materials_raw.replace("£", "").replace(",", "").strip())
            except (ValueError, AttributeError):
                materials_cost = 0.0

            # Boolean toggles
            bill_ll = 1 if cols.get("boolean_mm0phkaq", "") == "checked" else 0
            emergency = 1 if cols.get("boolean2hbqq7ey", "") == "checked" else 0

            # Reporter info
            reporter_name = cols.get("short_textcvckh2h3", "")
            reporter_email = cols.get("emailzit7svgb", "")

            # File paths (photo evidence + contractor invoices)
            photo_paths = _parse_photo_paths(cols)
            invoice_paths = _parse_invoice_paths(cols)

            # Check if this item already exists in local DB
            existing = db.execute(
                "SELECT id, status, priority, type, address, contractor, "
                "labour_cost, materials_cost, bill_ll, emergency, "
                "reporter_name, reporter_email, photo_paths, invoice_paths, "
                "location, description, team_notes "
                "FROM maintenance_jobs WHERE monday_id = ?",
                [monday_id],
            ).fetchone()

            if existing:
                # ── UPDATE existing row ──
                # Compare key fields to decide if an update is needed
                changed = False
                updates = {}
                compare_map = {
                    "title": title,
                    "status": status,
                    "priority": priority,
                    "type": maint_type,
                    "address": address,
                    "contractor": contractor,
                    "location": location,
                    "labour_cost": labour_cost,
                    "materials_cost": materials_cost,
                    "bill_ll": bill_ll,
                    "emergency": emergency,
                    "reporter_name": reporter_name,
                    "reporter_email": reporter_email,
                    "photo_paths": photo_paths,
                    "invoice_paths": invoice_paths,
                }
                for field, new_val in compare_map.items():
                    old_val = existing[field]
                    if old_val is None:
                        old_val = ""
                    # Normalise types for comparison
                    if isinstance(old_val, float) or isinstance(new_val, float):
                        if abs(float(old_val or 0) - float(new_val or 0)) > 0.001:
                            updates[field] = new_val
                            changed = True
                    elif str(old_val).strip() != str(new_val).strip():
                        updates[field] = new_val
                        changed = True

                if changed:
                    updates["modified"] = "datetime('now')"
                    set_clause = ", ".join(f"{k} = ?" for k in updates)
                    values = list(updates.values())
                    values.append(existing["id"])
                    db.execute(
                        f"UPDATE maintenance_jobs SET {set_clause} WHERE id = ?",
                        values,
                    )
                    updated += 1
                    pending += 1
                else:
                    unchanged += 1
            else:
                # ── INSERT new row ──
                db.execute(
                    """INSERT INTO maintenance_jobs
                       (monday_id, title, status, priority, type, address,
                        contractor, location, labour_cost, materials_cost,
                        bill_ll, emergency, reporter_name, reporter_email,
                        photo_paths, invoice_paths, source)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'monday')""",
                    [
                        monday_id,
                        title,
                        status,
                        priority,
                        maint_type,
                        address,
                        contractor,
                        location,
                        labour_cost,
                        materials_cost,
                        bill_ll,
                        emergency,
                        reporter_name,
                        reporter_email,
                        photo_paths,
                        invoice_paths,
                    ],
                )
                inserted += 1
                pending += 1

            # Commit in batches so the write lock is released periodically,
            # letting concurrent user saves through instead of waiting on one
            # giant transaction spanning every item.
            if pending >= 50:
                db.commit()
                pending = 0

        db.commit()
        return json_success(
            {
                "inserted": inserted,
                "updated": updated,
                "unchanged": unchanged,
                "total_on_monday": len(all_items),
            }
        )
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2. PROPERTIES
# ═══════════════════════════════════════════════

def _ensure_landlord_link(db, data):
    """Guarantee every property is linked to a landlord row in property_owners.

    Resolves ``data['property_owner_id']`` to a valid ``property_owners.id`` and
    keeps ``property_owner_name`` in sync. Auto-creates the landlord when needed
    so a property can never exist without a linked landlord. Mutates + returns data.
    """
    raw_id = data.get("property_owner_id")
    oid = str(raw_id).strip() if raw_id not in (None, "") else ""
    oname = (data.get("property_owner_name") or "").strip()

    # 1. An id was supplied
    if oid:
        row = db.execute(
            "SELECT id, name FROM property_owners WHERE id = CAST(? AS INTEGER)", (oid,)
        ).fetchone()
        if row:
            data["property_owner_id"] = str(row["id"])
            if not oname:
                data["property_owner_name"] = row["name"]
            return data
        # id given but no matching landlord row — create it, preserving the id
        nm = oname or f"Landlord {oid}"
        db.execute(
            "INSERT INTO property_owners (id, name, status) VALUES (CAST(? AS INTEGER), ?, 'active')",
            (oid, nm),
        )
        data["property_owner_id"] = oid
        data["property_owner_name"] = nm
        return data

    # 2. Only a name was supplied — match existing (case-insensitive) or create
    if oname:
        row = db.execute(
            "SELECT id, name FROM property_owners WHERE LOWER(name) = LOWER(?)", (oname,)
        ).fetchone()
        if row:
            data["property_owner_id"] = str(row["id"])
            data["property_owner_name"] = row["name"]
            return data
        cur = db.execute(
            "INSERT INTO property_owners (name, status) VALUES (?, 'active')", (oname,)
        )
        data["property_owner_id"] = str(cur.lastrowid)
        data["property_owner_name"] = oname
        return data

    # 3. Nothing to link — caller should have validated already
    raise ValueError("A landlord is required — every property must be linked to a landlord.")


@banksia_os_bp.route("/properties", methods=["GET", "POST"])
def api_properties():
    if request.method == "POST":
        return api_create_property()
    all_mode = request.args.get("all", "").strip().lower() in ("1", "true", "yes")

    page = int_param(request.args.get("page")) if not all_mode else 1
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE) if not all_mode else 9999
    search = request.args.get("search", "").strip()
    occ_filter = request.args.get("occupancy", "").strip().lower()
    rent_min = float_param(request.args.get("rent_min"))
    rent_max = float_param(request.args.get("rent_max"))
    sort_field = request.args.get("sort_field", "").strip()
    sort_direction = request.args.get("sort_direction", "asc").strip().lower()

    # Category scoping (Norbert, 2026-08-03). The properties page is entered through
    # three cards, and they deliberately OVERLAP: a managed property is still an
    # active property, and an inactive one can be a managed one too. So these are
    # three ways of asking about the same 45 rows, not a partition of them.
    #   active     — trading properties (is_active = 1, not archived)
    #   management — let on a management fee rather than fixed rent
    #   inactive   — stood down or archived, whatever the fee basis
    category = request.args.get("category", "").strip().lower()
    not_archived = "(status IS NULL OR status <> 'archived')"
    if category == "inactive":
        # Archived properties are NOT inactive properties (Norbert, 2026-08-03): he saw
        # 59 Claylands Road and 95 Wheat Sheaf Close here and said both should be
        # archived — they already were. Archived stock lives on the Archive page and
        # nowhere else, so this category is "stood down but still on the books".
        scope_clause = f"COALESCE(is_active, 1) = 0 AND {not_archived}"
    elif category == "all":
        # "I still want the option to see all the properties in 1" (Norbert). Archived
        # rows stay out — the Archive page is where those live.
        scope_clause = not_archived
    elif category == "management":
        scope_clause = f"management_type = 'Management Fee' AND {not_archived}"
    elif category == "active":
        scope_clause = f"COALESCE(is_active, 1) = 1 AND {not_archived}"
    else:
        # No category — the whole list, archived excluded (they live in Archive).
        scope_clause = not_archived

    base_where = scope_clause
    base_params = []

    if search:
        search_clause, search_params = build_search_clause(
            ["name", "ref", "address_line_1", "city", "postcode"], search
        )
        base_where = f"({search_clause}) AND {scope_clause}"
        base_params = search_params

    # Occupancy filter
    # ── Safety: apply occupancy filter as inner WHERE (not outer) to avoid
    #    parse errors from concurrent write transactions in WAL mode ──
    occ_inner = ""
    if occ_filter == "vacant":
        occ_inner = " AND 0 = (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id AND u2.unit_vacant = 0)"
    elif occ_filter == "full":
        occ_inner = " AND (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id AND u2.unit_vacant = 0) = (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id)"
    elif occ_filter == "partial":
        occ_inner = " AND (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id AND u2.unit_vacant = 0) > 0 AND (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id AND u2.unit_vacant = 0) < (SELECT COUNT(*) FROM units u2 WHERE u2.property_id = properties.id)"

    # Rent range filter — apply as inner WHERE too
    rent_inner = ""
    if rent_min is not None:
        rent_inner = f" AND (SELECT COALESCE(SUM(t2.rent_amount), 0) FROM tenancies t2 WHERE t2.property_id = properties.id AND t2.status IN ('Current','current','Periodic','periodic','Active','active')) >= {rent_min}"
        if rent_max is not None and rent_max < float('inf'):
            rent_inner += f" AND (SELECT COALESCE(SUM(t2.rent_amount), 0) FROM tenancies t2 WHERE t2.property_id = properties.id AND t2.status IN ('Current','current','Periodic','periodic','Active','active')) <= {rent_max}"

    # Management type filter — safe-string column comparison, no injection
    mgmt_filter = request.args.get("management_type", "").strip()
    mgmt_inner = ""
    if mgmt_filter and mgmt_filter in ("Fixed Rent", "Management Fee"):
        mgmt_inner = f" AND management_type = '{mgmt_filter}'"

    inner_query = f"SELECT *, " \
        f"(SELECT COUNT(*) FROM units u WHERE u.property_id = properties.id) AS actual_units, " \
        f"(SELECT COUNT(*) FROM units u WHERE u.property_id = properties.id AND u.unit_vacant = 0) AS occupied_units, " \
        f"(SELECT COALESCE(NULLIF(ref, ''), NULLIF(address_line_1, ''), name)) AS sort_name, " \
        f"COALESCE((SELECT SUM(t.rent_amount) FROM tenancies t WHERE t.property_id = properties.id AND t.status IN ('Current','current','Periodic','periodic','Active','active')), 0) AS monthly_rent, " \
        f"CASE WHEN (SELECT COUNT(*) FROM units u WHERE u.property_id = properties.id AND u.unit_vacant = 0) > 0 THEN 'Active' ELSE 'Vacant' END AS property_status, " \
        f"(SELECT po.name FROM property_owners po WHERE po.id = CAST(properties.property_owner_id AS INTEGER) LIMIT 1) AS owner_display_name, " \
        f"(SELECT po.id FROM property_owners po WHERE po.id = CAST(properties.property_owner_id AS INTEGER) LIMIT 1) AS owner_display_id " \
        f"FROM properties WHERE {base_where}{occ_inner}{rent_inner}{mgmt_inner}"

    # Build ORDER BY — use only column names from the inner SELECT, no subquery references in ORDER
    safe_sort_fields = {
        "name": "sort_name",
        "ref": "ref",
        "type": "property_type",
        "city": "city",
        "status": "property_status",
        "units": "actual_units",
        "occupied": "occupied_units",
        "vacant": "(total_units - occupied_units)",
        "rent": "monthly_rent",
        "owner": "owner_display_name",
        "management_type": "management_type",
    }
    order_col = safe_sort_fields.get(sort_field, "ref")
    order_dir = "DESC" if sort_direction == "desc" else "ASC"
    order_clause = f"ORDER BY {order_col} {order_dir}"

    # Simple wrapper query — no outer WHERE clause (filters are all in the inner query)
    base_query = f"SELECT * FROM ({inner_query}) sub {order_clause}"
    count_query = f"SELECT COUNT(*) AS cnt FROM ({inner_query}) sub"

    rows, total = paginate(base_query, count_query, base_params, page, per_page)

    # Real totals from DB — respect filters (also use the inner query directly)
    totals_query = f"SELECT COUNT(*) AS props_cnt, COALESCE(SUM(actual_units),0) AS units_cnt, COALESCE(SUM(occupied_units),0) AS occ_cnt FROM ({inner_query}) sub"
    db2 = get_dict_db()
    try:
        totals_row = db2.execute(totals_query, base_params).fetchone()
        real_props_count = totals_row["props_cnt"]
        real_units = totals_row["units_cnt"]
        real_occupied = totals_row["occ_cnt"]
        # Active/inactive counts — unfiltered (from full DB, not current page)
        active_cnt = db2.execute("SELECT COUNT(*) AS cnt FROM properties WHERE COALESCE(is_active, 1) = 1 AND (status IS NULL OR status <> 'archived')").fetchone()["cnt"]
        inactive_cnt = db2.execute("SELECT COUNT(*) AS cnt FROM properties WHERE COALESCE(is_active, 1) = 0 AND (status IS NULL OR status <> 'archived')").fetchone()["cnt"]
        all_cnt = db2.execute("SELECT COUNT(*) AS cnt FROM properties WHERE status IS NULL OR status <> 'archived'").fetchone()["cnt"]
        # Card counts. Unfiltered on purpose: the landing cards state how big each
        # category is, so they must not shrink because someone left a search in the box.
        management_cnt = db2.execute("SELECT COUNT(*) AS cnt FROM properties WHERE management_type = 'Management Fee' AND (status IS NULL OR status <> 'archived')").fetchone()["cnt"]
    finally:
        db2.close()

    return json_success({
        "items": rows,
        "totals": {"properties": real_props_count, "units": real_units, "occupied": real_occupied,
                   "active": active_cnt, "inactive": inactive_cnt, "management": management_cnt,
                   "all": all_cnt},
    }, total, page, per_page)


def api_create_property():
    """POST handler for creating a new property with onboarding details."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")
    required = ["ref", "name"]
    for r in required:
        if not data.get(r):
            return json_error(f"'{r}' is required")

    # Every property must be linked to a landlord.
    if not data.get("property_owner_id") and not (data.get("property_owner_name") or "").strip():
        return json_error("A landlord is required — every property must be linked to a landlord.")

    db = get_dict_db()
    try:
        data = _ensure_landlord_link(db, data)
        cols = ["ref", "name", "address_line_1", "address_line_2", "city", "county", "postcode", "country",
                "property_type", "total_units", "bedrooms", "bathrooms", "council_tax_band",
                "council_account_no", "property_owner_id", "property_owner_name", "features", "notes"]
        ins_cols = [c for c in cols if c in data]
        ins_vals = [data[c] for c in ins_cols]
        placeholders = ",".join(["?"] * len(ins_cols))
        cursor = db.execute(
            f"INSERT INTO properties ({','.join(ins_cols)}) VALUES ({placeholders})",
            ins_vals
        )
        db.commit()
        new_id = cursor.lastrowid
        user_info = session.get("user", {})
        user_name = user_info.get("username", "Unknown")
        record_change(user_name, 'created', 'property', str(new_id), data.get("name", ""))
        return json_success({"id": new_id, "message": "Property created"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/properties/sync-monday-list", methods=["POST"])
def api_sync_monday_property_list():
    """Fetch Monday.com Property List board and update properties.custom_fields."""
    import requests as req_lib
    from pathlib import Path

    token_path = Path("/root/.hermes/secrets/monday_token.txt")
    if not token_path.exists():
        return json_error("Monday token not found", 500)
    token = token_path.read_text().strip()

    headers = {"Authorization": token, "Content-Type": "application/json"}

    # 1. Fetch all items from Property List board
    query = """{ boards(ids: 5930667891) { items_page(limit: 500) { items { id name column_values { id text } } } } }"""
    try:
        resp = req_lib.post("https://api.monday.com/v2", json={"query": query}, headers=headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return json_error(safe_error(e), 502)

    items = data.get("data", {}).get("boards", [{}])[0].get("items_page", {}).get("items", [])
    if not items:
        return json_error("No items found on Monday board", 404)

    # Column ID to custom_fields key mapping
    col_map = {
        "status": "legal_type",
        "numeric_mks0z366": "rent_to_ll_monday",
        "numeric_mm0psp3q": "rent_received_monday",
        "status7": "operating_as",
        "color_mm06cjy9": "model_monday",
        "check__1": "has_keys",
        "numbers2__1": "smart_tvs",
        "numbers7__1": "ring_doorbells",
        "numbers__1": "rooms_monday",
        "laundry__1": "laundry_code",
        "letterbox__1": "letterbox_code",
        "numbers4__1": "beds_monday",
        "label__1": "boiler_location",
        "numbers5__1": "keybox_code_monday",
        "numbers47__1": "smart_locks",
        "text__1": "alias_name",
        "text_mky1a405": "joint_venture",
        "date4": "start_date_monday",
        "date_mknagbz": "expiry_date_monday",
        "link": "maps_link",
    }

    import re as re_lib

    def normalize(s):
        s = s.lower().strip()
        s = re_lib.sub(r'[,\.]', '', s)
        s = re_lib.sub(r'\s+', ' ', s)
        return s.strip()

    def extract_street_num(s):
        nums = set(re_lib.findall(r'\d+', s))
        return nums

    # 2. Get local properties
    db = get_dict_db()
    try:
        local_props = db.execute(
            "SELECT id, name, address_line_1 FROM properties WHERE (status IS NULL OR status = '' OR status = 'Active')"
        ).fetchall()
    except Exception as e:
        db.close()
        return json_error(safe_error(e), 500)

    updated = 0
    matched = 0
    unmatched_monday = []

    for item in items:
        mname = item.get("name", "")
        mname_norm = normalize(mname)
        m_numbers = extract_street_num(mname)

        # Extract postcode from Monday name
        pc_match = re_lib.search(r'([A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2})', mname, re_lib.IGNORECASE)
        m_postcode = pc_match.group(1).strip().upper() if pc_match else ""

        # Build custom_fields dict from Monday columns
        custom = {}
        for cv in item.get("column_values", []):
            cid = cv.get("id", "")
            ckey = col_map.get(cid)
            if ckey:
                val = cv.get("text", "").strip()
                if val:
                    custom[ckey] = val

        if not custom:
            continue

        custom["_monday_name"] = mname
        custom["_monday_id"] = item.get("id", "")

        # Try matching to a local property
        found = False
        for lp in local_props:
            lp_id = lp["id"]
            lp_name_norm = normalize(lp["name"])
            lp_addr_norm = normalize(lp.get("address_line_1", ""))
            lp_numbers = extract_street_num(lp["name"] + " " + (lp.get("address_line_1", "") or ""))

            # Exact match
            if mname_norm == lp_name_norm or mname_norm == lp_addr_norm:
                found = True
            else:
                # Match by postcode + street number overlap
                common_nums = m_numbers & lp_numbers
                if m_postcode and common_nums:
                    # Check street name overlap
                    m_words = set(re_lib.sub(r'\d+', '', mname_norm).split())
                    lp_words = set(re_lib.sub(r'\d+', '', lp_name_norm + " " + lp_addr_norm).split())
                    common_words = m_words & lp_words
                    if len(common_words) >= 2 or (len(common_words) >= 1 and len([w for w in common_words if len(w) > 3]) >= 1):
                        found = True
                elif common_nums:
                    m_words = set(re_lib.sub(r'\d+', '', mname_norm).split())
                    lp_words = set(re_lib.sub(r'\d+', '', lp_name_norm + " " + lp_addr_norm).split())
                    common_words = m_words & lp_words
                    if len(common_words) >= 3 or (len(common_nums) >= 1 and len(common_words) >= 2):
                        found = True

            if found:
                try:
                    existing = db.execute("SELECT custom_fields FROM properties WHERE id=?", (lp_id,)).fetchone()
                    existing_json = {}
                    if existing and existing.get("custom_fields"):
                        try:
                            existing_json = json.loads(existing["custom_fields"])
                        except (json.JSONDecodeError, TypeError):
                            existing_json = {}
                    # Merge: Monday data overwrites, but preserve non-Monday custom fields
                    # Only update the Monday-specific keys
                    custom_fields = {k: v for k, v in existing_json.items() if not k.startswith("_monday") and k not in col_map.values()}
                    custom_fields.update(custom)

                    db.execute("UPDATE properties SET custom_fields=? WHERE id=?", (json.dumps(custom_fields), lp_id))
                    updated += 1
                    matched += 1
                except Exception as _e:
                    current_app.logger.error(f"Error in line ~1865: {_e}")
                    pass
                break

        if not found:
            unmatched_monday.append(mname)

        # Batch commits so the write lock is released periodically rather than
        # held for the entire loop.
        if updated and updated % 50 == 0:
            db.commit()

    db.commit()
    db.close()

    return json_success({
        "updated": updated,
        "total_monday_items": len(items),
        "matched": matched,
        "unmatched_monday_count": len(unmatched_monday),
    })


@banksia_os_bp.route("/properties/<int:prop_id>", methods=["GET", "PATCH"])
def api_property(prop_id):
    if request.method == "PATCH":
        return _api_patch_property(prop_id)
    db = get_dict_db()
    try:
        prop = db.execute("SELECT * FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        # Resolve owner display info from property_owners table
        owner_info = db.execute("SELECT id, name, company_name FROM property_owners WHERE id = CAST(? AS INTEGER) LIMIT 1",
                                (prop.get("property_owner_id", "0"),)).fetchone()
        if owner_info:
            prop["owner_display_name"] = owner_info["name"]
            prop["owner_display_id"] = owner_info["id"]
            prop["owner_company"] = owner_info.get("company_name", "")
        else:
            prop["owner_display_name"] = prop.get("property_owner_name", "")
            prop["owner_display_id"] = None

        # ── Units (existing) ──
        units = db.execute("SELECT * FROM units WHERE property_id = ? AND (is_active IS NULL OR is_active = 1) ORDER BY sort_order ASC, CAST(SUBSTR(unit_ref, 2) AS INTEGER) ASC, unit_ref ASC", (prop_id,)).fetchall()
        computed_rent = 0
        for u in units:
            tn = db.execute("SELECT * FROM tenancies WHERE unit_id=? AND status IN ('Active','active','Periodic','periodic') ORDER BY id DESC LIMIT 1", (u["id"],)).fetchone()
            if tn:
                u["tenant_name"] = tn.get("main_tenant_name") or ""
                u["tenancy_id"] = tn["id"]
                u['tenancy_rent'] = tn.get('rent_amount', 0)
                u['tenancy_status'] = tn.get('status', '')
                u['deposit_amount'] = tn.get('deposit_registered_amount', 0) or 0
                computed_rent += (tn.get('rent_amount', 0) or 0)
                cnt = db.execute("SELECT COUNT(*) AS cnt FROM tenants WHERE tenancy_id=?", (tn["id"],)).fetchone()
                u["occupant_count"] = cnt["cnt"] if cnt else 0
                first_t = db.execute("SELECT id FROM tenants WHERE tenancy_id=? LIMIT 1", (tn["id"],)).fetchone()
                u["tenant_id"] = first_t["id"] if first_t else None
            else:
                u["tenant_name"] = ""
                u["tenancy_id"] = None
                u["tenancy_rent"] = 0
                u["tenancy_status"] = ""
                u["occupant_count"] = 0
                u["tenant_id"] = None
        prop["units"] = units
        prop["monthly_rent"] = computed_rent or prop.get("monthly_rent", 0) or prop.get("monthly_property_rent", 0)
        prop["occupied_units"] = len([u for u in units if not u.get("unit_vacant")])
        prop["vacant_units"] = len([u for u in units if u.get("unit_vacant")])

        # ── Tenancies ──
        tenancies = db.execute(
            "SELECT t.*, u.unit_ref FROM tenancies t LEFT JOIN units u ON t.unit_id=u.id "
            "WHERE t.property_id=? AND t.status IN ('Active','active','Periodic','periodic') ORDER BY t.id DESC",
            (prop_id,)
        ).fetchall()
        prop["tenancies"] = [clean_none(dict(r)) for r in tenancies]

        # ── Tenants ──
        tenants = db.execute(
            "SELECT t.id, t.first_name, t.last_name, t.email, t.mobile, t.phone_home, "
            "t.title, t.date_of_birth, t.gender, t.main_tenant, t.tenancy_id, t.property_id, "
            "u.unit_ref, tn.start_date, tn.end_date, tn.rent_amount "
            "FROM tenants t "
            "JOIN tenancies tn ON t.tenancy_id=tn.id "
            "LEFT JOIN units u ON t.unit_id=u.id "
            "WHERE tn.property_id=? "
            "AND (t.main_tenant=1 OR t.status='active' OR t.status='Active') "
            "ORDER BY COALESCE(t.last_name,t.first_name)",
            (prop_id,)
        ).fetchall()
        # tenants use first_name/last_name — compose name
        for t in tenants:
            if not t.get("name") and t.get("first_name"):
                t["name"] = (t.get("first_name", "") or "") + " " + (t.get("last_name", "") or "")
        prop["tenants"] = [clean_none(dict(r)) for r in tenants]

        # ── Maintenance jobs ──
        maint = db.execute(
            "SELECT * FROM maintenance_jobs WHERE property_id=? ORDER BY created DESC LIMIT 20",
            (prop_id,)
        ).fetchall()
        prop["maintenance"] = [clean_none(dict(r)) for r in maint]

        # ── Documents ──
        docs = db.execute(
            "SELECT * FROM documents WHERE related_to='property' AND related_id=? ORDER BY created DESC LIMIT 20",
            (str(prop_id),)
        ).fetchall()
        prop["documents"] = [clean_none(dict(r)) for r in docs]

        # ── Activity ──
        activity = db.execute(
            "SELECT * FROM activity_log WHERE entity_type='property' AND entity_id=? ORDER BY created DESC LIMIT 20",
            (prop_id,)
        ).fetchall()
        prop["activity"] = [clean_none(dict(r)) for r in activity]

        # ── Images ──
        images = db.execute(
            "SELECT pi.id, pi.property_id, pi.unit_id, pi.image_url, pi.caption, "
            "pi.category, pi.sort_order, pi.created_at, "
            "u.unit_ref "
            "FROM property_images pi "
            "LEFT JOIN units u ON pi.unit_id = u.id "
            "WHERE pi.property_id = ? "
            "ORDER BY pi.sort_order ASC, pi.id ASC",
            (prop_id,)
        ).fetchall()
        img_list = []
        for img in images:
            d = dict(img)
            d["url"] = d.pop("image_url")
            d["thumbnail_url"] = d["url"]
            img_list.append(d)
        prop["images"] = img_list

        return json_success(prop)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Allowed fields for PATCH on properties (mass-assignment protection) ──
# All editable columns from the properties table except: id, arthur_id, created, modified,
# sync_dirty, local_modified, pushed_at, sync_origin, total_units, rentable_units.
ALLOWED_PROPERTY_FIELDS = {
    "name", "ref", "address_line_1", "address_line_2", "city", "county",
    "postcode", "country", "lat", "lng", "property_type",
    "property_owner_id", "property_owner_name",
    "max_occupancy", "bathrooms", "bedrooms", "kitchens", "floors",
    "council_tax_band", "council_account_no", "main_image_url", "image_urls", "epc_urls",
    "floor_plan_urls", "thumbnail_urls", "features", "notes", "tags",
    "custom_fields",
    # Extended HMO onboarding fields
    "status", "property_ref", "acquisition_date", "owner_company",
    "management_type", "monthly_property_rent", "management_fee",
    "contract_start", "contract_end", "notice_period_days",
    "deposit_paid_to_landlord", "responsible_manager",
    "is_hmo", "licence_number", "licence_expiry",
    "heating_type", "boiler_details", "utility_suppliers", "wifi_provider",
    "main_door_instructions", "keybox_location", "keybox_code",
    "smart_lock_provider", "smart_lock_code", "intercom_details",
    "alarm_details", "emergency_access_notes",
    "description", "internal_notes", "is_active",
}

SYNC_TABLES = {"properties", "units", "tenancies", "tenants", "applicants"}


def _api_patch_property(prop_id):
    """PATCH endpoint for properties with concurrency protection and audit logging."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    # Separate concurrency token from payload
    provided_modified = data.pop("modified", None)

    db = get_dict_db()
    try:
        # 1. Fetch current property state
        prop = db.execute("SELECT * FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        # Status set to/from 'archived' triggers the full cascade (units,
        # tenancies, deposits) instead of just writing the status text — this is
        # what makes archiving actually take effect from the front-end.
        _incoming = (data.get("status") or "").strip().lower()
        _current = (prop.get("status") or "").strip().lower()
        if _incoming == "archived" and _current != "archived":
            _archive_property_cascade(db, prop_id)
            _create_activity_log(db, "property_archived", prop_id,
                "Property '" + (prop.get("name") or prop.get("ref") or "") + "' archived via status change")
            db.commit()
            return json_success({"updated": True, "id": prop_id, "status": "archived", "cascade": True})
        if _current == "archived" and _incoming and _incoming != "archived":
            _restore_property_cascade(db, prop_id)
            _create_activity_log(db, "property_restored", prop_id,
                "Property '" + (prop.get("name") or prop.get("ref") or "") + "' restored via status change")
            db.commit()
            return json_success({"updated": True, "id": prop_id, "status": _incoming, "cascade": True})

        # 2. Optimistic concurrency check
        if provided_modified is not None:
            current_modified = prop.get("modified")
            if current_modified and current_modified != provided_modified:
                return json_error({
                    "message": "Property was modified by another user. Please refresh and try again.",
                    "code": "CONCURRENCY_CONFLICT",
                    "current_modified": current_modified,
                    "your_modified": provided_modified,
                }, 409)

        # 3. Filter to allowed fields only (mass-assignment protection)
        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(properties)").fetchall()}
        protected = {"id", "sync_dirty", "local_modified", "sync_origin", "pushed_at", "arthur_id"}

        updates = {}
        ignored = []
        for key, val in data.items():
            if key in protected:
                continue
            if key not in ALLOWED_PROPERTY_FIELDS:
                ignored.append(key)
                continue
            if key not in real_cols:
                continue
            updates[key] = val

        if not updates:
            return json_error(f"No valid fields to update (ignored: {', '.join(ignored) or 'none'})", 400)

        # 4. Build per-field activity descriptions and track changes
        activity_entries = []
        for key, val in updates.items():
            old_val = prop.get(key)
            # Convert to string for logging comparison
            old_str = str(old_val) if old_val is not None else None
            new_str = str(val) if val is not None else None
            if old_str != new_str:
                activity_entries.append({
                    "field_changed": key,
                    "old_value": old_str,
                    "new_value": new_str,
                })

        now = datetime.now(timezone.utc).isoformat()

        # 5. Apply update with sync tracking
        set_parts = [f"{k} = ?" for k in updates]
        params = list(updates.values())

        set_parts.append("modified = ?")
        params.append(now)

        if "properties" in SYNC_TABLES:
            set_parts.append("sync_dirty = ?")
            params.append(1)
            set_parts.append("local_modified = ?")
            params.append(now)
            set_parts.append("sync_origin = ?")
            params.append("banksia_os")

        params.append(prop_id)
        db.execute(
            f"UPDATE properties SET {', '.join(set_parts)} WHERE id = ?",
            params
        )
        db.commit()

        user_info = session.get("user", {})
        user_name = user_info.get("username", "Unknown")
        record_change(user_name, 'updated', 'property', str(prop_id), prop.get("name", ""), details=json.dumps(activity_entries) if activity_entries else None)

        # 6. Log activity for each changed field
        user_name = getattr(request, "current_user", {}).get("username", "system")
        for entry in activity_entries:
            _log_activity(
                entity_type="property",
                entity_id=prop_id,
                action="update",
                field_changed=entry["field_changed"],
                old_value=entry["old_value"],
                new_value=entry["new_value"],
                notes=f"Property '{prop.get('name', '') or prop.get('ref', '')}' updated",
                db=db,
            )
        db.commit()

        return json_success({
            "updated": True,
            "id": prop_id,
            "fields": list(updates.keys()),
            "ignored": ignored,
            "modified": now,
        })
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2b. PROPERTY DEPENDENCIES — for archive/delete UI
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/<int:prop_id>/dependencies", methods=["GET"])
def api_property_dependencies(prop_id):
    """Return counts of linked entities for archive/delete pre-checks."""
    db = get_dict_db()
    try:
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        units = db.execute("SELECT COUNT(*) AS cnt FROM units WHERE property_id = ?", (prop_id,)).fetchone()["cnt"]

        active_statuses = ("'Current','current','Periodic','periodic','Active','active'")
        active_tenancies = db.execute(
            f"SELECT COUNT(*) AS cnt FROM tenancies WHERE property_id = ? AND status IN ({active_statuses})",
            (prop_id,)
        ).fetchone()["cnt"]

        total_tenancies = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies WHERE property_id = ?",
            (prop_id,)
        ).fetchone()["cnt"]

        tenants = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenants WHERE property_id = ?",
            (prop_id,)
        ).fetchone()["cnt"]

        # applicants don't have a property_id column — skip them
        applicants = 0

        documents = db.execute(
            "SELECT COUNT(*) AS cnt FROM documents WHERE related_to = 'property' AND related_id = ?",
            (str(prop_id),)
        ).fetchone()["cnt"]

        maintenance_jobs = db.execute(
            "SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE property_id = ?",
            (prop_id,)
        ).fetchone()["cnt"]

        # Also check for images
        images = db.execute(
            "SELECT COUNT(*) AS cnt FROM property_images WHERE property_id = ?",
            (prop_id,)
        ).fetchone()["cnt"]

        return json_success({
            "units": units,
            "active_tenancies": active_tenancies,
            "total_tenancies": total_tenancies,
            "tenants": tenants,
            "applicants": applicants,
            "documents": documents,
            "maintenance_jobs": maintenance_jobs,
            "images": images,
            "has_active_tenancies": active_tenancies > 0,
            "has_any_records": any([
                units, total_tenancies, tenants, applicants,
                documents, maintenance_jobs, images
            ]),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2c. CREATE FULL PROPERTY — transactional multi-step
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/create-full", methods=["POST"])
def api_create_property_full():
    """Complete multi-step property creation in one transactional request.
    Creates the property record, optional units, access records, and property info.
    Rolls back entirely on any failure."""
    data = request.get_json(force=True, silent=True) or {}
    if not data:
        return json_error("No data provided")

    # Validate required fields
    required = ["name", "address_line_1", "postcode"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return json_error(f"Missing required fields: {', '.join(missing)}")

    # A property cannot exist without a landlord. Accept either a linked
    # landlord id or a landlord name (which we auto-create + link below).
    if not data.get("property_owner_id") and not (data.get("property_owner_name") or "").strip():
        return json_error("A landlord is required — every property must be linked to a landlord.")

    db = get_dict_db()
    try:
        # ── 0. Resolve/auto-create landlord so the property is always linked ──
        data = _ensure_landlord_link(db, data)
        # ── 1. Build property insert ──
        property_fields = [
            "name", "address_line_1", "address_line_2", "city", "county",
            "postcode", "country", "property_ref", "property_type", "status",
            "acquisition_date", "property_owner_id", "owner_company",
            "management_type", "monthly_property_rent", "management_fee",
            "contract_start", "contract_end", "notice_period_days",
            "deposit_paid_to_landlord", "responsible_manager",
            "bedrooms", "bathrooms", "kitchens", "floors", "max_occupancy",
            "is_hmo", "licence_number", "licence_expiry",
            "notes",
        ]
        ins_cols = []
        ins_vals = []
        for f in property_fields:
            if f in data and data[f] is not None:
                val = data[f]
                if f == "is_hmo":
                    val = 1 if val else 0
                ins_cols.append(f)
                ins_vals.append(val)

        placeholders = ",".join(["?"] * len(ins_cols))
        cursor = db.execute(
            f"INSERT INTO properties ({','.join(ins_cols)}) VALUES ({placeholders})",
            ins_vals
        )
        property_id = cursor.lastrowid

        # ── 2. Create units if provided ──
        units_data = data.get("units", [])
        if units_data and isinstance(units_data, list):
            unit_fields = [
                "unit_ref", "unit_type", "floor", "bedrooms", "capacity",
                "market_rent", "furnished", "status", "unit_vacant", "sort_order",
            ]
            for u_data in units_data:
                u_ins_cols = ["property_id"]
                u_ins_vals = [property_id]
                for f in unit_fields:
                    if f in u_data and u_data[f] is not None:
                        val = u_data[f]
                        if f == "furnished":
                            val = 1 if val else 0
                        if f == "unit_vacant":
                            val = 1 if val else 0
                        u_ins_cols.append(f)
                        u_ins_vals.append(val)
                u_placeholders = ",".join(["?"] * len(u_ins_cols))
                db.execute(
                    f"INSERT INTO units ({','.join(u_ins_cols)}) VALUES ({u_placeholders})",
                    u_ins_vals
                )

        # ── 3. Create access record if provided ──
        access_data = data.get("access")
        if access_data and isinstance(access_data, dict):
            # Map the frontend fields to DB access_records columns
            access_field_map = {
                "main_door_instructions": "label",
                "keybox_location": "identifier",
                "keybox_code": "notes",
                "smart_lock_provider": "notes",
                "smart_lock_code": "notes",
                "intercom_details": "label",
                "alarm_details": "label",
                "keys_count": "notes",
                "key_holder": "assigned_to",
                "emergency_access_notes": "notes",
            }
            # Build a combined notes string and label from access data
            access_parts = []
            for k, v in access_data.items():
                if v and isinstance(v, str) and v.strip():
                    label = k.replace("_", " ").title()
                    access_parts.append(f"{label}: {v}")
            combined_notes = "; ".join(access_parts)

            db.execute(
                "INSERT INTO access_records (property_id, type, label, notes) VALUES (?, 'property_access', ?, ?)",
                [property_id, "Main Access", combined_notes]
            )

        # ── 4. Store property_info as notes if provided ──
        info_data = data.get("property_info")
        if info_data and isinstance(info_data, dict):
            info_parts = []
            for k, v in info_data.items():
                if v and isinstance(v, str) and v.strip():
                    label = k.replace("_", " ").title()
                    info_parts.append(f"{label}: {v}")
            if info_parts:
                info_str = "; ".join(info_parts)
                existing_notes = db.execute(
                    "SELECT notes FROM properties WHERE id = ?", (property_id,)
                ).fetchone()
                current_notes = existing_notes["notes"] or "" if existing_notes else ""
                if current_notes:
                    info_str = current_notes + "\n\n" + info_str
                db.execute(
                    "UPDATE properties SET notes = ? WHERE id = ?",
                    [info_str, property_id]
                )

        # ── 5. Create activity log ──
        _create_activity_log(db, "property_created", property_id,
                             f"Property '{data.get('name', '')}' created with {len(units_data) if units_data else 0} units")

        db.commit()
        return json_success({
            "id": property_id,
            "message": "Property created successfully",
            "units_created": len(units_data) if units_data else 0,
        }), 201

    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


def _create_activity_log(db, action, resource_id, description, user=None):
    """Create an activity log entry. Writes to the activity_log table
    (used by the PATCH endpoint) as well as the legacy activity table."""
    try:
        if has_request_context():
            g._audit_detailed = True
    except Exception:
        pass
    if user is None:
        user = getattr(request, "current_user", {}).get("username", "system")
    try:
        now = datetime.now(timezone.utc).isoformat()

        # Write to activity_log table (primary, used by activity endpoints)
        try:
            db.execute(
                "INSERT INTO activity_log (entity_type, entity_id, action, user_name, notes, created) "
                "VALUES ('property', ?, ?, ?, ?, ?)",
                [resource_id, action, user, description, now]
            )
        except Exception:
            pass  # activity_log table may not exist

        # Legacy: Try activity table if it exists
        try:
            db.execute(
                "INSERT INTO activity (action, resource_type, resource_id, description, user, created_at) "
                "VALUES (?, 'property', ?, ?, ?, ?)",
                [action, resource_id, description, user, now]
            )
        except Exception:
            pass  # activity table doesn't exist — benign

        # Also try notifications table as a last fallback
        try:
            db.execute(
                "INSERT INTO notifications (type, title, message, created_at) "
                "VALUES ('activity', ?, ?, ?)",
                [f"{action}: {description}", f"Property #{resource_id}", now]
            )
        except Exception:
            pass  # No logging table at all — benign
    except Exception:
        pass  # Never let logging crash the main operation


# ── Sensitive fields that must be redacted from activity logs ──
SENSITIVE_FIELDS = {
    "keybox_code",
    "smart_lock_code",
    "alarm_code",
    "wifi_password",
    "alarm_details",
    "emergency_access_notes",
    "passport_number",
    "ni_number",
    "visa_number",
    "bank_name",
    "bank_account",
    "bank_sort_code",
}


def _redact_if_sensitive(val):
    """Return '[REDACTED]' if val looks sensitive or is one of the sensitive fields."""
    return "[REDACTED]" if val is not None else None


def _log_activity(entity_type, entity_id, action, field_changed=None,
                  old_value=None, new_value=None, notes=None, db=None):
    """Log an activity entry to the activity_log table.

    Automatically redacts sensitive field values.
    If no db connection is provided, creates one (for use outside request context).
    
    For property updates, also creates notifications for the responsible_manager
    and all super_admin users.
    """
    try:
        if has_request_context():
            g._audit_detailed = True
    except Exception:
        pass
    # Redact sensitive fields
    if field_changed and field_changed in SENSITIVE_FIELDS:
        old_value = _redact_if_sensitive(old_value)
        new_value = _redact_if_sensitive(new_value)

    user_name = getattr(request, "current_user", {}).get("username", "system")
    own_conn = db is None
    if own_conn:
        db = get_dict_db()
    try:
        db.execute(
            "INSERT INTO activity_log (entity_type, entity_id, action, field_changed, "
            "old_value, new_value, user_name, notes) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            [entity_type, entity_id, action, field_changed, old_value, new_value, user_name, notes]
        )
        
        # For property updates, create notifications for responsible_manager + super_admins
        if entity_type == "property" and action == "update" and field_changed:
            try:
                now = datetime.now(timezone.utc).isoformat()
                prop = db.execute(
                    "SELECT name, ref, responsible_manager FROM properties WHERE id = ?",
                    (entity_id,)
                ).fetchone()
                if prop:
                    prop_label = prop.get("name") or prop.get("ref") or f"#{entity_id}"
                    rm = prop.get("responsible_manager") or ""
                    message = (f"{user_name} updated {field_changed} on property "
                               f"'{prop_label}' ({_format_value(old_value)} → {_format_value(new_value)})")
                    link = f"/banksia-os?entity=properties&id={entity_id}"
                    
                    # Notify responsible_manager
                    notified = set()
                    if rm and rm.strip():
                        db.execute(
                            "INSERT INTO notifications (username, message, link, read, created) "
                            "VALUES (?, ?, ?, 0, ?)",
                            (rm.strip(), message, link, now)
                        )
                        notified.add(rm.strip())
                    
                    # Notify super_admins (Sami, Roo, Norbert, Sadman) who aren't the updater
                    super_admins = ["Sami", "Roo", "Norbert", "Sadman"]
                    for sa in super_admins:
                        if sa not in notified and sa != user_name:
                            db.execute(
                                "INSERT INTO notifications (username, message, link, read, created) "
                                "VALUES (?, ?, ?, 0, ?)",
                                (sa, message, link, now)
                            )
                            notified.add(sa)
            except Exception:
                pass  # Never let notification creation crash logging
        
        if own_conn:
            db.commit()
    except Exception:
        pass  # Never let logging crash the main operation
    finally:
        if own_conn:
            db.close()


def _format_value(val):
    """Format a value for notification messages — truncate and clean up."""
    if val is None:
        return "∅"
    s = str(val)
    if len(s) > 60:
        s = s[:57] + "..."
    return s


# ═══════════════════════════════════════════════
# 2d. ARCHIVE PROPERTY — soft delete with dependency check
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/archive-bulk", methods=["POST"])
def api_archive_properties_bulk():
    """Archive multiple properties in one request. Accepts JSON body with 'ids' array.
    Returns per-property results and any blockers found."""
    data = request.get_json(silent=True) or {}
    prop_ids = data.get("ids", [])
    if not prop_ids or not isinstance(prop_ids, list):
        return json_error("Provide an 'ids' array of property IDs")

    db = get_dict_db()
    try:
        results = []
        archived = []
        for pid in prop_ids:
            prop = db.execute("SELECT * FROM properties WHERE id = ?", (pid,)).fetchone()
            if not prop:
                results.append({"id": pid, "archived": False, "error": "Not found"})
                continue

            active_statuses = ("'Current','current','Periodic','periodic','Active','active'")
            active_tenancies = db.execute(
                f"SELECT COUNT(*) AS cnt FROM tenancies WHERE property_id=? AND status IN ({active_statuses})",
                (pid,)
            ).fetchone()["cnt"]
            active_jobs = db.execute(
                "SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE property_id=? AND status NOT IN ('COMPLETED','CANCELLED','No Invoice Found')",
                (pid,)
            ).fetchone()["cnt"]

            if active_tenancies > 0 or active_jobs > 0:
                results.append({"id": pid, "archived": False, "blockers": {"active_tenancies": active_tenancies, "active_maintenance_jobs": active_jobs}})
                continue

            db.execute("UPDATE properties SET status='archived', modified=? WHERE id=?",
                       [datetime.now(timezone.utc).isoformat(), pid])

            # Cascade: archive all units under this property
            db.execute(
                "UPDATE units SET is_active = 0, modified = ? WHERE property_id = ? AND (is_active IS NULL OR is_active = 1)",
                [datetime.now(timezone.utc).isoformat(), pid]
            )

            _create_activity_log(db, "property_archived", pid,
                                 f"Property '{prop.get('name','') or prop.get('ref','')}' archived (bulk)")
            archived.append(pid)
            results.append({"id": pid, "archived": True})

        db.commit()
        return json_success({
            "archived": archived,
            "results": results,
            "total_requested": len(prop_ids),
            "total_archived": len(archived),
        })
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/properties/<int:prop_id>/archive", methods=["POST"])
def api_archive_property(prop_id):
    """Archive a property and CASCADE to everything connected to it — units,
    tenancies and deposits — so nothing tied to it keeps showing on live
    figures/financials. Unlike the old behaviour this does NOT block when
    active tenancies exist (an archived property is one returned to the
    landlord / no longer operated), it archives them too. Fully reversible
    via /restore. Returns the dependency counts it acted on."""
    db = get_dict_db()
    try:
        prop = db.execute("SELECT * FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)
        if prop.get("status") == "archived":
            return json_error("Property is already archived", 409)

        user = (getattr(request, "current_user", None) or session.get("user", {}) or {}).get("username", "system")
        counts = _archive_property_cascade(db, prop_id)

        _create_activity_log(
            db, "property_archived", prop_id,
            f"Property '{prop.get('name','') or prop.get('ref','')}' archived (cascade: "
            f"{counts['tenancies']} tenancies, {counts['units']} units, {counts['deposits']} deposits)"
        )
        db.commit()
        return json_success({
            "id": prop_id,
            "status": "archived",
            "message": "Property and all connected records archived",
            "cascade": counts,
        })
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


def _archive_property_cascade(db, prop_id):
    """Set property + all connected units/tenancies/deposits to archived.
    Preserves each record's live status so restore can put it back exactly."""
    now = datetime.now(timezone.utc).isoformat()
    user = (getattr(request, "current_user", None) or session.get("user", {}) or {}).get("username", "system")

    db.execute(
        "UPDATE properties SET status='archived', is_active=0, archived_at=?, archived_by=?, modified=? WHERE id=?",
        [now, user, now, prop_id]
    )
    u = db.execute(
        "UPDATE units SET is_active=0, archived_at=?, modified=? WHERE property_id=? AND (is_active IS NULL OR is_active=1)",
        [now, now, prop_id]
    ).rowcount
    # Tenancies: stash live status into pre_archive_status, flip to 'Archived'
    # (every live query filters status IN Active/Periodic/Current, so this drops
    # them from rent roll, arrears, occupancy automatically).
    t = db.execute(
        "UPDATE tenancies SET pre_archive_status=status, status='Archived', archived_at=?, modified=? "
        "WHERE property_id=? AND status <> 'Archived'",
        [now, now, prop_id]
    ).rowcount
    # Deposits: stash current_status, flip to 'archived' (drops from held-deposit
    # figures + the deposits register automatically).
    d = db.execute(
        "UPDATE deposits SET pre_archive_status=current_status, current_status='archived', archived_at=?, modified=? "
        "WHERE property_id=? AND current_status <> 'archived'",
        [now, now, prop_id]
    ).rowcount
    return {"units": u, "tenancies": t, "deposits": d}


def _restore_property_cascade(db, prop_id):
    """Reverse _archive_property_cascade — put every record back to its
    pre-archive live status."""
    now = datetime.now(timezone.utc).isoformat()
    db.execute(
        "UPDATE properties SET status='active', is_active=1, archived_at=NULL, archived_by=NULL, modified=? WHERE id=?",
        [now, prop_id]
    )
    u = db.execute(
        "UPDATE units SET is_active=1, archived_at=NULL, modified=? WHERE property_id=? AND is_active=0",
        [now, prop_id]
    ).rowcount
    t = db.execute(
        "UPDATE tenancies SET status=COALESCE(NULLIF(pre_archive_status,''),'Active'), pre_archive_status=NULL, "
        "archived_at=NULL, modified=? WHERE property_id=? AND status='Archived'",
        [now, prop_id]
    ).rowcount
    d = db.execute(
        "UPDATE deposits SET current_status=COALESCE(NULLIF(pre_archive_status,''),'held'), pre_archive_status=NULL, "
        "archived_at=NULL, modified=? WHERE property_id=? AND current_status='archived'",
        [now, prop_id]
    ).rowcount
    return {"units": u, "tenancies": t, "deposits": d}


@banksia_os_bp.route("/properties/<int:prop_id>/restore", methods=["POST"])
def api_restore_property(prop_id):
    """Restore an archived property (and everything cascaded) to live."""
    db = get_dict_db()
    try:
        prop = db.execute("SELECT * FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)
        if prop.get("status") != "archived":
            return json_error(f"Property is not archived (status: '{prop.get('status')}')", 409)

        counts = _restore_property_cascade(db, prop_id)
        _create_activity_log(
            db, "property_restored", prop_id,
            f"Property '{prop.get('name','') or prop.get('ref','')}' restored from archive (cascade: "
            f"{counts['tenancies']} tenancies, {counts['units']} units, {counts['deposits']} deposits)"
        )
        db.commit()
        return json_success({
            "id": prop_id, "status": "active",
            "message": "Property and all connected records restored",
            "cascade": counts,
        })
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ─── Universal archive / Records module ───

def _ensure_archive_columns(db):
    """Idempotently add the archive-tracking columns used by the universal
    archive/Records module. Safe to call repeatedly (each ALTER is guarded)."""
    for s in (
        "ALTER TABLE units ADD COLUMN archived_by TEXT",
        "ALTER TABLE tenancies ADD COLUMN archived_by TEXT",
        "ALTER TABLE applicants ADD COLUMN archived_at TEXT",
        "ALTER TABLE applicants ADD COLUMN archived_by TEXT",
        "ALTER TABLE applicants ADD COLUMN pre_archive_status TEXT",
        "ALTER TABLE guarantors ADD COLUMN archived_at TEXT",
        "ALTER TABLE guarantors ADD COLUMN archived_by TEXT",
        # Guarantors are held on the tenant row; this is their own status,
        # set by hand and overriding the one derived from the tenancy.
        "ALTER TABLE tenants ADD COLUMN guarantor_status TEXT",
    ):
        try:
            db.execute(s)
        except Exception:
            pass  # column already present
    db.commit()


def _archive_actor():
    return (getattr(request, "current_user", None) or session.get("user", {}) or {}).get("username", "system")


@banksia_os_bp.route("/archive", methods=["GET"])
def api_archive_list():
    """Universal archive / Records module. Returns every archived record across
    all entity types: properties (with their cascade counts), plus tenancies,
    units, deposits, applicants and guarantors that were archived in their own
    right. Cascade children of an archived property are summarised under that
    property rather than listed twice. Every row is restorable. Searchable."""
    search = request.args.get("search", "").strip()
    lv = "%" + search + "%"
    db = get_dict_db()
    try:
        _ensure_archive_columns(db)

        def like(cols):
            if not search:
                return "", []
            clause = " OR ".join("COALESCE(%s,'') LIKE ?" % c for c in cols)
            return " AND (" + clause + ")", [lv] * len(cols)

        # Properties (cascade parents)
        pc, pp = like(["p.ref", "p.name", "p.address_line_1"])
        properties = db.execute(
            "SELECT p.id, p.ref, p.name, p.address_line_1, p.archived_at, p.archived_by, "
            "(SELECT COUNT(*) FROM units u WHERE u.property_id=p.id) AS units_count, "
            "(SELECT COUNT(*) FROM tenancies t WHERE t.property_id=p.id) AS tenancies_count, "
            "(SELECT COUNT(*) FROM deposits d WHERE d.property_id=p.id) AS deposits_count, "
            "(SELECT COALESCE(SUM(d.amount),0) FROM deposits d WHERE d.property_id=p.id) AS deposits_total "
            "FROM properties p WHERE p.status='archived'" + pc + " ORDER BY p.archived_at DESC, p.id DESC",
            pp
        ).fetchall()

        arch = "(SELECT id FROM properties WHERE status='archived')"

        # Tenancies archived independently (parent property not archived)
        tc, tp = like(["t.ref", "t.main_tenant_name", "pr.name", "pr.ref"])
        tenancies = db.execute(
            "SELECT t.id, t.ref, t.main_tenant_name, t.status, t.pre_archive_status, t.archived_at, t.archived_by, "
            "t.start_date, t.rent_amount, pr.name AS property_name, pr.ref AS property_ref "
            "FROM tenancies t LEFT JOIN properties pr ON pr.id=t.property_id "
            "WHERE t.status='Archived' AND (t.property_id IS NULL OR t.property_id NOT IN " + arch + ")" + tc +
            " ORDER BY t.archived_at DESC, t.id DESC",
            tp
        ).fetchall()

        # Units archived independently
        uc, up = like(["u.unit_ref", "pr.name", "pr.ref"])
        units = db.execute(
            "SELECT u.id, u.unit_ref AS ref, u.archived_at, u.archived_by, "
            "pr.name AS property_name, pr.ref AS property_ref "
            "FROM units u LEFT JOIN properties pr ON pr.id=u.property_id "
            "WHERE u.status='archived' AND (u.property_id IS NULL OR u.property_id NOT IN " + arch + ")" + uc +
            " ORDER BY u.archived_at DESC, u.id DESC",
            up
        ).fetchall()

        # Deposits archived independently
        dc, dp = like(["dep.protection_reference", "pr.name", "pr.ref"])
        deposits = db.execute(
            "SELECT dep.id, dep.protection_reference AS ref, dep.amount, dep.archived_at, dep.pre_archive_status, "
            "pr.name AS property_name, pr.ref AS property_ref "
            "FROM deposits dep LEFT JOIN properties pr ON pr.id=dep.property_id "
            "WHERE dep.current_status='archived' AND (dep.property_id IS NULL OR dep.property_id NOT IN " + arch + ")" + dc +
            " ORDER BY dep.archived_at DESC, dep.id DESC",
            dp
        ).fetchall()

        # Applicants archived
        ac, ap = like(["a.first_name", "a.last_name", "a.email"])
        applicants = db.execute(
            "SELECT a.id, a.first_name, a.last_name, a.email, a.archived_at, a.archived_by, a.pre_archive_status "
            "FROM applicants a WHERE a.archived_at IS NOT NULL AND a.archived_at<>''" + ac +
            " ORDER BY a.archived_at DESC, a.id DESC",
            ap
        ).fetchall()

        # Guarantors archived
        gc, gp = like(["g.first_name", "g.last_name", "g.email"])
        guarantors = db.execute(
            "SELECT g.id, g.first_name, g.last_name, g.email, g.archived_at, g.archived_by, "
            "a.first_name AS applicant_first, a.last_name AS applicant_last "
            "FROM guarantors g LEFT JOIN applicants a ON a.id=g.applicant_id "
            "WHERE g.archived_at IS NOT NULL AND g.archived_at<>''" + gc +
            " ORDER BY g.archived_at DESC, g.id DESC",
            gp
        ).fetchall()

        counts = {
            "properties": len(properties), "tenancies": len(tenancies),
            "units": len(units), "deposits": len(deposits),
            "applicants": len(applicants), "guarantors": len(guarantors),
        }
        return json_success({
            "properties": properties, "tenancies": tenancies, "units": units,
            "deposits": deposits, "applicants": applicants, "guarantors": guarantors,
            "counts": counts, "total": sum(counts.values()),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tid>/archive", methods=["POST"])
def api_archive_tenancy(tid):
    db = get_dict_db()
    try:
        _ensure_archive_columns(db)
        t = db.execute("SELECT * FROM tenancies WHERE id=?", (tid,)).fetchone()
        if not t:
            return json_error("Tenancy not found", 404)
        if t.get("status") == "Archived":
            return json_error("Tenancy is already archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE tenancies SET pre_archive_status=status, status='Archived', archived_at=?, archived_by=?, modified=? WHERE id=?",
                   (now, _archive_actor(), now, tid))
        db.commit()
        _log_activity("tenancy", tid, "archived", notes="Tenancy '" + str(t.get("ref", "")) + "' archived", db=db)
        return json_success({"id": tid, "status": "Archived", "message": "Tenancy archived"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tid>/restore", methods=["POST"])
def api_restore_tenancy(tid):
    db = get_dict_db()
    try:
        t = db.execute("SELECT * FROM tenancies WHERE id=?", (tid,)).fetchone()
        if not t:
            return json_error("Tenancy not found", 404)
        if t.get("status") != "Archived":
            return json_error("Tenancy is not archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE tenancies SET status=COALESCE(NULLIF(pre_archive_status,''),'Active'), pre_archive_status=NULL, archived_at=NULL, archived_by=NULL, modified=? WHERE id=?",
                   (now, tid))
        db.commit()
        _log_activity("tenancy", tid, "restored", notes="Tenancy '" + str(t.get("ref", "")) + "' restored from archive", db=db)
        return json_success({"id": tid, "message": "Tenancy restored"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/units/<int:unit_id>/restore", methods=["POST"])
def api_restore_unit(unit_id):
    db = get_dict_db()
    try:
        u = db.execute("SELECT * FROM units WHERE id=?", (unit_id,)).fetchone()
        if not u:
            return json_error("Unit not found", 404)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE units SET status='', is_active=1, unit_vacant=0, archived_at=NULL, archived_by=NULL, modified=? WHERE id=?",
                   (now, unit_id))
        db.commit()
        _log_activity("unit", unit_id, "restored", notes="Unit '" + str(u.get("unit_ref", "")) + "' restored from archive", db=db)
        return json_success({"id": unit_id, "message": "Unit restored"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:dep_id>/restore", methods=["POST"])
def api_restore_deposit(dep_id):
    db = get_dict_db()
    try:
        d = db.execute("SELECT * FROM deposits WHERE id=?", (dep_id,)).fetchone()
        if not d:
            return json_error("Deposit not found", 404)
        if d.get("current_status") != "archived":
            return json_error("Deposit is not archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE deposits SET current_status=COALESCE(NULLIF(pre_archive_status,''),'held'), pre_archive_status=NULL, archived_at=NULL, modified=? WHERE id=?",
                   (now, dep_id))
        db.commit()
        _log_activity("deposit", dep_id, "restored", notes="Deposit restored from archive", db=db)
        return json_success({"id": dep_id, "message": "Deposit restored"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:aid>", methods=["DELETE"])
def api_delete_applicant(aid):
    db = get_dict_db()
    try:
        a = db.execute("SELECT * FROM applicants WHERE id=?", (aid,)).fetchone()
        if not a:
            return json_error("Applicant not found", 404)
        name = (str(a.get("first_name", "")) + " " + str(a.get("last_name", ""))).strip() or "Applicant #%s" % aid
        db.execute("DELETE FROM applicants WHERE id=?", (aid,))
        db.commit()
        _log_activity("applicant", aid, "deleted", notes="Applicant '%s' deleted" % name, db=db)
        return json_success({"id": aid, "message": "Applicant deleted"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/applicants/<int:aid>/archive", methods=["POST"])
def api_archive_applicant(aid):
    db = get_dict_db()
    try:
        _ensure_archive_columns(db)
        a = db.execute("SELECT * FROM applicants WHERE id=?", (aid,)).fetchone()
        if not a:
            return json_error("Applicant not found", 404)
        if a.get("archived_at"):
            return json_error("Applicant is already archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE applicants SET pre_archive_status=status, status='Archived', archived_at=?, archived_by=? WHERE id=?",
                   (now, _archive_actor(), aid))
        db.commit()
        _log_activity("applicant", aid, "archived", notes="Applicant '" + str(a.get("first_name", "")) + " " + str(a.get("last_name", "")) + "' archived", db=db)
        return json_success({"id": aid, "message": "Applicant archived"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:aid>/restore", methods=["POST"])
def api_restore_applicant(aid):
    db = get_dict_db()
    try:
        a = db.execute("SELECT * FROM applicants WHERE id=?", (aid,)).fetchone()
        if not a:
            return json_error("Applicant not found", 404)
        if not a.get("archived_at"):
            return json_error("Applicant is not archived", 409)
        db.execute("UPDATE applicants SET status=COALESCE(NULLIF(pre_archive_status,''),'Active'), pre_archive_status=NULL, archived_at=NULL, archived_by=NULL WHERE id=?",
                   (aid,))
        db.commit()
        _log_activity("applicant", aid, "restored", notes="Applicant '" + str(a.get("first_name", "")) + " " + str(a.get("last_name", "")) + "' restored from archive", db=db)
        return json_success({"id": aid, "message": "Applicant restored"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/guarantors/<int:gid>/archive", methods=["POST"])
def api_archive_guarantor(gid):
    db = get_dict_db()
    try:
        _ensure_archive_columns(db)
        g = db.execute("SELECT * FROM guarantors WHERE id=?", (gid,)).fetchone()
        if not g:
            return json_error("Guarantor not found", 404)
        if g.get("archived_at"):
            return json_error("Guarantor is already archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE guarantors SET archived_at=?, archived_by=?, modified=? WHERE id=?",
                   (now, _archive_actor(), now, gid))
        db.commit()
        _log_activity("guarantor", gid, "archived", notes="Guarantor '" + str(g.get("first_name", "")) + " " + str(g.get("last_name", "")) + "' archived", db=db)
        return json_success({"id": gid, "message": "Guarantor archived"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/guarantors/<int:gid>/restore", methods=["POST"])
def api_restore_guarantor(gid):
    db = get_dict_db()
    try:
        g = db.execute("SELECT * FROM guarantors WHERE id=?", (gid,)).fetchone()
        if not g:
            return json_error("Guarantor not found", 404)
        if not g.get("archived_at"):
            return json_error("Guarantor is not archived", 409)
        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE guarantors SET archived_at=NULL, archived_by=NULL, modified=? WHERE id=?", (now, gid))
        db.commit()
        _log_activity("guarantor", gid, "restored", notes="Guarantor '" + str(g.get("first_name", "")) + " " + str(g.get("last_name", "")) + "' restored from archive", db=db)
        return json_success({"id": gid, "message": "Guarantor restored"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2f. DELETE PROPERTY — permanent removal (super_admin only)
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/<int:prop_id>/delete", methods=["POST"])
def api_delete_property(prop_id):
    """Permanently delete a property. Requires super_admin role,
    `?confirm=PROPERTY-NAME` query parameter, and no remaining
    operational records."""
    user = session.get("user", {})
    if user.get("role") != "super_admin":
        return json_error("Only super admins can permanently delete properties", 403)

    db = get_dict_db()
    try:
        prop = db.execute("SELECT * FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        # Require confirmation via ?confirm=<property-name>
        confirm = request.args.get("confirm", "")
        expected_name = prop.get("name", "") or prop.get("ref", "")
        if not confirm or confirm.strip() != expected_name:
            return json_error({
                "message": "Confirmation required. Pass ?confirm=<property-name> matching the property name.",
                "code": "CONFIRMATION_REQUIRED",
                "expected": expected_name,
                "provided": confirm,
            }, 400)

        # Check for any remaining records
        dependencies = {}

        units_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM units WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["units"] = units_count

        tenancies_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["tenancies"] = tenancies_count

        tenants_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenants WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["tenants"] = tenants_count

        # applicants don't have a property_id column — skip them
        dependencies["applicants"] = 0

        documents_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM documents WHERE related_to = 'property' AND related_id = ?",
            (str(prop_id),)
        ).fetchone()["cnt"]
        dependencies["documents"] = documents_count

        maintenance_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["maintenance_jobs"] = maintenance_count

        images_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM property_images WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["images"] = images_count

        access_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM access_records WHERE property_id = ?", (prop_id,)
        ).fetchone()["cnt"]
        dependencies["access_records"] = access_count

        # Check deposits
        deposits_count = 0
        try:
            deposits_count = db.execute(
                "SELECT COUNT(*) AS cnt FROM deposits WHERE property_id = ?", (prop_id,)
            ).fetchone()["cnt"]
        except Exception as _e:
            current_app.logger.error(f"Error in line ~2817: {_e}")
            pass
        dependencies["deposits"] = deposits_count

        # Check transactions
        try:
            transactions_count = db.execute(
                "SELECT COUNT(*) AS cnt FROM transactions WHERE property_id = ?", (prop_id,)
            ).fetchone()["cnt"]
        except Exception:
            transactions_count = 0
        dependencies["transactions"] = transactions_count

        has_history = any(v > 0 for v in dependencies.values())
        if has_history:
            return json_error({
                "message": "Cannot delete property with operational history",
                "code": "HAS_OPERATIONAL_HISTORY",
                "dependencies": {k: v for k, v in dependencies.items() if v > 0},
            }, 409)

        # Delete related records explicitly (CASCADE may not be configured)
        for table in ("units", "access_records", "property_images", "documents"):
            try:
                if table == "documents":
                    db.execute(
                        f"DELETE FROM {table} WHERE related_to = 'property' AND related_id = ?",
                        (str(prop_id),)
                    )
                elif table == "units":
                    db.execute(f"DELETE FROM {table} WHERE property_id = ?", (prop_id,))
                elif table == "access_records":
                    db.execute(f"DELETE FROM {table} WHERE property_id = ?", (prop_id,))
                else:
                    db.execute(f"DELETE FROM {table} WHERE property_id = ?", (prop_id,))
            except Exception:
                pass  # table may not exist

        # Delete the property
        db.execute("DELETE FROM properties WHERE id = ?", (prop_id,))

        _create_activity_log(db, "property_deleted", prop_id,
                             f"Property '{prop.get('name', '') or prop.get('ref', '')}' permanently deleted by {user.get('username', 'unknown')}")

        db.commit()
        return json_success({
            "id": prop_id,
            "message": "Property permanently deleted",
        })

    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2g. PROPERTY ACTIVITY LOG — activity history for a specific property
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/<int:prop_id>/activity", methods=["GET"])
def api_property_activity(prop_id):
    """Return paginated activity log entries for a specific property."""
    page = int_param(request.args.get("page"), default=1)
    limit = int_param(request.args.get("per_page", 50), default=50, max_val=MAX_PAGE_SIZE)
    offset = (page - 1) * limit

    db = get_dict_db()
    try:
        # Verify property exists
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        total = db.execute(
            "SELECT COUNT(*) AS cnt FROM activity_log "
            "WHERE entity_type = 'property' AND entity_id = ?",
            (prop_id,)
        ).fetchone()["cnt"]

        rows = db.execute(
            "SELECT * FROM activity_log "
            "WHERE entity_type = 'property' AND entity_id = ? "
            "ORDER BY created DESC LIMIT ? OFFSET ?",
            (prop_id, limit, offset)
        ).fetchall()

        return json_success(rows, total=total, page=page, per_page=limit)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2h. ACTIVITY LOG — query and create entries
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/activity", methods=["GET"])
def api_get_activity():
    """Query activity log. Supports filtering by entity_type + entity_id and pagination."""
    entity_type = request.args.get("entity_type", "")
    entity_id_str = request.args.get("entity_id", "")
    entity_id = int(entity_id_str) if entity_id_str and entity_id_str.isdigit() else None
    limit = int_param(request.args.get("limit", 50), default=50, max_val=MAX_PAGE_SIZE)
    page = int_param(request.args.get("page", 1), default=1)
    offset = (page - 1) * limit

    db = get_dict_db()
    try:
        if entity_type and entity_id is not None:
            total = db.execute(
                "SELECT COUNT(*) AS cnt FROM activity_log WHERE entity_type = ? AND entity_id = ?",
                (entity_type, entity_id)
            ).fetchone()["cnt"]
            rows = db.execute(
                "SELECT * FROM activity_log WHERE entity_type = ? AND entity_id = ? "
                "ORDER BY created DESC LIMIT ? OFFSET ?",
                (entity_type, entity_id, limit, offset)
            ).fetchall()
        else:
            total = db.execute("SELECT COUNT(*) AS cnt FROM activity_log").fetchone()["cnt"]
            rows = db.execute(
                "SELECT * FROM activity_log ORDER BY created DESC LIMIT ? OFFSET ?",
                (limit, offset)
            ).fetchall()

        return json_success(rows, total=total, page=page, per_page=limit)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/activity", methods=["POST"])
def api_create_activity():
    """Create an activity log entry programmatically."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    entity_type = data.get("entity_type")
    entity_id = data.get("entity_id")
    action = data.get("action")

    if not all([entity_type, entity_id, action]):
        return json_error("entity_type, entity_id, and action are required", 400)

    field_changed = data.get("field_changed")
    old_value = data.get("old_value")
    new_value = data.get("new_value")
    notes = data.get("notes")

    _log_activity(
        entity_type=entity_type,
        entity_id=entity_id,
        action=action,
        field_changed=field_changed,
        old_value=old_value,
        new_value=new_value,
        notes=notes,
    )

    return json_success({"message": "Activity logged"})


# ═══════════════════════════════════════════════
# 2h. UNIT CRUD — create, edit, archive, delete, bulk-create
# ═══════════════════════════════════════════════

# The only three unit statuses (Norbert, 2026-08-03). Was a free-text mix of
# Let / Available To Let / Available / Unavailable to Let / Vacant.
UNIT_STATUSES = {"Occupied", "Vacant", "Inactive"}

ALLOWED_UNIT_FIELDS = {
    "unit_ref", "unit_type", "floor", "bedrooms", "capacity",
    "market_rent", "furnished", "status", "notes",
    "unit_status", "max_occupancy", "bathrooms",
}


@banksia_os_bp.route("/properties/<int:prop_id>/units", methods=["GET", "POST"])
def api_units_for_property(prop_id):
    if request.method == "GET":
        return api_list_units_for_property(prop_id)
    return api_create_unit_for_property(prop_id)


def api_list_units_for_property(prop_id):
    """GET /api/banksia-os/properties/{prop_id}/units — list units for a property."""
    db = get_dict_db()
    try:
        units = db.execute(
            "SELECT id, unit_ref AS ref, unit_type, unit_status, market_rent, floor, bedrooms, "
            "capacity, max_occupancy, furnished, status, notes, created, modified "
            "FROM units WHERE property_id = ? ORDER BY unit_ref",
            (prop_id,)
        ).fetchall()
        return json_success(units)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


def api_create_unit_for_property(prop_id):
    """POST /api/banksia-os/properties/{prop_id}/units — create a unit."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    unit_ref = data.get("unit_ref")
    unit_type = data.get("unit_type")
    if not unit_ref or not unit_type:
        return json_error("'unit_ref' and 'unit_type' are required", 400)

    db = get_dict_db()
    try:
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        now = datetime.now(timezone.utc).isoformat()
        ins_cols = ["property_id", "unit_ref", "unit_type", "created", "modified"]
        ins_vals = [prop_id, unit_ref, unit_type, now, now]

        optional_fields = ["floor", "bedrooms", "capacity", "market_rent", "furnished", "status", "notes",
                           "unit_status", "max_occupancy", "bathrooms"]
        for f in optional_fields:
            if f in data:
                ins_cols.append(f)
                ins_vals.append(data[f])

        placeholders = ",".join(["?"] * len(ins_cols))
        cursor = db.execute(
            f"INSERT INTO units ({','.join(ins_cols)}) VALUES ({placeholders})",
            ins_vals
        )
        new_id = cursor.lastrowid
        db.commit()

        _log_activity("unit", new_id, "created",
                      notes=f"Unit '{unit_ref}' created on property #{prop_id}",
                      db=db)

        created = db.execute("SELECT * FROM units WHERE id = ?", (new_id,)).fetchone()
        return json_success(clean_none(dict(created))), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/units/<int:unit_id>", methods=["PATCH"])
def api_update_unit(unit_id):
    """PATCH /api/banksia-os/units/{unit_id} — edit a unit."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    db = get_dict_db()
    try:
        unit = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        now = datetime.now(timezone.utc).isoformat()
        updates = {}
        for key, val in data.items():
            if key in ALLOWED_UNIT_FIELDS:
                updates[key] = val

        if not updates:
            return json_error("No valid fields to update", 400)

        # Unit status is one of exactly three values (Norbert, 2026-08-03), and it is
        # the same fact as unit_vacant — so the two are written together. Letting them
        # drift is how a room reads "Vacant" on this page and "occupied" in every count
        # that reads unit_vacant instead.
        if "unit_status" in updates:
            wanted = str(updates["unit_status"] or "").strip().title()
            if wanted not in UNIT_STATUSES:
                return json_error("Status must be one of: %s" % ", ".join(sorted(UNIT_STATUSES)), 400)
            updates["unit_status"] = wanted

            # Status has to agree with the tenancy (Norbert, 2026-08-03):
            #   Occupied -> there must be a tenant
            #   Vacant   -> there must not be
            #   Inactive -> there must not be
            # Enforced here rather than in the UI because the tenancy is the fact and the
            # status is a label on it. A room marked empty while someone lives in it is
            # how rent stops being chased and a re-let gets advertised over a tenant.
            live = db.execute(
                "SELECT main_tenant_name, status FROM tenancies WHERE unit_id = ? "
                "AND LOWER(COALESCE(status,'')) IN ('current','active','periodic') "
                "ORDER BY start_date DESC LIMIT 1", (unit_id,)
            ).fetchone()
            who = (live or {}).get("main_tenant_name") or "a tenant"

            if wanted == "Occupied" and not live:
                return json_error(
                    "This unit has no tenant, so it cannot be marked Occupied. "
                    "Create the tenancy first and the unit becomes occupied with it.", 409)
            if wanted in ("Vacant", "Inactive") and live:
                return json_error(
                    "%s still has a live tenancy on this unit, so it cannot be marked %s. "
                    "End the tenancy first." % (who, wanted), 409)

            if wanted == "Vacant":
                updates["unit_vacant"] = 1
            elif wanted == "Occupied":
                updates["unit_vacant"] = 0
            # Inactive says nothing about whether anyone is in there, so unit_vacant
            # is left exactly as it was rather than being invented.

        set_parts = [f"{k} = ?" for k in updates]
        params = list(updates.values())
        set_parts.append("modified = ?")
        params.append(now)
        params.append(unit_id)

        db.execute(f"UPDATE units SET {', '.join(set_parts)} WHERE id = ?", params)

        # Logged before the commit — _log_activity does not commit when it is handed a
        # connection, so an audit row written after db.commit() is thrown away.
        if "unit_status" in updates:
            _log_activity("unit", unit_id, "update", field_changed="unit_status",
                          old_value=unit.get("unit_status") or "",
                          new_value=updates["unit_status"],
                          notes=f"Unit '{unit.get('unit_ref', '')}' status changed",
                          db=db)
        else:
            _log_activity("unit", unit_id, "update",
                          notes=f"Unit '{unit.get('unit_ref', '')}' updated",
                          db=db)

        db.commit()

        updated = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        return json_success(clean_none(dict(updated)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/units/<int:unit_id>/archive", methods=["POST"])
def api_archive_unit(unit_id):
    """POST /api/banksia-os/units/{unit_id}/archive — soft-delete a unit."""
    db = get_dict_db()
    try:
        unit = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        # Check for active tenancies
        active = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies WHERE unit_id=? AND status IN ('Active','active','Periodic','periodic','Current','current')",
            (unit_id,)
        ).fetchone()["cnt"]
        if active > 0:
            return json_error(f"Cannot archive unit with {active} active tenanc{'y' if active == 1 else 'ies'}", 409)

        # Soft-archive by setting status to 'archived'
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "UPDATE units SET status='archived', unit_vacant=1, is_active=0, archived_at=?, archived_by=?, modified=? WHERE id=?",
            (now, _archive_actor(), now, unit_id)
        )
        db.commit()

        _log_activity("unit", unit_id, "archived",
                      notes=f"Unit '{unit.get('unit_ref', '')}' archived",
                      db=db)

        return json_success({"id": unit_id, "message": "Unit archived successfully", "status": "archived"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/units/<int:unit_id>/delete", methods=["POST"])
def api_delete_unit(unit_id):
    """POST /api/banksia-os/units/{unit_id}/delete — permanently delete a unit."""
    db = get_dict_db()
    try:
        unit = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        # Check for active tenancies
        active = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenancies WHERE unit_id=? AND status IN ('Active','active','Periodic','periodic','Current','current')",
            (unit_id,)
        ).fetchone()["cnt"]
        if active > 0:
            return json_error(f"Cannot delete unit with {active} active tenanc{'y' if active == 1 else 'ies'}", 409)

        # Delete related records first
        db.execute("UPDATE tenants SET unit_id=NULL, property_id=NULL WHERE unit_id=?", (unit_id,))
        db.execute("UPDATE tenancies SET unit_id=NULL WHERE unit_id=?", (unit_id,))
        db.execute("DELETE FROM units WHERE id=?", (unit_id,))
        db.commit()

        _log_activity("unit", unit_id, "deleted",
                      notes=f"Unit '{unit.get('unit_ref', '')}' permanently deleted",
                      db=db)

        return json_success({"id": unit_id, "message": "Unit permanently deleted"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/properties/<int:prop_id>/units/bulk", methods=["POST"])
def api_bulk_create_units(prop_id):
    """POST /api/banksia-os/properties/{prop_id}/units/bulk — bulk-create units."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    units_data = data.get("units", [])
    if not units_data or not isinstance(units_data, list):
        return json_error("'units' must be a non-empty array", 400)

    db = get_dict_db()
    try:
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        now = datetime.now(timezone.utc).isoformat()
        created_ids = []
        errors = []

        for idx, unit_data in enumerate(units_data):
            unit_ref = unit_data.get("unit_ref")
            unit_type = unit_data.get("unit_type")
            if not unit_ref or not unit_type:
                errors.append({"index": idx, "error": "'unit_ref' and 'unit_type' are required"})
                continue

            try:
                ins_cols = ["property_id", "unit_ref", "unit_type", "created", "modified"]
                ins_vals = [prop_id, unit_ref, unit_type, now, now]

                optional_fields = ["floor", "bedrooms", "capacity", "market_rent", "furnished", "status", "notes",
                                   "unit_status", "max_occupancy", "bathrooms"]
                for f in optional_fields:
                    if f in unit_data:
                        ins_cols.append(f)
                        ins_vals.append(unit_data[f])

                placeholders = ",".join(["?"] * len(ins_cols))
                cursor = db.execute(
                    f"INSERT INTO units ({','.join(ins_cols)}) VALUES ({placeholders})",
                    ins_vals
                )
                created_ids.append(cursor.lastrowid)
            except Exception as e:
                errors.append({"index": idx, "error": str(e)})

        db.commit()

        return json_success({
            "created_ids": created_ids,
            "count": len(created_ids),
            "errors": errors if errors else None,
        }), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 2h. PROPERTY IMAGES — list images for a property
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/<int:prop_id>/images", methods=["GET"])
def api_property_images(prop_id):
    db = get_dict_db()
    try:
        # Check property exists
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        # Return images with unit_ref joined in
        images = db.execute(
            "SELECT pi.id, pi.property_id, pi.unit_id, pi.image_url, pi.caption, "
            "pi.category, pi.sort_order, pi.created_at, "
            "u.unit_ref "
            "FROM property_images pi "
            "LEFT JOIN units u ON pi.unit_id = u.id "
            "WHERE pi.property_id = ? "
            "ORDER BY pi.sort_order ASC, pi.id ASC",
            (prop_id,)
        ).fetchall()

        result = []
        for img in images:
            d = dict(img)
            d["url"] = d.pop("image_url")
            d["thumbnail_url"] = d["url"]
            result.append(d)

        return json_success(result)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 3. UNITS
# ═══════════════════════════════════════════════

def _next_months(from_month, count):
    """The next `count` calendar month keys after `from_month` ("YYYY-MM")."""
    y, m = int(from_month[:4]), int(from_month[5:7])
    out = []
    for _ in range(count):
        m += 1
        if m > 12:
            m, y = 1, y + 1
        out.append("%04d-%02d" % (y, m))
    return out


@banksia_os_bp.route("/units", methods=["GET", "POST"])
def api_units():
    if request.method == "POST":
        return api_create_unit()
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    status_filter = request.args.get("status", "").strip()
    property_id = request.args.get("property_id", "").strip()
    search = request.args.get("search", "").strip()
    sort_by = request.args.get("sort_by", "").strip()
    sort_dir = request.args.get("sort_dir", "asc").strip().lower()

    # Allowed sort columns mapped to SQL expressions
    sortable_columns = {
        "unit_ref": "u.unit_ref",
        "property_name": "property_name",
        "unit_type": "u.unit_type",
        "unit_vacant": "u.unit_vacant",
        "tenant_name": "tenant_name",
        "tenancy_rent": "tenancy_rent",
        "tenancy_deposit": "tenancy_deposit",
        "tenancy_start_date": "tenancy_start_date",
        "tenancy_end_date": "tenancy_end_date",
    }

    where_parts = ["(u.is_active IS NULL OR u.is_active = 1)"]
    params = []

    # Category scoping (Norbert, 2026-08-03) — the units page is entered by category,
    # the same shape as Properties. Unlike the property categories these do NOT overlap:
    # a unit in a stood-down property is not "vacant" in any useful sense (you cannot let
    # it), so Inactive takes it out of both Occupied and Vacant rather than sitting
    # alongside them.
    #   occupied — someone is living there
    #   vacant   — empty and lettable, optionally narrowed to when it frees up
    #   inactive — the property it belongs to is inactive, so the unit is too
    category = request.args.get("category", "").strip().lower()
    vacant_when = request.args.get("vacant_when", "").strip().lower()

    _this_month = datetime.now(timezone.utc).strftime("%Y-%m")
    # Buckets start at the CURRENT month so a tenancy ending this month cannot fall
    # through the gap between "now" and the first future month.
    _upcoming_months = [_this_month] + _next_months(_this_month, 4)

    # "Vacant now" means EMPTY NOW, nothing else (Norbert, 2026-08-03, after 5 Radford
    # House E3 turned up here while Jessica Tamang was living in it).
    NOW_CLAUSE = "u.unit_vacant = 1"

    # A unit is "coming vacant in month X" only when a live tenancy is actually ENDING
    # then. available_from is not evidence on its own: it holds the date the unit was
    # last advertised, so E3 carried 2026-08-31 against an open-ended periodic tenancy.
    # It is used only as a fallback for a unit with no live tenancy at all, where there
    # is nothing better to go on. Takes the month twice.
    MONTH_CLAUSE = (
        "(COALESCE(u.unit_vacant, 0) = 0 AND ("
        # Bucketed by the day AFTER the tenancy ends — that is the day the room can be
        # let. A tenancy ending 31 October makes the room a NOVEMBER vacancy.
        # An APPROVED notice to quit is the better evidence when there is one: it
        # carries the agreed vacate date, which is the first day the room can be let.
        # Falls back to the day after the tenancy end date when there is no notice.
        "substr(COALESCE((SELECT COALESCE(n.move_out_date, date(t.end_date, '+1 day')) "
        "FROM tenancies t LEFT JOIN tenancy_notice n ON n.tenancy_id = t.id AND n.status = 'approved' "
        "WHERE t.unit_id = u.id "
        "AND LOWER(COALESCE(t.status,'')) IN ('current','active','periodic') "
        "ORDER BY t.start_date DESC LIMIT 1), ''), 1, 7) = ? "
        "OR (NOT EXISTS (SELECT 1 FROM tenancies t2 WHERE t2.unit_id = u.id "
        "AND LOWER(COALESCE(t2.status,'')) IN ('current','active','periodic')) "
        "AND substr(COALESCE(u.available_from, ''), 1, 7) = ?)))"
    )

    prop_inactive = ("EXISTS (SELECT 1 FROM properties p2 WHERE p2.id = u.property_id "
                     "AND COALESCE(p2.is_active, 1) = 0)")
    prop_archived = ("EXISTS (SELECT 1 FROM properties p3 WHERE p3.id = u.property_id "
                     "AND p3.status = 'archived')")

    if category:
        # Archived properties live on the Archive page, so their units never appear here.
        where_parts.append("NOT " + prop_archived)

    # The status column is now the answer, except that a stood-down property still
    # forces its units inactive whatever their own row says — that rule is Norbert's
    # and it has to survive someone flipping a single unit back to Occupied.
    if category == "all":
        pass  # every live unit, no further narrowing
    elif category == "inactive":
        where_parts.append("(%s OR LOWER(COALESCE(u.unit_status,'')) = 'inactive')" % prop_inactive)
    elif category == "occupied":
        where_parts.append("NOT " + prop_inactive)
        where_parts.append("LOWER(COALESCE(u.unit_status,'')) <> 'inactive'")
        where_parts.append("(u.unit_vacant = 0 OR u.unit_vacant IS NULL)")
    elif category == "vacant":
        where_parts.append("NOT " + prop_inactive)
        where_parts.append("LOWER(COALESCE(u.unit_status,'')) <> 'inactive'")
        if vacant_when and vacant_when != "now":
            # A named month: units still let today that free up in that month.
            # Already-empty units belong under Now, not under a future month.
            where_parts.append(MONTH_CLAUSE)
            params.extend([vacant_when[:7], vacant_when[:7]])
        elif vacant_when == "now":
            where_parts.append(NOW_CLAUSE)
        else:
            # No month chosen: the whole vacancy pipeline, which is exactly the union of
            # the buckets shown on screen — so the card count and this list always agree.
            month_ors = " OR ".join([MONTH_CLAUSE] * len(_upcoming_months))
            where_parts.append("(%s OR %s)" % (NOW_CLAUSE, month_ors))
            for _m in _upcoming_months:
                params.extend([_m, _m])

    if status_filter:
        if status_filter.lower() == 'vacant':
            where_parts.append("unit_vacant = 1")
        elif status_filter.lower() == 'occupied':
            where_parts.append("(unit_vacant = 0 OR unit_vacant IS NULL)")
        else:
            where_parts.append("unit_status = ?")
            params.append(status_filter)

    if property_id:
        try:
            where_parts.append("property_id = ?")
            params.append(int(property_id))
        except ValueError:
            pass

    if search:
        search_clause, search_params = build_search_clause(
            ["unit_ref", "full_address", "unit_type", "owner_name"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    where = " AND ".join(where_parts)

    # Build ORDER BY clause
    if sort_by in sortable_columns:
        direction = "DESC" if sort_dir == "desc" else "ASC"
        order_clause = f"{sortable_columns[sort_by]} {direction}"
    else:
        order_clause = "sort_order ASC, CAST(SUBSTR(unit_ref, 2) AS INTEGER) ASC, unit_ref ASC"

    rows, total = paginate(
        f"SELECT u.*, "
        f"(SELECT COALESCE(NULLIF(NULLIF(p.name,'multi'),'single'), p.address_line_1, p.name) FROM properties p WHERE p.id = u.property_id) AS property_name, "
        f"(SELECT t.main_tenant_name FROM tenancies t WHERE t.unit_id = u.id AND t.status IN ('Current','current','Periodic','periodic','Active','active') ORDER BY t.start_date DESC LIMIT 1) AS tenant_name, "
        f"(SELECT t.rent_amount FROM tenancies t WHERE t.unit_id = u.id AND t.status IN ('Current','current','Periodic','periodic','Active','active') ORDER BY t.start_date DESC LIMIT 1) AS tenancy_rent, "
        f"(SELECT t.deposit_registered_amount FROM tenancies t WHERE t.unit_id = u.id AND t.status IN ('Current','current','Periodic','periodic','Active','active') ORDER BY t.start_date DESC LIMIT 1) AS tenancy_deposit, "
        f"(SELECT t.start_date FROM tenancies t WHERE t.unit_id = u.id AND t.status IN ('Current','current','Periodic','periodic','Active','active') ORDER BY t.start_date DESC LIMIT 1) AS tenancy_start_date, "
        f"(SELECT t.end_date FROM tenancies t WHERE t.unit_id = u.id AND t.status IN ('Current','current','Periodic','periodic','Active','active') ORDER BY t.start_date DESC LIMIT 1) AS tenancy_end_date "
        f"FROM units u WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM units u WHERE {where}",
        params, page, per_page
    )

    # Convert unit_vacant to bool
    vac_count = 0
    occ_count = 0
    for r in rows:
        bool_fields(r, "unit_vacant")
        if r.get("unit_vacant"):
            vac_count += 1
        else:
            occ_count += 1

    # Get real totals from DB (not just the current page).
    # NOTE: the params must be passed — this counted with a bare WHERE before, which
    # threw as soon as any filter carried a placeholder.
    db2 = get_dict_db()
    try:
        real_total = db2.execute(f"SELECT COUNT(*) AS cnt FROM units u WHERE {where}", params).fetchone()["cnt"]
        real_vacant = db2.execute(f"SELECT COUNT(*) AS cnt FROM units u WHERE {where} AND unit_vacant = 1", params).fetchone()["cnt"]
        real_occupied = real_total - real_vacant

        # ── Card counts for the category landing ──
        # Unfiltered by category on purpose: the cards say how big each category is, so
        # they must not change according to which one you are standing in.
        live = ("(u.is_active IS NULL OR u.is_active = 1) "
                "AND NOT EXISTS (SELECT 1 FROM properties p3 WHERE p3.id = u.property_id AND p3.status = 'archived')")
        active_prop = ("NOT EXISTS (SELECT 1 FROM properties p2 WHERE p2.id = u.property_id "
                       "AND COALESCE(p2.is_active, 1) = 0) "
                       "AND LOWER(COALESCE(u.unit_status,'')) <> 'inactive'")
        cnt = lambda clause, args=(): db2.execute(
            f"SELECT COUNT(*) AS cnt FROM units u WHERE {live} AND {clause}", args).fetchone()["cnt"]

        cat_occupied = cnt(f"{active_prop} AND (u.unit_vacant = 0 OR u.unit_vacant IS NULL)")
        cat_inactive = cnt("(EXISTS (SELECT 1 FROM properties p2 WHERE p2.id = u.property_id "
                           "AND COALESCE(p2.is_active, 1) = 0) "
                           "OR LOWER(COALESCE(u.unit_status,'')) = 'inactive')")

        # The four months roll forward from today rather than being pinned, so the board
        # cannot go stale. Today that renders exactly as Norbert listed it:
        # September / October / November / December 2026.
        buckets = [{"key": "now", "label": "Now",
                    "count": cnt(f"{active_prop} AND {NOW_CLAUSE}")}]
        for key in _upcoming_months:
            buckets.append({
                "key": key,
                "label": datetime(int(key[:4]), int(key[5:7]), 1).strftime("%B %Y"),
                "count": cnt(f"{active_prop} AND {MONTH_CLAUSE}", (key, key)),
            })
        cat_vacant = sum(b["count"] for b in buckets)
        # Every live unit in one list, the way the page worked before it gained
        # categories — the categories answer different questions and do not have
        # to add up to this.
        cat_all = cnt("1 = 1")
    finally:
        db2.close()

    return json_success({
        "items": rows,
        "totals": {"total": real_total, "occupied": real_occupied, "vacant": real_vacant,
                   "cat_occupied": cat_occupied, "cat_vacant": cat_vacant,
                   "cat_inactive": cat_inactive, "cat_all": cat_all,
                   "vacancy_buckets": buckets},
    }, total, page, per_page)


def api_create_unit():
    """POST handler for creating a new unit with room/fixture details."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")
    if not data.get("property_id"):
        return json_error("'property_id' is required")

    db = get_dict_db()
    try:
        cols = ["property_id", "unit_ref", "unit_type", "unit_status", "unit_vacant",
                "full_address", "market_rent", "market_rent_frequency", "deposit_amount",
                "short_description", "furnished", "bedrooms", "bathrooms", "max_occupancy",
                "council_tax_band", "features", "owner_name", "notes"]
        ins_cols = [c for c in cols if c in data]
        ins_vals = [data[c] for c in ins_cols]
        placeholders = ",".join(["?"] * len(ins_cols))
        cursor = db.execute(
            f"INSERT INTO units ({','.join(ins_cols)}) VALUES ({placeholders})",
            ins_vals
        )
        db.commit()
        new_id = cursor.lastrowid
        user_info = session.get("user", {})
        user_name = user_info.get("username", "Unknown")
        record_change(user_name, 'created', 'unit', str(new_id), data.get("unit_ref", ""))
        return json_success({"id": new_id, "message": "Unit created"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/units/<int:unit_id>", methods=["GET", "PATCH"])
def api_unit(unit_id):
    if request.method == "PATCH":
        return api_update_resource("units", unit_id)
    db = get_dict_db()
    try:
        unit = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        bool_fields(unit, "unit_vacant")

        # Linked tenancy (current active)
        tenancy = db.execute(
            "SELECT * FROM tenancies WHERE unit_id = ? AND status IN ('Active', 'active', 'Periodic', 'periodic') "
            "ORDER BY start_date DESC LIMIT 1",
            (unit_id,)
        ).fetchone()
        unit["current_tenancy"] = tenancy

        # Current tenant(s) for the unit
        if tenancy:
            tenants = db.execute(
                "SELECT * FROM tenants WHERE tenancy_id = ? ORDER BY main_tenant DESC",
                (tenancy["id"],)
            ).fetchall()
            unit["current_tenants"] = tenants
        else:
            unit["current_tenants"] = []

        return json_success(unit)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 4. TENANCIES
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/tenancies")
def api_tenancies():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    status_filter = request.args.get("status", "").strip()
    search = request.args.get("search", "").strip()

    where_parts = ["1=1"]
    params = []

    if status_filter:
        if status_filter.lower() == 'all':
            pass  # No status filter — show everything including past
        else:
            where_parts.append("status = ?")
            params.append(status_filter)
    else:
        # Default: live tenancies plus prospective ones. A signed application creates
        # a Prospective tenancy (see tenant_application_api._create_prospective_tenancy),
        # and hiding those by default would mean the hand-off the team is meant to act
        # on never appears on this page.
        where_parts.append(
            "status IN ('Current', 'current', 'Periodic', 'periodic', 'Active', 'active', "
            "'Prospective', 'prospective')"
        )

    if search:
        search_clause, search_params = build_search_clause(
            ["ref", "full_address", "main_tenant_name"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    where = " AND ".join(where_parts)

    order_clause = build_order_by({
        "ref": "t.ref", "full_address": "t.full_address",
        "main_tenant_name": "t.main_tenant_name", "status": "t.status",
        "start_date": "t.start_date", "end_date": "t.end_date",
        "rent_amount": "t.rent_amount", "property_name": "property_name",
    }, "t.start_date DESC")

    rows, total = paginate(
        f"SELECT t.*, "
        f"(SELECT COALESCE(NULLIF(p.name, 'multi'), p.address_line_1) FROM properties p WHERE p.id = t.property_id) AS property_name, "
        f"(SELECT p.address_line_1 FROM properties p WHERE p.id = t.property_id) AS property_address, "
        f"(SELECT u.unit_ref FROM units u WHERE u.id = t.unit_id) AS unit_ref, "
        f"(SELECT u.unit_type FROM units u WHERE u.id = t.unit_id) AS unit_type_name, "
        f"t.deposit_registered_amount AS deposit_amount, "
        f"(SELECT n.status FROM tenancy_notice n WHERE n.tenancy_id = t.id) AS notice_status, "
        f"(SELECT n.move_out_date FROM tenancy_notice n WHERE n.tenancy_id = t.id) AS notice_move_out "
        f"FROM tenancies t WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM tenancies t WHERE {where}",
        params, page, per_page
    )

    for r in rows:
        bool_fields(r, "deposit_registered", "section_21_served", "is_renewed")
        # Only an APPROVED notice earns the tag — a pending or declined one must not
        # read as though the room is going.
        r["notice_tag"] = None
        if r.get("notice_status") == "approved" and r.get("notice_move_out"):
            try:
                r["notice_tag"] = "Notice (%s)" % datetime.strptime(
                    r["notice_move_out"][:10], "%Y-%m-%d").strftime("%d/%m/%Y")
            except (ValueError, TypeError):
                r["notice_tag"] = "Notice"

    return json_success(rows, total, page, per_page)


@banksia_os_bp.route("/tenancies/<int:ten_id>", methods=["GET", "PATCH"])
def api_tenancy(ten_id):
    if request.method == "PATCH":
        # An end date is a consequence of an approved notice, never a free field
        # (Norbert, 2026-08-03). Enforced here rather than in the form so it holds
        # however the change arrives.
        _data = request.get_json(silent=True) or {}
        if "end_date" in _data:
            _db = get_dict_db()
            try:
                _cur = _db.execute("SELECT end_date FROM tenancies WHERE id = ?", (ten_id,)).fetchone()
                if not _cur:
                    return json_error("Tenancy not found", 404)
                _new = (str(_data.get("end_date") or "").strip())[:10]
                _old = (str(dict(_cur).get("end_date") or "").strip())[:10]
                if _new and _new != _old:
                    _agreed = _approved_move_out(_db, ten_id)
                    if not _agreed:
                        return json_error(
                            "An end date can only be set once the tenant has served notice and "
                            "that notice has been approved. Record it on the Notice tab first.", 400)
                    if _new != str(_agreed)[:10]:
                        return json_error(
                            "The end date must match the approved notice, which gives up the room on "
                            "%s. Change the notice date if that is wrong." % str(_agreed)[:10], 400)
            finally:
                _db.close()
        return api_update_resource("tenancies", ten_id)
    db = get_dict_db()
    try:
        ten = db.execute("SELECT * FROM tenancies WHERE id = ?", (ten_id,)).fetchone()
        if not ten:
            return json_error("Tenancy not found", 404)

        bool_fields(ten, "deposit_registered", "section_21_served", "is_renewed")

        # Tenant info — from tenants table (NOT JSON string), with full detail
        tenant_rows = db.execute(
            "SELECT id, first_name, last_name, email, mobile, phone_home, phone_work, "
            "date_of_birth, gender, citizen, ni_number, passport_number, "
            "main_tenant, status, has_guarantor, "
            "guarantor_first_name, guarantor_last_name, guarantor_email, guarantor_mobile, "
            "employment_company, employment_salary, student_status, university, "
            "move_in_date, move_out_date, applicant_note, manager_note, "
            "created, modified "
            "FROM tenants WHERE tenancy_id = ? ORDER BY main_tenant DESC",
            (ten_id,)
        ).fetchall()
        for t in tenant_rows:
            bool_fields(t, "main_tenant", "has_guarantor")
        ten["tenants"] = tenant_rows

        # Transactions
        transactions = db.execute(
            "SELECT * FROM transactions WHERE tenancy_id = ? ORDER BY date DESC",
            (ten_id,)
        ).fetchall()
        for t in transactions:
            bool_fields(t, "is_overdue", "is_outstanding")
        ten["transactions"] = transactions

        # Deposits (from authoritative deposits table)
        deposit_rows = db.execute(
            "SELECT * FROM deposits WHERE tenancy_id = ? ORDER BY created DESC",
            (ten_id,)
        ).fetchall()
        ten["deposits"] = deposit_rows

        # Rent charges
        rent_charge_rows = db.execute(
            "SELECT * FROM rent_charges WHERE tenancy_id = ? ORDER BY month DESC",
            (ten_id,)
        ).fetchall()
        ten["rent_charges"] = rent_charge_rows

        # Invoices
        invoice_rows = db.execute(
            "SELECT * FROM invoices WHERE tenancy_id = ? ORDER BY due_date DESC",
            (ten_id,)
        ).fetchall()
        ten["invoices"] = invoice_rows

        # Property info — with address display name
        if ten.get("property_id"):
            prop = db.execute(
                "SELECT id, ref, name, address_line_1, address_line_2, city, postcode, property_type FROM properties WHERE id = ?",
                (ten["property_id"],)
            ).fetchone()
            if prop:
                # Use address_line_1 if name is 'multi'
                display_name = prop["address_line_1"] if prop["name"] == "multi" else prop["name"]
                ten["property"] = {**prop, "display_name": display_name}

        # Unit info
        if ten.get("unit_id"):
            unit = db.execute(
                "SELECT id, unit_ref, unit_type, full_address FROM units WHERE id = ?",
                (ten["unit_id"],)
            ).fetchone()
            ten["unit"] = unit

        # Linked maintenance jobs (by property_id)
        maintenance_jobs = db.execute(
            "SELECT id, reference, title, status, priority, type AS category, created, "
            "contractor AS assigned_to, total_cost, "
            "(SELECT COUNT(*) FROM ll_communications WHERE job_id = maintenance_jobs.id) AS ll_comms_count "
            "FROM maintenance_jobs "
            "WHERE property_id = ? "
            "ORDER BY created DESC LIMIT 20",
            (ten.get("property_id"),)
        ).fetchall()
        ten["maintenance_jobs"] = maintenance_jobs

        # Notice to quit. Lives in its own table rather than tenancies.tags, which
        # the Arthur sync rewrites wholesale on every pull.
        ten["notice"] = _notice_payload(_get_notice(db, ten_id))

        return json_success(ten)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/ending-soon")
def api_tenancies_ending_soon():
    """Tenancies ending within 30 days."""
    from datetime import date, timedelta
    today = date.today()
    end_date = today + timedelta(days=30)

    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT * FROM tenancies WHERE end_date >= ? AND end_date <= ? "
            "AND status IN ('Active', 'active', 'Periodic', 'periodic') "
            "ORDER BY end_date ASC",
            (today.isoformat(), end_date.isoformat())
        ).fetchall()

        for r in rows:
            bool_fields(r, "deposit_registered", "section_21_served", "is_renewed")

        return json_success(rows, total=len(rows), page=1, per_page=len(rows))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/moving-in-this-month")
def api_tenancies_moving_in():
    db = get_dict_db()
    try:
        now = datetime.now(timezone.utc)
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
        if now.month == 12:
            next_month = now.replace(year=now.year + 1, month=1, day=1)
        else:
            next_month = now.replace(month=now.month + 1, day=1)
        month_end = next_month.isoformat()

        rows = db.execute(
            "SELECT t.*, "
            "(SELECT COALESCE(NULLIF(p.name, 'multi'), p.address_line_1) FROM properties p WHERE p.id = t.property_id) AS property_name, "
            "(SELECT p.address_line_1 FROM properties p WHERE p.id = t.property_id) AS property_address, "
            "(SELECT u.unit_ref FROM units u WHERE u.id = t.unit_id) AS unit_ref, "
            "(SELECT u.unit_type FROM units u WHERE u.id = t.unit_id) AS unit_type_name, "
            "t.deposit_registered_amount AS deposit_amount "
            "FROM tenancies t WHERE move_in_date >= ? AND move_in_date < ? "
            "ORDER BY move_in_date ASC",
            (month_start, month_end)
        ).fetchall()

        for r in rows:
            bool_fields(r, "deposit_registered", "section_21_served", "is_renewed")

        return json_success(rows, total=len(rows), page=1, per_page=len(rows))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/moving-out-this-month")
def api_tenancies_moving_out():
    db = get_dict_db()
    try:
        now = datetime.now(timezone.utc)
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
        if now.month == 12:
            next_month = now.replace(year=now.year + 1, month=1, day=1)
        else:
            next_month = now.replace(month=now.month + 1, day=1)
        month_end = next_month.isoformat()

        rows = db.execute(
            "SELECT t.*, "
            "(SELECT COALESCE(NULLIF(p.name, 'multi'), p.address_line_1) FROM properties p WHERE p.id = t.property_id) AS property_name, "
            "(SELECT p.address_line_1 FROM properties p WHERE p.id = t.property_id) AS property_address, "
            "(SELECT u.unit_ref FROM units u WHERE u.id = t.unit_id) AS unit_ref, "
            "(SELECT u.unit_type FROM units u WHERE u.id = t.unit_id) AS unit_type_name, "
            "t.deposit_registered_amount AS deposit_amount "
            "FROM tenancies t WHERE move_out_date >= ? AND move_out_date < ? "
            "ORDER BY move_out_date ASC",
            (month_start, month_end)
        ).fetchall()

        for r in rows:
            bool_fields(r, "deposit_registered", "section_21_served", "is_renewed")

        return json_success(rows, total=len(rows), page=1, per_page=len(rows))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 4a. HMO OCCUPANCY / PROPERTY TENANCY OVERVIEW
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/tenancies/property/<int:property_id>")
def api_tenancies_by_property(property_id):
    """All active tenancies + vacant units for a property. HMO room-level overview."""
    db = get_dict_db()
    try:
        active_statuses = ("'Current', 'current', 'Periodic', 'periodic', 'Active', 'active'")

        # Get property
        prop = db.execute(
            "SELECT id, name, address_line_1, address_line_2, city, postcode, property_type, "
            "total_units, rentable_units, max_occupancy, bathrooms, bedrooms "
            "FROM properties WHERE id = ?", (property_id,)
        ).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        display_name = prop["address_line_1"] if prop["name"] == "multi" else prop["name"]
        prop_data = dict(prop)
        prop_data["display_name"] = display_name

        # All units at this property
        units = db.execute(
            "SELECT id, unit_ref, unit_type, unit_status, unit_vacant, max_occupancy, "
            "market_rent, deposit_amount, short_description, features "
            "FROM units WHERE property_id = ? ORDER BY sort_order ASC, unit_ref", (property_id,)
        ).fetchall()

        # Active tenancies with tenant info
        tenancies = db.execute(f"""
            SELECT t.id, t.ref, t.unit_id, t.main_tenant_name, t.rent_amount, t.rent_frequency,
                   t.status, t.start_date, t.end_date, t.move_in_date, t.move_out_date,
                   t.deposit_registered_amount, t.deposit_registered, t.notice_period,
                   t.break_clause_date, t.section_21_served
            FROM tenancies t
            WHERE t.property_id = ? AND t.status IN ({active_statuses})
            ORDER BY t.unit_id, t.start_date DESC
        """, (property_id,)).fetchall()

        for t in tenancies:
            bool_fields(t, "deposit_registered", "section_21_served")
            # Get tenants for this tenancy
            t["tenant_list"] = db.execute(
                "SELECT id, first_name, last_name, email, mobile, main_tenant, status, "
                "date_of_birth, employment_company, student_status "
                "FROM tenants WHERE tenancy_id = ? ORDER BY main_tenant DESC",
                (t["id"],)
            ).fetchall()

        # Enrich units with active tenancy info
        unit_data = []
        for u in units:
            unit_dict = dict(u)
            tenancy = next((t for t in tenancies if t["unit_id"] == u["id"]), None)
            unit_dict["active_tenancy"] = tenancy
            unit_dict["occupied"] = tenancy is not None
            unit_data.append(unit_dict)

        # Vacant unit count
        occupied = sum(1 for u in unit_data if u["occupied"])
        vacant = sum(1 for u in unit_data if not u["occupied"])

        # Summary
        total_rent = sum(t["rent_amount"] or 0 for t in tenancies)

        return json_success({
            "property": prop_data,
            "units": unit_data,
            "tenancies": tenancies,
            "summary": {
                "total_units": len(unit_data),
                "occupied": occupied,
                "vacant": vacant,
                "occupancy_pct": round(occupied / len(unit_data) * 100, 1) if unit_data else 0,
                "active_tenancies": len(tenancies),
                "total_monthly_rent": round(total_rent, 2),
            }
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 4b. SHORT-ALIAS ROUTES (flat paths for Next.js frontend)
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/financials")
def api_financials():
    db = get_dict_db()
    try:
        # Active statuses matching dashboard logic
        active_statuses = ("'Current', 'current', 'Periodic', 'periodic', 'Active', 'active'")
        # Total monthly rent
        rent_row = db.execute(
            f"SELECT COALESCE(SUM(rent_amount),0) as total FROM tenancies WHERE status IN ({active_statuses})"
        ).fetchone()
        # Arrears — from transactions outstanding (same as dashboard)
        arrears_row = db.execute(
            "SELECT COALESCE(SUM(amount_outstanding), 0) AS total FROM transactions WHERE is_outstanding = 1"
        ).fetchone()
        # Count tenancies with outstanding transactions
        arrears_count = db.execute(
            "SELECT COUNT(DISTINCT t.id) as cnt FROM tenancies t WHERE EXISTS (SELECT 1 FROM transactions tx WHERE tx.tenancy_id = t.arthur_id AND tx.is_outstanding = 1 AND tx.amount_outstanding > 0)"
        ).fetchone()
        # Deposits
        dep_count = db.execute("SELECT COUNT(*) as c FROM tenancies WHERE deposit_registered = 1").fetchone()
        dep_total = db.execute("SELECT COALESCE(SUM(deposit_registered_amount),0) as t FROM tenancies WHERE deposit_registered = 1").fetchone()
        dep_unreg = db.execute("SELECT COUNT(*) as c FROM tenancies WHERE deposit_registered = 0 AND deposit_registered IS NOT NULL").fetchone()
        
        # Active tenancy count
        active_count = db.execute(
            f"SELECT COUNT(*) as cnt FROM tenancies WHERE status IN ({active_statuses})"
        ).fetchone()
        total_count = db.execute("SELECT COUNT(*) as cnt FROM tenancies").fetchone()
        
        return json_success({
            "monthly_rent_income": rent_row["total"],
            "monthly_rent_roll": rent_row["total"],
            "total_arrears": arrears_row["total"],
            "total_deposits": dep_total["t"],
            "total_deposits_held": dep_total["t"],
            "deposits_registered": dep_count["c"],
            "deposits_total": dep_total["t"],
            "deposits_unregistered": dep_unreg["c"],
            "tenancies_in_arrears_count": arrears_count["cnt"],
            "unit_occupancy_rate": round(active_count["cnt"] / total_count["cnt"] * 100, 1) if total_count["cnt"] else 0,
            "payment_dates": [],
            "rent_collected": 0,
            "rent_outstanding": 0,
            "metrics": {
                "total_tenancies": total_count["cnt"],
                "active_tenancies": active_count["cnt"],
                "in_arrears": arrears_count["cnt"]
            }
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/arrears")
def api_arrears():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    db = get_dict_db()
    try:
        # Get tenancies with outstanding transactions (tenancy_id in transactions
        # maps to arthur_id in tenancies, not the local id)
        rows, total = paginate(
            "SELECT t.id, t.ref, t.full_address, t.main_tenant_name, t.rent_amount, t.rent_frequency, "
            "COALESCE((SELECT SUM(amount_outstanding) FROM transactions tx WHERE tx.tenancy_id = t.arthur_id AND tx.is_outstanding = 1), 0) as arrears_amount, "
            "t.status, (SELECT p.name FROM properties p WHERE p.id = t.property_id) AS property_name "
            "FROM tenancies t WHERE EXISTS (SELECT 1 FROM transactions tx WHERE tx.tenancy_id = t.arthur_id AND tx.is_outstanding = 1 AND tx.amount_outstanding > 0) "
            "ORDER BY arrears_amount DESC",
            "SELECT COUNT(DISTINCT t.id) as cnt FROM tenancies t WHERE EXISTS (SELECT 1 FROM transactions tx WHERE tx.tenancy_id = t.arthur_id AND tx.is_outstanding = 1 AND tx.amount_outstanding > 0)",
            [], page, per_page
        )
        return json_success(rows, total, page, per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()



@banksia_os_bp.route("/maintenance")
def api_maintenance_list():
    db = get_dict_db()
    try:
        page = int_param(request.args.get("page"))
        per_page = int_param(request.args.get("per_page"), 50, max_val=MAX_PAGE_SIZE)
        search = (request.args.get("search") or "").strip()
        status_filter = request.args.get("status", "")
        type_filter = request.args.get("type", "")
        priority_filter = request.args.get("priority", "")
        contractor_filter = request.args.get("contractor", "")

        where = ["1=1"]
        params = []

        if status_filter:
            where.append("mj.status = ?")
            params.append(status_filter)
        if type_filter:
            where.append("mj.type = ?")
            params.append(type_filter)
        if priority_filter:
            where.append("mj.priority = ?")
            params.append(priority_filter)
        if contractor_filter:
            where.append("mj.contractor = ?")
            params.append(contractor_filter)
        if search:
            where.append("(mj.title LIKE ? OR mj.description LIKE ? OR mj.address LIKE ? OR mj.reference LIKE ? OR mj.contractor LIKE ? OR mj.type LIKE ? OR mj.reporter_name LIKE ? OR mj.team_notes LIKE ?)")
            s = f"%{search}%"
            params.extend([s, s, s, s, s, s, s, s])

        where_clause = " AND ".join(where)

        total = db.execute(
            f"SELECT COUNT(*) AS cnt FROM maintenance_jobs mj WHERE {where_clause}",
            params
        ).fetchone()["cnt"]

        offset = (page - 1) * per_page
        rows = db.execute(
            f"""SELECT mj.*, COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN '' ELSE p.name END, ''), p.address_line_1, p.ref, p.name) AS property_name,
                       p.management_type AS property_management_type
                FROM maintenance_jobs mj
                LEFT JOIN properties p ON mj.property_id = p.id
                WHERE {where_clause}
                ORDER BY
                    CASE mj.priority
                        WHEN 'Emergency' THEN 0 WHEN 'Critical' THEN 1
                        WHEN 'High' THEN 2 WHEN 'Medium' THEN 3 WHEN 'Low' THEN 4
                    END,
                    mj.created DESC
                LIMIT ? OFFSET ?""",
            params + [per_page, offset]
        ).fetchall()

        for r in rows:
            r["bill_ll"] = bool(r["bill_ll"])
            r["emergency"] = bool(r["emergency"])
            r["ll_informed"] = bool(r["ll_informed"])
            o = db.execute(
                "SELECT COUNT(*) AS cnt FROM maintenance_orders WHERE job_id = ?",
                [r["id"]]
            ).fetchone()
            r["order_count"] = o["cnt"] if o else 0

        counts = {}
        for s in MAINT_STATUSES:
            c = db.execute("SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE status = ?", [s]).fetchone()
            counts[s] = c["cnt"] if c else 0

        return json_success(rows, total=total, page=page, per_page=per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/<int:job_id>")
def api_maintenance_detail(job_id):
    db = get_dict_db()
    try:
        job = db.execute(
            """SELECT mj.*, COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN '' ELSE p.name END, ''), p.address_line_1, p.ref, p.name) AS property_name
               FROM maintenance_jobs mj
               LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.id = ?""",
            [job_id]
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)
        orders = db.execute(
            "SELECT * FROM maintenance_orders WHERE job_id = ? ORDER BY created DESC",
            [job_id]
        ).fetchall()
        ll_comms = db.execute(
            "SELECT * FROM ll_communications WHERE job_id = ? ORDER BY sent_at DESC",
            [job_id]
        ).fetchall()
        result = dict(job)
        result["orders"] = [dict(o) for o in orders]
        result["ll_comms"] = [dict(c) for c in ll_comms]
        result["bill_ll"] = bool(result["bill_ll"])
        result["emergency"] = bool(result["emergency"])
        result["ll_informed"] = bool(result["ll_informed"])
        return json_success(result)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/<int:job_id>/emails")
def api_maintenance_job_emails(job_id):
    """Return Missive emails related to a maintenance job by matching property address."""
    import subprocess, json
    db = get_dict_db()
    try:
        job = db.execute(
            """SELECT mj.*, p.name as property_name, p.address_line_1, p.address_line_2,
                      p.city, p.postcode, p.property_owner_name
               FROM maintenance_jobs mj
               LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.id = ?""",
            [job_id]
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)

        # Call the fetch_job_emails helper script
        script_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts", "fetch_job_emails.py")
        if os.path.exists(script_path):
            try:
                result = subprocess.run(
                    [sys.executable, script_path, str(job_id)],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout:
                    data = json.loads(result.stdout)
                    if data.get("success"):
                        return json_success(data["data"])
                    else:
                        # Script returned error — degrade to empty
                        return json_success({
                            "emails": [],
                            "property_owner": job.get("property_owner_name"),
                            "property_address": job.get("address_line_1"),
                            "total_matched": 0,
                            "error": data.get("error")
                        })
                else:
                    # Script failed — degrade gracefully
                    return json_success({
                        "emails": [],
                        "property_owner": job.get("property_owner_name"),
                        "property_address": job.get("address_line_1"),
                        "total_matched": 0,
                        "error": result.stderr[:200] if result.stderr else "Script execution failed"
                    })
            except (subprocess.TimeoutExpired, Exception) as e:
                return json_success({
                    "emails": [],
                    "property_owner": job.get("property_owner_name"),
                    "property_address": job.get("address_line_1"),
                    "total_matched": 0,
                    "error": str(e)[:200]
                })
        else:
            return json_success({
                "emails": [],
                "property_owner": job.get("property_owner_name"),
                "property_address": job.get("address_line_1"),
                "total_matched": 0
            })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Order Emails (Missive Orders inbox) ─────────────────────────
@banksia_os_bp.route("/maintenance/<int:job_id>/order-emails")
def api_maintenance_job_order_emails(job_id):
    """Return Missive Orders inbox emails matched to this job."""
    import subprocess, json
    script_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts", "fetch_job_orders.py")
    try:
        result = subprocess.run(
            [sys.executable, script_path, str(job_id)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            return json_success({"orders": [], "total_matched": 0, "error": result.stderr.strip()})
        parsed = json.loads(result.stdout)
        if not parsed.get("success"):
            return json_success({"orders": [], "total_matched": 0, "error": parsed.get("error", "Script failed")})
        return json_success(parsed["data"])
    except subprocess.TimeoutExpired:
        return json_success({"orders": [], "total_matched": 0, "error": "Timed out"})
    except Exception as e:
        return json_success({"orders": [], "total_matched": 0, "error": str(e)})


# ═══════════════════════════════════════════════
#  CONVERSATION TIMELINE — WhatsApp contractor chats per job
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/maintenance/<int:job_id>/conversations")
def api_maintenance_job_conversations(job_id):
    """Return WhatsApp group conversations related to a maintenance job."""
    db = get_dict_db()
    try:
        # Check job exists
        job = db.execute(
            "SELECT id, reference, title, contractor, address FROM maintenance_jobs WHERE id = ?",
            (job_id,)
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)

        # Get conversation timeline entries
        rows = db.execute("""
            SELECT id, sender_name, body, message_timestamp, source_group_name,
                   linked_contractor
            FROM conversation_timeline
            WHERE job_id = ?
            ORDER BY message_timestamp ASC
        """, (job_id,)).fetchall()

        conversations = []
        for r in rows:
            conversations.append({
                "id": r["id"],
                "sender": r["sender_name"],
                "body": r["body"],
                "timestamp": r["message_timestamp"],
                "source_group": r["source_group_name"],
                "contractor": r["linked_contractor"],
            })

        return json_success({
            "job": {
                "id": job["id"],
                "reference": job["reference"],
                "title": job["title"],
                "contractor": job["contractor"],
                "address": job["address"],
            },
            "conversations": conversations,
            "count": len(conversations),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/orders/scan")
def api_scan_orders_inbox():
    """Scan Orders inbox and return all non-marketing order emails."""
    import subprocess, json
    script_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts", "fetch_job_orders.py")
    try:
        result = subprocess.run(
            [sys.executable, script_path, "0"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            return json_success({"orders": [], "total_matched": 0, "error": result.stderr.strip()})
        parsed = json.loads(result.stdout)
        if not parsed.get("success"):
            return json_success({"orders": [], "total_matched": 0, "error": parsed.get("error", "Script failed")})
        return json_success(parsed["data"])
    except subprocess.TimeoutExpired:
        return json_success({"orders": [], "total_matched": 0, "error": "Timed out"})
    except Exception as e:
        return json_success({"orders": [], "total_matched": 0, "error": str(e)})


@banksia_os_bp.route("/activity")
def api_activity():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 30, max_val=MAX_PAGE_SIZE)
    db = get_dict_db()
    try:
        rows = db.execute("""
            SELECT 'tenant' as type, id, first_name || ' ' || last_name as title, modified as timestamp, 'Modified' as action FROM tenants WHERE modified IS NOT NULL
            UNION ALL
            SELECT 'tenancy' as type, id, ref as title, modified as timestamp, 'Modified' as action FROM tenancies WHERE modified IS NOT NULL
            UNION ALL
            SELECT 'applicant' as type, id, first_name || ' ' || last_name as title, modified as timestamp, 'Modified' as action FROM applicants WHERE modified IS NOT NULL
            ORDER BY timestamp DESC LIMIT ? OFFSET ?
        """, [per_page, (page-1)*per_page]).fetchall()
        total = db.execute("SELECT COUNT(*) as cnt FROM (SELECT modified FROM tenants WHERE modified IS NOT NULL UNION ALL SELECT modified FROM tenancies WHERE modified IS NOT NULL UNION ALL SELECT modified FROM applicants WHERE modified IS NOT NULL)").fetchone()["cnt"]
        return json_success(rows, total, page, per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 5. TENANTS
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/tenants")
def api_tenants():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    search = request.args.get("search", "").strip()
    property_id = request.args.get("property_id", "").strip()
    status_filter = request.args.get("status", "").strip()
    has_email = request.args.get("has_email", "").strip()
    has_phone = request.args.get("has_phone", "").strip()

    where_parts = ["1=1"]
    params = []

    if search:
        search_clause, search_params = build_search_clause(
            ["first_name", "last_name", "email", "mobile", "full_address"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    if property_id:
        try:
            where_parts.append("property_id = ?")
            params.append(int(property_id))
        except ValueError:
            pass

    if status_filter:
        where_parts.append("LOWER(COALESCE(status,'')) LIKE ?")
        params.append(f"%{status_filter.lower()}%")

    if has_email:
        where_parts.append("email IS NOT NULL AND email != ''")

    if has_phone:
        where_parts.append("(mobile IS NOT NULL AND mobile != '') OR (phone_home IS NOT NULL AND phone_home != '')")

    where = " AND ".join(where_parts)

    order_clause = build_order_by({
        "last_name": "tn.last_name", "first_name": "tn.first_name",
        "email": "tn.email", "status": "tn.status",
        "property_name": "property_name", "rent_amount": "rent_amount",
        "arrears": "arrears", "created": "tn.created",
    }, "tn.last_name ASC, tn.first_name ASC")

    rows, total = paginate(
        f"SELECT tn.id, tn.arthur_id, tn.arthur_person_id, tn.tenancy_id, tn.unit_id, tn.property_id, "
        f"tn.full_address, tn.title, tn.first_name, tn.last_name, tn.date_of_birth, tn.gender, tn.citizen, "
        f"tn.email, tn.phone_home, tn.phone_work, tn.mobile AS phone, tn.passport_number, tn.visa_number, tn.visa_type, "
        f"tn.visa_years, tn.country_of_origin, tn.ni_number, tn.main_tenant, tn.status, tn.has_guarantor, "
        f"tn.guarantor_first_name, tn.guarantor_last_name, tn.guarantor_email, tn.guarantor_mobile, "
        f"tn.employment_company, tn.student_status, tn.university, "
        f"tn.bank_name, tn.latest_credit_score, tn.latest_credit_description, "
        f"tn.applicant_note, tn.manager_note, tn.move_in_date, tn.move_out_date, tn.modified, tn.created, "
        f"COALESCE((SELECT COUNT(*) FROM tenancies t2 WHERE t2.id = tn.tenancy_id), 0) AS tenancy_count, "
        f"COALESCE((SELECT COALESCE(NULLIF(p2.name, 'multi'), p2.address_line_1) FROM properties p2 WHERE p2.id = tn.property_id), '') AS property_name, "
        f"COALESCE((SELECT u2.unit_ref FROM units u2 WHERE u2.id = tn.unit_id), '') AS unit_ref, "
        f"COALESCE((SELECT t2.rent_amount FROM tenancies t2 WHERE t2.id = tn.tenancy_id LIMIT 1), 0) AS rent_amount, "
        f"COALESCE((SELECT SUM(x.amount_outstanding) FROM transactions x WHERE x.tenancy_id = tn.tenancy_id AND x.is_outstanding = 1), 0) AS arrears "
        f"FROM tenants tn WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM tenants WHERE {where}",
        params, page, per_page
    )

    for r in rows:
        bool_fields(r, "main_tenant", "has_guarantor")

    return json_success(rows, total, page, per_page)


@banksia_os_bp.route("/tenants/<int:tenant_id>", methods=["GET", "PATCH"])
def api_tenant(tenant_id):
    if request.method == "PATCH":
        return api_update_resource("tenants", tenant_id)
    db = get_dict_db()
    try:
        tenant = db.execute("SELECT * FROM tenants WHERE id = ?", (tenant_id,)).fetchone()
        if not tenant:
            return json_error("Tenant not found", 404)

        bool_fields(tenant, "main_tenant", "has_guarantor", "guarantor_home_owner")

        # Linked tenancy
        if tenant.get("tenancy_id"):
            tenancy = db.execute("SELECT * FROM tenancies WHERE id = ?", (tenant["tenancy_id"],)).fetchone()
            if tenancy:
                bool_fields(tenancy, "deposit_registered", "section_21_served", "is_renewed")
            tenant["tenancy"] = tenancy
        else:
            tenant["tenancy"] = None

        # Linked property — robust path first: tenant -> tenancy.property_id -> properties.id
        # (all clean integer FKs). Fall back to the legacy Arthur-ID text match only for
        # tenants that have no linked tenancy, so nothing regresses.
        prop_cols = (
            "SELECT id, ref, name, address_line_1, address_line_2, city, postcode, property_type, "
            "COALESCE(NULLIF(name, 'multi'), address_line_1) AS display_name FROM properties "
        )
        prop = None
        linked_tenancy = tenant.get("tenancy")
        if linked_tenancy and linked_tenancy.get("property_id"):
            prop = db.execute(prop_cols + "WHERE id = ?", (linked_tenancy["property_id"],)).fetchone()
        if prop is None and tenant.get("property_id"):
            prop = db.execute(
                prop_cols + "WHERE arthur_id = CAST(? AS TEXT) OR id = ?",
                (tenant["property_id"], tenant["property_id"])
            ).fetchone()
        tenant["property"] = prop

        # Linked unit
        if tenant.get("unit_id"):
            unit = db.execute(
                "SELECT id, unit_ref, unit_type, full_address FROM units WHERE id = ?",
                (tenant["unit_id"],)
            ).fetchone()
            tenant["unit"] = unit

        # Linked esignature requests via tenancy or referencing form
        tenant["esignatures"] = []
        tenancy_id = tenant.get("tenancy_id")
        if tenancy_id:
            tenant["esignatures"] = db.execute(
                "SELECT id, document_type, document_title, status, created_for, "
                "created, sent_at, signed_at, completed_at "
                "FROM esignature_requests WHERE tenancy_id = ? "
                "ORDER BY created DESC LIMIT 5",
                [tenancy_id]
            ).fetchall()

        return json_success(tenant)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 5b. GUARANTORS (from tenants table)
# ═══════════════════════════════════════════════

# Guarantors live on the tenant row, not in a table of their own, so a
# guarantor "is" the tenant record that names them.
GUARANTOR_BASE_WHERE = (
    "(t.has_guarantor = 1 OR (t.guarantor_first_name IS NOT NULL AND t.guarantor_first_name != ''))"
)

# A guarantor is active while the tenancy they stand behind is running. Anyone
# can override that by hand, and the override wins. There is no third state:
# the tenant's own portal status (invited/pending) says nothing about the
# guarantee, which is why it is not used here.
GUARANTOR_STATUS_SQL = (
    "LOWER(COALESCE(NULLIF(TRIM(t.guarantor_status), ''), "
    "CASE WHEN (SELECT tn.status FROM tenancies tn WHERE tn.id = t.tenancy_id) "
    "IN ('Current', 'Periodic') THEN 'active' ELSE 'inactive' END))"
)

GUARANTOR_CATEGORIES = ("active", "inactive")


@banksia_os_bp.route("/guarantors")
def api_guarantors():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    search = request.args.get("search", "").strip()
    category = request.args.get("category", "").strip().lower()

    where_parts = [GUARANTOR_BASE_WHERE]
    params = []

    if search:
        search_clause, search_params = build_search_clause(
            ["t.guarantor_first_name", "t.guarantor_last_name",
             "t.guarantor_email", "t.guarantor_mobile"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    # Scoped server-side so paging and counts are right; a client-side filter
    # would only ever filter the page you happen to be looking at.
    if category in GUARANTOR_CATEGORIES:
        where_parts.append(f"{GUARANTOR_STATUS_SQL} = ?")
        params.append(category)

    where = " AND ".join(where_parts)

    base_cols = (
        "t.id, "
        "t.guarantor_first_name AS first_name, "
        "t.guarantor_last_name AS last_name, "
        "t.guarantor_email AS email, "
        "COALESCE(NULLIF(t.guarantor_mobile, ''), t.guarantor_phone) AS phone, "
        "t.guarantor_mobile AS mobile, "
        "t.guarantor_relation AS relationship, "
        f"{GUARANTOR_STATUS_SQL} AS status, "
        "t.guarantor_status AS status_override, "
        "(SELECT tn.status FROM tenancies tn WHERE tn.id = t.tenancy_id) AS tenancy_status, "
        "t.first_name || ' ' || t.last_name AS linked_applicant_name, "
        "t.first_name || ' ' || t.last_name AS linked_tenant_name, "
        "COALESCE("
        "(SELECT COALESCE(NULLIF(p.address_line_1, ''), p.name) FROM properties p "
        "JOIN tenancies tn ON tn.property_id = p.id WHERE tn.id = t.tenancy_id), "
        "(SELECT COALESCE(NULLIF(p.address_line_1, ''), p.name) FROM properties p "
        "WHERE p.arthur_id = CAST(t.property_id AS TEXT))"
        ") AS property_name, "
        "t.employment_company AS employer_name"
    )

    order_clause = build_order_by({
        "guarantor_last_name": "t.guarantor_last_name",
        "guarantor_first_name": "t.guarantor_first_name",
        "guarantor_email": "t.guarantor_email",
        "status": "status",
    }, "t.guarantor_last_name ASC, t.guarantor_first_name ASC")

    rows, total = paginate(
        f"SELECT {base_cols} FROM tenants t WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM tenants t WHERE {where}",
        params, page, per_page
    )

    # Card counts are deliberately unfiltered, so they don't shrink when
    # somebody leaves a search in the box.
    db = get_dict_db()
    try:
        counts = {r["s"]: r["cnt"] for r in db.execute(
            f"SELECT {GUARANTOR_STATUS_SQL} AS s, COUNT(*) AS cnt FROM tenants t "
            f"WHERE {GUARANTOR_BASE_WHERE} GROUP BY s"
        ).fetchall()}
    except Exception:
        counts = {}
    totals = {
        "active": counts.get("active", 0),
        "inactive": counts.get("inactive", 0),
        "all": sum(counts.values()),
    }

    return json_success({"items": rows, "totals": totals}, total, page, per_page)


GUARANTOR_FIELDS = (
    "guarantor_first_name", "guarantor_last_name", "guarantor_email",
    "guarantor_phone", "guarantor_mobile", "guarantor_relation",
    "guarantor_address", "guarantor_city", "guarantor_postcode",
    "guarantor_country", "guarantor_date_of_birth", "guarantor_profession",
    "guarantor_home_owner", "guarantor_status",
)


@banksia_os_bp.route("/guarantors/tenant/<int:tenant_id>", methods=["DELETE"])
def api_delete_guarantor(tenant_id):
    """Remove a guarantor.

    A guarantor is not a record of its own — it is the guarantor_* fields on a
    tenant row. So deleting one clears the guarantee and leaves the tenant
    completely alone. The details that were removed are written into the
    activity log so the entry can be put back by hand if it was a mistake.
    """
    db = get_dict_db()
    try:
        t = db.execute("SELECT * FROM tenants WHERE id=?", (tenant_id,)).fetchone()
        if not t:
            return json_error("Guarantor not found", 404)
        name = f"{t.get('guarantor_first_name','') or ''} {t.get('guarantor_last_name','') or ''}".strip()
        tenant_name = f"{t.get('first_name','') or ''} {t.get('last_name','') or ''}".strip()
        if not name and not t.get("has_guarantor"):
            return json_error("There is no guarantor on this tenant to remove.", 409)

        snapshot = {k: t.get(k) for k in GUARANTOR_FIELDS if t.get(k)}
        now = datetime.now(timezone.utc).isoformat()
        sets = ", ".join(f"{c} = NULL" for c in GUARANTOR_FIELDS)
        db.execute(
            f"UPDATE tenants SET {sets}, has_guarantor = 0, modified = ?, "
            "sync_dirty = 1, local_modified = ?, sync_origin = 'banksia_os' WHERE id = ?",
            (now, now, tenant_id))
        db.commit()
        _log_activity(
            "guarantor", tenant_id, "deleted",
            notes=(f"Guarantor '{name or 'unnamed'}' removed from tenant "
                   f"{tenant_name or tenant_id}. Details were: "
                   + json.dumps(snapshot, default=str)),
            db=db)
        db.commit()
        return json_success({"deleted": True, "name": name, "tenant": tenant_name})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/guarantors/tenant/<int:tenant_id>/status", methods=["PATCH"])
def api_set_guarantor_status(tenant_id):
    """Set a guarantor's status by hand. Active and Inactive are the only two."""
    data = request.get_json() or {}
    new_status = str(data.get("status", "")).strip().lower()
    if new_status not in GUARANTOR_CATEGORIES:
        return json_error("Status must be Active or Inactive.", 400)
    db = get_dict_db()
    try:
        t = db.execute("SELECT * FROM tenants WHERE id=?", (tenant_id,)).fetchone()
        if not t:
            return json_error("Guarantor not found", 404)
        old = db.execute(
            f"SELECT {GUARANTOR_STATUS_SQL} AS s FROM tenants t WHERE t.id=?", (tenant_id,)
        ).fetchone()["s"]
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "UPDATE tenants SET guarantor_status=?, modified=?, sync_dirty=1, "
            "local_modified=?, sync_origin='banksia_os' WHERE id=?",
            (new_status, now, now, tenant_id))
        db.commit()
        name = f"{t.get('guarantor_first_name','')} {t.get('guarantor_last_name','')}".strip()
        _log_activity("guarantor", tenant_id, "update", field_changed="status",
                      old_value=old, new_value=new_status,
                      notes=f"Guarantor status: {name or 'unnamed'}", db=db)
        db.commit()
        return json_success({"id": tenant_id, "status": new_status})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# NOTE: there used to be a /guarantors/<id> route here that redirected to
# itself, so every guarantor row on the list 302'd forever and could not be
# opened. Removed — the real handler is api_get_guarantor further down.


# ═══════════════════════════════════════════════
# 5c. REFERENCING
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/referencing")
def api_referencing():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    search = request.args.get("search", "").strip()
    status_filter = request.args.get("status", "").strip()

    where_parts = ["1=1"]
    params = []

    if search:
        search_clause, search_params = build_search_clause(
            ["rf.first_name", "rf.last_name", "rf.email"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    if status_filter:
        # Map frontend 'new' status to DB 'draft'
        db_status = status_filter
        if db_status == "new":
            db_status = "draft"
        # Allow comma-separated status filter
        statuses = [s.strip() for s in db_status.split(",")]
        placeholders = ",".join("?" * len(statuses))
        where_parts.append(f"rf.status IN ({placeholders})")
        params.extend(statuses)

    where = " AND ".join(where_parts)

    order_clause = build_order_by({
        "last_name": "rf.last_name", "first_name": "rf.first_name",
        "email": "rf.email", "status": "rf.status",
        "created": "rf.created", "property_name": "property_name",
    }, "rf.created DESC")

    query = (
        "SELECT rf.id, rf.first_name, rf.last_name, rf.email, rf.status, "
        "rf.created, rf.submitted_at, "
        "COALESCE(p.name, '') AS property_name, "
        "COALESCE(u.unit_ref, '') AS unit_ref, "
        "'' AS assigned_to "
        "FROM referencing_forms rf "
        "LEFT JOIN applicants a ON rf.applicant_id = a.id "
        "LEFT JOIN properties p ON a.property_id = p.id "
        "LEFT JOIN units u ON a.unit_id = u.id "
        f"WHERE {where} ORDER BY {order_clause}"
    )

    count_query = f"SELECT COUNT(*) AS cnt FROM referencing_forms rf WHERE {where}"

    db = get_dict_db()
    try:
        total = db.execute(count_query, params).fetchone()["cnt"]
        offset = (page - 1) * per_page
        rows = db.execute(query + " LIMIT ? OFFSET ?", params + [per_page, offset]).fetchall()

        # Map DB status 'draft' to frontend 'new' in rows
        for row in rows:
            if row["status"] == "draft":
                row["status"] = "new"

        stats = db.execute("""
            SELECT
                SUM(CASE WHEN status = 'draft' THEN 1 ELSE 0 END) as new,
                SUM(CASE WHEN status = 'submitted' THEN 1 ELSE 0 END) as submitted,
                SUM(CASE WHEN status = 'under_review' THEN 1 ELSE 0 END) as under_review,
                SUM(CASE WHEN status = 'approved' THEN 1 ELSE 0 END) as approved,
                SUM(CASE WHEN status = 'rejected' THEN 1 ELSE 0 END) as rejected
            FROM referencing_forms
        """).fetchone()
        stats["tenancy_created"] = 0
        stats["total"] = stats["new"] + stats["submitted"] + stats["under_review"] + stats["approved"] + stats["rejected"]

        return json_success({"items": rows, "stats": stats}, total=total, page=page, per_page=per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 7. FINANCE
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/finance/overview")
def api_finance_overview():
    db = get_dict_db()
    try:
        # Monthly rent roll
        monthly_rent_roll = db.execute(
            "SELECT COALESCE(SUM(rent_amount), 0) AS total FROM tenancies "
            "WHERE status IN ('Active', 'active', 'Periodic', 'periodic')"
        ).fetchone()["total"]

        # Total arrears
        total_arrears = db.execute(
            "SELECT COALESCE(SUM(amount_outstanding), 0) AS total FROM transactions "
            "WHERE is_outstanding = 1"
        ).fetchone()["total"]

        # Overdue transactions count & total
        overdue = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount_outstanding), 0) AS total FROM transactions "
            "WHERE is_overdue = 1"
        ).fetchone()

        # Deposit summary from deposits table
        currently_held = db.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total FROM deposits WHERE current_status = 'held'"
        ).fetchone()["total"]

        protected = db.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total FROM deposits WHERE protection_status = 'protected' AND current_status = 'held'"
        ).fetchone()["total"]

        unprotected = currently_held - protected

        deposit_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM deposits WHERE current_status = 'held'"
        ).fetchone()["cnt"]

        protected_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM deposits WHERE protection_status = 'protected' AND current_status = 'held'"
        ).fetchone()["cnt"]

        # Total transactions
        total_transactions = db.execute(
            "SELECT COUNT(*) AS cnt FROM transactions"
        ).fetchone()["cnt"]

        # Total collected (amount_paid sum)
        total_collected = db.execute(
            "SELECT COALESCE(SUM(amount_paid), 0) AS total FROM transactions"
        ).fetchone()["total"]

        return json_success({
            "monthly_rent_roll": round(monthly_rent_roll, 2),
            "monthly_rent_income": round(monthly_rent_roll, 2),
            "monthly_income": round(monthly_rent_roll, 2),
            "total_expected_monthly": round(monthly_rent_roll, 2),
            "total_collected_monthly": round(total_collected, 2),
            "total_arrears": round(total_arrears, 2),
            "overdue_count": overdue["cnt"],
            "overdue_total": round(overdue["total"], 2),
            "total_collected": round(total_collected, 2),
            "total_transactions": total_transactions,
            "total_deposits_held": round(currently_held, 2),
            "total_deposits": round(currently_held, 2),
            "deposits": {
                "currently_held_count": deposit_count,
                "currently_held_total": round(currently_held, 2),
                "protected_count": protected_count,
                "protected_total": round(protected, 2),
                "unprotected_total": round(unprotected, 2),
            },
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/transactions")
def api_transactions():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    is_overdue = request.args.get("is_overdue", "").strip()
    is_outstanding = request.args.get("is_outstanding", "").strip()
    property_id = request.args.get("property_id", "").strip()

    where_parts = ["1=1"]
    params = []

    if is_overdue.lower() in ("1", "true", "yes"):
        where_parts.append("is_overdue = 1")
    elif is_overdue.lower() in ("0", "false", "no"):
        where_parts.append("is_overdue = 0")

    if is_outstanding.lower() in ("1", "true", "yes"):
        where_parts.append("is_outstanding = 1")
    elif is_outstanding.lower() in ("0", "false", "no"):
        where_parts.append("is_outstanding = 0")

    if property_id:
        try:
            where_parts.append("property_id = ?")
            params.append(int(property_id))
        except ValueError:
            pass

    where = " AND ".join(where_parts)

    order_clause = build_order_by({
        "date": "date", "amount": "amount", "type": "type",
        "status": "status", "description": "description",
        "amount_outstanding": "amount_outstanding",
    }, "date DESC")

    rows, total = paginate(
        f"SELECT * FROM transactions WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM transactions WHERE {where}",
        params, page, per_page
    )

    for r in rows:
        bool_fields(r, "is_overdue", "is_outstanding")

    return json_success(rows, total, page, per_page)


@banksia_os_bp.route("/finance/transactions/<int:txn_id>")
def api_transaction(txn_id):
    db = get_dict_db()
    try:
        txn = db.execute("SELECT * FROM transactions WHERE id = ?", (txn_id,)).fetchone()
        if not txn:
            return json_error("Transaction not found", 404)

        bool_fields(txn, "is_overdue", "is_outstanding")

        # Related entities
        if txn.get("tenancy_id"):
            ten = db.execute(
                "SELECT id, ref, status FROM tenancies WHERE id = ?",
                (txn["tenancy_id"],)
            ).fetchone()
            txn["tenancy"] = ten

        if txn.get("property_id"):
            prop = db.execute(
                "SELECT id, ref, name FROM properties WHERE id = ?",
                (txn["property_id"],)
            ).fetchone()
            txn["property"] = prop

        if txn.get("payee_tenant_id"):
            payee = db.execute(
                "SELECT id, first_name, last_name FROM tenants WHERE id = ?",
                (txn["payee_tenant_id"],)
            ).fetchone()
            if payee:
                txn["payee"] = f"{payee['first_name']} {payee['last_name']}".strip()

        return json_success(txn)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/deposits")
def api_deposits():
    """Deposit list from the authoritative deposits table (unified data source).
    Returns a flat list matching the old frontend format for backward compatibility."""
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT d.*, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "t.main_tenant_name, t.ref AS tenancy_ref "
            "FROM deposits d "
            "LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            "LEFT JOIN properties p ON d.property_id = p.id "
            "ORDER BY d.created DESC "
            "LIMIT 200"
        ).fetchall()

        all_deposits = []
        for r in rows:
            tenant = r.get("main_tenant_name") or ""
            if not tenant and r.get("first_name"):
                tenant = f"{r.get('first_name','')} {r.get('last_name','')}"
            all_deposits.append({
                "id": r["id"],
                "tenant_name": tenant or "—",
                "property_name": r.get("property_name") or "—",
                "amount": r.get("amount") or 0,
                "scheme": r.get("scheme") or "—",
                "registered": r.get("protection_status") == "protected",
                "ref": r.get("tenancy_ref") or "",
                "protection_status": r.get("protection_status") or "unprotected",
                "current_status": r.get("current_status") or "held",
                "tenancy_id": r.get("tenancy_id"),
                "property_id": r.get("property_id"),
            })
        return json_success(all_deposits)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# DEPOSITS — Authoritative deposits table endpoints
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/deposits", methods=["GET"])
def api_banksia_deposits():
    """Deposit register, entered by category (Norbert, 2026-08-04).

    Three questions, in the order the money actually moves:

      to_register   the tenancy is running and the deposit is not protected
                    yet. The clock is the statutory 30 days from the start of
                    the tenancy (Housing Act 2004 s.213 as amended by the
                    Localism Act 2011) — not the 28 this endpoint used to
                    count, which was stricter than the law.
      registered    the tenancy is running and the deposit is protected.
      unprotected   the tenancy has ended, so the deposit is no longer held
                    against a live tenancy and has to be returned or resolved.

    A deposit answers exactly one of them, so unlike the property categories
    (which deliberately overlap) these partition the register.

    "Ended" is read off the tenancy status first and the end date second. A
    periodic tenancy rolls on past its fixed-term end date by definition, so
    an end date on its own must never retire a deposit; but once an agreed end
    date has actually passed, the deposit belongs in `unprotected` whether or
    not anyone has got round to changing the status.

    Archived deposits stay out entirely — archived stock lives on the Archive
    page and nowhere else.

    `?tenancy_id=` returns a flat list for a single tenancy instead of the
    split, which is what the tenancy detail page asks for.
    """
    LIMIT_DAYS = 30
    ENDED_STATUSES = {"past", "ended", "terminated", "expired", "closed", "completed",
                      # A tenancy that was rejected, cancelled or withdrawn never
                      # became a let. Now that signing an application opens a
                      # tenancy (and with it a deposit), those have to fall out of
                      # the 30-day queue rather than be chased forever.
                      "rejected", "cancelled", "canceled", "withdrawn"}

    search = request.args.get("search", "").strip()
    tenancy_filter = request.args.get("tenancy_id", "").strip()

    where_parts = ["d.current_status = 'held'"]
    params = []
    if tenancy_filter:
        where_parts.append("d.tenancy_id = ?")
        params.append(tenancy_filter)
    if search:
        where_parts.append("(COALESCE(t.main_tenant_name,'') LIKE ? OR COALESCE(p.ref,'') LIKE ? OR COALESCE(p.address_line_1,'') LIKE ? OR COALESCE(t.ref,'') LIKE ?)")
        like_val = f"%{search}%"
        params.extend([like_val, like_val, like_val, like_val])
    where = " AND ".join(where_parts)

    def _days_between(a, b):
        """Whole days from date-string a to date-string b (b - a). None if unparseable."""
        if not a or not b:
            return None
        try:
            da = datetime.fromisoformat(str(a)[:10])
            db_ = datetime.fromisoformat(str(b)[:10])
            return (db_ - da).days
        except Exception:
            return None

    db = get_dict_db()
    try:
        rows = db.execute(
            f"SELECT d.id, d.tenancy_id, d.amount, d.registered_amount, d.scheme, d.protection_status, "
            f"d.protection_reference, d.date_received, d.date_protected, d.current_status, "
            f"d.unprotected_at, d.unprotected_by, d.unprotected_reason, d.date_returned, d.amount_returned, d.deductions, "
            f"t.deposit_held_by, "
            f"t.ref AS tenancy_ref, t.main_tenant_name, t.status AS tenancy_status, "
            f"NULLIF(t.end_date,'') AS tenancy_end, "
            f"COALESCE(NULLIF(t.start_date,''), NULLIF(t.move_in_date,''), d.date_received) AS commencement, "
            f"COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS full_address, "
            f"u.unit_ref "
            f"FROM deposits d "
            f"LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            f"LEFT JOIN properties p ON d.property_id = p.id "
            f"LEFT JOIN units u ON d.unit_id = u.id "
            f"WHERE {where} ORDER BY t.start_date DESC, d.id DESC",
            params
        ).fetchall()

        # Attachments (deposit protection certificates, prescribed information,
        # bank proof) live in entity_documents like every other file in the OS,
        # so they also surface in the Documents module. One query for the lot.
        attachments = {}
        for a in db.execute(
            "SELECT id, entity_id, original_filename, file_type, file_size, created "
            "FROM entity_documents WHERE entity_type = 'deposit' ORDER BY created DESC"
        ).fetchall():
            attachments.setdefault(a["entity_id"], []).append({
                "id": a["id"],
                "filename": a["original_filename"] or "attachment",
                "file_type": a["file_type"] or "",
                "file_size": a["file_size"] or 0,
                "uploaded_at": a["created"] or "",
            })

        today = datetime.now(timezone.utc).date().isoformat()
        to_register, registered, unprotected, refunded = [], [], [], []
        amt_to_register = amt_registered = amt_unprotected = amt_refunded = 0.0
        late_count = overdue_count = never_protected_count = refunded_count = 0

        for r in rows:
            amt = r.get("registered_amount") or r.get("amount") or 0
            addr = r.get("full_address") or ""
            if r.get("unit_ref"):
                addr = f"{addr} ({r['unit_ref']})" if addr else r["unit_ref"]

            t_status = (r.get("tenancy_status") or "").strip()
            end_date = str(r.get("tenancy_end") or "")[:10]
            # A deposit can also be stood down by hand (Norbert, 2026-08-04) —
            # released back, disputed, written off — without waiting for the
            # tenancy to end. That decision is a person's, so it is recorded on
            # the row with who and when, and it is reversible.
            marked = str(r.get("unprotected_at") or "")[:10]
            ended = (bool(marked)
                     or t_status.lower() in ENDED_STATUSES
                     or (bool(end_date) and end_date < today))
            protected = r.get("protection_status") == "protected"
            reg_date = r.get("date_protected") or ""

            base = {
                "id": r["id"],
                "tenancy_id": r.get("tenancy_id"),
                "ref": r.get("tenancy_ref") or "",
                "full_address": addr,
                "main_tenant_name": r.get("main_tenant_name") or "",
                "deposit_held_by": r.get("deposit_held_by") or "",
                "deposit_scheme": r.get("scheme") or "",
                "protection_reference": r.get("protection_reference") or "",
                "deposit_amount": round(amt, 2),
                "tenancy_start": r.get("commencement") or "",
                "tenancy_status": t_status,
                "tenancy_end": end_date,
                "date_registered": reg_date,
                "unprotected_at": marked,
                "unprotected_by": r.get("unprotected_by") or "",
                "unprotected_reason": r.get("unprotected_reason") or "",
                "manually_unprotected": bool(marked),
                "date_returned": r.get("date_returned") or "",
                "amount_returned": round(r.get("amount_returned") or 0, 2),
                "deductions": round(r.get("deductions") or 0, 2),
                "refunded": bool(r.get("date_returned")),
                "attachments": attachments.get(r["id"], []),
            }

            if ended:
                # The tenancy is over. Whether the deposit was ever protected
                # still matters — one is a deposit to release, the other is a
                # deposit that was never protected at all — so it is carried
                # through rather than collapsed into a single state.
                # But first: if the deposit has been returned to the tenant,
                # it goes into refunded, not unprotected (Norbert, 2026-08-04).
                is_refunded = bool(r.get("date_returned"))
                if is_refunded:
                    refunded_count += 1
                    refunded.append({
                        **base,
                        "category": "refunded",
                        "was_protected": protected,
                        "never_protected": not protected,
                        "date_returned": r.get("date_returned") or "",
                        "amount_returned": round(r.get("amount_returned") or 0, 2),
                        "deductions": round(r.get("deductions") or 0, 2),
                    })
                    amt_refunded += amt
                else:
                    never_protected = not protected
                    if never_protected:
                        never_protected_count += 1
                    unprotected.append({
                        **base,
                        "category": "unprotected",
                        "was_protected": protected,
                        "never_protected": never_protected,
                        "days_since_end": _days_between(end_date, today) if end_date else None,
                    })
                    # (base already carries manually_unprotected / unprotected_by,
                    # so the row can say whether a person stood it down or the
                    # tenancy simply ran out.)
                    amt_unprotected += amt
            elif protected:
                dtr = _days_between(r.get("commencement"), reg_date) if reg_date else None
                # A registration date in the future cannot have happened, so it
                # is a keying error, not a late registration — counting it as
                # either on-time or late would be inventing a fact.
                future_dated = bool(reg_date) and str(reg_date)[:10] > today
                late = (not future_dated) and dtr is not None and dtr > LIMIT_DAYS
                if late:
                    late_count += 1
                registered.append({
                    **base,
                    "category": "registered",
                    "days_to_register": dtr,
                    "late": late,
                    "future_dated": future_dated,
                    # Protecting a deposit before the tenancy commences is
                    # normal (it is taken at holding stage), so a negative
                    # count is not an error — it just is not a countdown.
                    "before_start": dtr is not None and dtr < 0,
                    "date_missing": not bool(reg_date),
                })
                amt_registered += amt
            else:
                dss = _days_between(r.get("commencement"), today)
                overdue_by = (dss - LIMIT_DAYS) if (dss is not None and dss > LIMIT_DAYS) else 0
                if overdue_by > 0:
                    overdue_count += 1
                to_register.append({
                    **base,
                    "category": "to_register",
                    "days_since_start": dss,
                    "days_left": (LIMIT_DAYS - dss) if (dss is not None and dss <= LIMIT_DAYS) else None,
                    "days_overdue": overdue_by,
                    "overdue": overdue_by > 0,
                })
                amt_to_register += amt

        # A single tenancy asked for its own deposits: give it the flat list it
        # expects rather than the register's category split. This call used to
        # be ignored — the whole register came back as an object, the tenancy
        # page read `.length` off it, and so no tenancy has ever shown its
        # deposit. The aliases below are the names that page reads.
        if tenancy_filter:
            flat = []
            for item in to_register + registered + unprotected + refunded:
                flat.append({
                    **item,
                    "amount": item["deposit_amount"],
                    "scheme": item["deposit_scheme"],
                    "held_by": item["deposit_held_by"],
                    "registered": item["category"] == "registered",
                    "date_paid": item["date_registered"] or item["tenancy_start"],
                    "status": {
                        "to_register": "Awaiting registration",
                        "registered": "Registered",
                        "unprotected": "Awaiting return",
                        "refunded": "Refunded",
                    }[item["category"]],
                })
            return json_success(flat)

        return json_success({
            "to_register": to_register,
            "registered": registered,
            "unprotected": unprotected,
            "refunded": refunded,
            "counts": {
                "to_register": len(to_register),
                "registered": len(registered),
                "unprotected": len(unprotected),
                "refunded": len(refunded),
                "all": len(to_register) + len(registered) + len(unprotected) + len(refunded),
            },
            "amounts": {
                "to_register": round(amt_to_register, 2),
                "registered": round(amt_registered, 2),
                "unprotected": round(amt_unprotected, 2),
                "refunded": round(amt_refunded, 2),
                "all": round(amt_to_register + amt_registered + amt_unprotected + amt_refunded, 2),
            },
            "overdue_count": overdue_count,
            "late_registered_count": late_count,
            "never_protected_count": never_protected_count,
            "refunded_count": refunded_count,
            "limit_days": LIMIT_DAYS,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/tenancy-lookup", methods=["GET"])
def api_deposit_tenancy_lookup():
    """Find a tenancy to hang a new deposit on (Norbert, 2026-08-04).

    Everything the deposit register shows is already held on the tenancy, so a
    new deposit is created by picking the tenancy rather than by re-typing the
    tenant, the property and the start date. This returns exactly the fields
    the Add Deposit dialog puts on screen, plus what is already on file:

      suggested_amount   the deposit figure recorded on the tenancy itself, so
                         the amount is prefilled rather than guessed at.
      existing_deposits  deposits already held against that tenancy. Not a
                         block — a top-up is legitimate — but the dialog says
                         so out loud, because a duplicate deposit is otherwise
                         invisible until it turns up twice in the register.

    Live tenancies come first: those are the ones a new deposit belongs to.
    Ended and archived ones are still findable, because a deposit can be keyed
    in after the fact, but they are labelled as such.
    """
    q = request.args.get("search", "").strip()
    limit = int_param(request.args.get("limit"), 25, max_val=50)

    where_parts = ["1=1"]
    params = []
    if q:
        like = f"%{q}%"
        where_parts.append(
            "(COALESCE(t.main_tenant_name,'') LIKE ? OR COALESCE(t.ref,'') LIKE ? "
            "OR COALESCE(t.full_address,'') LIKE ? OR COALESCE(p.ref,'') LIKE ? "
            "OR COALESCE(p.address_line_1,'') LIKE ?)"
        )
        params.extend([like] * 5)
    where = " AND ".join(where_parts)

    db = get_dict_db()
    try:
        rows = db.execute(
            f"SELECT t.id, t.ref, t.status, t.main_tenant_name, t.start_date, t.end_date, "
            f"t.move_in_date, t.deposit_registered_amount, t.rent_amount, "
            f"COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS property_label, "
            f"u.unit_ref, "
            f"(SELECT COUNT(*) FROM deposits d WHERE d.tenancy_id = t.id AND d.current_status = 'held') AS existing_deposits, "
            f"(SELECT COALESCE(SUM(COALESCE(d.registered_amount, d.amount)),0) FROM deposits d "
            f" WHERE d.tenancy_id = t.id AND d.current_status = 'held') AS existing_amount "
            f"FROM tenancies t "
            f"LEFT JOIN properties p ON t.property_id = p.id "
            f"LEFT JOIN units u ON t.unit_id = u.id "
            f"WHERE {where} "
            f"ORDER BY CASE WHEN LOWER(COALESCE(t.status,'')) IN ('current','periodic','active','prospective') "
            f"         THEN 0 ELSE 1 END, t.start_date DESC, t.id DESC "
            f"LIMIT ?",
            params + [limit]
        ).fetchall()

        out = []
        for r in rows:
            addr = r.get("property_label") or ""
            if r.get("unit_ref"):
                addr = f"{addr} ({r['unit_ref']})" if addr else r["unit_ref"]
            status = (r.get("status") or "").strip()
            out.append({
                "id": r["id"],
                "ref": r.get("ref") or "",
                "status": status,
                "live": status.lower() in ("current", "periodic", "active", "prospective"),
                "main_tenant_name": r.get("main_tenant_name") or "",
                "full_address": addr,
                "tenancy_start": r.get("start_date") or r.get("move_in_date") or "",
                "tenancy_end": r.get("end_date") or "",
                "suggested_amount": round(float(r.get("deposit_registered_amount") or 0), 2),
                "rent_amount": round(float(r.get("rent_amount") or 0), 2),
                "existing_deposits": r.get("existing_deposits") or 0,
                "existing_amount": round(float(r.get("existing_amount") or 0), 2),
            })
        return json_success({"tenancies": out, "count": len(out)})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits", methods=["POST"])
def api_create_deposit():
    """Add a deposit to the register against an existing tenancy.

    The tenancy is the only thing asked for: tenant, property, unit and start
    date are read off it rather than typed again, so a deposit cannot end up
    describing a property its tenancy is not on.

    Amount is required and must be a real, positive figure — a £0 deposit is
    almost always a half-finished entry, and it would sit in the register
    looking protected-but-worthless. A registration date can be supplied at
    the same time, which files it straight into Registered; without one it
    lands in To Be Registered with the 30-day clock running from the start of
    the tenancy.
    """
    data = request.get_json(silent=True) or {}

    tenancy_id = data.get("tenancy_id")
    if not tenancy_id:
        return json_error("Pick a tenancy for this deposit")

    raw_amount = data.get("amount")
    try:
        amount = round(float(raw_amount), 2)
    except (TypeError, ValueError):
        return json_error("Enter the deposit amount")
    if amount <= 0:
        return json_error("A deposit has to be more than £0")

    date_protected = (data.get("date_protected") or "").strip()[:10]
    date_received = (data.get("date_received") or "").strip()[:10]
    scheme = (data.get("scheme") or "").strip()
    reference = (data.get("protection_reference") or "").strip()

    today = datetime.now(timezone.utc).date().isoformat()
    if date_protected and date_protected > today:
        return json_error("The registration date cannot be in the future")

    db = get_dict_db()
    try:
        ten = db.execute(
            "SELECT t.id, t.ref, t.property_id, t.unit_id, t.main_tenant_name, t.status, "
            "t.start_date, t.move_in_date, t.deposit_held_by, "
            "COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS property_label "
            "FROM tenancies t LEFT JOIN properties p ON t.property_id = p.id "
            "WHERE t.id = ?", (tenancy_id,)).fetchone()
        if not ten:
            return json_error("That tenancy no longer exists", 404)

        # The tenant row is the person, the tenancy is the agreement. Link both
        # where we can so the deposit is reachable from either side; a missing
        # tenant row is not a reason to refuse the deposit.
        tenant = db.execute(
            "SELECT id FROM tenants WHERE tenancy_id = ? ORDER BY id LIMIT 1",
            (tenancy_id,)).fetchone()

        if not date_received:
            date_received = ten.get("start_date") or ten.get("move_in_date") or today

        protection_status = "protected" if date_protected else "unprotected"

        cur = db.execute(
            "INSERT INTO deposits (tenancy_id, tenant_id, unit_id, property_id, amount, "
            "registered_amount, deposit_type, scheme, protection_status, protection_reference, "
            "date_received, date_protected, current_status, source) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (tenancy_id, tenant["id"] if tenant else None, ten.get("unit_id"), ten.get("property_id"),
             amount, amount, data.get("deposit_type") or "cash", scheme or None,
             protection_status, reference or None,
             date_received, date_protected or None, "held", "banksia-os"))
        deposit_id = cur.lastrowid
        db.commit()

        _log_activity(
            "deposit", deposit_id, "created",
            notes=("Deposit of £%s added for %s (%s), tenancy %s. %s"
                   % (amount,
                      ten.get("main_tenant_name") or "no tenant on record",
                      ten.get("property_label") or "no property on record",
                      ten.get("ref") or tenancy_id,
                      ("Registered on %s." % date_protected) if date_protected
                      else "Not registered yet.")),
            db=db)
        db.commit()

        # Tell the caller where it landed, so the page can drop the user into
        # the category the deposit actually went to rather than the one they
        # happened to be looking at.
        ended = (ten.get("status") or "").strip().lower() in (
            "past", "ended", "terminated", "expired", "closed", "completed")
        return json_success({
            "id": deposit_id,
            "tenancy_id": tenancy_id,
            "amount": amount,
            "category": "unprotected" if ended else ("registered" if date_protected else "to_register"),
            "main_tenant_name": ten.get("main_tenant_name") or "",
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:deposit_id>/unprotect", methods=["POST"])
def api_unprotect_deposit(deposit_id):
    """Stand a deposit down as unprotected by hand (Norbert, 2026-08-04).

    Until now a deposit only reached the Unprotected category when its tenancy
    ended. That is not the only way it happens: a deposit can be released back
    to the tenant, disputed, or written off while the tenancy is still running.
    This records that decision on the row, with who made it and when, and moves
    it into Unprotected.

    Nothing is destroyed. `protection_status` and the registration date are left
    exactly as they are, so the register still knows the deposit *was* protected
    and can say "to release" rather than "never registered". That also means
    Restore puts it back where it came from with one click.
    """
    data = request.get_json(silent=True) or {}
    reason = (data.get("reason") or "").strip()[:300]

    db = get_dict_db()
    try:
        dep = db.execute(
            "SELECT d.id, d.amount, d.registered_amount, d.unprotected_at, "
            "t.main_tenant_name, "
            "COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS property_label "
            "FROM deposits d "
            "LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            "LEFT JOIN properties p ON d.property_id = p.id "
            "WHERE d.id = ?", (deposit_id,)).fetchone()
        if not dep:
            return json_error("Deposit not found", 404)
        if dep.get("unprotected_at"):
            return json_error("This deposit is already marked unprotected", 409)

        who = getattr(request, "current_user", {}).get("username", "system")
        when = datetime.now(timezone.utc).isoformat()
        db.execute(
            "UPDATE deposits SET unprotected_at = ?, unprotected_by = ?, unprotected_reason = ?, "
            "modified = ? WHERE id = ?",
            (when, who, reason or None, when, deposit_id))
        db.commit()

        amount = dep.get("registered_amount") or dep.get("amount") or 0
        _log_activity(
            "deposit", deposit_id, "unprotected",
            field_changed="unprotected_at", old_value="", new_value=when,
            notes=("Deposit of £%s for %s (%s) marked unprotected.%s"
                   % (round(float(amount or 0), 2),
                      dep.get("main_tenant_name") or "no tenant on record",
                      dep.get("property_label") or "no property on record",
                      (" Reason: %s" % reason) if reason else "")),
            db=db)
        db.commit()
        return json_success({"id": deposit_id, "unprotected_at": when, "unprotected_by": who,
                             "category": "unprotected"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:deposit_id>/restore-protection", methods=["POST"])
def api_restore_deposit_protection(deposit_id):
    """Undo a manual unprotect and send the deposit back to where it was.

    Only clears the manual mark. A deposit whose tenancy has genuinely ended
    stays in Unprotected afterwards, because that is not a decision anyone made
    and it is not this route's to reverse — it says so rather than appearing to
    work and changing nothing.
    """
    db = get_dict_db()
    try:
        dep = db.execute(
            "SELECT d.id, d.amount, d.registered_amount, d.unprotected_at, t.main_tenant_name "
            "FROM deposits d LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            "WHERE d.id = ?", (deposit_id,)).fetchone()
        if not dep:
            return json_error("Deposit not found", 404)
        if not dep.get("unprotected_at"):
            return json_error(
                "Nothing to undo — nobody marked this deposit unprotected. If it is "
                "sitting in Unprotected it is because its tenancy has ended, and that "
                "is not something this can reverse", 409)

        was = dep["unprotected_at"]
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "UPDATE deposits SET unprotected_at = NULL, unprotected_by = NULL, "
            "unprotected_reason = NULL, modified = ? WHERE id = ?", (now, deposit_id))
        db.commit()

        amount = dep.get("registered_amount") or dep.get("amount") or 0
        _log_activity(
            "deposit", deposit_id, "protection_restored",
            field_changed="unprotected_at", old_value=was, new_value="",
            notes=("Deposit of £%s for %s put back on the register."
                   % (round(float(amount or 0), 2),
                      dep.get("main_tenant_name") or "no tenant on record")),
            db=db)
        db.commit()
        return json_success({"id": deposit_id, "restored": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:deposit_id>", methods=["DELETE"])
def api_delete_deposit(deposit_id):
    """Delete a deposit outright.

    Arthur does not own the deposits table — it syncs tenancy-level deposit
    fields, not these rows — so a deleted deposit stays deleted and will not
    reappear on the next pull.

    Everything the row held is written into the activity log as JSON before it
    goes, so a mistaken delete can be keyed straight back. Any files attached
    to the deposit go with it (rows and files both), because an attachment
    hanging off a record that no longer exists is invisible everywhere.
    """
    db = get_dict_db()
    try:
        dep = db.execute(
            "SELECT d.*, t.main_tenant_name, t.ref AS tenancy_ref, "
            "COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS property_label "
            "FROM deposits d "
            "LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            "LEFT JOIN properties p ON d.property_id = p.id "
            "WHERE d.id = ?", (deposit_id,)).fetchone()
        if not dep:
            return json_error("Deposit not found", 404)

        docs = db.execute(
            "SELECT id, file_path, original_filename FROM entity_documents "
            "WHERE entity_type = 'deposit' AND entity_id = ?", (deposit_id,)).fetchall()

        snapshot = {k: v for k, v in dict(dep).items() if v not in (None, "")}
        tenant = (dep.get("main_tenant_name") or "").strip()
        where = (dep.get("property_label") or "").strip()
        amount = dep.get("registered_amount") or dep.get("amount") or 0

        db.execute("DELETE FROM entity_documents WHERE entity_type = 'deposit' AND entity_id = ?",
                   (deposit_id,))
        db.execute("DELETE FROM deposits WHERE id = ?", (deposit_id,))
        db.commit()

        # Files only after the row is safely gone, so a disk error cannot
        # leave a deleted-looking deposit still in the table.
        for d in docs:
            try:
                if d["file_path"] and os.path.exists(d["file_path"]):
                    os.remove(d["file_path"])
            except Exception:
                pass  # a stray file is harmless; the record is what matters

        _log_activity(
            "deposit", deposit_id, "deleted",
            notes=("Deposit of \u00a3%s for %s (%s) deleted%s. Record was: %s"
                   % (round(float(amount or 0), 2), tenant or "no tenant on record",
                      where or "no property on record",
                      (", along with %d attachment(s)" % len(docs)) if docs else "",
                      json.dumps(snapshot, default=str))),
            db=db)
        db.commit()
        return json_success({
            "deleted": True,
            "id": deposit_id,
            "tenant": tenant,
            "attachments_removed": len(docs),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/reconciliation", methods=["GET"])
def api_deposits_reconciliation():
    """Returns the reconciliation report for deposits."""
    db = get_dict_db()
    try:
        # Currently held
        currently_held = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits WHERE current_status = 'held'"
        ).fetchone()
        protected = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits WHERE protection_status = 'protected' AND current_status = 'held'"
        ).fetchone()
        unprotected = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits WHERE protection_status = 'unprotected' AND current_status = 'held'"
        ).fetchone()

        # Historic
        historic = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits WHERE current_status IN ('returned', 'deducted')"
        ).fetchone()

        # Orphans — deposits without a linked tenancy
        orphans = db.execute(
            "SELECT d.*, COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name "
            "FROM deposits d "
            "LEFT JOIN properties p ON d.property_id = p.id "
            "WHERE d.tenancy_id IS NULL OR d.tenancy_id NOT IN (SELECT id FROM tenancies)"
        ).fetchall()

        # Tenancies without a deposit record
        tenancies_without_deposit = db.execute(
            "SELECT t.id, t.ref, t.main_tenant_name, t.status, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "t.deposit_registered_amount "
            "FROM tenancies t "
            "LEFT JOIN properties p ON t.property_id = p.id "
            "WHERE t.id NOT IN (SELECT tenancy_id FROM deposits WHERE tenancy_id IS NOT NULL) "
            "AND (t.deposit_registered_amount IS NOT NULL AND t.deposit_registered_amount > 0)"
        ).fetchall()

        # Mismatches — tenancy deposit_registered_amount != deposit record amount
        mismatches = db.execute(
            "SELECT t.id AS tenancy_id, t.ref AS tenancy_ref, t.main_tenant_name, "
            "t.deposit_registered_amount AS tenancy_amount, "
            "d.id AS deposit_id, d.amount AS deposit_amount, "
            "ABS(COALESCE(t.deposit_registered_amount, 0) - COALESCE(d.amount, 0)) AS difference "
            "FROM tenancies t "
            "JOIN deposits d ON d.tenancy_id = t.id "
            "WHERE ABS(COALESCE(t.deposit_registered_amount, 0) - COALESCE(d.amount, 0)) > 0.01"
        ).fetchall()

        # Totals
        total_all_time = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits"
        ).fetchone()

        return json_success({
            "currently_held": {
                "count": currently_held["cnt"],
                "total": round(currently_held["total"], 2),
                "protected": {
                    "count": protected["cnt"],
                    "total": round(protected["total"], 2),
                },
                "unprotected": {
                    "count": unprotected["cnt"],
                    "total": round(unprotected["total"], 2),
                },
            },
            "historic": {
                "count": historic["cnt"],
                "total": round(historic["total"], 2),
            },
            "orphans": orphans,
            "tenancies_without_deposit": tenancies_without_deposit,
            "mismatches": mismatches,
            "total_all_time": total_all_time["cnt"],
            "corrected_total": round(total_all_time["total"], 2),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/migrate", methods=["POST"])
def api_deposits_migrate():
    """One-time migration: populate deposits from existing tenancy data.
    Idempotent — skips tenancy_ids that already have a deposit record.

    Requires super_admin role. Guarded by migration_log to prevent re-runs.
    """
    import hashlib
    from datetime import datetime, timezone

    # ── Super admin check ──
    user = session.get("user", {})
    if user.get("role") != "super_admin":
        return json_error("Only super admins can run deposit migration", 403)

    db = get_dict_db()
    log_id = None
    try:
        # ── Migration log guard ──
        existing = db.execute(
            "SELECT id, status, checksum FROM migration_log WHERE name = ?",
            ("deposit_migration_v1",)
        ).fetchone()
        if existing and existing["status"] == "completed":
            return json_error(
                f"Migration already completed (id={existing['id']}, checksum={existing['checksum']})",
                409
            )

        # ── Create or resume migration log entry ──
        now_iso = datetime.now(timezone.utc).isoformat()
        requester = user.get("username", "unknown")
        if existing and existing["status"] == "failed":
            # Reset a previously failed migration
            db.execute(
                "UPDATE migration_log SET status='in_progress', notes=?, start_time=? WHERE id=?",
                ("Retry after failure", now_iso, existing["id"])
            )
            log_id = existing["id"]
        elif not existing:
            cursor = db.execute(
                "INSERT INTO migration_log (name, version, start_time, user_process, status) "
                "VALUES (?, ?, ?, ?, 'in_progress')",
                ("deposit_migration_v1", "1.0", now_iso, requester)
            )
            log_id = cursor.lastrowid
        else:
            # Already 'in_progress' — continue
            log_id = existing["id"]

        # ── Run the migration ──
        active_statuses = ("'Active', 'active', 'Periodic', 'periodic', 'Current', 'current'")
        tenancies_to_migrate = db.execute(
            f"SELECT t.id, t.unit_id, t.property_id, t.deposit_registered_amount, "
            f"t.deposit_scheme, t.deposit_registered, t.main_tenant_name, "
            f"t.ref, t.start_date, t.status "
            f"FROM tenancies t "
            f"WHERE t.id NOT IN (SELECT tenancy_id FROM deposits WHERE tenancy_id IS NOT NULL) "
            f"AND t.deposit_registered_amount IS NOT NULL AND t.deposit_registered_amount > 0 "
            f"AND t.status IN ({active_statuses})"
        ).fetchall()

        total_reviewed = len(tenancies_to_migrate)
        inserted = 0
        skipped = 0
        errors = 0

        for t in tenancies_to_migrate:
            try:
                tenancy_id = t["id"]
                amount = t["deposit_registered_amount"] or 0
                scheme = t["deposit_scheme"]
                protection_status = "protected" if t["deposit_registered"] else "unprotected"
                deposit_type = "cash"

                primary_tenant = db.execute(
                    "SELECT id FROM tenants WHERE tenancy_id = ? AND main_tenant = 1 LIMIT 1",
                    (tenancy_id,)
                ).fetchone()
                tenant_id = primary_tenant["id"] if primary_tenant else None

                db.execute(
                    "INSERT INTO deposits (tenancy_id, tenant_id, unit_id, property_id, "
                    "amount, deposit_type, scheme, protection_status, date_received, "
                    "current_status, source) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'held', 'migration')",
                    (tenancy_id, tenant_id, t["unit_id"], t["property_id"],
                     amount, deposit_type, scheme, protection_status, t["start_date"])
                )
                inserted += 1
                # Batch commits so a several-hundred-row migration doesn't hold
                # the write lock as one transaction, blocking live user saves.
                if inserted % 50 == 0:
                    db.commit()
            except Exception:
                errors += 1
                continue

        db.commit()

        # ── Update migration log on success ──
        completion_iso = datetime.now(timezone.utc).isoformat()
        # Compute a simple checksum over the deposits table
        checksum_data = db.execute(
            "SELECT COUNT(*) AS cnt, COALESCE(SUM(amount), 0) AS total FROM deposits"
        ).fetchone()
        checksum_input = f"deposit_migration_v1|{checksum_data['cnt']}|{checksum_data['total']}|{completion_iso}"
        checksum = hashlib.sha256(checksum_input.encode()).hexdigest()[:16]

        db.execute(
            "UPDATE migration_log SET "
            "status='completed', completion_time=?, records_reviewed=?, "
            "records_inserted=?, records_skipped=?, errors=?, checksum=? "
            "WHERE id=?",
            (completion_iso, total_reviewed, inserted, skipped, errors, checksum, log_id)
        )
        db.commit()

        return json_success({
            "message": f"Migration complete. Inserted {inserted} deposit records, skipped {skipped} (already present), errors {errors}.",
            "inserted": inserted,
            "skipped": skipped,
            "errors": errors,
            "log_id": log_id,
            "checksum": checksum,
        })

    except Exception as e:
        # ── Update migration log on failure ──
        try:
            if log_id is not None:
                db.execute(
                    "UPDATE migration_log SET status='failed', completion_time=?, notes=? WHERE id=?",
                    (datetime.now(timezone.utc).isoformat(), f"Error: {str(e)}", log_id)
                )
                db.commit()
        except Exception as _e:
            current_app.logger.error(f"Error in line ~4955: {_e}")
            pass
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:deposit_id>", methods=["GET"])
def api_deposits_detail(deposit_id):
    """Single deposit detail with linked tenancy/tenant/property info."""
    db = get_dict_db()
    try:
        deposit = db.execute(
            "SELECT d.*, "
            "t.main_tenant_name, t.ref AS tenancy_ref, t.status AS tenancy_status, "
            "t.start_date AS tenancy_start_date, t.end_date AS tenancy_end_date, "
            "t.rent_amount, t.rent_frequency, "
            "tn.first_name AS tenant_first_name, tn.last_name AS tenant_last_name, "
            "tn.email AS tenant_email, tn.mobile AS tenant_mobile, "
            "COALESCE(NULLIF(p.ref, ''), NULLIF(p.address_line_1, ''), p.name) AS property_name, "
            "p.address_line_1, p.address_line_2, p.city, p.postcode, "
            "u.unit_ref, u.unit_type "
            "FROM deposits d "
            "LEFT JOIN tenancies t ON d.tenancy_id = t.id "
            "LEFT JOIN tenants tn ON d.tenant_id = tn.id "
            "LEFT JOIN properties p ON d.property_id = p.id "
            "LEFT JOIN units u ON d.unit_id = u.id "
            "WHERE d.id = ?",
            (deposit_id,)
        ).fetchone()

        if not deposit:
            return json_error("Deposit not found", 404)

        return json_success(deposit)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/deposits/<int:deposit_id>", methods=["PATCH"])
def api_update_deposit(deposit_id):
    """PATCH /api/banksia-os/deposits/{id} — update deposit fields."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    db = get_dict_db()
    try:
        deposit = db.execute("SELECT * FROM deposits WHERE id = ?", (deposit_id,)).fetchone()
        if not deposit:
            return json_error("Deposit not found", 404)

        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(deposits)").fetchall()}
        protected_keys = {"id", "created", "modified"}

        set_parts = []
        params = []
        for key, val in data.items():
            if key in protected_keys or key not in real_cols:
                continue
            set_parts.append(f"{key} = ?")
            params.append(val)

        if not set_parts:
            return json_error("No valid fields to update", 400)

        now = datetime.now(timezone.utc).isoformat()
        set_parts.append("modified = ?")
        params.append(now)
        params.append(deposit_id)

        db.execute(f"UPDATE deposits SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()
        # -- Audit: log before/after for every changed deposit field --
        try:
            _old = dict(deposit) if deposit else {}
            for key, val in data.items():
                if key in protected_keys or key not in real_cols:
                    continue
                if str(_old.get(key) if _old.get(key) is not None else "") != str(val if val is not None else ""):
                    _log_activity("deposit", deposit_id, "update", key, _old.get(key), val, db=db)
            db.commit()
        except Exception:
            pass

        updated = db.execute("SELECT * FROM deposits WHERE id = ?", (deposit_id,)).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# RENT CHARGES — Per-month editable schedule
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/finance/rent-charges/<int:tenancy_id>", methods=["GET"])
def api_get_rent_charges(tenancy_id):
    """Get all monthly rent charges for a tenancy."""
    db = get_dict_db()
    try:
        charges = db.execute(
            "SELECT id, month, rent_amount, paid_amount, status, notes, created, modified "
            "FROM rent_charges WHERE tenancy_id = ? ORDER BY month ASC",
            (tenancy_id,)
        ).fetchall()
        return json_success(charges)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/rent-charges/generate/<int:tenancy_id>", methods=["POST"])
def api_generate_rent_charges(tenancy_id):
    """Generate monthly rent charges for a tenancy based on its start/end dates and rent_amount.
    Existing charges are preserved; only missing months are added."""
    db = get_dict_db()
    try:
        tenancy = db.execute("SELECT id, start_date, end_date, rent_amount, rent_frequency FROM tenancies WHERE id = ?",
                             (tenancy_id,)).fetchone()
        if not tenancy:
            return json_error("Tenancy not found", 404)

        start = tenancy["start_date"]
        end = tenancy["end_date"] or (datetime.now(timezone.utc).replace(day=1) + timedelta(days=365)).isoformat()[:10]
        rent = float(tenancy["rent_amount"] or 0)
        freq = (tenancy["rent_frequency"] or "pcm").lower()

        # Generate from start to end (or 24 months max)
        try:
            cur = datetime.strptime(start[:7], "%Y-%m") if start else datetime.now(timezone.utc).replace(day=1)
        except:
            cur = datetime.now(timezone.utc).replace(day=1)
        try:
            end_dt = datetime.strptime(end[:7], "%Y-%m")
        except:
            end_dt = cur + timedelta(days=365)

        max_months = 24
        count = 0
        while cur <= end_dt and count < max_months:
            month_str = cur.strftime("%Y-%m")
            existing = db.execute("SELECT id FROM rent_charges WHERE tenancy_id = ? AND month = ?",
                                  (tenancy_id, month_str)).fetchone()
            if not existing:
                db.execute(
                    "INSERT INTO rent_charges (tenancy_id, month, rent_amount, status, created) "
                    "VALUES (?, ?, ?, 'due', ?)",
                    (tenancy_id, month_str, rent, datetime.now(timezone.utc).isoformat())
                )
            count += 1
            # Advance by frequency
            if freq in ("pw", "week", "weekly"):
                cur += timedelta(weeks=4)
            else:
                if cur.month == 12:
                    cur = cur.replace(year=cur.year + 1, month=1)
                else:
                    cur = cur.replace(month=cur.month + 1)

        db.commit()
        total_charges = db.execute("SELECT COUNT(*) AS c FROM rent_charges WHERE tenancy_id = ?",
                                   (tenancy_id,)).fetchone()["c"]
        return json_success({"generated": count, "total_charges": total_charges, "tenancy_id": tenancy_id})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/rent-charges/<int:charge_id>", methods=["PATCH"])
def api_update_rent_charge(charge_id):
    """Update a specific month's rent charge (amount, paid_amount, status, notes)."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")
    db = get_dict_db()
    try:
        charge = db.execute("SELECT id, tenancy_id FROM rent_charges WHERE id = ?", (charge_id,)).fetchone()
        if not charge:
            return json_error("Charge not found", 404)
        set_parts = ["modified = ?"]
        params = [datetime.now(timezone.utc).isoformat()]
        for key in ("rent_amount", "paid_amount", "status", "notes"):
            if key in data:
                set_parts.append(f"{key} = ?")
                params.append(data[key])
        params.append(charge_id)
        db.execute(f"UPDATE rent_charges SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()

        # Recalculate tenancy financial summary
        tenancy_id = charge["tenancy_id"]
        totals = db.execute(
            "SELECT COALESCE(SUM(rent_amount),0) AS total_expected, "
            "COALESCE(SUM(paid_amount),0) AS total_paid "
            "FROM rent_charges WHERE tenancy_id = ?",
            (tenancy_id,)
        ).fetchone()
        return json_success({
            "updated": True,
            "charge_id": charge_id,
            "total_expected": round(totals["total_expected"], 2),
            "total_paid": round(totals["total_paid"], 2),
            "balance": round(totals["total_expected"] - totals["total_paid"], 2)
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/recalculate", methods=["POST"])
def api_recalculate_finances():
    """Recalculate all financial KPI summary data from rent_charges."""
    db = get_dict_db()
    try:
        tenancy_counts = db.execute("SELECT COUNT(DISTINCT tenancy_id) AS c FROM rent_charges").fetchone()["c"]
        total_expected = db.execute("SELECT COALESCE(SUM(rent_amount),0) AS t FROM rent_charges").fetchone()["t"]
        total_paid = db.execute("SELECT COALESCE(SUM(paid_amount),0) AS t FROM rent_charges").fetchone()["t"]
        overdue = db.execute("SELECT COALESCE(SUM(rent_amount - paid_amount),0) AS t FROM rent_charges WHERE status IN ('due','overdue')").fetchone()["t"]
        monthly = db.execute(
            "SELECT COALESCE(SUM(rc.rent_amount),0) AS t FROM rent_charges rc "
            "JOIN tenancies t ON rc.tenancy_id = t.id "
            "WHERE rc.month = strftime('%Y-%m', 'now') AND t.status IN ('Active','Periodic','active','periodic')"
        ).fetchone()["t"]
        return json_success({
            "tenancies_with_charges": tenancy_counts,
            "total_expected": round(total_expected, 2),
            "total_paid": round(total_paid, 2),
            "total_outstanding": round(total_expected - total_paid, 2),
            "current_month_rent": round(monthly, 2),
            "overdue_estimated": round(overdue, 2)
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 8. SEARCH
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/search")
def api_search():
    q = request.args.get("q", "").strip()
    if not q:
        return json_error("Query parameter 'q' is required")

    like_val = f"%{q}%"
    db = get_dict_db()
    try:
        # Properties
        properties = db.execute(
            "SELECT id, ref, name, address_line_1, city, postcode, main_image_url, "
            "'property' AS result_type FROM properties "
            "WHERE name LIKE ? OR ref LIKE ? OR address_line_1 LIKE ? OR city LIKE ? OR postcode LIKE ? "
            "LIMIT 10",
            [like_val] * 5
        ).fetchall()

        # Units
        units = db.execute(
            "SELECT id, unit_ref, full_address, unit_type, unit_status, "
            "'unit' AS result_type FROM units "
            "WHERE unit_ref LIKE ? OR full_address LIKE ? OR unit_type LIKE ? OR owner_name LIKE ? "
            "LIMIT 10",
            [like_val] * 4
        ).fetchall()
        for u in units:
            bool_fields(u, "unit_vacant") if "unit_vacant" in u else None

        # Tenancies
        tenancies = db.execute(
            "SELECT id, ref, full_address, status, main_tenant_name, rent_amount, start_date, end_date, "
            "'tenancy' AS result_type FROM tenancies "
            "WHERE ref LIKE ? OR full_address LIKE ? OR main_tenant_name LIKE ? "
            "LIMIT 10",
            [like_val] * 3
        ).fetchall()

        # Tenants
        tenants = db.execute(
            "SELECT id, first_name, last_name, email, mobile, full_address, status, "
            "'tenant' AS result_type FROM tenants "
            "WHERE first_name LIKE ? OR last_name LIKE ? OR email LIKE ? OR mobile LIKE ? OR full_address LIKE ? "
            "LIMIT 10",
            [like_val] * 5
        ).fetchall()

        # Applicants
        applicants = db.execute(
            "SELECT id, first_name, last_name, email, mobile, full_address, status, "
            "'applicant' AS result_type FROM applicants "
            "WHERE first_name LIKE ? OR last_name LIKE ? OR email LIKE ? OR mobile LIKE ? OR full_address LIKE ? "
            "LIMIT 10",
            [like_val] * 5
        ).fetchall()

        # Maintenance jobs
        maintenance = db.execute(
            "SELECT mj.id, mj.title, mj.reference, mj.address AS full_address, mj.status, "
            "p.name AS property_name, 'maintenance' AS result_type "
            "FROM maintenance_jobs mj LEFT JOIN properties p ON mj.property_id = p.id "
            "WHERE mj.title LIKE ? OR mj.reference LIKE ? OR mj.address LIKE ? "
            "LIMIT 10",
            [like_val] * 3
        ).fetchall()

        # Documents
        documents = db.execute(
            "SELECT id, filename, filename AS name, file_type, category, related_to AS entity_type, "
            "'document' AS result_type FROM documents "
            "WHERE (filename IS NOT NULL AND filename LIKE ?) OR (category IS NOT NULL AND category LIKE ?) OR (notes IS NOT NULL AND notes LIKE ?) "
            "LIMIT 10",
            [like_val] * 3
        ).fetchall()

        results = {
            "properties": properties,
            "units": units,
            "tenancies": tenancies,
            "tenants": tenants,
            "applicants": applicants,
            "maintenance": maintenance,
            "documents": documents,
            "total_count": len(properties) + len(units) + len(tenancies)
            + len(tenants) + len(applicants) + len(maintenance) + len(documents),
        }

        return json_success(results)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 9. WRITE ENDPOINTS — Tenancies, Applicants, Tenants
# ═══════════════════════════════════════════════

def ensure_deposit_for_tenancy(db, tenancy_id, amount=None, origin=None):
    """Give every new tenancy its deposit row, in To Be Registered (Norbert, 2026-08-04).

    A deposit that only exists once somebody remembers to add it is a deposit that
    gets protected late. The register's whole job is the statutory 30 days from the
    start of the tenancy, and that clock starts the day the tenancy starts whether
    or not anyone has keyed the money in — so the row is created with the tenancy
    and the countdown is on screen from day one.

    Idempotent by tenancy. If the tenancy already holds a deposit this returns that
    one and writes nothing, because a single let passes through more than one
    creation path: signing an application opens a Prospective tenancy, conversion
    then promotes it. Two deposit rows for one let is a double count in the
    register's money as well as its counts.

    Deliberately NOT called from arthur_sync. An Arthur pull imports history, and
    inventing deposits for 298 tenancies that ended years ago would be fiction.

    The amount is whatever the tenancy already records, and a tenancy with no
    deposit figure still gets its row: "we hold a deposit and have not keyed the
    amount" is a job, and hiding it until someone types a number is how the 30
    days get missed. The amount is editable on the register.

    The caller owns the transaction — nothing is committed here, and the activity
    line is written before the caller's commit because _log_activity(db=db) only
    commits when it opened the connection itself.

    Returns {"id", "created", "amount"}, or None if the tenancy has gone.
    """
    ten = db.execute(
        "SELECT t.id, t.ref, t.property_id, t.unit_id, t.main_tenant_name, t.status, "
        "t.start_date, t.move_in_date, t.deposit_registered_amount, t.deposit_scheme, "
        "COALESCE(NULLIF(p.ref,''), NULLIF(p.address_line_1,''), p.name) AS property_label "
        "FROM tenancies t LEFT JOIN properties p ON t.property_id = p.id "
        "WHERE t.id = ?", (tenancy_id,)).fetchone()
    if not ten:
        return None

    existing = db.execute(
        "SELECT id, amount FROM deposits WHERE tenancy_id = ? AND archived_at IS NULL "
        "ORDER BY id LIMIT 1", (tenancy_id,)).fetchone()
    if existing:
        return {"id": existing["id"], "created": False,
                "amount": float(existing.get("amount") or 0)}

    def _num(v):
        # Some of these figures arrive as typed text ("£1,200"), because the
        # tenant application stores what the applicant was quoted rather than a
        # number. A pound sign is not a reason to record the deposit as zero.
        if v is None:
            return 0.0
        try:
            return round(float(str(v).replace("\u00a3", "").replace(",", "").strip()), 2)
        except (TypeError, ValueError):
            return 0.0

    value = _num(amount) if amount not in (None, "") else 0.0
    if value <= 0:
        value = _num(ten.get("deposit_registered_amount"))
    if value < 0:
        value = 0.0

    # The tenant row is the person, the tenancy is the agreement. Link both where
    # we can so the deposit is reachable from either side; a tenancy whose tenant
    # row has not been written yet is not a reason to skip the deposit.
    tenant = db.execute(
        "SELECT id FROM tenants WHERE tenancy_id = ? ORDER BY main_tenant DESC, id LIMIT 1",
        (tenancy_id,)).fetchone()

    today = datetime.now(timezone.utc).date().isoformat()
    date_received = (ten.get("start_date") or ten.get("move_in_date") or today)
    date_received = str(date_received)[:10]

    cur = db.execute(
        "INSERT INTO deposits (tenancy_id, tenant_id, unit_id, property_id, amount, "
        "registered_amount, deposit_type, scheme, protection_status, date_received, "
        "current_status, source) "
        "VALUES (?,?,?,?,?,?,?,?,'unprotected',?,'held','auto-tenancy')",
        (tenancy_id, tenant["id"] if tenant else None, ten.get("unit_id"),
         ten.get("property_id"), value, value, "cash",
         (ten.get("deposit_scheme") or None), date_received))
    deposit_id = cur.lastrowid

    _log_activity(
        "deposit", deposit_id, "created",
        notes=("Deposit opened automatically with %s for %s (%s). %s Awaiting "
               "registration — 30 days from %s.%s"
               % (("tenancy %s" % ten.get("ref")) if ten.get("ref") else "the tenancy",
                  ten.get("main_tenant_name") or "no tenant on record",
                  ten.get("property_label") or "no property on record",
                  ("Amount \u00a3%s." % value) if value > 0
                  else "Amount not recorded on the tenancy, so it needs keying.",
                  date_received,
                  (" Source: %s." % origin) if origin else "")),
        db=db)

    return {"id": deposit_id, "created": True, "amount": value}


@banksia_os_bp.route("/tenancies", methods=["POST"])
def api_create_tenancy():
    """Create a new tenancy."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("Request body must be JSON")

    property_id = data.get("property_id")
    unit_id = data.get("unit_id")
    start_date = data.get("start_date")
    rent_amount = data.get("rent_amount")
    rent_frequency = data.get("rent_frequency", "pcm")
    deposit_amount = data.get("deposit_amount")
    main_tenant_name = data.get("main_tenant_name")
    tenant_email = (data.get("tenant_email") or "").strip()
    send_agreement = bool(data.get("send_agreement"))

    if not all([property_id, unit_id, start_date, rent_amount]):
        return json_error("Missing required fields: property_id, unit_id, start_date, rent_amount")
    if send_agreement and not tenant_email:
        return json_error("A tenant email is required to send the tenancy agreement for signature")

    db = get_dict_db()
    try:
        # Verify unit exists
        unit = db.execute("SELECT id, unit_ref, owner_name FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        # Generate a ref
        now = datetime.now(timezone.utc)
        ref = f"TEN-{now.strftime('%Y%m')}-{db.execute('SELECT COALESCE(MAX(id),0)+1 FROM tenancies').fetchone()['COALESCE(MAX(id),0)+1']}"

        full_address = unit.get("owner_name") or f"Unit {unit.get('unit_ref')}"
        property_id_val = int(property_id)
        rent_amount_val = float(rent_amount)
        deposit_amount_val = float(deposit_amount) if deposit_amount else 0
        now_iso = now.isoformat()

        db.execute(
            """INSERT INTO tenancies
               (property_id, unit_id, ref, full_address, status, start_date,
                rent_amount, rent_frequency, deposit_registered_amount,
                main_tenant_name, modified, created)
               VALUES (?, ?, ?, ?, 'Active', ?, ?, ?, ?, ?, ?, ?)""",
            (property_id_val, int(unit_id), ref, full_address, start_date,
             rent_amount_val, rent_frequency, deposit_amount_val,
             main_tenant_name or "", now_iso, now_iso)
        )
        new_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]

        # The deposit is part of opening a tenancy, not a separate errand, so it
        # is written in the same transaction: if it cannot be written the tenancy
        # rolls back with it rather than leaving a let with no deposit on record.
        deposit = ensure_deposit_for_tenancy(
            db, new_id, amount=deposit_amount_val, origin="new tenancy")
        db.commit()

        # Mark unit as not vacant
        db.execute("UPDATE units SET unit_vacant = 0, unit_status = 'Let' WHERE id = ?", (int(unit_id),))
        db.commit()

        esign = None
        delivery = None
        if send_agreement:
            from referencing_api import generate_token, _current_username, _deliver_esignature
            signer_token = generate_token()
            expires_at = (now + timedelta(days=14)).isoformat()
            db.execute(
                """INSERT INTO esignature_requests
                   (tenancy_id, document_type, document_title, status,
                    created_for, created_for_email, signer_token, expires_at, created_by)
                   VALUES (?, 'tenancy_agreement', ?, 'draft', ?, ?, ?, ?, ?)""",
                (new_id, f"Tenancy Agreement — {ref}", main_tenant_name or tenant_email,
                 tenant_email, signer_token, expires_at, _current_username()),
            )
            eid = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
            db.execute(
                "INSERT INTO esignature_audit_log (request_id, event_type, event_detail, ip_address, user_agent) "
                "VALUES (?, 'created', ?, ?, ?)",
                (eid, f"Auto-created with tenancy {ref}", request.remote_addr or "", request.headers.get("User-Agent", "")),
            )
            ereq = db.execute("SELECT * FROM esignature_requests WHERE id = ?", (eid,)).fetchone()
            ok, delivery = _deliver_esignature(db, ereq, actual_send=True)
            esign = db.execute("SELECT * FROM esignature_requests WHERE id = ?", (eid,)).fetchone()
            db.commit()

        return json_success({"id": new_id, "ref": ref, "esignature": esign,
                             "delivery": delivery, "deposit": deposit})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:ten_id>/end", methods=["POST"])
def api_end_tenancy(ten_id):
    """End a tenancy — set end_date, move_out_date, and update status."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("Request body must be JSON")

    end_date = data.get("end_date")
    move_out_date = data.get("move_out_date")

    if not end_date:
        return json_error("end_date is required")

    db = get_dict_db()
    try:
        tenancy = db.execute("SELECT * FROM tenancies WHERE id = ?", (ten_id,)).fetchone()
        if not tenancy:
            return json_error("Tenancy not found", 404)

        db.execute(
            "UPDATE tenancies SET status = 'Ended', end_date = ?, move_out_date = ?, modified = ? WHERE id = ?",
            (end_date, move_out_date or end_date, datetime.now(timezone.utc).isoformat(), ten_id)
        )
        db.commit()

        # Mark unit as vacant
        if tenancy.get("unit_id"):
            db.execute("UPDATE units SET unit_vacant = 1, unit_status = 'Available' WHERE id = ?",
                       (tenancy["unit_id"],))
            db.commit()

        return json_success({"id": ten_id, "status": "Ended"})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenants", methods=["POST"])
def api_create_tenant():
    """Create a new tenant."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("Request body must be JSON")

    first_name = data.get("first_name", "").strip()
    last_name = data.get("last_name", "").strip()
    email = data.get("email", "").strip()
    mobile = data.get("mobile", "").strip()
    tenancy_id = data.get("tenancy_id")

    if not first_name or not last_name:
        return json_error("first_name and last_name are required")

    db = get_dict_db()
    try:
        # Look up tenancy for property/unit info
        tenancy = None
        property_id = None
        unit_id = None
        full_address = ""
        if tenancy_id:
            tenancy = db.execute(
                "SELECT id, property_id, unit_id, full_address FROM tenancies WHERE id = ?",
                (tenancy_id,)
            ).fetchone()
            if tenancy:
                property_id = tenancy["property_id"]
                unit_id = tenancy["unit_id"]
                full_address = tenancy["full_address"] or ""

        now_iso = datetime.now(timezone.utc).isoformat()
        db.execute(
            """INSERT INTO tenants
               (first_name, last_name, email, mobile, tenancy_id, property_id, unit_id,
                full_address, main_tenant, modified, created)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?, ?)""",
            (first_name, last_name, email, mobile, tenancy_id,
             property_id, unit_id, full_address, now_iso, now_iso)
        )
        new_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
        db.commit()
        return json_success({"id": new_id})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 10. DOCUMENT GENERATION
# ═══════════════════════════════════════════════

from document_engine import generate_document, save_template, list_templates, delete_template, list_generated_documents, record_generated_document, get_template_info
import mammoth
import fitz  # PyMuPDF


@banksia_os_bp.route("/documents/templates", methods=["GET"])
def api_list_templates():
    templates = list_templates()
    return json_success(templates)


@banksia_os_bp.route("/documents/templates", methods=["POST"])
def api_upload_template():
    if "file" not in request.files:
        return json_error("No file uploaded")
    file = request.files["file"]
    name = request.form.get("name", file.filename or "Untitled")
    description = request.form.get("description", "")
    tid, err = save_template(file, name, description)
    if err:
        return json_error(err)
    return json_success({"id": tid, "name": name})


@banksia_os_bp.route("/documents/templates/<template_id>", methods=["DELETE"])
def api_delete_template(template_id):
    if delete_template(template_id):
        return json_success({"deleted": True})
    return json_error("Template not found", 404)


@banksia_os_bp.route("/documents/templates/<template_id>/download")
def api_download_template(template_id):
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    path = os.path.join(os.path.dirname(__file__), "documents", "templates", info["filename"])
    if not os.path.exists(path):
        return json_error("File not found", 404)
    from flask import send_file
    return send_file(path, as_attachment=True, download_name=info["filename"])


@banksia_os_bp.route("/documents/generate", methods=["POST"])
def api_generate_document():
    data = request.get_json(silent=True) or {}
    template_id = data.get("template_id")
    tenancy_id = data.get("tenancy_id")
    if not template_id or not tenancy_id:
        return json_error("template_id and tenancy_id are required")
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    template_path = os.path.join(os.path.dirname(__file__), "documents", "templates", info["filename"])
    output_path, err = generate_document(template_path, tenancy_id)
    if err:
        return json_error(err)
    doc_id = record_generated_document(output_path, info["name"], tenancy_id, "Tenant")
    return json_success({"id": doc_id, "filename": os.path.basename(output_path)})


@banksia_os_bp.route("/documents/generated", methods=["GET"])
def api_list_generated():
    docs = list_generated_documents()
    return json_success(docs)


@banksia_os_bp.route("/documents/generated/<doc_id>/download")
def api_download_generated(doc_id):
    docs = list_generated_documents()
    info = next((d for d in docs if d["id"] == doc_id), None)
    if not info:
        return json_error("Document not found", 404)
    path = os.path.join(os.path.dirname(__file__), "documents", "generated", info["filename"])
    if not os.path.exists(path):
        return json_error("File not found", 404)
    from flask import send_file
    return send_file(path, as_attachment=True, download_name=info["filename"])


# ═══════════════════════════════════════════════
# 10b. TEMPLATE EDITOR API (Phase 2 — Visual Editor)
# ═══════════════════════════════════════════════

DOCUMENTS_TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "documents", "templates")


@banksia_os_bp.route("/documents/templates/<template_id>/preview", methods=["GET"])
def api_template_preview(template_id):
    """Convert a .docx template to HTML + paragraph structure for the visual editor."""
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    path = os.path.join(DOCUMENTS_TEMPLATES_DIR, info["filename"])
    if not os.path.exists(path):
        return json_error("Template file not found on disk", 404)
    try:
        # Mammoth HTML conversion
        with open(path, "rb") as f:
            result = mammoth.convert_to_html(f)
            html = result.value

        # python-docx paragraph extraction for the editor structure
        from docx import Document
        doc = Document(path)
        paragraphs = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                style_name = p.style.name if p.style else "Normal"
                paragraphs.append({
                    "index": i,
                    "text": text[:200],  # first 200 chars for preview
                    "style": style_name,
                })

        # Also extract table preview
        tables = []
        for ti, table in enumerate(doc.tables):
            rows = []
            for row in table.rows[:3]:  # first 3 rows preview
                cells = [cell.text.strip()[:50] for cell in row.cells]
                rows.append(cells)
            tables.append({"index": ti, "rows": rows, "total_rows": len(table.rows)})

        return json_success({
            "html": html,
            "paragraphs": paragraphs,
            "tables": tables,
            "filename": info["filename"],
            "name": info["name"],
        })
    except Exception as e:
        return json_error(safe_error(e), 500)


@banksia_os_bp.route("/documents/templates/<template_id>/layout", methods=["GET"])
def api_get_template_layout(template_id):
    """Load the stored field layout for a template."""
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    layout_path = os.path.join(DOCUMENTS_TEMPLATES_DIR, f"{template_id}.layout.json")
    if not os.path.exists(layout_path):
        return json_success({"fields": [], "signature_blocks": [], "page_width": 0, "page_height": 0})
    try:
        with open(layout_path) as f:
            layout = json.load(f)
        return json_success(layout)
    except Exception as e:
        return json_error(safe_error(e), 500)


@banksia_os_bp.route("/documents/templates/<template_id>/layout", methods=["POST"])
def api_save_template_layout(template_id):
    """Save the field layout for a template."""
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    data = request.get_json(silent=True) or {}
    layout_path = os.path.join(DOCUMENTS_TEMPLATES_DIR, f"{template_id}.layout.json")
    try:
        layout = {
            "fields": data.get("fields", []),
            "signature_blocks": data.get("signature_blocks", []),
            "page_width": data.get("page_width", 595),  # A4 default pts
            "page_height": data.get("page_height", 842),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        with open(layout_path, "w") as f:
            json.dump(layout, f, indent=2)
        return json_success({"saved": True})
    except Exception as e:
        return json_error(safe_error(e), 500)


@banksia_os_bp.route("/documents/merge-fields", methods=["GET"])
def api_get_merge_fields():
    """Return all available merge fields grouped by category."""
    fields = {
        "Tenant Info": {
            "TenantName": {"label": "Tenant Name(s)", "description": "Full names of all tenants", "type": "text"},
            "MainTenantName": {"label": "Main Tenant Name", "description": "Primary tenant full name", "type": "text"},
            "TenantFirstName": {"label": "Tenant First Name", "description": "Primary tenant first name", "type": "text"},
            "TenantLastName": {"label": "Tenant Last Name", "description": "Primary tenant last name", "type": "text"},
            "TenantEmail": {"label": "Tenant Email", "description": "Primary tenant email address", "type": "text"},
            "TenantPhone": {"label": "Tenant Phone", "description": "Primary tenant phone number", "type": "text"},
            "TenantDOB": {"label": "Tenant Date of Birth", "description": "Primary tenant date of birth", "type": "text"},
            "TenantEmployer": {"label": "Tenant Employer", "description": "Primary tenant employer/company", "type": "text"},
            "TenantNI": {"label": "Tenant NI Number", "description": "National Insurance number", "type": "text"},
            "TenantPassport": {"label": "Tenant Passport", "description": "Passport number", "type": "text"},
            "TenantSignatureBlock": {"label": "Tenant Signature(s)", "description": "Signature lines for every tenant (one line per tenant)", "type": "signature"},
        },
        "Guarantor": {
            "GuarantorName": {"label": "Guarantor Name", "description": "Guarantor full name", "type": "text"},
            "GuarantorEmail": {"label": "Guarantor Email", "description": "Guarantor email address", "type": "text"},
        },
        "Property": {
            "PropertyName": {"label": "Property Name", "description": "Property name/ref", "type": "text"},
            "PropertyRef": {"label": "Property Ref", "description": "Property reference code", "type": "text"},
            "PropertyAddress": {"label": "Full Address", "description": "Full property address", "type": "text"},
            "PropertyAddressLine1": {"label": "Address Line 1", "description": "First line of address", "type": "text"},
            "PropertyCity": {"label": "City", "description": "Property city/town", "type": "text"},
            "PropertyPostcode": {"label": "Postcode", "description": "Property postcode", "type": "text"},
            "CouncilTaxBand": {"label": "Council Tax Band", "description": "Council tax band", "type": "text"},
        },
        "Unit": {
            "UnitRef": {"label": "Unit Ref", "description": "Unit/room reference", "type": "text"},
            "UnitType": {"label": "Unit Type", "description": "Studio/1-bed/2-bed etc", "type": "text"},
            "UnitAddress": {"label": "Unit Address", "description": "Full unit address line", "type": "text"},
            "UnitBedrooms": {"label": "Bedrooms", "description": "Number of bedrooms", "type": "text"},
            "UnitMaxOccupancy": {"label": "Max Occupancy", "description": "Maximum occupants", "type": "text"},
        },
        "Rent & Deposit": {
            "RentAmount": {"label": "Rent Amount (formatted)", "description": "Rent amount with GBP symbol", "type": "text"},
            "RentAmountNumeric": {"label": "Rent Amount (numeric)", "description": "Rent amount as number only", "type": "text"},
            "RentFrequency": {"label": "Rent Frequency", "description": "pcm/pw/etc", "type": "text"},
            "RentReviewDate": {"label": "Rent Review Date", "description": "Rent review date", "type": "text"},
            "DepositAmount": {"label": "Deposit Amount", "description": "Deposit with GBP symbol", "type": "text"},
            "DepositScheme": {"label": "Deposit Scheme", "description": "DPS/TDS/MyDeposits", "type": "text"},
            "DepositHeldBy": {"label": "Deposit Held By", "description": "Who holds the deposit", "type": "text"},
        },
        "Dates": {
            "TenancyStartDate": {"label": "Start Date", "description": "Tenancy start date", "type": "text"},
            "TenancyEndDate": {"label": "End Date", "description": "Tenancy end date", "type": "text"},
            "RenewalStart": {"label": "Renewal Start", "description": "Renewal period start", "type": "text"},
            "RenewalEnd": {"label": "Renewal End", "description": "Renewal period end", "type": "text"},
            "BreakClauseDate": {"label": "Break Clause Date", "description": "Break clause date", "type": "text"},
            "MoveInDate": {"label": "Move In Date", "description": "Move in date", "type": "text"},
            "MoveOutDate": {"label": "Move Out Date", "description": "Move out date", "type": "text"},
            "NoticePeriod": {"label": "Notice Period", "description": "Notice period required", "type": "text"},
        },
        "Landlord / Agent": {
            "LandlordName": {"label": "Landlord Name", "description": "Property owner/landlord", "type": "text"},
            "LandlordAddress": {"label": "Landlord Address", "description": "Landlord's address", "type": "text"},
            "AgentName": {"label": "Agent Name", "description": "Managing agent name", "type": "text"},
            "AgentAddress": {"label": "Agent Address", "description": "Agent office address", "type": "text"},
            "AgentEmail": {"label": "Agent Email", "description": "Agent contact email", "type": "text"},
            "AgentPhone": {"label": "Agent Phone", "description": "Agent contact phone", "type": "text"},
        },
        "Today": {
            "Date": {"label": "Date (dd/mm/yyyy)", "description": "Today's date short format", "type": "text"},
            "DateLong": {"label": "Date (long format)", "description": "Today's date long format", "type": "text"},
            "Year": {"label": "Year", "description": "Current year", "type": "text"},
            "Month": {"label": "Month", "description": "Current month name", "type": "text"},
            "Day": {"label": "Day", "description": "Current day number", "type": "text"},
        },
    }

    return json_success(fields)


# ─── Inline document editor: read/write template body as ordered blocks ───────

def _iter_block_items(doc):
    """Yield ('p', Paragraph) and ('tbl', Table) in true document order."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag.endswith('}p'):
            yield ('p', Paragraph(child, doc))
        elif tag.endswith('}tbl'):
            yield ('tbl', Table(child, doc))


def _set_paragraph_text(paragraph, text):
    """Replace a paragraph's text while keeping its first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


@banksia_os_bp.route("/documents/templates/<template_id>/document", methods=["GET"])
def api_template_document(template_id):
    """Return the template body as an ordered list of editable text blocks.
    PDF/non-docx templates come back non-editable with a preview URL."""
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    path = os.path.join(DOCUMENTS_TEMPLATES_DIR, info["filename"])
    if not os.path.exists(path):
        return json_error("Template file not found on disk", 404)
    try:
        with open(path, "rb") as fh:
            head = fh.read(5)
    except Exception as e:
        return json_error(safe_error(e), 500)

    if head[:4] == b"%PDF":
        return json_success({
            "editable": False, "format": "pdf",
            "name": info["name"], "filename": info["filename"],
            "preview_url": "/api/banksia-os/documents/templates/%s/download" % template_id,
            "blocks": [],
            "reason": "This template is a PDF, which cannot hold inline Word variables. Upload a Word (.docx) version to edit variables inline.",
        })

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        return json_success({
            "editable": False, "format": "unknown",
            "name": info["name"], "filename": info["filename"],
            "preview_url": "/api/banksia-os/documents/templates/%s/download" % template_id,
            "blocks": [],
            "reason": "This file is not a valid Word document, so variables cannot be edited inline. Upload a Word (.docx) version. (%s)" % safe_error(e),
        })

    blocks = []
    pcount = 0
    tcount = 0
    for kind, item in _iter_block_items(doc):
        if kind == 'p':
            style = item.style.name if item.style else "Normal"
            blocks.append({
                "id": "p%d" % pcount, "kind": "paragraph",
                "style": style, "text": item.text,
            })
            pcount += 1
        else:
            rows = []
            for ri, row in enumerate(item.rows):
                cells = []
                for ci, cell in enumerate(row.cells):
                    cells.append({"id": "t%dr%dc%d" % (tcount, ri, ci), "text": cell.text})
                rows.append(cells)
            blocks.append({"id": "t%d" % tcount, "kind": "table", "rows": rows})
            tcount += 1

    return json_success({
        "editable": True, "format": "docx",
        "name": info["name"], "filename": info["filename"],
        "blocks": blocks,
    })


@banksia_os_bp.route("/documents/templates/<template_id>/document", methods=["POST"])
def api_save_template_document(template_id):
    """Write edited block texts back into the .docx, preserving each
    paragraph/cell's formatting. Only touches blocks whose text changed."""
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)
    path = os.path.join(DOCUMENTS_TEMPLATES_DIR, info["filename"])
    if not os.path.exists(path):
        return json_error("Template file not found on disk", 404)
    with open(path, "rb") as fh:
        if fh.read(4) == b"%PDF":
            return json_error("This template is a PDF and cannot be edited inline. Upload a Word (.docx) version.", 400)

    data = request.get_json(silent=True) or {}
    incoming = data.get("blocks", [])
    if not isinstance(incoming, list):
        return json_error("blocks must be a list")

    para_edits, cell_edits = {}, {}
    for b in incoming:
        bid = b.get("id", "")
        if bid.startswith("t") and "r" in bid and "c" in bid:
            cell_edits[bid] = b.get("text", "")
        elif bid.startswith("p"):
            para_edits[bid] = b.get("text", "")

    try:
        from docx import Document
        doc = Document(path)
        changed = 0
        pcount = 0
        tcount = 0
        for kind, item in _iter_block_items(doc):
            if kind == 'p':
                key = "p%d" % pcount
                if key in para_edits and para_edits[key] != item.text:
                    _set_paragraph_text(item, para_edits[key])
                    changed += 1
                pcount += 1
            else:
                for ri, row in enumerate(item.rows):
                    for ci, cell in enumerate(row.cells):
                        key = "t%dr%dc%d" % (tcount, ri, ci)
                        if key in cell_edits and cell_edits[key] != cell.text:
                            paras = cell.paragraphs
                            if paras:
                                _set_paragraph_text(paras[0], cell_edits[key])
                                for extra in paras[1:]:
                                    _set_paragraph_text(extra, "")
                            changed += 1
                tcount += 1

        if changed:
            try:
                shutil.copy2(path, path + ".bak")
            except Exception:
                pass
            doc.save(path)
        return json_success({"saved": True, "changed": changed})
    except Exception as e:
        return json_error(safe_error(e), 500)


@banksia_os_bp.route("/documents/templates/<template_id>/generate-with-layout", methods=["POST"])
def api_generate_with_layout(template_id):
    """Generate document using layout-based field positions + merge data."""
    data = request.get_json(silent=True) or {}
    tenancy_id = data.get("tenancy_id")
    if not tenancy_id:
        return json_error("tenancy_id is required")

    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)

    template_path = os.path.join(DOCUMENTS_TEMPLATES_DIR, info["filename"])
    if not os.path.exists(template_path):
        return json_error("Template file not found", 404)

    # Load layout from request or stored file
    layout = data.get("layout")
    if not layout:
        layout_path = os.path.join(DOCUMENTS_TEMPLATES_DIR, f"{template_id}.layout.json")
        if os.path.exists(layout_path):
            try:
                with open(layout_path) as f:
                    layout = json.load(f)
            except:
                layout = {}
        else:
            layout = {}

    # Generate the base document
    output_path, err = generate_document(template_path, tenancy_id)
    if err:
        return json_error(err)

    # Convert to PDF for signature overlay
    try:
        doc_pdf_path = output_path.replace(".docx", ".pdf")
        # Use python-docx to save then fitz to overlay
        from docx import Document as DocxDocument
        doc = DocxDocument(output_path)

        # Use subprocess to convert via LibreOffice if available
        import subprocess
        pdf_ok = False
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
                 os.path.dirname(doc_pdf_path), output_path],
                capture_output=True, timeout=30, check=False,
            )
            if os.path.exists(doc_pdf_path):
                pdf_ok = True
        except:
            pass

        if pdf_ok:
            # Add signature placeholder overlays
            sig_blocks = layout.get("signature_blocks", [])
            if sig_blocks:
                pdf_doc = fitz.open(doc_pdf_path)
                for page in pdf_doc:
                    for block in sig_blocks:
                        x = block.get("x", 50)
                        y = block.get("y", 700)
                        w = block.get("width", 200)
                        h = block.get("height", 80)
                        label = block.get("type", "signature")
                        # Draw signature placeholder rectangle
                        rect = fitz.Rect(x, y, x + w, y + h)
                        page.draw_rect(rect, color=(0.2, 0.2, 0.2), width=0.5)
                        # Add label
                        sig_label = "Tenant Signature" if label == "tenant" else "Banksia Authorised Signatory"
                        page.insert_text(
                            fitz.Point(x + 5, y + 15),
                            sig_label,
                            fontsize=8,
                            color=(0.4, 0.4, 0.4),
                        )
                        page.insert_text(
                            fitz.Point(x + 5, y + h - 5),
                            "_________________________",
                            fontsize=8,
                            color=(0.6, 0.6, 0.6),
                        )
                out_sig_path = doc_pdf_path.replace(".pdf", "_signed.pdf")
                pdf_doc.save(out_sig_path)
                pdf_doc.close()
                doc_pdf_path = out_sig_path

        # Record and return
        doc_id = record_generated_document(
            output_path, info["name"], tenancy_id, "Tenant"
        )
        return json_success({
            "id": doc_id,
            "filename": os.path.basename(output_path),
            "pdf_path": doc_pdf_path if pdf_ok else None,
            "has_layout": len(sig_blocks) > 0 if sig_blocks else False,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)


@banksia_os_bp.route("/documents/generate-template-preview", methods=["POST"])
def api_generate_template_preview():
    """Generate a PDF and return base64 PNG preview pages."""
    data = request.get_json(silent=True) or {}
    template_id = data.get("template_id")
    tenancy_id = data.get("tenancy_id")
    if not template_id or not tenancy_id:
        return json_error("template_id and tenancy_id are required")

    # First generate the document
    info = get_template_info(template_id)
    if not info:
        return json_error("Template not found", 404)

    template_path = os.path.join(DOCUMENTS_TEMPLATES_DIR, info["filename"])
    output_path, err = generate_document(template_path, tenancy_id)
    if err:
        return json_error(err)

    # Convert to PDF
    import subprocess
    doc_pdf_path = output_path.replace(".docx", ".pdf")
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             os.path.dirname(doc_pdf_path), output_path],
            capture_output=True, timeout=30, check=False,
        )
    except:
        pass

    pages = []
    if os.path.exists(doc_pdf_path):
        pdf_doc = fitz.open(doc_pdf_path)
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_data = pix.tobytes("png")
            import base64
            b64 = base64.b64encode(img_data).decode()
            pages.append(f"data:image/png;base64,{b64}")
        pdf_doc.close()

    doc_id = record_generated_document(
        output_path, info["name"], tenancy_id, "Tenant"
    )

    return json_success({
        "pages": pages,
        "doc_id": doc_id,
        "filename": os.path.basename(output_path),
        "total_pages": len(pages),
    })


# ═══════════════════════════════════════════════
# UPLOADED DOCUMENTS STORAGE
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/documents/upload", methods=["POST"])
def api_upload_document():
    """Upload a document file and associate it with a tenancy/property/tenant."""
    if "file" not in request.files:
        return json_error("No file provided")
    file = request.files["file"]
    if file.filename == "":
        return json_error("Empty filename")

    docs_dir = os.path.join(os.path.dirname(__file__), "documents", "uploads")
    os.makedirs(docs_dir, exist_ok=True)

    category = request.form.get("category", "general")
    related_to = request.form.get("related_to", "")
    related_id = request.form.get("related_id", "")
    notes = request.form.get("notes", "")

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_name = f"{ts}_{file.filename}"
    save_path = os.path.join(docs_dir, safe_name)
    file.save(save_path)

    db = get_dict_db()
    try:
        # ── Auto-match: parse filename to find tenant/tenancy ──
        matched_to = None
        auto_match = request.form.get("auto_match", "false") == "true"
        if auto_match:
            fn = file.filename.lower()
            # Try to extract tenancy ref pattern (TE followed by digits)
            import re
            ref_match = re.search(r'[Tt][Ee]\d+', fn)
            if ref_match:
                ref = ref_match.group().upper()
                tenancy = db.execute(
                    "SELECT id, ref, main_tenant_name, full_address FROM tenancies WHERE ref LIKE ? LIMIT 1",
                    (f"%{ref}%",)
                ).fetchone()
                if tenancy:
                    related_to = "tenancy"
                    related_id = str(tenancy["id"])
                    matched_to = f"Tenancy {ref} ({tenancy.get('main_tenant_name','')[:30]})"
            if not matched_to:
                # Try tenant name match
                name_parts = fn.replace("_", " ").replace("-", " ").split()
                for name in name_parts:
                    if len(name) > 3:
                        tenant = db.execute(
                            "SELECT id, first_name, last_name FROM tenants WHERE first_name LIKE ? OR last_name LIKE ? LIMIT 1",
                            (f"%{name}%", f"%{name}%")
                        ).fetchone()
                        if tenant:
                            related_to = "tenant"
                            related_id = str(tenant["id"])
                            matched_to = f"Tenant {tenant['first_name']} {tenant['last_name']}"
                            break
            if not matched_to:
                # Try tenancy ID in filename
                id_match = re.search(r'\b(\d{3,5})\b', fn)
                if id_match:
                    tid = id_match.group(1)
                    tenancy = db.execute(
                        "SELECT id, ref, main_tenant_name FROM tenancies WHERE id LIKE ? OR ref LIKE ? LIMIT 1",
                        (f"%{tid}%", f"%{tid}%")
                    ).fetchone()
                    if tenancy:
                        related_to = "tenancy"
                        related_id = str(tenancy["id"])
                        matched_to = f"Tenancy {tenancy.get('ref','')}"

        db.execute(
            "INSERT INTO documents (filename, file_path, file_type, category, related_to, related_id, notes, created) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (file.filename, save_path, os.path.splitext(file.filename)[1].lower().lstrip("."),
             category, related_to, related_id, notes, datetime.now(timezone.utc).isoformat())
        )
        db.commit()
        result = {"id": db.lastrowid, "filename": file.filename}
        if matched_to:
            result["matched_to"] = matched_to
        return json_success(result)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/documents/uploaded", methods=["GET"])
def api_list_uploaded():
    db = get_dict_db()
    try:
        docs = db.execute(
            "SELECT id, filename, file_type, category, related_to, related_id, notes, created "
            "FROM documents ORDER BY created DESC"
        ).fetchall()
        return json_success(docs)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/documents/uploaded/<int:doc_id>/download")
def api_download_uploaded(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, filename, file_path FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if not os.path.exists(doc["file_path"]):
            return json_error("File not found on disk", 404)
        from flask import send_file
        return send_file(doc["file_path"], as_attachment=True, download_name=doc["filename"])
    finally:
        db.close()


@banksia_os_bp.route("/documents/uploaded/<int:doc_id>", methods=["DELETE"])
def api_delete_uploaded(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, file_path FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if os.path.exists(doc["file_path"]):
            os.remove(doc["file_path"])
        db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        db.commit()
        return json_success({"deleted": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════

@banksia_os_bp.route("/compliance", methods=["GET"])

def api_compliance_list():
    """Return all compliance records, optionally filtered by group."""
    db = get_dict_db()
    try:
        group = request.args.get("group", "")
        if group:
            rows = db.execute(
                "SELECT * FROM compliance WHERE monday_group = ? ORDER BY property_name",
                [group]
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT * FROM compliance ORDER BY monday_group, property_name"
            ).fetchall()
        return jsonify({"success": True, "data": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


@banksia_os_bp.route("/compliance/groups", methods=["GET"])

def api_compliance_groups():
    """Return compliance group names and counts."""
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT monday_group, COUNT(*) as cnt FROM compliance GROUP BY monday_group ORDER BY monday_group"
        ).fetchall()
        return jsonify({"success": True, "data": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


# Compliance certificates live on Monday as file assets; the board stores only the
# asset id per certificate column. Files replaced from inside Banksia OS are stored
# locally instead, referenced as "local:<filename>".
_COMPLIANCE_DOC_FIELDS = {
    "gas": "gas_doc",
    "electric": "electrical_doc",
    "epc": "epc_doc",
    "fire-alarm": "fire_alarm_doc",
    "emergency-lighting": "emergency_lighting_doc",
    "fra": "fra_doc",
    "floor-plan": "floor_plan_doc",
    "fire-doors": "fire_doors_doc",
    "fire-blanket": "fire_blanket_doc",
    "co2-alarm": "co2_alarm_doc",
}

# Only the dated certificates have an expiry; Floor Plan is a Yes/No status.
_COMPLIANCE_DATE_FIELDS = {
    "gas": "gas_date",
    "electric": "electrical_date",
    "epc": "epc_date",
    "fire-alarm": "fire_alarm_date",
    "emergency-lighting": "emergency_lighting_date",
    "fra": "fra_date",
    "fire-doors": "fire_doors_date",
    "fire-blanket": "fire_blanket_date",
    "co2-alarm": "co2_alarm_date",
}

COMPLIANCE_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "media", "compliance")
os.makedirs(COMPLIANCE_UPLOAD_DIR, exist_ok=True)

_CERT_ALLOWED_EXT = {".pdf", ".jpg", ".jpeg", ".png", ".webp", ".heic", ".doc", ".docx"}
_CERT_MAX_BYTES = 25 * 1024 * 1024


def _compliance_row(row_id, field):
    db = get_dict_db()
    try:
        return db.execute(
            "SELECT id, property_name, %s AS ref FROM compliance WHERE id = ?" % field,
            (row_id,)
        ).fetchone()
    finally:
        db.close()


def _safe_cert_filename(name):
    name = os.path.basename(name or "").strip()
    return re.sub(r"[^A-Za-z0-9._ -]", "_", name)[:120] or "certificate"


@banksia_os_bp.route("/compliance/<int:row_id>/certificate/<cert>", methods=["GET"])
def api_compliance_certificate(row_id, cert):
    """Stream a property's compliance certificate so it opens in the browser.

    Two storage shapes are served here: a file uploaded through Banksia OS
    ("local:<filename>" on disk) and the original Monday asset id. Monday's signed
    download url expires after an hour, so it is resolved per request and the bytes
    are proxied — the viewer needs no Monday access of their own, only a Banksia OS
    session (enforced by this blueprint's before_request).
    """
    import mimetypes, urllib.request
    from flask import Response, send_file

    field = _COMPLIANCE_DOC_FIELDS.get(cert)
    if not field:
        return json_error("Unknown certificate type", 404)

    row = _compliance_row(row_id, field)
    if not row:
        return json_error("Compliance record not found", 404)

    ref = (row["ref"] or "").strip()
    if not ref:
        return json_error("No certificate on record for this property", 404)

    if ref.startswith("local:"):
        path = os.path.join(COMPLIANCE_UPLOAD_DIR, ref[len("local:"):])
        if not os.path.exists(path):
            return json_error("Certificate file missing on disk", 404)
        return send_file(path, as_attachment=False)

    # Tolerate either a bare asset id or a full Monday resource url
    m = re.search(r"/resources/(\d+)/", ref)
    asset_id = m.group(1) if m else ref
    if not asset_id.isdigit():
        return json_error("Certificate reference is not a Monday asset", 422)

    mtok = get_monday_token()
    if not mtok:
        return json_error("Monday credentials unavailable on this server", 502)

    try:
        data = _monday_graphql(mtok, "{assets(ids:[%s]){name file_extension public_url}}" % asset_id)
        assets = ((data or {}).get("data") or {}).get("assets") or []
        if not assets or not assets[0].get("public_url"):
            return json_error("Certificate file no longer available on Monday", 404)
        asset = assets[0]
        with urllib.request.urlopen(asset["public_url"], timeout=45) as r:
            payload = r.read()
    except Exception as e:
        return json_error(safe_error(e), 502)

    filename = asset.get("name") or ("certificate" + (asset.get("file_extension") or ""))
    mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    resp = Response(payload, mimetype=mime)
    # inline so PDFs and images open in a tab rather than downloading
    resp.headers["Content-Disposition"] = 'inline; filename="%s"' % filename.replace('"', "")
    resp.headers["Cache-Control"] = "private, max-age=300"
    return resp


@banksia_os_bp.route("/compliance/<int:row_id>/certificate/<cert>", methods=["POST"])
def api_compliance_certificate_upload(row_id, cert):
    """Replace a property's certificate with an uploaded file.

    The new file is stored on this server; the Monday asset it replaces is only
    dereferenced, never deleted, so the original stays recoverable on the board.
    """
    field = _COMPLIANCE_DOC_FIELDS.get(cert)
    if not field:
        return json_error("Unknown certificate type", 404)
    if "file" not in request.files:
        return json_error("No file provided (use field 'file')")

    file = request.files["file"]
    if not file.filename:
        return json_error("Empty filename")

    safe_name = _safe_cert_filename(file.filename)
    ext = os.path.splitext(safe_name)[1].lower()
    if ext not in _CERT_ALLOWED_EXT:
        return json_error("Unsupported file type %s — allowed: %s" % (ext, ", ".join(sorted(_CERT_ALLOWED_EXT))), 415)

    row = _compliance_row(row_id, field)
    if not row:
        return json_error("Compliance record not found", 404)

    payload = file.read()
    if not payload:
        return json_error("Uploaded file is empty")
    if len(payload) > _CERT_MAX_BYTES:
        return json_error("File is larger than 25MB", 413)

    stored = "%s_%s_%s" % (row_id, cert, safe_name)
    with open(os.path.join(COMPLIANCE_UPLOAD_DIR, stored), "wb") as fh:
        fh.write(payload)

    old_ref = (row["ref"] or "").strip()
    db = get_dict_db()
    try:
        db.execute(
            "UPDATE compliance SET %s = ?, updated_at = datetime('now') WHERE id = ?" % field,
            ("local:" + stored, row_id)
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", field, old_ref, "local:" + stored,
                  notes="%s certificate replaced for %s" % (cert, row["property_name"]))
    return json_success({"certificate": stored, "filename": safe_name})


@banksia_os_bp.route("/compliance/<int:row_id>/certificate/<cert>", methods=["DELETE"])
def api_compliance_certificate_delete(row_id, cert):
    """Remove the certificate reference from a property.

    A locally uploaded file is deleted from disk; a Monday asset is only
    dereferenced here — nothing is removed from the Monday board.
    """
    field = _COMPLIANCE_DOC_FIELDS.get(cert)
    if not field:
        return json_error("Unknown certificate type", 404)

    row = _compliance_row(row_id, field)
    if not row:
        return json_error("Compliance record not found", 404)

    old_ref = (row["ref"] or "").strip()
    if not old_ref:
        return json_error("No certificate on record for this property", 404)

    if old_ref.startswith("local:"):
        path = os.path.join(COMPLIANCE_UPLOAD_DIR, old_ref[len("local:"):])
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass  # reference still clears; a stray file is harmless

    db = get_dict_db()
    try:
        db.execute(
            "UPDATE compliance SET %s = '', updated_at = datetime('now') WHERE id = ?" % field,
            (row_id,)
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", field, old_ref, "",
                  notes="%s certificate removed for %s" % (cert, row["property_name"]))
    return json_success({"removed": True})


@banksia_os_bp.route("/compliance/<int:row_id>", methods=["PATCH"])
def api_compliance_update(row_id):
    """Edit a compliance record — a certificate's expiry date, or the landlord.

    Body: {"cert": "gas", "date": "2027-02-20"}  ("" clears the date)
       or: {"landlord": "Jane Smith"}             ("" clears the name)
    Changes are local to Banksia OS; they are not written back to the Monday board.
    """
    data = request.get_json(silent=True) or {}

    if "property_name" in data:
        # The address is the only human handle on a compliance row, so it is
        # editable — but never blank, and never silently truncated.
        name = " ".join(str(data.get("property_name", "")).split())
        if not name:
            return json_error("The property address cannot be empty", 422)
        if len(name) > 200:
            return json_error("Property address is too long (200 characters max)", 422)
        row = _compliance_row(row_id, "property_name")
        if not row:
            return json_error("Compliance record not found", 404)
        old_name = (row["property_name"] or "").strip()
        db = get_dict_db()
        try:
            db.execute(
                "UPDATE compliance SET property_name = ?, updated_at = datetime('now') WHERE id = ?",
                (name, row_id)
            )
            db.commit()
        finally:
            db.close()
        _log_activity("compliance", row_id, "update", "property_name", old_name, name,
                      notes="property address changed from %s" % old_name)
        return json_success({"property_name": name})

    if "returned" in data:
        # Handing a property back is a real business state, not a display choice:
        # it exempts the row from every compliance calculation. The trigger on
        # monday_group keeps automation_exempt in step either way.
        row = _compliance_row(row_id, "monday_group")
        if not row:
            return json_error("Compliance record not found", 404)
        old_group = (row["ref"] or "").strip()
        new_group = COMPLIANCE_RETURNED_GROUP if data.get("returned") else COMPLIANCE_LIVE_GROUP
        if old_group.upper() == new_group.upper():
            return json_success({"monday_group": old_group})
        db = get_dict_db()
        try:
            db.execute(
                "UPDATE compliance SET monday_group = ?, updated_at = datetime('now') WHERE id = ?",
                (new_group, row_id)
            )
            db.commit()
        finally:
            db.close()
        _log_activity("compliance", row_id, "update", "monday_group", old_group, new_group,
                      notes="%s moved to %s" % (row["property_name"], new_group))
        return json_success({"monday_group": new_group})

    if "co2_alarm_status" in data:
        # The CO2 alarm is not a certificate with an expiry — it is either fitted
        # or it is not, evidenced by a photo. It rides on the Gas board because
        # that is the visit where anyone would check it (Norbert, 2026-08-04).
        raw = str(data.get("co2_alarm_status", "")).strip().lower()
        mapped = {"yes": "Yes", "no": "No", "": ""}.get(raw)
        if mapped is None:
            return json_error("CO2 alarm must be Yes or No.", 422)
        row = _compliance_row(row_id, "co2_alarm_status")
        if not row:
            return json_error("Compliance record not found", 404)
        old = (row["ref"] or "").strip()
        db = get_dict_db()
        try:
            db.execute(
                "UPDATE compliance SET co2_alarm_status = ?, updated_at = datetime('now') WHERE id = ?",
                (mapped, row_id)
            )
            db.commit()
        finally:
            db.close()
        _log_activity("compliance", row_id, "update", "co2_alarm_status", old, mapped,
                      notes="CO2 alarm set to %s for %s" % (mapped or "blank", row["property_name"]))
        return json_success({"co2_alarm_status": mapped})

    if "landlord" in data:
        name = str(data.get("landlord", "")).strip()
        if len(name) > 120:
            return json_error("Landlord name is too long (120 characters max)", 422)
        row = _compliance_row(row_id, "landlord")
        if not row:
            return json_error("Compliance record not found", 404)
        old_name = (row["ref"] or "").strip()
        db = get_dict_db()
        try:
            db.execute(
                "UPDATE compliance SET landlord = ?, updated_at = datetime('now') WHERE id = ?",
                (name, row_id)
            )
            db.commit()
        finally:
            db.close()
        _log_activity("compliance", row_id, "update", "landlord", old_name, name,
                      notes="landlord changed for %s" % row["property_name"])
        return json_success({"landlord": name})

    cert = str(data.get("cert", "")).strip()
    field = _COMPLIANCE_DATE_FIELDS.get(cert)
    if not field:
        return json_error("Unknown or undated certificate type", 422)

    value = str(data.get("date", "")).strip()
    if value:
        try:
            parsed = datetime.strptime(value, "%Y-%m-%d")
        except ValueError:
            return json_error("Date must be YYYY-MM-DD", 422)
        # A partially typed year (0002, 0020...) is never a real certificate date.
        # Rejecting it here means no client bug can quietly corrupt the record.
        if not (1900 <= parsed.year <= 2100):
            return json_error("Year %d is not a plausible certificate date" % parsed.year, 422)

    row = _compliance_row(row_id, field)
    if not row:
        return json_error("Compliance record not found", 404)

    old_value = (row["ref"] or "").strip()
    db = get_dict_db()
    try:
        db.execute(
            "UPDATE compliance SET %s = ?, updated_at = datetime('now') WHERE id = ?" % field,
            (value, row_id)
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", field, old_value, value,
                  notes="%s expiry date changed for %s" % (cert, row["property_name"]))
    return json_success({"cert": cert, "date": value})


# ── Custom groups on the /compliance-test board ───────────────────────────────
# The board's own sections (Expired, Due for Renew, Active, Not applicable) are
# derived from each certificate's expiry date and cannot be assigned by hand — a
# property sits in Expired because its date has passed, not because someone put it
# there. Custom groups exist alongside them for states the date cannot express,
# "To be arranged" being the first (Norbert, 2026-07-30).
#
# A group belongs to ONE certificate (Norbert, 2026-08-06): creating, filling or
# deleting a group on Fire Alarm must leave Emergency Lighting alone. Groups were
# board-wide between 2026-07-30 and 2026-08-06 — a property dragged in on Gas
# appeared in that group on every certificate — because a state like "to be
# arranged" reads as a fact about the property. In practice it is a fact about one
# certificate's renewal, so the scope moved back onto the certificate. Both tables
# already carried a cert_key, so this was a change to the writers and to a new
# compliance_groups.cert_key, not to the board's reads.

COMPLIANCE_RETURNED_GROUP = "PROPERTY RETURNED"
COMPLIANCE_LIVE_GROUP = "VERV COMPLIANCE CERTIFICATES"
COMPLIANCE_CERT_KEYS = {
    "gas", "electric", "epc", "fire-alarm", "emergency-lighting", "fra", "floor-plan",
    # Added by Norbert 2026-07-30. Unlike the seven above these come from nobody --
    # there is no Monday column behind them, so they start undated on every property.
    "fire-doors", "fire-blanket", "co2-alarm",
}


# Floor Plan is scored on the document held rather than an expiry date, so its
# board splits On file / Not on file. Everything else splits on the date.
COMPLIANCE_PRESENCE_CERTS = {"floor-plan"}


def _compliance_cert_kind(cert):
    return "presence" if cert in COMPLIANCE_PRESENCE_CERTS else "date"


def _ensure_compliance_group_tables():
    """Created on demand so the board works on a database that predates it."""
    db = get_dict_db()
    try:
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_groups ("
            "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "  name TEXT NOT NULL,"
            "  colour TEXT DEFAULT '#2563eb',"
            "  position INTEGER DEFAULT 0,"
            "  created TEXT DEFAULT (datetime('now'))"
            ")"
        )
        # Which certificate the group belongs to (Norbert, 2026-08-06). Additive, so
        # a database written before the change still opens; groups made while they
        # were board-wide would carry '' and are treated as belonging to no
        # certificate rather than to all of them.
        cols = {c["name"] for c in db.execute("PRAGMA table_info(compliance_groups)").fetchall()}
        if "cert_key" not in cols:
            db.execute("ALTER TABLE compliance_groups ADD COLUMN cert_key TEXT NOT NULL DEFAULT ''")
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_group_members ("
            "  group_id INTEGER NOT NULL,"
            "  compliance_id INTEGER NOT NULL,"
            "  cert_key TEXT NOT NULL,"
            "  created TEXT DEFAULT (datetime('now')),"
            "  PRIMARY KEY (compliance_id, cert_key)"
            ")"
        )
        # The order the sections are shown in, set by hand on the board (Norbert,
        # 2026-07-30). Stored as an ordered list of section keys rather than a
        # number on each section, because the built-in sections (Expired, Active,
        # ...) have no row of their own to carry a position. The `kind` column now
        # holds the CERTIFICATE key, one order per board (Norbert, 2026-08-06) —
        # it held 'date'/'presence' while groups were shared across certificates.
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_section_order ("
            "  kind TEXT PRIMARY KEY,"
            "  keys TEXT NOT NULL,"
            "  updated TEXT DEFAULT (datetime('now'))"
            ")"
        )
        # Migrate the two shared orders onto every certificate of that kind, so a
        # board someone had already arranged by hand keeps the arrangement.
        legacy = db.execute(
            "SELECT kind, keys FROM compliance_section_order WHERE kind IN ('date', 'presence')"
        ).fetchall()
        for row in legacy:
            for cert in sorted(COMPLIANCE_CERT_KEYS):
                if _compliance_cert_kind(cert) != row["kind"]:
                    continue
                db.execute(
                    "INSERT INTO compliance_section_order (kind, keys, updated)"
                    " VALUES (?, ?, datetime('now'))"
                    " ON CONFLICT(kind) DO NOTHING",
                    (cert, row["keys"])
                )
            db.execute("DELETE FROM compliance_section_order WHERE kind = ?", (row["kind"],))
        db.commit()
    finally:
        db.close()


# Section keys are the board's own, not database ids: the built-in sections are
# derived from dates and exist only in the UI. Custom groups key as "g:<id>".
COMPLIANCE_SECTION_KEYS = {
    "date": {"expired", "due", "active", "undated", "returned"},
    "presence": {"present-no", "present-yes", "returned"},
}


def _compliance_section_order(db):
    """{cert_key: [key, ...]} — certificates absent until reordered by hand."""
    out = {}
    try:
        for row in db.execute("SELECT kind, keys FROM compliance_section_order").fetchall():
            if row["kind"] in COMPLIANCE_CERT_KEYS:
                out[row["kind"]] = [k for k in str(row["keys"] or "").split(",") if k]
    except Exception:
        pass  # table predates this feature on an old database
    return out


@banksia_os_bp.route("/compliance/board-groups", methods=["GET"])
def api_compliance_board_groups():
    """Custom groups plus every property currently sitting in one.

    Every certificate's groups come back in one response and the board picks out
    the ones for the page it is showing — the payload is tiny and it saves a fetch
    on every certificate switch.
    """
    _ensure_compliance_group_tables()
    db = get_dict_db()
    try:
        groups = db.execute(
            "SELECT id, name, colour, position, cert_key FROM compliance_groups ORDER BY position, id"
        ).fetchall()
        members = db.execute(
            "SELECT group_id, compliance_id, cert_key FROM compliance_group_members"
        ).fetchall()
        order = _compliance_section_order(db)
    finally:
        db.close()
    return json_success({"groups": groups, "members": members, "order": order})


@banksia_os_bp.route("/compliance/section-order", methods=["PUT"])
def api_compliance_section_order_save():
    """Save the order the board's sections are shown in.

    Body: {"cert": "fire-alarm", "keys": ["expired", "g:3", "due", ...]}

    One order per certificate (Norbert, 2026-08-06): arranging the Fire Alarm board
    must not rearrange Emergency Lighting. `kind` is still accepted as the field
    name so an open tab from before the change does not 422 on its next drag.
    """
    _ensure_compliance_group_tables()
    data = request.get_json(silent=True) or {}
    cert = str(data.get("cert") or data.get("kind") or "").strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)
    kind = _compliance_cert_kind(cert)
    raw = data.get("keys")
    if not isinstance(raw, list) or not raw:
        return json_error("Give the order to save", 422)

    db = get_dict_db()
    try:
        live = {
            "g:%d" % g["id"] for g in
            db.execute("SELECT id FROM compliance_groups WHERE cert_key = ?", (cert,)).fetchall()
        }
        allowed = COMPLIANCE_SECTION_KEYS[kind] | live
        keys, seen = [], set()
        for k in raw:
            k = str(k).strip()
            # Unknown or repeated keys are dropped rather than rejected: a group
            # deleted in another tab must not wedge the whole board's order.
            if k in allowed and k not in seen:
                seen.add(k)
                keys.append(k)
        if not keys:
            return json_error("None of those sections exist any more", 422)
        before = ",".join(_compliance_section_order(db).get(cert, []))
        db.execute(
            "INSERT INTO compliance_section_order (kind, keys, updated)"
            " VALUES (?, ?, datetime('now'))"
            " ON CONFLICT(kind) DO UPDATE SET keys = excluded.keys, updated = excluded.updated",
            (cert, ",".join(keys))
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", 0, "update", "section_order_" + cert, before, ",".join(keys),
                  notes="compliance board section order changed")
    return json_success({"cert": cert, "kind": kind, "keys": keys})


@banksia_os_bp.route("/compliance/board-groups", methods=["POST"])
def api_compliance_board_group_create():
    """Create a custom group on ONE certificate's board (Norbert, 2026-08-06)."""
    _ensure_compliance_group_tables()
    data = request.get_json(silent=True) or {}
    cert = str(data.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)
    name = " ".join(str(data.get("name", "")).split())
    if not name:
        return json_error("Give the group a name", 422)
    if len(name) > 60:
        return json_error("Group name is too long (60 characters max)", 422)
    colour = str(data.get("colour", "")).strip() or "#2563eb"
    if not re.fullmatch(r"#[0-9a-fA-F]{6}", colour):
        colour = "#2563eb"
    db = get_dict_db()
    try:
        # Scoped to the certificate: the same name on Gas and Fire Alarm is now two
        # separate groups, which is the point of the change.
        clash = db.execute(
            "SELECT id FROM compliance_groups WHERE cert_key = ? AND LOWER(name) = LOWER(?)",
            (cert, name)
        ).fetchone()
        if clash:
            return json_error("A group called %s already exists on this certificate" % name, 409)
        nxt = db.execute(
            "SELECT COALESCE(MAX(position), 0) + 1 AS p FROM compliance_groups WHERE cert_key = ?",
            (cert,)
        ).fetchone()
        cur = db.execute(
            "INSERT INTO compliance_groups (name, colour, position, cert_key) VALUES (?, ?, ?, ?)",
            (name, colour, nxt["p"], cert)
        )
        db.commit()
        gid = cur.lastrowid
    finally:
        db.close()
    _log_activity("compliance", 0, "create", "board_group", "", name,
                  notes="compliance board group %s created on %s" % (name, cert))
    return json_success({
        "id": gid, "name": name, "colour": colour, "position": nxt["p"], "cert_key": cert,
    })


@banksia_os_bp.route("/compliance/board-groups/<int:group_id>", methods=["DELETE"])
def api_compliance_board_group_delete(group_id):
    """Delete a custom group. Its properties are not touched — they simply fall
    back to the section their expiry date puts them in, so nothing is lost."""
    _ensure_compliance_group_tables()
    db = get_dict_db()
    try:
        row = db.execute(
            "SELECT name, cert_key FROM compliance_groups WHERE id = ?", (group_id,)
        ).fetchone()
        if not row:
            return json_error("Group not found", 404)
        freed = db.execute(
            "SELECT COUNT(DISTINCT compliance_id) AS n FROM compliance_group_members WHERE group_id = ?",
            (group_id,)
        ).fetchone()["n"]
        db.execute("DELETE FROM compliance_group_members WHERE group_id = ?", (group_id,))
        db.execute("DELETE FROM compliance_groups WHERE id = ?", (group_id,))
        db.commit()
    finally:
        db.close()
    _log_activity("compliance", 0, "delete", "board_group", row["name"], "",
                  notes="compliance board group %s deleted from %s (%d propert%s released)"
                        % (row["name"], row["cert_key"] or "no certificate", freed,
                           "y" if freed == 1 else "ies"))
    return json_success({"deleted": group_id, "released": freed, "cert_key": row["cert_key"]})


@banksia_os_bp.route("/compliance/<int:row_id>/group", methods=["PUT"])
def api_compliance_set_group(row_id):
    """Put one property into a custom group on ONE certificate, or take it out.

    Body: {"cert": "gas", "group_id": 3}   group_id null/absent removes it.

    The cert scopes the change (Norbert, 2026-08-06): a property dragged into a
    group on Fire Alarm moves on Fire Alarm and nowhere else. It used to write a
    row for every certificate key, which is what made one drag land on ten boards.
    """
    _ensure_compliance_group_tables()
    data = request.get_json(silent=True) or {}
    cert = str(data.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)
    row = _compliance_row(row_id, "property_name")
    if not row:
        return json_error("Compliance record not found", 404)

    raw = data.get("group_id")
    group_id = None if raw in (None, "", 0) else int(raw)

    db = get_dict_db()
    try:
        before = db.execute(
            "SELECT g.name FROM compliance_group_members m JOIN compliance_groups g ON g.id = m.group_id"
            " WHERE m.compliance_id = ? AND m.cert_key = ?", (row_id, cert)
        ).fetchone()
        old_name = before["name"] if before else ""
        if group_id is None:
            db.execute(
                "DELETE FROM compliance_group_members WHERE compliance_id = ? AND cert_key = ?",
                (row_id, cert)
            )
            new_name = ""
        else:
            grp = db.execute(
                "SELECT name, cert_key FROM compliance_groups WHERE id = ?", (group_id,)
            ).fetchone()
            if not grp:
                return json_error("Group not found", 404)
            # A group belongs to one certificate, so dropping a property into it
            # from another board would file the property somewhere it cannot be
            # seen. Refuse rather than write an invisible row.
            if grp["cert_key"] and grp["cert_key"] != cert:
                return json_error("That group belongs to a different certificate", 422)
            db.execute(
                "INSERT INTO compliance_group_members (group_id, compliance_id, cert_key)"
                " VALUES (?, ?, ?)"
                " ON CONFLICT(compliance_id, cert_key) DO UPDATE SET group_id = excluded.group_id",
                (group_id, row_id, cert)
            )
            new_name = grp["name"]
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", "board_group", old_name, new_name,
                  notes="%s moved to %s on the %s board only"
                        % (row["property_name"], new_name or "its expiry-date section", cert))
    return json_success({"cert": cert, "group_id": group_id, "board_wide": False})

# -- Landlord renewal emails ---------------------------------------------------
# Step 1 of the renewal process (Norbert, 2026-07-30): fifteen days before a
# certificate expires the landlord is emailed with the quote for the work. This is
# the manual half -- a button on the row opens the drafted email, the address and
# wording can be changed, and Send puts it out. Every send is logged so the board
# can show "Emailed" on the row and nobody sends the same reminder twice.
#
# Transport is Missive's drafts API with send:true, the same one the Verv dashboard
# already uses for notification mail. The Banksia backend's smtp_config.json is a
# stub -- empty username -- so SMTP is not a working path here.

MISSIVE_DRAFTS_URL = "https://public.missiveapp.com/v1/drafts"
# Who the reminder comes from. Norbert, 2026-07-30: this is a choice, not a
# constant -- some landlords are dealt with as Banksia and some as Verv, and the
# person sending picks per email.
#
# Both addresses below were confirmed as live Missive senders by an actual send,
# which is the only test that means anything: Missive validates from_field only
# when send is true, so a draft probe returns 201 for any address at all.
# compliance@, info@ and admin@banksialondon.com are NOT connected and are
# rejected. Adding a sender here is safe; guessing one is not.
COMPLIANCE_EMAIL_SENDERS = [
    {"address": "team@banksialondon.com", "name": "Banksia London"},
    {"address": "admin@vervrooms.com", "name": "Verv Rooms"},
]
COMPLIANCE_EMAIL_FROM = COMPLIANCE_EMAIL_SENDERS[0]["address"]
COMPLIANCE_EMAIL_FROM_NAME = COMPLIANCE_EMAIL_SENDERS[0]["name"]


def _sender_for(address):
    """The chosen sender, or None if it is not one we are allowed to send as.

    An allowlist rather than a free-text field on purpose. The From line decides
    who the landlord thinks they are talking to, and it cannot be corrected after
    the email has gone.
    """
    want = str(address or "").strip().lower()
    if not want:
        return COMPLIANCE_EMAIL_SENDERS[0]
    for s in COMPLIANCE_EMAIL_SENDERS:
        if s["address"].lower() == want:
            return s
    return None


def _last_used_sender():
    """The sender used on the most recent reminder, so the choice sticks between
    emails instead of resetting to the top of the list every time."""
    db = get_dict_db()
    try:
        row = db.execute(
            "SELECT from_email FROM compliance_emails"
            " WHERE from_email <> '' ORDER BY id DESC LIMIT 1"
        ).fetchone()
    except Exception:
        row = None
    finally:
        db.close()
    return _sender_for(row["from_email"]) if row else None

# A drafted email always carries this where the price goes. Sending is refused
# while it is still there, because "[QUOTE]" landing in a landlord's inbox is
# worse than a blocked send.
COMPLIANCE_QUOTE_TOKEN = "[QUOTE]"


def _vat_mentioned(body):
    """Norbert, 2026-08-06: no email to a landlord mentions VAT.

    Enforced at send time rather than only kept out of the templates, because the
    body is editable -- the templates have never carried it, so a template-only
    rule would guard the one place the word was never going to come from. Matched
    on word boundaries so 'private' and similar are not caught.
    """
    return bool(re.search(r"\bvat\b", str(body or ""), re.IGNORECASE))

COMPLIANCE_CERT_LABELS = {
    "gas": "gas safety certificate",
    "electric": "electrical certificate (EICR)",
    "epc": "EPC",
    "fire-alarm": "fire alarm certificate",
    "emergency-lighting": "emergency lighting certificate",
    "fra": "fire risk assessment",
    "floor-plan": "floor plan",
    "fire-doors": "fire door inspection",
    "fire-blanket": "fire blanket check",
    "co2-alarm": "CO2 alarm check",
}


def _missive_token():
    for path in ("/root/banksia-backend/missive.json", "/root/.hermes/secrets/missive.json"):
        if os.path.exists(path):
            try:
                return (json.load(open(path)) or {}).get("token")
            except Exception:
                pass
    return None


def _ensure_compliance_email_table():
    db = get_dict_db()
    try:
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_emails ("
            "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "  compliance_id INTEGER NOT NULL,"
            "  cert_key TEXT NOT NULL,"
            "  expiry_date TEXT DEFAULT '',"
            "  to_email TEXT NOT NULL,"
            "  subject TEXT DEFAULT '',"
            "  body TEXT DEFAULT '',"
            "  sent_by TEXT DEFAULT '',"
            "  sent_at TEXT DEFAULT (datetime('now')),"
            "  from_email TEXT DEFAULT '',"
            "  conversation_id TEXT DEFAULT ''"
            ")"
        )
        # Added after the table shipped, so an existing database needs it bolting on.
        cols = {c["name"] for c in db.execute("PRAGMA table_info(compliance_emails)").fetchall()}
        if "from_email" not in cols:
            db.execute("ALTER TABLE compliance_emails ADD COLUMN from_email TEXT DEFAULT ''")
        # Missive's thread id for this email. Stored from the first send onwards so a
        # landlord's reply can be found later and shown against the right property --
        # a thread that was never recorded cannot be matched back afterwards.
        if "conversation_id" not in cols:
            db.execute("ALTER TABLE compliance_emails ADD COLUMN conversation_id TEXT DEFAULT ''")
        # Which of the two emails this was (Norbert, 2026-08-06). 'renewal' is the
        # 15-day chase that asks the landlord to approve the work; 'certificate'
        # sends the finished certificate once it has been uploaded. They are
        # separate events on the same row, so the board must not read one as the
        # other -- the renewal chip in particular means "we have asked", and a
        # certificate email would wrongly satisfy it.
        if "kind" not in cols:
            db.execute("ALTER TABLE compliance_emails ADD COLUMN kind TEXT NOT NULL DEFAULT 'renewal'")
        db.execute(
            "CREATE INDEX IF NOT EXISTS idx_compliance_emails_row"
            " ON compliance_emails (compliance_id, cert_key)"
        )
        db.commit()
    finally:
        db.close()


def _uk_date(value):
    """2026-08-14 -> 14 August 2026. Anything unparseable comes back untouched."""
    try:
        return datetime.strptime(str(value)[:10], "%Y-%m-%d").strftime("%d %B %Y").lstrip("0")
    except Exception:
        return str(value or "")


def _compliance_email_draft(row, cert, sender=None):
    """The email as it is first shown. Everything in it can be edited before sending."""
    label = COMPLIANCE_CERT_LABELS.get(cert, cert.replace("-", " "))
    prop = (row["property_name"] or "").strip()
    landlord = (row["landlord"] or "").strip()
    field = _COMPLIANCE_DATE_FIELDS.get(cert)
    expiry = (row[field] or "").strip() if field else ""
    greeting = landlord if landlord else "Sir or Madam"

    subject = "%s due for renewal - %s" % (label[:1].upper() + label[1:], prop)
    # Deliberately no claim about what the law requires: this one wording goes out
    # for nine different certificates whose legal footing is not the same, and a
    # landlord who checks and finds it overstated will trust none of the rest.
    lines = [
        "Dear %s," % greeting,
        "",
        "We look after the compliance certificates for %s on your behalf, and one of "
        "them is now coming up for renewal." % prop,
        "",
        "The %s expires on %s. We would like to arrange the renewal before that date so "
        "there is no gap in the property's compliance."
        % (label, _uk_date(expiry) or "its recorded date"),
        "",
        "We can book the work for you. Our quote is %s." % COMPLIANCE_QUOTE_TOKEN,
        "",
        "If you are happy to go ahead, just reply to confirm. We will book the contractor, "
        "arrange access with the tenants where it is needed, and send you the new "
        "certificate as soon as it has been issued.",
        "",
        "If you would rather arrange it yourself, that is no problem — please let us know, "
        "and send us a copy of the new certificate so our records stay up to date.",
        "",
        "Kind regards,",
        (sender or COMPLIANCE_EMAIL_SENDERS[0])["name"],
    ]
    return {"subject": subject, "body": "\n".join(lines), "expiry_date": expiry}


# Letterhead per sending mailbox (Norbert, 2026-08-06: a bare email risks the spam
# folder). Banksia's postal details are taken from their own letterhead on the
# guarantor form, not invented -- an address nobody can verify in a footer is
# worse than no address. Verv Rooms has no postal address on record here, so it
# shows none rather than borrowing Banksia's.
EMAIL_BRANDS = {
    "team@banksialondon.com": {
        "name": "Banksia London",
        "accent": "#0f766e",
        "address": "Banksia Limited, 29-31 Adelaide Road, London NW3 7BB",
        "contact": "team@banksialondon.com",
    },
    "admin@vervrooms.com": {
        "name": "Verv Rooms",
        "accent": "#52BA31",
        "address": "",
        "contact": "admin@vervrooms.com",
    },
}


def _email_html(body_text, sender=None):
    """Plain text in, email-safe HTML out.

    The editor is a plain textarea on purpose -- these are short letters, and a
    rich editor would only invite formatting that renders differently in every
    mail client. The letterhead is added here instead, so every email carries a
    sender name, a footer and a reason for existing: a message that is little more
    than an attachment is the shape spam filters are built to catch (Norbert,
    2026-08-06).
    """
    brand = EMAIL_BRANDS.get(
        ((sender or {}).get("address") or "").lower(),
        EMAIL_BRANDS["team@banksialondon.com"],
    )
    esc = (body_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    paragraphs = [p.strip() for p in esc.split("\n\n")]
    html = "".join(
        "<p style=\"margin:0 0 14px;\">%s</p>" % p.replace("\n", "<br>")
        for p in paragraphs if p
    )
    footer_address = (
        "<div style=\"margin-top:4px;\">%s</div>" % brand["address"]
    ) if brand["address"] else ""
    return (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">"
        "</head><body style=\"margin:0;padding:0;background:#f2f4f7;\">"
        "<table role=\"presentation\" width=\"100%%\" cellpadding=\"0\" cellspacing=\"0\">"
        "<tr><td align=\"center\" style=\"padding:26px 12px;\">"
        "<table role=\"presentation\" width=\"100%%\" style=\"max-width:600px;background:#ffffff;"
        "border-radius:12px;overflow:hidden;font-family:Arial,Helvetica,sans-serif;"
        "color:#1a2233;font-size:15px;line-height:1.6;"
        "border:1px solid #e3e8ef;\">"
        # Letterhead
        "<tr><td style=\"background:%(accent)s;padding:18px 32px;\">"
        "<div style=\"color:#ffffff;font-size:17px;font-weight:bold;letter-spacing:0.3px;\">"
        "%(name)s</div>"
        "<div style=\"color:#ffffff;opacity:0.85;font-size:12.5px;margin-top:2px;\">"
        "Property compliance</div>"
        "</td></tr>"
        # Letter
        "<tr><td style=\"padding:28px 32px 8px;\">%(body)s</td></tr>"
        # Footer
        "<tr><td style=\"padding:14px 32px 24px;\">"
        "<div style=\"border-top:1px solid #e3e8ef;padding-top:14px;"
        "font-size:11.5px;line-height:1.55;color:#6b7688;\">"
        "This email was sent by %(name)s regarding a property we manage or maintain "
        "compliance records for. If anything here looks wrong, reply to this email "
        "and a person will read it."
        "<div style=\"margin-top:8px;\">%(contact)s</div>"
        "%(address)s"
        "</div></td></tr>"
        "</table></td></tr></table></body></html>"
    ) % {
        "accent": brand["accent"],
        "name": brand["name"],
        "contact": brand["contact"],
        "address": footer_address,
        "body": html,
    }


def _compliance_cert_bytes(row_id, cert):
    """(filename, bytes, error) for a property's certificate, whichever way it is
    stored -- uploaded here ("local:<name>" on disk) or still a Monday asset id.

    The same two shapes api_compliance_certificate streams to the browser. Pulled
    out into a helper so emailing a certificate cannot drift from viewing one.
    """
    import urllib.request

    field = _COMPLIANCE_DOC_FIELDS.get(cert)
    if not field:
        return None, None, "Unknown certificate type"
    row = _compliance_row(row_id, field)
    if not row:
        return None, None, "Compliance record not found"
    ref = (row["ref"] or "").strip()
    if not ref:
        return None, None, "There is no certificate on record to attach"

    if ref.startswith("local:"):
        name = ref[len("local:"):]
        path = os.path.join(COMPLIANCE_UPLOAD_DIR, name)
        if not os.path.exists(path):
            return None, None, "The certificate file is missing on disk"
        with open(path, "rb") as fh:
            return name, fh.read(), None

    m = re.search(r"/resources/(\d+)/", ref)
    asset_id = m.group(1) if m else ref
    if not asset_id.isdigit():
        return None, None, "The certificate reference is not a Monday asset"
    mtok = get_monday_token()
    if not mtok:
        return None, None, "Monday credentials unavailable on this server"
    try:
        data = _monday_graphql(mtok, "{assets(ids:[%s]){name file_extension public_url}}" % asset_id)
        assets = ((data or {}).get("data") or {}).get("assets") or []
        if not assets or not assets[0].get("public_url"):
            return None, None, "The certificate file is no longer available on Monday"
        asset = assets[0]
        with urllib.request.urlopen(asset["public_url"], timeout=45) as r:
            return (asset.get("name") or "certificate"), r.read(), None
    except Exception as e:
        return None, None, safe_error(e)


def _send_via_missive(to_email, subject, body_text, sender=None, attachments=None):
    """Send it. Returns (ok, error, conversation_id) -- the thread id is what makes
    a later reply findable, so it is carried back even though nothing reads it yet.

    `attachments` is [(filename, bytes)]. Missive takes them base64-encoded on the
    draft itself; verified 2026-08-06 by sending a 58KB PDF internally and reading
    it back off the sent message at the same byte length.
    """
    # Imported here rather than at module level to match _monday_graphql above,
    # the only other outbound HTTP call in this file. A module-level import was
    # skipped by the earlier patch because that local import already matched the
    # 'is urllib imported' check.
    import base64
    import urllib.request
    import urllib.error
    token = _missive_token()
    if not token:
        return False, "No Missive token on this server - email is not configured", ""
    sender = sender or COMPLIANCE_EMAIL_SENDERS[0]
    payload = {
        "drafts": {
            "subject": subject,
            "body": _email_html(body_text, sender),
            "to_fields": [{"address": to_email}],
            "from_field": {"address": sender["address"], "name": sender["name"]},
            "send": True,
        }
    }
    if attachments:
        payload["drafts"]["attachments"] = [
            {"base64_data": base64.b64encode(blob).decode(), "filename": name}
            for name, blob in attachments
        ]
    try:
        req = urllib.request.Request(
            MISSIVE_DRAFTS_URL,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Authorization": "Bearer " + token, "Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=25) as res:
            if 200 <= res.status < 300:
                conv = ""
                try:
                    conv = str(((json.loads(res.read().decode("utf-8"))
                                 or {}).get("drafts") or {}).get("conversation") or "")
                except Exception:
                    pass  # the email went; a missing thread id is not worth failing over
                return True, None, conv
            return False, "Missive returned %s" % res.status, ""
    except urllib.error.HTTPError as e:
        detail = ""
        try:
            detail = e.read().decode("utf-8", "replace")[:200]
        except Exception:
            pass
        return False, "Missive rejected the email (%s) %s" % (e.code, detail), ""
    except Exception as e:
        return False, "Could not reach Missive: %s" % str(e)[:150], ""


@banksia_os_bp.route("/compliance/emails", methods=["GET"])
def api_compliance_emails():
    """Every reminder already sent, so the board can mark the rows."""
    _ensure_compliance_email_table()
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT compliance_id, cert_key, expiry_date, to_email, from_email, sent_by,"
            " MAX(sent_at) AS sent_at"
            " FROM compliance_emails WHERE kind = 'renewal'"
            " GROUP BY compliance_id, cert_key, expiry_date"
        ).fetchall()
    finally:
        db.close()
    return json_success(rows)


@banksia_os_bp.route("/compliance/<int:row_id>/email", methods=["GET"])
def api_compliance_email_draft(row_id):
    """The drafted reminder for one property's certificate."""
    _ensure_compliance_email_table()
    cert = str(request.args.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)

    # Whatever went out last time, so the choice is not made again every email.
    sender = _last_used_sender() or COMPLIANCE_EMAIL_SENDERS[0]

    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        if not row:
            return json_error("Compliance record not found", 404)
        draft = _compliance_email_draft(row, cert, sender)
        prev = db.execute(
            "SELECT to_email, from_email, sent_at, sent_by FROM compliance_emails"
            " WHERE compliance_id = ? AND cert_key = ? AND expiry_date = ? AND kind = 'renewal'"
            " ORDER BY id DESC LIMIT 1",
            (row_id, cert, draft["expiry_date"])
        ).fetchone()
    finally:
        db.close()

    return json_success({
        "property_name": row["property_name"],
        "landlord": row["landlord"] or "",
        "to": (row["landlord_email"] or "").strip(),
        "subject": draft["subject"],
        "body": draft["body"],
        "expiry_date": draft["expiry_date"],
        "quote_token": COMPLIANCE_QUOTE_TOKEN,
        "from": sender["address"],
        "senders": COMPLIANCE_EMAIL_SENDERS,
        "sign_off": sender["name"],
        "already_sent": prev,
    })


@banksia_os_bp.route("/compliance/<int:row_id>/email", methods=["POST"])
def api_compliance_email_send(row_id):
    """Send the reminder and write it down.

    Body: {"cert": "gas", "to": "...", "subject": "...", "body": "...", "confirm": true}

    `confirm` is only needed to send a second time for the same expiry date -- the
    first send is what the board marks as done, and a duplicate reminder to a
    landlord is the thing Norbert asked to be protected from.
    """
    _ensure_compliance_email_table()
    data = request.get_json(silent=True) or {}
    cert = str(data.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)

    sender = _sender_for(data.get("from"))
    if sender is None:
        return json_error("That is not one of the addresses this board can send from", 422)

    to_email = str(data.get("to", "")).strip()
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", to_email):
        return json_error("That is not an email address", 422)
    subject = " ".join(str(data.get("subject", "")).split())
    if not subject:
        return json_error("The email needs a subject", 422)
    body = str(data.get("body", "")).strip()
    if not body:
        return json_error("The email is empty", 422)
    if COMPLIANCE_QUOTE_TOKEN in body:
        return json_error("Replace %s with the quote before sending" % COMPLIANCE_QUOTE_TOKEN, 422)
    if _vat_mentioned(body):
        return json_error("These emails do not mention VAT — please take it out of the wording", 422)

    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        if not row:
            return json_error("Compliance record not found", 404)
        field = _COMPLIANCE_DATE_FIELDS.get(cert)
        expiry = (row[field] or "").strip() if field else ""
        prev = db.execute(
            "SELECT sent_at FROM compliance_emails"
            " WHERE compliance_id = ? AND cert_key = ? AND expiry_date = ? AND kind = 'renewal'"
            " ORDER BY id DESC LIMIT 1",
            (row_id, cert, expiry)
        ).fetchone()
    finally:
        db.close()

    if prev and not data.get("confirm"):
        return json_error(
            "This reminder was already sent on %s. Send it again?" % prev["sent_at"], 409)

    ok, err, conversation_id = _send_via_missive(to_email, subject, body, sender)
    if not ok:
        return json_error(err or "The email could not be sent", 502)

    actor = _archive_actor()
    db = get_dict_db()
    try:
        db.execute(
            "INSERT INTO compliance_emails (compliance_id, cert_key, expiry_date, to_email,"
            " from_email, subject, body, sent_by, conversation_id, kind)"
            " VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'renewal')",
            (row_id, cert, expiry, to_email, sender["address"], subject, body, actor,
             conversation_id)
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", "renewal_email_" + cert, "", to_email,
                  notes="%s renewal reminder emailed to %s from %s for %s"
                        % (cert, to_email, sender["address"], row["property_name"]))
    return json_success({"cert": cert, "to": to_email, "from": sender["address"],
                         "resent": bool(prev)})


# -- Certificate issued: send it to the landlord -------------------------------
# Norbert, 2026-08-06: the last step of the renewal. Once a certificate has been
# uploaded, the landlord gets it -- their property, their document, and they
# should not have to ask for it.
#
# It is NOT sent automatically. Uploading opens the drafted email and somebody
# presses Send, which is what Norbert asked for and is also the safeguard: an
# upload is a filing action, and filing a document should never put mail in
# somebody's inbox on its own. Every field can be edited first.
#
# admin@vervrooms.com is the default sender here (Norbert's choice) rather than
# team@banksialondon.com, which is the default on the renewal chase.

CERTIFICATE_EMAIL_SENDER = "admin@vervrooms.com"


def _certificate_email_draft(row, cert, sender=None):
    """The email as first shown. Name, address and certificate type are what Norbert
    asked to see in it; the certificate itself rides along as the attachment."""
    label = COMPLIANCE_CERT_LABELS.get(cert, cert.replace("-", " "))
    prop = (row["property_name"] or "").strip()
    landlord = (row["landlord"] or "").strip()
    field = _COMPLIANCE_DATE_FIELDS.get(cert)
    expiry = (row[field] or "").strip() if field else ""
    # 3 of 62 properties have no landlord name on record. "Dear ," is worse than a
    # formal greeting, so an unknown name falls back rather than printing an empty.
    greeting = landlord if landlord else "Sir or Madam"

    lines = [
        "Dear %s," % greeting,
        "",
        "We arrange and keep track of the compliance certificates for %s on your behalf." % prop,
        "",
        "The %s has now been carried out and the certificate is attached to this email. "
        "Please keep a copy for your records — it is the document you would need to produce "
        "if it were ever asked for." % label,
    ]
    if expiry:
        lines += [
            "",
            "The certificate is valid until %s. We hold the renewal date on file and will "
            "contact you before it runs out, so there is nothing you need to do now."
            % _uk_date(expiry),
        ]
    else:
        lines += [
            "",
            "There is nothing you need to do — we hold this on file and will let you know "
            "when it next needs attention.",
        ]
    lines += [
        "",
        "If anything in the certificate looks wrong, or you have a question about the work "
        "that was done, just reply to this email and we will look into it.",
        "",
        "Kind regards,",
        (sender or _sender_for(CERTIFICATE_EMAIL_SENDER))["name"],
    ]
    return {
        "subject": "%s - %s" % (label[:1].upper() + label[1:], prop),
        "body": "\n".join(lines),
        "expiry_date": expiry,
    }


@banksia_os_bp.route("/compliance/<int:row_id>/certificate-email", methods=["GET"])
def api_compliance_certificate_email_draft(row_id):
    """The drafted 'here is your certificate' email, for the confirmation step.

    Query: ?cert=gas
    """
    _ensure_compliance_email_table()
    cert = str(request.args.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)

    sender = _sender_for(CERTIFICATE_EMAIL_SENDER) or COMPLIANCE_EMAIL_SENDERS[0]
    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        if not row:
            return json_error("Compliance record not found", 404)
        draft = _certificate_email_draft(row, cert, sender)
        prev = db.execute(
            "SELECT to_email, from_email, sent_at, sent_by FROM compliance_emails"
            " WHERE compliance_id = ? AND cert_key = ? AND kind = 'certificate'"
            " ORDER BY id DESC LIMIT 1",
            (row_id, cert)
        ).fetchone()
    finally:
        db.close()

    # Resolved now, not at send time, so the modal can say what is attached -- and
    # so a missing file is discovered before anybody has written an email.
    filename, blob, err = _compliance_cert_bytes(row_id, cert)
    return json_success({
        "property_name": row["property_name"],
        "landlord": row["landlord"] or "",
        "to": (row["landlord_email"] or "").strip(),
        "subject": draft["subject"],
        "body": draft["body"],
        "from": sender["address"],
        "senders": COMPLIANCE_EMAIL_SENDERS,
        "sign_off": sender["name"],
        "attachment": filename,
        "attachment_size": len(blob) if blob else 0,
        "attachment_error": err,
        "already_sent": prev,
    })


@banksia_os_bp.route("/compliance/<int:row_id>/certificate-email", methods=["POST"])
def api_compliance_certificate_email_send(row_id):
    """Send the certificate to the landlord, with the certificate attached.

    Body: {"cert": "gas", "to": "...", "subject": "...", "body": "...", "confirm": true}

    Refuses outright if there is no certificate to attach: an email that says "the
    certificate is attached" and carries nothing is worse than no email at all.
    """
    _ensure_compliance_email_table()
    data = request.get_json(silent=True) or {}
    cert = str(data.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)

    sender = _sender_for(data.get("from") or CERTIFICATE_EMAIL_SENDER)
    if sender is None:
        return json_error("That is not one of the addresses this board can send from", 422)

    to_email = str(data.get("to", "")).strip()
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", to_email):
        return json_error("That is not an email address", 422)
    subject = " ".join(str(data.get("subject", "")).split())
    if not subject:
        return json_error("The email needs a subject", 422)
    body = str(data.get("body", "")).strip()
    if not body:
        return json_error("The email is empty", 422)
    if _vat_mentioned(body):
        return json_error("These emails do not mention VAT — please take it out of the wording", 422)

    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        if not row:
            return json_error("Compliance record not found", 404)
        # Property Returned is exempt from every automation (Norbert, 2026-07-30) --
        # emailing a certificate to the landlord of a property we handed back is
        # exactly the kind of thing that rule exists to stop.
        if int(row["automation_exempt"] or 0) == 1:
            return json_error("This property has been returned to the landlord", 422)
        field = _COMPLIANCE_DATE_FIELDS.get(cert)
        expiry = (row[field] or "").strip() if field else ""
        prev = db.execute(
            "SELECT sent_at FROM compliance_emails"
            " WHERE compliance_id = ? AND cert_key = ? AND expiry_date = ? AND kind = 'certificate'"
            " ORDER BY id DESC LIMIT 1",
            (row_id, cert, expiry)
        ).fetchone()
    finally:
        db.close()

    # Keyed on the expiry date like the renewal log, so next year's certificate for
    # the same property is a new send rather than a blocked duplicate.
    if prev and not data.get("confirm"):
        return json_error(
            "This certificate was already sent on %s. Send it again?" % prev["sent_at"], 409)

    filename, blob, err = _compliance_cert_bytes(row_id, cert)
    if err or not blob:
        return json_error(err or "There is no certificate on record to attach", 422)

    ok, send_err, conversation_id = _send_via_missive(
        to_email, subject, body, sender, attachments=[(filename, blob)])
    if not ok:
        return json_error(send_err or "The email could not be sent", 502)

    actor = _archive_actor()
    db = get_dict_db()
    try:
        db.execute(
            "INSERT INTO compliance_emails (compliance_id, cert_key, expiry_date, to_email,"
            " from_email, subject, body, sent_by, conversation_id, kind)"
            " VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'certificate')",
            (row_id, cert, expiry, to_email, sender["address"], subject, body, actor,
             conversation_id)
        )
        db.commit()
    finally:
        db.close()

    _log_activity("compliance", row_id, "update", "certificate_email_" + cert, "", to_email,
                  notes="%s certificate (%s) emailed to %s from %s for %s"
                        % (cert, filename, to_email, sender["address"], row["property_name"]))
    return json_success({"cert": cert, "to": to_email, "from": sender["address"],
                         "attachment": filename, "resent": bool(prev)})


# -- Landlord replies ----------------------------------------------------------
# Norbert, 2026-07-30: "if the Landlord replies, means someone needs to go on
# Missive and check the email". They should not have to. The board sent the
# email, so the board shows the answer.
#
# Read-only on purpose. Missive stays the system of record -- nothing is moved
# out of it, nothing is deleted, and replying from the board is deliberately NOT
# built while the standing instruction is that no landlord is emailed.
#
# The link between a property and a thread is `compliance_emails.conversation_id`,
# written at send time. It cannot be reconstructed afterwards, which is why it
# went in before any of this.

MISSIVE_API = "https://public.missiveapp.com/v1"

# Missive has no webhook that can be created through the API (/v1/hooks is a 404 --
# it is set up by hand in Missive's own integration settings). So replies are
# pulled, not pushed. Pulling on every board load would hammer their API for no
# reason, so a refresh does nothing if one ran in the last minute.
_REPLY_REFRESH_MIN_GAP = 60
_reply_refresh_last = [0.0]


def _ensure_compliance_reply_table():
    db = get_dict_db()
    try:
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_email_replies ("
            "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "  compliance_id INTEGER NOT NULL,"
            "  cert_key TEXT NOT NULL,"
            "  conversation_id TEXT NOT NULL,"
            "  message_id TEXT NOT NULL UNIQUE,"
            "  from_email TEXT DEFAULT '',"
            "  from_name TEXT DEFAULT '',"
            "  subject TEXT DEFAULT '',"
            "  body TEXT DEFAULT '',"
            "  attachments TEXT DEFAULT '',"
            "  received_at TEXT DEFAULT '',"
            "  seen INTEGER DEFAULT 0,"
            "  created TEXT DEFAULT (datetime('now'))"
            ")"
        )
        db.execute(
            "CREATE INDEX IF NOT EXISTS idx_compliance_replies_row"
            " ON compliance_email_replies (compliance_id, cert_key)"
        )
        db.commit()
    finally:
        db.close()


def _missive_get(path):
    import urllib.request
    import urllib.error
    token = _missive_token()
    if not token:
        return None
    try:
        req = urllib.request.Request(
            MISSIVE_API + path,
            headers={"Authorization": "Bearer " + token},
            method="GET",
        )
        with urllib.request.urlopen(req, timeout=20) as res:
            return json.loads(res.read().decode("utf-8"))
    except Exception:
        return None


_TAG_RE = re.compile(r"<[^>]+>")
_STYLE_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.S | re.I)


def _html_to_text(html):
    """A reply arrives as HTML written by someone outside the business. It is
    reduced to plain text here rather than rendered, so nothing a landlord (or
    anyone forging their address) puts in an email can execute on our board."""
    v = str(html or "")
    v = _STYLE_RE.sub(" ", v)
    v = re.sub(r"<br\s*/?>", "\n", v, flags=re.I)
    v = re.sub(r"</p\s*>", "\n\n", v, flags=re.I)
    v = _TAG_RE.sub("", v)
    for a, b in (("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"),
                 ("&quot;", '"'), ("&#39;", "'")):
        v = v.replace(a, b)
    v = re.sub(r"[ \t]+", " ", v)
    # Marketing and portal mail is built out of nested tables, which leaves a line of
    # whitespace per cell. Trimming each line first is what makes the blank-run
    # collapse below actually work -- " \n \n " is not an empty run until it is.
    lines = [ln.strip() for ln in v.split("\n")]
    out = []
    for ln in lines:
        if not ln and out and not out[-1]:
            continue
        out.append(ln)
    return "\n".join(out).strip()


# A reply usually carries the whole message it is answering underneath it. The
# board already shows what we sent directly above, so the quoted copy is pure
# repetition -- it is cut here rather than in the page, so the stored body is
# clean for anything else that reads it later.
_QUOTE_CUTS = (
    # Some clients put the quoted body on the same line as the attribution
    # ("...wrote:Dear Norbert"), so this searches rather than matching a line.
    re.compile(r"On .{0,200}?wrote:", re.I),
    re.compile(r"-{2,}\s*Original Message\s*-{2,}", re.I),
    re.compile(r"^_{5,}$", re.M),
    re.compile(r"^From: .+$", re.I | re.M),
    re.compile(r"^Sent from my .*$", re.I | re.M),
    re.compile(r"^>", re.M),
)


def _strip_quoted_reply(text):
    v = str(text or "")
    cut = None
    for rx in _QUOTE_CUTS:
        m = rx.search(v)
        if m and (cut is None or m.start() < cut):
            cut = m.start()
    if cut is None:
        return v.strip()
    head = v[:cut].strip()
    # A reply that is nothing but a quote is better shown whole than shown blank.
    return head or v.strip()


def _our_addresses():
    return {x["address"].lower() for x in COMPLIANCE_EMAIL_SENDERS}


def _ingest_thread(thread, known, ours):
    """Read one Missive conversation and store anything on it that is not ours.

    Pulled out of the sweep so the webhook can do a single thread rather than all
    of them. `known` is mutated so one run never stores a message twice; the
    UNIQUE on message_id is the backstop for two runs racing (a webhook and a
    board load can land in the same second).
    """
    data = _missive_get("/conversations/%s/messages" % thread["conversation_id"])
    if not data:
        return 0
    found = 0
    for m in (data.get("messages") or []):
        mid = str(m.get("id") or "")
        frm = ((m.get("from_field") or {}).get("address") or "").lower()
        # Ours went out; anything else on the thread is the answer.
        if not mid or mid in known or frm in ours:
            continue
        full = _missive_get("/messages/%s" % mid) or {}
        full = full.get("messages") or {}
        body = _strip_quoted_reply(_html_to_text(full.get("body") or m.get("preview") or ""))
        atts = [a.get("filename") or a.get("name") or "attachment"
                for a in (full.get("attachments") or [])]
        stamp = full.get("delivered_at") or m.get("delivered_at") or m.get("created_at")
        try:
            stamp = datetime.fromtimestamp(int(stamp), timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            stamp = ""
        db = get_dict_db()
        try:
            db.execute(
                "INSERT OR IGNORE INTO compliance_email_replies (compliance_id, cert_key,"
                " conversation_id, message_id, from_email, from_name, subject, body,"
                " attachments, received_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (thread["compliance_id"], thread["cert_key"], thread["conversation_id"], mid, frm,
                 (m.get("from_field") or {}).get("name") or "",
                 full.get("subject") or m.get("subject") or "", body,
                 ", ".join(atts), stamp)
            )
            db.commit()
        finally:
            db.close()
        known.add(mid)
        found += 1
    return found


def _refresh_compliance_replies(force=False):
    """Pull anything new on the threads we started. Returns how many replies landed."""
    import time
    if not force and (time.time() - _reply_refresh_last[0]) < _REPLY_REFRESH_MIN_GAP:
        return None  # too soon -- say nothing rather than pretend a check happened
    _reply_refresh_last[0] = time.time()

    _ensure_compliance_email_table()
    _ensure_compliance_reply_table()
    db = get_dict_db()
    try:
        threads = db.execute(
            "SELECT DISTINCT conversation_id, compliance_id, cert_key FROM compliance_emails"
            " WHERE conversation_id <> ''"
        ).fetchall()
        known = {r["message_id"] for r in
                 db.execute("SELECT message_id FROM compliance_email_replies").fetchall()}
    finally:
        db.close()

    ours = _our_addresses()
    found = 0
    for t in threads:
        found += _ingest_thread(t, known, ours)
    return found


# -- Missive webhook -----------------------------------------------------------
# Sami and Norbert, 2026-07-30: "We have to add a webhook into Missive". Missive
# has no API that creates one, so the rule is added by hand in Missive under
# Settings > Rules with a Webhook action pointed at this route. This end is the
# half that can be built, and it is built.
#
# The posted body is deliberately NOT trusted as the reply. A valid signature
# only tells us something happened on a thread; the message is then read back
# from Missive's own API exactly as a board load reads it. So the worst a forged
# or replayed post can do is make us re-read a conversation we already own.
_MISSIVE_HOOK_PATH = "/api/banksia-os/hooks/missive"
_MISSIVE_HOOK_SECRET_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                         "missive_hook_secret")


def _missive_hook_secret():
    try:
        with open(_MISSIVE_HOOK_SECRET_FILE) as fh:
            return fh.read().strip()
    except Exception:
        return ""


@banksia_os_bp.route("/hooks/missive", methods=["POST"])
def api_missive_hook():
    """Missive tells us a thread moved; we go and read it.

    Unsigned posts are accepted and ignored, because Missive pings the endpoint
    to validate it before the rule can be saved and that ping carries no
    signature. Ignoring costs nothing -- nothing is read and nothing is written.
    A signature that is present but wrong is refused outright.
    """
    import hashlib
    import hmac as _hmac

    raw = request.get_data() or b""
    sent = request.headers.get("X-Hook-Signature", "")
    if not sent:
        return json_success({"signed": False, "matched": False, "new": 0})

    secret = _missive_hook_secret()
    if not secret:
        return json_error("Webhook secret not configured on this server", 503)
    mine = "sha256=" + _hmac.new(secret.encode("utf-8"), raw, hashlib.sha256).hexdigest()
    if not _hmac.compare_digest(sent, mine):
        return json_error("Bad signature", 401)

    payload = request.get_json(silent=True) or {}
    conv = str((payload.get("conversation") or {}).get("id") or "")

    _ensure_compliance_email_table()
    _ensure_compliance_reply_table()
    db = get_dict_db()
    try:
        if conv:
            threads = db.execute(
                "SELECT DISTINCT conversation_id, compliance_id, cert_key FROM compliance_emails"
                " WHERE conversation_id = ?", (conv,)).fetchall()
        else:
            threads = db.execute(
                "SELECT DISTINCT conversation_id, compliance_id, cert_key FROM compliance_emails"
                " WHERE conversation_id <> ''").fetchall()
        known = {r["message_id"] for r in
                 db.execute("SELECT message_id FROM compliance_email_replies").fetchall()}
    finally:
        db.close()

    # A hook for a conversation this board never started is the normal case --
    # the rule watches a whole mailbox, most of which is nothing to do with
    # compliance. Do nothing, quietly.
    if not threads:
        return json_success({"signed": True, "matched": False, "new": 0})

    ours = _our_addresses()
    found = 0
    for t in threads:
        found += _ingest_thread(t, known, ours)
    return json_success({"signed": True, "matched": True, "new": found})


@banksia_os_bp.route("/compliance/emails/refresh", methods=["POST"])
def api_compliance_replies_refresh():
    """Check the threads for new answers. Safe to call often -- it throttles itself."""
    found = _refresh_compliance_replies(force=bool((request.get_json(silent=True) or {}).get("force")))
    if found is None:
        return json_success({"checked": False, "new": 0})
    return json_success({"checked": True, "new": found})


@banksia_os_bp.route("/compliance/replies", methods=["GET"])
def api_compliance_replies():
    """One line per row that has been answered, so the board can mark it."""
    _ensure_compliance_reply_table()
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT compliance_id, cert_key, COUNT(*) AS reply_count,"
            " MAX(received_at) AS last_reply, MIN(seen) AS all_seen"
            " FROM compliance_email_replies GROUP BY compliance_id, cert_key"
        ).fetchall()
    finally:
        db.close()
    return json_success(rows)


@banksia_os_bp.route("/compliance/<int:row_id>/thread", methods=["GET"])
def api_compliance_thread(row_id):
    """The whole conversation for one property's certificate: what we sent and
    what came back, oldest first."""
    _ensure_compliance_email_table()
    _ensure_compliance_reply_table()
    cert = str(request.args.get("cert", "")).strip().lower()
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate type", 422)

    db = get_dict_db()
    try:
        prop = db.execute("SELECT property_name, landlord FROM compliance WHERE id = ?",
                          (row_id,)).fetchone()
        if not prop:
            return json_error("Compliance record not found", 404)
        sent = db.execute(
            "SELECT to_email, from_email, subject, body, sent_at, sent_by"
            " FROM compliance_emails WHERE compliance_id = ? AND cert_key = ?"
            " ORDER BY sent_at",
            (row_id, cert)
        ).fetchall()
        replies = db.execute(
            "SELECT id, from_email, from_name, subject, body, attachments, received_at"
            " FROM compliance_email_replies WHERE compliance_id = ? AND cert_key = ?"
            " ORDER BY received_at",
            (row_id, cert)
        ).fetchall()
        if replies:
            db.execute(
                "UPDATE compliance_email_replies SET seen = 1"
                " WHERE compliance_id = ? AND cert_key = ?", (row_id, cert))
            db.commit()
    finally:
        db.close()

    messages = []
    for m in sent:
        messages.append({
            "direction": "out", "who": m["from_email"], "name": "",
            "to": m["to_email"], "subject": m["subject"], "body": m["body"],
            "at": m["sent_at"], "by": m["sent_by"], "attachments": "",
        })
    for m in replies:
        messages.append({
            "direction": "in", "who": m["from_email"], "name": m["from_name"],
            "to": "", "subject": m["subject"], "body": m["body"],
            "at": m["received_at"], "by": "", "attachments": m["attachments"],
        })
    messages.sort(key=lambda x: x["at"] or "")

    return json_success({
        "property_name": prop["property_name"],
        "landlord": prop["landlord"] or "",
        "messages": messages,
    })


# -- Contractors on the /compliance-test board ---------------------------------
# Norbert, 2026-07-30: the board needs to know who does what, so the renewal
# chaser can message the right trade's WhatsApp group 15 days before a
# certificate expires. Full name, WhatsApp group id, and the certificates they
# cover. This table is the single source of truth for that automation -- it used
# to be a hard-coded map in the script, which meant every new contractor was a
# code change.


def _ensure_compliance_contractor_table():
    """Created on demand so the board works on a database that predates it."""
    db = get_dict_db()
    try:
        db.execute(
            "CREATE TABLE IF NOT EXISTS compliance_contractors ("
            "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "  name TEXT NOT NULL,"
            "  group_id TEXT DEFAULT '',"
            "  certs TEXT DEFAULT '',"
            "  preferred_certs TEXT DEFAULT '',"
            "  created TEXT DEFAULT (datetime('now')),"
            "  updated TEXT DEFAULT (datetime('now'))"
            ")"
        )
        cols = [r["name"] for r in db.execute("PRAGMA table_info(compliance_contractors)").fetchall()]
        if "preferred_certs" not in cols:
            db.execute("ALTER TABLE compliance_contractors ADD COLUMN preferred_certs TEXT DEFAULT ''")
        # What trade they are, which is not the same question as which
        # certificates they sign: a painter belongs on this list and will never
        # hold one (Norbert, 2026-08-07).
        if "trades" not in cols:
            db.execute("ALTER TABLE compliance_contractors ADD COLUMN trades TEXT DEFAULT ''")
        db.commit()
    finally:
        db.close()


def _clean_group_id(raw):
    """Normalise a WhatsApp group id, or explain why it is not one.

    Getting this wrong means messaging strangers, so a half-typed id is rejected
    rather than stored. Both `1203...@g.us` and the bare digits are accepted --
    people copy it either way -- and the suffix is added back on.
    """
    v = str(raw or "").strip()
    if not v:
        return "", None
    if v.endswith("@g.us"):
        digits = v[:-5]
    else:
        digits = v
    if not digits.isdigit() or not (10 <= len(digits) <= 25):
        return None, "That is not a WhatsApp group id - it should look like 120363260519419014@g.us"
    return digits + "@g.us", None


def _clean_certs(raw):
    """Keep only certificates the board actually has, in the board's own order."""
    if isinstance(raw, str):
        items = [x.strip().lower() for x in raw.split(",")]
    elif isinstance(raw, (list, tuple)):
        items = [str(x).strip().lower() for x in raw]
    else:
        return ""
    order = ["gas", "electric", "epc", "fire-alarm", "emergency-lighting", "fra",
             "fire-doors", "fire-blanket", "co2-alarm", "floor-plan"]
    keep = [k for k in order if k in items]
    return ",".join(keep)


def _clean_trades(raw):
    """Keep only real categories, in the list's own order.

    Matched case-insensitively because the value arrives from a picker but may
    also arrive from a script, and "plumbing" should not create a second category
    alongside "Plumbing".
    """
    if isinstance(raw, str):
        items = [x.strip().lower() for x in raw.split(",")]
    elif isinstance(raw, (list, tuple)):
        items = [str(x).strip().lower() for x in raw]
    else:
        return ""
    return ",".join([t for t in MAINT_TRADES if t.lower() in items])


def _clean_preferred(raw, certs):
    """Certificates this contractor is the first call for.

    Three trades cover gas. Until now the chaser messaged whoever sorted first
    alphabetically, which is an accident rather than a decision (Sami,
    2026-08-06). A preference can only name a certificate they actually cover, so
    dropping the certificate drops the preference with it.
    """
    picked = _clean_certs(raw)
    covered = set(x for x in _clean_certs(certs).split(",") if x)
    return ",".join(k for k in picked.split(",") if k and k in covered)


def _contractor_row(db, cid):
    return db.execute(
        "SELECT id, name, group_id, certs, preferred_certs, trades FROM compliance_contractors WHERE id = ?",
        (cid,)
    ).fetchone()


@banksia_os_bp.route("/compliance/contractors", methods=["GET"])
def api_compliance_contractors():
    _ensure_compliance_contractor_table()
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT id, name, group_id, certs, preferred_certs, trades FROM compliance_contractors "
            "ORDER BY name COLLATE NOCASE"
        ).fetchall()
    finally:
        db.close()
    return json_success(rows)


@banksia_os_bp.route("/compliance/contractors", methods=["POST"])
def api_compliance_contractor_create():
    _ensure_compliance_contractor_table()
    data = request.get_json(silent=True) or {}
    name = " ".join(str(data.get("name", "")).split())
    if not name:
        return json_error("Give the contractor a name", 422)
    if len(name) > 120:
        return json_error("That name is too long (120 characters max)", 422)
    group_id, err = _clean_group_id(data.get("group_id"))
    if err:
        return json_error(err, 422)
    certs = _clean_certs(data.get("certs"))
    trades = _clean_trades(data.get("trades"))
    # Asked for at the point of adding, because a contractor with no category is
    # invisible to every "who does this?" question the boards ask.
    if not trades:
        return json_error("Choose at least one category — what does %s do?" % name, 422)

    db = get_dict_db()
    try:
        clash = db.execute(
            "SELECT id FROM compliance_contractors WHERE LOWER(name) = LOWER(?)", (name,)
        ).fetchone()
        if clash:
            return json_error("%s is already on the list" % name, 409)
        cur = db.execute(
            "INSERT INTO compliance_contractors (name, group_id, certs, trades) VALUES (?, ?, ?, ?)",
            (name, group_id, certs, trades)
        )
        db.commit()
        cid = cur.lastrowid
    finally:
        db.close()
    _log_activity("compliance", 0, "create", "contractor", "", name,
                  notes="contractor %s added" % name)
    return json_success({"id": cid, "name": name, "group_id": group_id, "certs": certs,
                         "preferred_certs": "", "trades": trades})


@banksia_os_bp.route("/compliance/contractors/<int:cid>", methods=["PATCH"])
def api_compliance_contractor_update(cid):
    _ensure_compliance_contractor_table()
    data = request.get_json(silent=True) or {}
    db = get_dict_db()
    try:
        row = _contractor_row(db, cid)
        if not row:
            return json_error("Contractor not found", 404)

        changes, notes = {}, []
        if "name" in data:
            name = " ".join(str(data.get("name", "")).split())
            if not name:
                return json_error("A contractor needs a name", 422)
            if len(name) > 120:
                return json_error("That name is too long (120 characters max)", 422)
            changes["name"] = name
            notes.append(("name", row["name"] or "", name))
        if "group_id" in data:
            group_id, err = _clean_group_id(data.get("group_id"))
            if err:
                return json_error(err, 422)
            changes["group_id"] = group_id
            notes.append(("group_id", row["group_id"] or "", group_id))
        if "trades" in data:
            trades = _clean_trades(data.get("trades"))
            if not trades:
                return json_error("A contractor needs at least one category.", 422)
            changes["trades"] = trades
            notes.append(("trades", row.get("trades") or "", trades))
            # Taking someone out of the Certificate category has to take their
            # certificates with it. Leaving them ticked but hidden would keep the
            # 15-day chaser messaging a painter about a gas safety check, which is
            # exactly the kind of thing nobody would find until it happened.
            if "certificate" not in [t.strip().lower() for t in trades.split(",")]:
                if (row.get("certs") or "").strip():
                    notes.append(("certs", row["certs"], ""))
                changes["certs"] = ""
                changes["preferred_certs"] = ""
        certs_now = row["certs"] or ""
        if "certs" in data:
            certs_now = _clean_certs(data.get("certs"))
            changes["certs"] = certs_now
            notes.append(("certs", row["certs"] or "", certs_now))
        # A preference is re-checked against the certificates they cover whenever
        # either side moves, so unticking gas cannot leave them preferred for it.
        if "preferred_certs" in data:
            pref = _clean_preferred(data.get("preferred_certs"), certs_now)
            # Saying nothing and quietly dropping it would read as the page
            # ignoring the click, so name the certificate they do not cover.
            asked = [x for x in _clean_certs(data.get("preferred_certs")).split(",") if x]
            spare = [x for x in asked if x not in pref.split(",")]
            if spare:
                return json_error(
                    "%s is not down as doing %s, so they cannot be first call for it"
                    % (row["name"], ", ".join(spare)), 422)
        elif "certs" in data:
            pref = _clean_preferred(row["preferred_certs"] or "", certs_now)
        else:
            pref = None
        if pref is not None and pref != (row["preferred_certs"] or ""):
            changes["preferred_certs"] = pref
            notes.append(("preferred_certs", row["preferred_certs"] or "", pref))
        if not changes:
            return json_error("Nothing to change", 422)

        sets = ", ".join("%s = ?" % k for k in changes)
        db.execute(
            "UPDATE compliance_contractors SET %s, updated = datetime('now') WHERE id = ?" % sets,
            tuple(changes.values()) + (cid,)
        )
        # One certificate, one first call: taking gas takes it off whoever held it,
        # otherwise the chaser is back to picking between two preferred trades.
        claimed = set(x for x in changes.get("preferred_certs", "").split(",") if x)
        if claimed:
            for other in db.execute(
                "SELECT id, preferred_certs FROM compliance_contractors WHERE id != ?", (cid,)
            ).fetchall():
                theirs = [x for x in (other["preferred_certs"] or "").split(",") if x]
                left = [x for x in theirs if x not in claimed]
                if len(left) != len(theirs):
                    db.execute(
                        "UPDATE compliance_contractors SET preferred_certs = ?, "
                        "updated = datetime('now') WHERE id = ?",
                        (",".join(left), other["id"])
                    )
        db.commit()
        after = dict(_contractor_row(db, cid))
    finally:
        db.close()

    for field, old, new in notes:
        _log_activity("compliance", 0, "update", "contractor_" + field, old, new,
                      notes="contractor %s %s changed" % (after.get("name"), field))
    return json_success(after)


@banksia_os_bp.route("/compliance/contractors/<int:cid>", methods=["DELETE"])
def api_compliance_contractor_delete(cid):
    _ensure_compliance_contractor_table()
    db = get_dict_db()
    try:
        row = _contractor_row(db, cid)
        if not row:
            return json_error("Contractor not found", 404)
        db.execute("DELETE FROM compliance_contractors WHERE id = ?", (cid,))
        db.commit()
    finally:
        db.close()
    _log_activity("compliance", 0, "delete", "contractor", row["name"] or "", "",
                  notes="contractor %s removed" % row["name"])
    return json_success({"deleted": cid})


# 11. POLYMORPHIC ENTITY DOCUMENTS (drag-and-drop file storage)
# ═══════════════════════════════════════════════

DOCUMENTS_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "media", "documents")
os.makedirs(DOCUMENTS_UPLOAD_DIR, exist_ok=True)

VALID_ENTITY_TYPES = {
    "tenant", "tenants", "guarantor", "guarantors",
    "applicant", "applicants", "property", "properties",
    "unit", "units", "tenancy", "tenancies",
    "referencing", "referencing_form", "referencing_forms",
    "maintenance_job", "maintenance_jobs",
    "property_owner", "property_owners",
    "deposit", "deposits",
}

def _normalise_entity_type(et: str) -> str:
    singular = {
        "tenants": "tenant", "tenancies": "tenancy",
        "properties": "property", "applicants": "applicant",
        "guarantors": "guarantor", "units": "unit",
        "referencing_form": "referencing", "referencing_forms": "referencing",
        "maintenance_job": "maintenance_job", "maintenance_jobs": "maintenance_job",
        "property_owners": "property_owner",
        "deposits": "deposit",
    }
    return singular.get(et, et)


def _validate_entity_exists(entity_type: str, entity_id: int) -> tuple:
    et = _normalise_entity_type(entity_type)
    table_map = {
        "tenant": ("tenants", "first_name", "last_name"),
        "guarantor": ("guarantors", "first_name", "last_name"),
        "applicant": ("applicants", "first_name", "last_name"),
        "property": ("properties", "name", "address_line_1"),
        "unit": ("units", "unit_ref", "full_address"),
        "tenancy": ("tenancies", "ref", "full_address"),
        "referencing": ("referencing_forms", "first_name", "last_name"),
        "maintenance_job": ("maintenance_jobs", "reference", "title"),
        "property_owner": ("property_owners", "name", "email"),
        # A deposit has no name of its own; both columns are usually empty,
        # so the label falls through to "Deposit #<id>".
        "deposit": ("deposits", "protection_reference", "notes"),
    }
    info = table_map.get(et)
    if not info:
        return False, "Unknown entity type"
    table, name_col, alt_col = info
    db = get_dict_db()
    try:
        row = db.execute(f"SELECT {name_col}, {alt_col} FROM {table} WHERE id = ?", (entity_id,)).fetchone()
        if row:
            label = f"{row[name_col] or ''} {row[alt_col] or ''}".strip()
            if not label:
                label = f"{et.capitalize()} #{entity_id}"
            return True, label
        return False, f"{table} #{entity_id} not found"
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/upload", methods=["POST"])
def api_entity_document_upload():
    if "file" not in request.files:
        return json_error("No file provided (use field 'file')")
    file = request.files["file"]
    if file.filename == "":
        return json_error("Empty filename")
    entity_type = request.form.get("entity_type", "").strip().lower()
    entity_id_str = request.form.get("entity_id", "").strip()
    category = request.form.get("category", "general").strip()
    notes = request.form.get("notes", "").strip()
    uploaded_by = request.form.get("uploaded_by", "team").strip()
    if entity_type not in VALID_ENTITY_TYPES:
        return json_error(f"Invalid entity_type. Must be one of: {', '.join(sorted(VALID_ENTITY_TYPES))}")
    if not entity_id_str:
        return json_error("entity_id is required")
    try:
        entity_id = int(entity_id_str)
    except ValueError:
        return json_error("entity_id must be an integer")
    et = _normalise_entity_type(entity_type)
    exists, label = _validate_entity_exists(et, entity_id)
    if not exists:
        return json_error(f"Entity not found: {label}", 404)
    
    # ── Security: validate file type ──
    ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".png", ".jpg", ".jpeg",
                          ".gif", ".webp", ".txt", ".csv", ".rtf", ".odt", ".ods",
                          ".msg", ".eml", ".heic", ".heif"}
    orig_name = (file.filename or "").strip()
    if not orig_name:
        return json_error("Empty filename")
    ext = os.path.splitext(orig_name)[1].lower() or ""
    if ext not in ALLOWED_EXTENSIONS:
        return json_error(f"File type '{ext}' not allowed. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}")
    
    # ── Security: validate actual MIME type ──
    mime_type = file.content_type or "application/octet-stream"
    ALLOWED_MIMES = {
        ".pdf": ["application/pdf"],
        ".doc": ["application/msword"],
        ".docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
        ".xls": ["application/vnd.ms-excel"],
        ".xlsx": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
        ".png": ["image/png"],
        ".jpg": ["image/jpeg"],
        ".jpeg": ["image/jpeg"],
        ".gif": ["image/gif"],
        ".webp": ["image/webp"],
        ".txt": ["text/plain"],
        ".csv": ["text/csv", "text/plain"],
        ".rtf": ["text/rtf", "application/rtf"],
        ".odt": ["application/vnd.oasis.opendocument.text"],
        ".ods": ["application/vnd.oasis.opendocument.spreadsheet"],
        ".msg": ["application/vnd.ms-outlook", "application/octet-stream"],
        ".eml": ["message/rfc822"],
        ".heic": ["image/heic"],
        ".heif": ["image/heif"],
    }
    expected_mimes = ALLOWED_MIMES.get(ext, [])
    if expected_mimes and mime_type not in expected_mimes and mime_type not in ("application/octet-stream",):
        return json_error(f"MIME type '{mime_type}' does not match extension '{ext}'")
    
    # ── Security: block path traversal in filename ──
    if ".." in orig_name or "/" in orig_name or "\\" in orig_name:
        return json_error("Invalid filename")
    
    # ── Security: enforce max file size (25MB) ──
    MAX_FILE_SIZE = 25 * 1024 * 1024
    file.seek(0, os.SEEK_END)
    real_size = file.tell()
    file.seek(0)
    if real_size > MAX_FILE_SIZE:
        return json_error(f"File too large ({real_size} bytes). Max 25MB.")
    
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    import hashlib
    hash_part = hashlib.md5(orig_name.encode()).hexdigest()[:8]
    stored_name = f"{et}_{entity_id}_{ts}_{hash_part}{ext}"
    entity_dir = os.path.join(DOCUMENTS_UPLOAD_DIR, et, str(entity_id))
    os.makedirs(entity_dir, exist_ok=True)
    file_path = os.path.join(entity_dir, stored_name)
    file.save(file_path)
    file_size = os.path.getsize(file_path)
    mime_type = file.content_type or "application/octet-stream"
    db = get_dict_db()
    try:
        db.execute(
            "INSERT INTO entity_documents "
            "(entity_type, entity_id, original_filename, stored_filename, file_path, "
            "file_type, file_size, mime_type, category, notes, uploaded_by) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (et, entity_id, orig_name, stored_name, file_path,
             ext.lstrip("."), file_size, mime_type, category, notes, uploaded_by))
        db.commit()
        doc_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
        doc = db.execute(
            "SELECT id, entity_type, entity_id, original_filename, stored_filename, "
            "file_type, file_size, mime_type, category, notes, uploaded_by, is_verified, created, original_filename AS filename, created AS uploaded_at, created AS created_at "
            "FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        return json_success(dict(doc))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<entity_type>/<int:entity_id>", methods=["GET"])
def api_list_entity_documents(entity_type, entity_id):
    et = _normalise_entity_type(entity_type)
    category_filter = request.args.get("category", "").strip()
    db = get_dict_db()
    try:
        sql = "SELECT id, entity_type, entity_id, original_filename, stored_filename, " \
              "file_type, file_size, mime_type, category, notes, uploaded_by, is_verified, created, original_filename AS filename, created AS uploaded_at, created AS created_at " \
              "FROM entity_documents WHERE entity_type = ? AND entity_id = ?"
        params = [et, entity_id]
        if category_filter:
            sql += " AND category = ?"
            params.append(category_filter)
        sql += " ORDER BY created DESC"
        docs = db.execute(sql, params).fetchall()
        return json_success(docs)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<int:doc_id>/download")
def api_download_entity_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, original_filename, file_path FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if not os.path.exists(doc["file_path"]):
            return json_error("File not found on disk", 404)
        from flask import send_file
        return send_file(doc["file_path"], as_attachment=True, download_name=doc["original_filename"])
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<int:doc_id>/preview")
def api_preview_entity_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, original_filename, file_path, mime_type FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if not os.path.exists(doc["file_path"]):
            return json_error("File not found on disk", 404)
        from flask import send_file
        return send_file(doc["file_path"], mimetype=doc["mime_type"] or "application/octet-stream", as_attachment=False)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<int:doc_id>", methods=["DELETE"])
def api_delete_entity_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, file_path FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if os.path.exists(doc["file_path"]):
            os.remove(doc["file_path"])
        db.execute("DELETE FROM entity_documents WHERE id = ?", (doc_id,))
        db.commit()
        return json_success({"deleted": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<int:doc_id>/verify", methods=["PATCH"])
def api_verify_entity_document(doc_id):
    data = request.get_json(silent=True) or {}
    is_verified = 1 if data.get("is_verified", True) else 0
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        db.execute("UPDATE entity_documents SET is_verified = ?, updated = datetime('now') WHERE id = ?", (is_verified, doc_id))
        db.commit()
        return json_success({"id": doc_id, "is_verified": bool(is_verified)})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/<int:doc_id>", methods=["PATCH"])
def api_update_entity_document(doc_id):
    data = request.get_json(silent=True) or {}
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id FROM entity_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        updates = []
        params = []
        if "category" in data:
            updates.append("category = ?")
            params.append(data["category"])
        if "notes" in data:
            updates.append("notes = ?")
            params.append(data["notes"])
        if not updates:
            return json_error("Nothing to update")
        updates.append("updated = datetime('now')")
        sql = f"UPDATE entity_documents SET {', '.join(updates)} WHERE id = ?"
        params.append(doc_id)
        db.execute(sql, params)
        db.commit()
        return json_success({"id": doc_id, "updated": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/all", methods=["GET"])
def api_list_all_documents():
    entity_filter = request.args.get("entity_type", "").strip()
    category_filter = request.args.get("category", "").strip()
    search = request.args.get("search", "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 50, type=int)
    db = get_dict_db()
    try:
        sql = "SELECT id, entity_type, entity_id, original_filename, stored_filename, " \
              "file_type, file_size, mime_type, category, notes, uploaded_by, is_verified, created, original_filename AS filename, created AS uploaded_at, created AS created_at " \
              "FROM entity_documents WHERE 1=1"
        count_sql = "SELECT COUNT(*) AS cnt FROM entity_documents WHERE 1=1"
        params = []
        if entity_filter:
            et = _normalise_entity_type(entity_filter)
            sql += " AND entity_type = ?"
            count_sql += " AND entity_type = ?"
            params.append(et)
        if category_filter:
            sql += " AND category = ?"
            count_sql += " AND category = ?"
            params.append(category_filter)
        if search:
            sql += " AND (original_filename LIKE ? OR notes LIKE ?)"
            count_sql += " AND (original_filename LIKE ? OR notes LIKE ?)"
            like = f"%{search}%"
            params.extend([like, like])
        total = db.execute(count_sql, params).fetchone()["cnt"]
        offset = (page - 1) * per_page
        sql += " ORDER BY created DESC LIMIT ? OFFSET ?"
        rows = db.execute(sql, params + [per_page, offset]).fetchall()
        results = []
        for row in rows:
            d = dict(row)
            _, label = _validate_entity_exists(d["entity_type"], d["entity_id"])
            d["entity_label"] = label
            results.append(d)
        return jsonify({
            "success": True, "data": results,
            "total": total, "page": page, "per_page": per_page,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/entity-documents/stats", methods=["GET"])
def api_document_stats():
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT entity_type, COUNT(*) AS count, SUM(file_size) AS total_bytes "
            "FROM entity_documents GROUP BY entity_type ORDER BY count DESC").fetchall()
        total = db.execute("SELECT COUNT(*) AS cnt FROM entity_documents").fetchone()["cnt"]
        total_bytes = db.execute("SELECT COALESCE(SUM(file_size), 0) AS s FROM entity_documents").fetchone()["s"]
        by_type = {}
        for r in rows:
            by_type[r["entity_type"]] = {"count": r["count"], "total_bytes": r["total_bytes"] or 0}
        return json_success({"total_count": total, "total_bytes": total_bytes, "by_type": by_type})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()




# ═══════════════════════════════════════════════
# UNPLACED DOCUMENTS  (Arthur migration manual-placement queue)
# Parked files that could not be auto-matched to a Banksia entity.
# Team reviews each one and allocates it to a Property / Tenancy / Tenant / Applicant.
# ═══════════════════════════════════════════════

UNPLACED_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "media", "unplaced")
os.makedirs(UNPLACED_UPLOAD_DIR, exist_ok=True)

# Placement targets offered in the UI (subset of VALID_ENTITY_TYPES).
UNPLACED_PLACE_TYPES = {"property", "unit", "tenancy", "tenant", "applicant", "guarantor"}


def _ensure_unplaced_table():
    db = get_dict_db()
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS unplaced_documents (
                id                 INTEGER PRIMARY KEY AUTOINCREMENT,
                arthur_asset_uuid  TEXT UNIQUE,
                source_url         TEXT,
                original_filename  TEXT,
                doc_name           TEXT,
                document_type      TEXT,
                relationship       TEXT,
                created_date       TEXT,
                raw_property_id    TEXT,
                raw_unit_id        TEXT,
                raw_tenancy_id     TEXT,
                reason             TEXT,          -- 'no_id' | 'id_not_in_banksia'
                stored_filename    TEXT,
                file_path          TEXT,
                file_type          TEXT,
                file_size          INTEGER,
                mime_type          TEXT,
                status             TEXT DEFAULT 'pending',   -- pending | placed | skipped
                placed_entity_type TEXT,
                placed_entity_id   INTEGER,
                placed_document_id INTEGER,
                placed_by          TEXT,
                placed_at          TEXT,
                created            TEXT DEFAULT (datetime('now'))
            )
        """)
        db.execute("CREATE INDEX IF NOT EXISTS idx_unplaced_status ON unplaced_documents(status)")
        db.commit()
    except Exception as e:
        try:
            app.logger.warning(f"unplaced table init: {e}")
        except Exception:
            pass
    finally:
        db.close()

_ensure_unplaced_table()


def _unplaced_category(document_type, relationship):
    rel = (relationship or "").strip()
    dt = (document_type or "").strip()
    if rel == "Unit":
        return "photo"
    if dt == "Certificate":
        return "certificate"
    if dt == "Contract":
        return "contract"
    if dt == "Reference":
        return "id"
    return "general"


@banksia_os_bp.route("/unplaced-documents", methods=["GET"])
def api_list_unplaced_documents():
    status = (request.args.get("status", "pending") or "").strip().lower()
    reason = (request.args.get("reason", "") or "").strip()
    search = (request.args.get("search", "") or "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 30, type=int)
    if per_page > 200:
        per_page = 200
    db = get_dict_db()
    try:
        where = "WHERE 1=1"
        params = []
        if status and status != "all":
            where += " AND status = ?"
            params.append(status)
        if reason:
            where += " AND reason = ?"
            params.append(reason)
        if search:
            where += " AND (doc_name LIKE ? OR original_filename LIKE ? OR document_type LIKE ? OR relationship LIKE ?)"
            like = f"%{search}%"
            params.extend([like, like, like, like])
        total = db.execute(f"SELECT COUNT(*) AS c FROM unplaced_documents {where}", params).fetchone()["c"]
        offset = (page - 1) * per_page
        rows = db.execute(
            f"""SELECT id, arthur_asset_uuid, doc_name, original_filename, document_type,
                       relationship, created_date, reason, file_type, file_size, mime_type,
                       status, placed_entity_type, placed_entity_id, placed_document_id,
                       placed_by, placed_at, raw_property_id, raw_unit_id, raw_tenancy_id
                FROM unplaced_documents {where}
                ORDER BY (file_path IS NULL) ASC, id ASC
                LIMIT ? OFFSET ?""",
            params + [per_page, offset]).fetchall()
        results = []
        for row in rows:
            d = dict(row)
            if d.get("placed_entity_type") and d.get("placed_entity_id"):
                _, label = _validate_entity_exists(d["placed_entity_type"], d["placed_entity_id"])
                d["placed_entity_label"] = label
            results.append(d)
        return jsonify({
            "success": True, "data": results,
            "total": total, "page": page, "per_page": per_page,
            "total_pages": max(1, (total + per_page - 1) // per_page),
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/stats", methods=["GET"])
def api_unplaced_stats():
    db = get_dict_db()
    try:
        rows = db.execute("SELECT status, COUNT(*) AS c FROM unplaced_documents GROUP BY status").fetchall()
        by_status = {r["status"]: r["c"] for r in rows}
        total = sum(by_status.values())
        downloaded = db.execute("SELECT COUNT(*) AS c FROM unplaced_documents WHERE file_path IS NOT NULL").fetchone()["c"]
        return json_success({
            "total": total,
            "pending": by_status.get("pending", 0),
            "placed": by_status.get("placed", 0),
            "skipped": by_status.get("skipped", 0),
            "downloaded": downloaded,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/<int:doc_id>/preview")
def api_preview_unplaced_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, doc_name, file_path, mime_type FROM unplaced_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if not doc["file_path"] or not os.path.exists(doc["file_path"]):
            return json_error("File not yet downloaded", 404)
        from flask import send_file
        return send_file(doc["file_path"], mimetype=doc["mime_type"] or "application/octet-stream", as_attachment=False)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/<int:doc_id>/download")
def api_download_unplaced_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, doc_name, original_filename, file_path FROM unplaced_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if not doc["file_path"] or not os.path.exists(doc["file_path"]):
            return json_error("File not yet downloaded", 404)
        from flask import send_file
        name = doc["original_filename"] or doc["doc_name"] or f"document_{doc_id}"
        return send_file(doc["file_path"], as_attachment=True, download_name=name)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/<int:doc_id>/place", methods=["POST"])
def api_place_unplaced_document(doc_id):
    """Allocate a parked file to a Banksia entity — attaches it into that entity's Documents tab."""
    data = request.get_json(silent=True) or {}
    entity_type_raw = (data.get("entity_type", "") or "").strip().lower()
    entity_id_str = str(data.get("entity_id", "")).strip()
    placed_by = (data.get("placed_by", "") or "team").strip() or "team"
    if not entity_type_raw or not entity_id_str:
        return json_error("entity_type and entity_id are required")
    et = _normalise_entity_type(entity_type_raw)
    if et not in UNPLACED_PLACE_TYPES:
        return json_error(f"Cannot place to '{et}'. Allowed: {', '.join(sorted(UNPLACED_PLACE_TYPES))}")
    try:
        entity_id = int(entity_id_str)
    except ValueError:
        return json_error("entity_id must be an integer")
    exists, label = _validate_entity_exists(et, entity_id)
    if not exists:
        return json_error(f"Entity not found: {label}", 404)

    db = get_dict_db()
    try:
        doc = db.execute("SELECT * FROM unplaced_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        doc = dict(doc)
        if doc.get("status") == "placed":
            return json_error("This document has already been placed")
        if not doc.get("file_path") or not os.path.exists(doc["file_path"]):
            return json_error("File not yet downloaded — cannot place")

        ext = "." + (doc.get("file_type") or "bin").lstrip(".")
        orig_name = doc.get("original_filename") or doc.get("doc_name") or f"document_{doc_id}"
        if os.path.splitext(orig_name)[1].lower() not in ("", None) and not os.path.splitext(orig_name)[1]:
            orig_name = orig_name + ext
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        import hashlib as _hl, shutil as _sh
        hash_part = _hl.md5((orig_name + str(doc_id)).encode()).hexdigest()[:8]
        stored = f"{et}_{entity_id}_{ts}_{hash_part}{ext}"
        edir = os.path.join(DOCUMENTS_UPLOAD_DIR, et, str(entity_id))
        os.makedirs(edir, exist_ok=True)
        dest = os.path.join(edir, stored)
        _sh.copy2(doc["file_path"], dest)
        size = os.path.getsize(dest)
        cat = _unplaced_category(doc.get("document_type"), doc.get("relationship"))
        notes = (f"Manually placed from Arthur parked queue | arthur-asset:{doc.get('arthur_asset_uuid')} | "
                 f"type:{doc.get('document_type') or '-'} rel:{doc.get('relationship') or '-'} | placed_by:{placed_by}")
        db.execute(
            "INSERT INTO entity_documents "
            "(entity_type, entity_id, original_filename, stored_filename, file_path, "
            "file_type, file_size, mime_type, category, notes, uploaded_by, is_verified) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,0)",
            (et, entity_id, orig_name, stored, dest, ext.lstrip("."), size,
             doc.get("mime_type") or "application/octet-stream", cat, notes, placed_by))
        new_doc_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
        db.execute(
            "UPDATE unplaced_documents SET status='placed', placed_entity_type=?, placed_entity_id=?, "
            "placed_document_id=?, placed_by=?, placed_at=datetime('now') WHERE id=?",
            (et, entity_id, new_doc_id, placed_by, doc_id))
        db.commit()
        return json_success({
            "placed": True, "unplaced_id": doc_id, "document_id": new_doc_id,
            "entity_type": et, "entity_id": entity_id, "entity_label": label, "category": cat,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/<int:doc_id>/skip", methods=["POST"])
def api_skip_unplaced_document(doc_id):
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, status FROM unplaced_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if doc["status"] == "placed":
            return json_error("Cannot skip a placed document")
        db.execute("UPDATE unplaced_documents SET status='skipped' WHERE id=?", (doc_id,))
        db.commit()
        return json_success({"id": doc_id, "status": "skipped"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/unplaced-documents/<int:doc_id>/reset", methods=["POST"])
def api_reset_unplaced_document(doc_id):
    """Return a skipped item to the pending queue (does not undo a real placement)."""
    db = get_dict_db()
    try:
        doc = db.execute("SELECT id, status FROM unplaced_documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return json_error("Document not found", 404)
        if doc["status"] == "placed":
            return json_error("Cannot reset a placed document")
        db.execute("UPDATE unplaced_documents SET status='pending' WHERE id=?", (doc_id,))
        db.commit()
        return json_success({"id": doc_id, "status": "pending"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# COMMENTS & NOTIFICATIONS (Monday.com-style updates)
# ═══════════════════════════════════════════════

# ── Comments table migration (idempotent) ──
try:
    _cmt_db = get_dict_db()
    # Check existing columns via PRAGMA
    existing_cols = {row["name"] for row in _cmt_db.execute("PRAGMA table_info(comments)").fetchall()}
    _cmt_migrated = False
    for col_name, col_def in [
        ("author_id", "TEXT"),
        ("parent_id", "INTEGER DEFAULT NULL"),
        ("media_paths", "TEXT DEFAULT '[]'"),
        ("is_edited", "INTEGER DEFAULT 0"),
        ("is_deleted", "INTEGER DEFAULT 0"),
        ("modified", "TEXT"),
    ]:
        if col_name not in existing_cols:
            _cmt_db.execute(f"ALTER TABLE comments ADD COLUMN {col_name} {col_def}")
            _cmt_migrated = True
    if _cmt_migrated:
        _cmt_db.commit()
    _cmt_db.close()
except Exception as _e:
    current_app.logger.error(f"Error in for col_name, col_def in [: {_e}")
    pass

# ── Ensure notifications table exists ──
try:
    _not_db = get_dict_db()
    _not_db.execute("""
        CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            message TEXT NOT NULL,
            link TEXT DEFAULT '',
            read INTEGER DEFAULT 0,
            created TEXT NOT NULL
        )
    """)
    _not_db.execute("CREATE INDEX IF NOT EXISTS idx_notifications_user ON notifications(username, read)")
    _not_db.commit()
    _not_db.close()
except Exception as _e:
    current_app.logger.error(f"Error initialising notifications table: {_e}")
    pass


# ── Helper: load users.json ──
def _load_comment_users():
    try:
        return json.load(open("/root/banksia-dashboard/users.json"))
    except Exception:
        return {}


# ── Helper: map entity type to display label ──
_ENTITY_LABELS = {
    "tenancies": "Tenancy", "tenancy": "Tenancy",
    "properties": "Property", "property": "Property",
    "tenants": "Tenant", "tenant": "Tenant",
    "applicants": "Applicant", "applicant": "Applicant",
    "units": "Unit", "unit": "Unit",
    "transactions": "Transaction", "transaction": "Transaction",
    "maintenance_jobs": "Maintenance Job", "maintenance_job": "Maintenance Job",
    "property_owners": "Property Owner", "property_owner": "Property Owner",
}


# ── Helper: get current user info from session ──
def _get_current_user():
    u = getattr(request, 'current_user', None) or session.get("user", {})
    if isinstance(u, dict):
        return u.get("username", "System"), u.get("role", "")
    return getattr(u, "username", "System"), getattr(u, "role", "")


# ═══════════════════════════════════════════════
# UNIVERSAL AUDIT HOOK
# Logs every successful mutating request that a detailed logger did not
# already record (guarded by g._audit_detailed). Guarantees no user action
# is invisible on /activity. Uses its own short-lived connection so it never
# interferes with the request's shared thread-local connection.
# ═══════════════════════════════════════════════

_AUDIT_SEGMENT_ENTITY = {
    "properties": "property", "property-owners": "property_owner", "owners": "property_owner",
    "units": "unit", "tenancies": "tenancy", "tenants": "tenant", "applicants": "applicant",
    "referencing": "referencing_form", "guarantors": "guarantor", "guarantor": "guarantor",
    "deposits": "deposit", "maintenance-jobs": "maintenance_job", "maintenance-orders": "maintenance_order",
    "maintenance": "maintenance_job", "invoices": "invoice", "threads": "message_thread",
    "messages": "message_thread", "tags": "tag", "users": "user", "documents": "document",
    "entity-documents": "document", "templates": "template", "comments": "comment",
    "access": "access", "rent-charges": "rent_charge", "images": "property_image",
    "ll-comms": "ll_comms", "company-settings": "company_settings",
}

_AUDIT_DENY = {
    "sync", "sync-from-monday", "sync-monday-property-list", "monday-sync",
    "import-monday-comments", "preview", "fingerprint", "read", "mark-read",
    "mark-all-read", "recalculate", "keepalive", "health", "me", "login",
    "logout", "session", "deposits-migrate",
}


def _audit_generic(entity, entity_id, action, username, notes):
    import sqlite3
    from banksia_os_db import DB_PATH
    conn = sqlite3.connect(DB_PATH, timeout=15)
    try:
        conn.execute(
            "INSERT INTO activity_log (entity_type, entity_id, action, user_name, notes) "
            "VALUES (?, ?, ?, ?, ?)",
            [entity, entity_id, action, username, notes],
        )
        conn.commit()
    finally:
        conn.close()


@banksia_os_bp.after_request
def _audit_mutations(response):
    try:
        if request.method not in ("POST", "PUT", "PATCH", "DELETE"):
            return response
        if not (200 <= response.status_code < 300):
            return response
        if getattr(g, "_audit_detailed", False):
            return response
        rel = (request.path or "").replace("/api/banksia-os", "", 1).strip("/")
        segs = [x for x in rel.split("/") if x]
        if not segs or any(x in _AUDIT_DENY for x in segs):
            return response
        entity, entity_id = "system", None
        for i, seg in enumerate(segs):
            if seg in _AUDIT_SEGMENT_ENTITY:
                entity = _AUDIT_SEGMENT_ENTITY[seg]
                entity_id = segs[i + 1] if (i + 1 < len(segs) and segs[i + 1].isdigit()) else None
        last = segs[-1]
        trailing = last if (not last.isdigit() and last not in _AUDIT_SEGMENT_ENTITY) else None
        if request.method == "DELETE":
            action = "deleted"
        elif request.method == "POST":
            action = trailing.replace("-", " ") if trailing else "created"
        else:
            action = trailing.replace("-", " ") if trailing else "updated"
        if entity_id is None:
            # New-record ids aren't in the URL; try the response body, else sentinel 0
            try:
                body = response.get_json(silent=True)
                if isinstance(body, dict):
                    src = body.get("data") if isinstance(body.get("data"), dict) else body
                    for k in ("id", "new_id", f"{entity}_id"):
                        v = src.get(k)
                        if isinstance(v, int):
                            entity_id = v
                            break
            except Exception:
                pass
        username, _role = _get_current_user()
        notes = f"{action} {entity.replace(chr(95), chr(32))}"
        _audit_generic(entity, entity_id if entity_id is not None else 0, action, username, notes)
    except Exception:
        pass
    return response


@banksia_os_bp.route("/me")
def api_get_current_user():
    username, role = _get_current_user()
    users_data = _load_comment_users()
    user_info = users_data.get(username, {})
    avatar_url = f"/static/uploads/avatars/{username}.jpg" if username and os.path.isfile(f"/root/banksia-dashboard/static/uploads/avatars/{username}.jpg") else None
    return json_success({
        "username": username,
        "name": user_info.get("display_name") or username,
        "role": role,
        "email": user_info.get("email", ""),
        "avatar_url": avatar_url,
    })


@banksia_os_bp.route("/comments/<entity_type>/<int:entity_id>", methods=["GET"])
def api_get_comments(entity_type, entity_id):
    # "compliance" added 2026-08-06 (Norbert) so the compliance board carries the
    # same per-row updates Monday has. Property-level: an update is about the
    # property, so it reads the same on all nine certificate pages.
    valid = {"tenancy","tenancies","property","properties","tenant","tenants",
             "applicant","applicants","unit","units","transaction","transactions",
             "maintenance_job","maintenance_jobs","property_owner","property_owners",
             "compliance"}
    if entity_type not in valid:
        return json_error("Invalid entity type", 400)
    sg_map = {"tenancy":"tenancies","property":"properties","applicant":"applicants",
              "transaction":"transactions","maintenance_job":"maintenance_jobs",
              "tenant":"tenants","unit":"units","property_owner":"property_owners"}
    etype = sg_map.get(entity_type, entity_type)
    current_user, current_role = _get_current_user()
    users_data = _load_comment_users()
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT id, author, author_id, body, mentions, media_paths, "
            "       is_edited, is_deleted, parent_id, created, modified "
            "FROM comments "
            "WHERE entity_type = ? AND entity_id = ? AND is_deleted = 0 "
            "ORDER BY created ASC",
            (etype, entity_id)
        ).fetchall()
        results = []
        for r in rows:
            author_id = r.get("author_id") or r.get("author", "")
            user_info = users_data.get(author_id, {})
            author_name = user_info.get("display_name") or author_id or r.get("author", "Unknown")
            avatar_url = f"/static/uploads/avatars/{author_id}.jpg" if author_id and os.path.isfile(f"/root/banksia-dashboard/static/uploads/avatars/{author_id}.jpg") else None
            can_delete = (current_user == author_id or current_role == "super_admin")
            try:
                mp = json.loads(r.get("media_paths") or "[]")
            except (json.JSONDecodeError, TypeError):
                mp = []
            try:
                ment = json.loads(r.get("mentions") or "[]")
            except (json.JSONDecodeError, TypeError):
                ment = []
            results.append({
                "id": r["id"],
                "author": r.get("author", ""),
                "author_id": author_id,
                "author_name": author_name,
                "avatar_url": avatar_url,
                "body": r["body"],
                "mentions": ment,
                "media_paths": mp,
                "is_edited": bool(r.get("is_edited")),
                "parent_id": r.get("parent_id"),
                "created": r["created"],
                "modified": r.get("modified"),
                "can_delete": can_delete,
            })
        return json_success(results)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/comments/<entity_type>/<int:entity_id>", methods=["POST"])
def api_add_comment(entity_type, entity_id):
    data = request.get_json()
    if not data or not data.get("body","").strip():
        return json_error("Comment body is required")
    body = data["body"].strip()
    etype = entity_type
    sg_map = {"tenancy":"tenancies","property":"properties","applicant":"applicants",
              "transaction":"transactions","maintenance_job":"maintenance_jobs",
              "tenant":"tenants","unit":"units","property_owner":"property_owners"}
    etype = sg_map.get(etype, etype)
    # This is a second, separately-worded allowlist from the one on GET (plurals
    # only, since the singular has already been mapped by here). Adding an entity
    # type means touching both -- "compliance" reads on one and 400s on the other
    # otherwise, which looks like the board silently swallowing updates.
    valid = {"tenancies","properties","tenants","applicants","units","transactions",
             "maintenance_jobs","property_owners","compliance"}
    if etype not in valid:
        return json_error("Invalid entity type", 400)
    current_user, _ = _get_current_user()
    author_id = current_user
    import re
    mentioned = list(set(re.findall(r'@(\w+)', body)))
    media_paths = data.get("media_paths", [])
    if not isinstance(media_paths, list):
        media_paths = []
    now_iso = datetime.now(timezone.utc).isoformat()
    db = get_dict_db()
    try:
        c = db.execute(
            "INSERT INTO comments (entity_type, entity_id, author, author_id, body, mentions, media_paths, created) "
            "VALUES (?,?,?,?,?,?,?,?)",
            (etype, entity_id, author_id, author_id, body, json.dumps(mentioned), json.dumps(media_paths), now_iso)
        )
        cid = c.lastrowid
        for u in mentioned:
            # Only notify if the mentioned user actually exists
            users_data = _load_comment_users()
            if u in users_data and u != author_id:
                db.execute(
                    "INSERT INTO notifications (username, message, link, read, created) VALUES (?,?,?,0,?)",
                    (u, f"{author_id} @mentioned you on {_ENTITY_LABELS.get(etype, etype[:-1])} #{entity_id}",
                     f"/banksia-os?entity={etype}&id={entity_id}",
                     now_iso)
                )
        db.commit()
        return json_success({
            "id": cid,
            "author": author_id,
            "author_id": author_id,
            "body": body,
            "mentions": mentioned,
            "media_paths": media_paths,
            "is_edited": False,
            "is_deleted": False,
            "parent_id": None,
            "created": now_iso,
            "modified": None,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/comments/<int:comment_id>", methods=["PUT"])
def api_edit_comment(comment_id):
    """Edit a comment: soft-deletes the old and inserts a clone with parent_id pointing to original."""
    current_user, current_role = _get_current_user()
    db = get_dict_db()
    try:
        old = db.execute(
            "SELECT * FROM comments WHERE id = ?", (comment_id,)
        ).fetchone()
        if not old:
            return json_error("Comment not found", 404)
        if old.get("is_deleted"):
            return json_error("Comment has been deleted", 400)
        author_id = old.get("author_id") or old.get("author", "")
        if current_user != author_id:
            return json_error("You can only edit your own comments", 403)
        data = request.get_json()
        if not data or not data.get("body","").strip():
            return json_error("Comment body is required")
        new_body = data["body"].strip()
        new_media = data.get("media_paths", [])
        if not isinstance(new_media, list):
            new_media = []
        import re
        mentioned = list(set(re.findall(r'@(\w+)', new_body)))
        now_iso = datetime.now(timezone.utc).isoformat()
        # Soft-delete the old version. It kept only is_edited until 2026-08-06,
        # but GET filters on is_deleted, so the original stayed visible alongside
        # the new one and a single edit showed the comment twice. The clone points
        # back here via parent_id, so hiding the original loses no history.
        # Never triggered in production -- nobody had used the in-app edit -- but
        # it would have shown up the first time anyone edited a compliance update.
        db.execute(
            "UPDATE comments SET is_edited = 1, is_deleted = 1, modified = ? WHERE id = ?",
            (now_iso, comment_id)
        )
        # Insert new version with parent_id pointing to original
        c = db.execute(
            "INSERT INTO comments (entity_type, entity_id, author, author_id, body, mentions, "
            "media_paths, parent_id, created) VALUES (?,?,?,?,?,?,?,?,?)",
            (old["entity_type"], old["entity_id"], author_id, author_id, new_body, json.dumps(mentioned),
             json.dumps(new_media), comment_id, now_iso)
        )
        new_id = c.lastrowid
        # Notify mentions in the edited comment
        users_data = _load_comment_users()
        for u in mentioned:
            if u in users_data and u != author_id:
                db.execute(
                    "INSERT INTO notifications (username, message, link, read, created) VALUES (?,?,?,0,?)",
                    (u, f"{author_id} @mentioned you in an edited comment on "
                         f"{_ENTITY_LABELS.get(old['entity_type'], old['entity_type'][:-1])} #{old['entity_id']}",
                     f"/banksia-os?entity={old['entity_type']}&id={old['entity_id']}",
                     now_iso)
                )
        db.commit()
        return json_success({
            "id": new_id,
            "author": author_id,
            "author_id": author_id,
            "body": new_body,
            "mentions": mentioned,
            "media_paths": new_media,
            "is_edited": False,
            "is_deleted": False,
            "parent_id": comment_id,
            "created": now_iso,
            "modified": None,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/comments/<int:comment_id>", methods=["DELETE"])
def api_delete_comment(comment_id):
    """Soft-delete a comment. Author or super_admin can delete."""
    current_user, current_role = _get_current_user()
    db = get_dict_db()
    try:
        old = db.execute(
            "SELECT * FROM comments WHERE id = ?", (comment_id,)
        ).fetchone()
        if not old:
            return json_error("Comment not found", 404)
        author_id = old.get("author_id") or old.get("author", "")
        if current_user != author_id and current_role != "super_admin":
            return json_error("You do not have permission to delete this comment", 403)
        now_iso = datetime.now(timezone.utc).isoformat()
        db.execute(
            "UPDATE comments SET is_deleted = 1, modified = ? WHERE id = ?",
            (now_iso, comment_id)
        )
        db.commit()
        return json_success({"deleted": True, "id": comment_id})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/comments/upload", methods=["POST"])
def api_upload_comment_media():
    """Upload a file for comment media attachments."""
    if "file" not in request.files:
        return json_error("No file provided")
    file = request.files["file"]
    if file.filename == "":
        return json_error("No file selected")
    # Secure the filename
    import uuid
    fname = file.filename or "upload.bin"
    ext = os.path.splitext(fname)[1].lower()
    safe_name = f"{uuid.uuid4().hex}{ext}"
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "uploads", "comments")
    os.makedirs(upload_dir, exist_ok=True)
    save_path = os.path.join(upload_dir, safe_name)
    file.save(save_path)
    url_path = f"/static/uploads/comments/{safe_name}"
    return json_success({"url": url_path, "filename": safe_name})


@banksia_os_bp.route("/comments/counts/<entity_type>", methods=["GET"])
def api_comment_counts(entity_type):
    """How many updates each row of one entity type has, plus when the last one
    landed. One request for the whole board -- a count per row would be 62 calls
    on a page that already loads in one.
    """
    etype = str(entity_type or "").strip().lower()
    if not re.fullmatch(r"[a-z_]+", etype):
        return json_error("Invalid entity type", 400)
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT entity_id, COUNT(*) AS n, MAX(created) AS last_at"
            " FROM comments WHERE entity_type = ? AND is_deleted = 0"
            " GROUP BY entity_id",
            (etype,)
        ).fetchall()
    finally:
        db.close()
    return json_success(rows)


@banksia_os_bp.route("/comments/recent")
def api_recent_comments():
    """Return the 20 most recent non-deleted comments across all entities."""
    limit = min(int_param(request.args.get("limit"), 20, max_val=MAX_PAGE_SIZE), 20)
    current_user, _ = _get_current_user()
    users_data = _load_comment_users()
    db = get_dict_db()
    try:
        rows = db.execute(
            "SELECT id, author, author_id, body, entity_type, entity_id, created "
            "FROM comments WHERE is_deleted = 0 "
            "ORDER BY id DESC LIMIT ?", (limit,)
        ).fetchall()
        results = []
        for r in rows:
            author_id = r.get("author_id") or r.get("author", "")
            user_info = users_data.get(author_id, {})
            author_name = user_info.get("display_name") or author_id or r.get("author", "Unknown")
            preview = (r["body"] or "")[:100]
            results.append({
                "id": r["id"],
                "author": r.get("author", ""),
                "author_id": author_id,
                "author_name": author_name,
                "body_preview": preview,
                "entity_type": r["entity_type"],
                "entity_type_label": _ENTITY_LABELS.get(r["entity_type"], r["entity_type"].replace("_", " ").title()),
                "entity_id": r["entity_id"],
                "created": r["created"],
            })
        return json_success(results)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/properties/compliance")
def api_properties_compliance():
    """Return compliance issues across all properties (missing certificates, etc)."""
    db = get_dict_db()
    try:
        issues = []
        props = db.execute("SELECT id, ref, name FROM properties ORDER BY name").fetchall()
        for p in props:
            # Check for missing council tax band
            if not p.get("council_tax_band"):
                issues.append({"property_id": p["id"], "property_name": p["ref"] or p["name"],
                               "issue": "Council Tax Band not set", "status": "missing"})
            # Check for missing EPC check based on tenancies
            tenancies = db.execute(
                "SELECT COUNT(*) AS cnt FROM tenancies WHERE property_id=? AND status IN ('Current','current','Periodic','periodic')",
                (p["id"],)
            ).fetchone()
            if tenancies and tenancies["cnt"] > 0:
                issues.append({"property_id": p["id"], "property_name": p["ref"] or p["name"],
                               "issue": f"{tenancies['cnt']} active tenancies — compliance review needed",
                               "status": "pending"})
        return json_success(issues)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/notifications", methods=["GET"])
def api_get_notifications():
    """Enhanced GET /notifications — returns full items + unread_count.
    
    Query params:
        mark_read=true|false  (default true) — mark returned notifications as read
        unread_only=true       — legacy compat, just returns unread_count
    """
    db = get_dict_db()
    try:
        u = getattr(request, 'current_user', None) or session.get("user", {})
        uname = u.get("username", "") if isinstance(u, dict) else getattr(u, "username", "")
        
        # Legacy compat: just return count
        if request.args.get("unread_only", "") == "true":
            cnt = db.execute(
                "SELECT COUNT(*) AS c FROM notifications WHERE username=? AND read=0",
                (uname,)
            ).fetchone()["c"]
            return json_success({"unread_count": cnt})
        
        # Fetch unread notifications (limit 20, ordered by created DESC)
        items = db.execute(
            "SELECT id, message, link, read, created FROM notifications "
            "WHERE username=? AND read=0 ORDER BY created DESC LIMIT 20",
            (uname,)
        ).fetchall()
        
        uc = db.execute(
            "SELECT COUNT(*) AS c FROM notifications WHERE username=? AND read=0",
            (uname,)
        ).fetchone()["c"]
        
        # Mark as read if requested (default true)
        mark_read = request.args.get("mark_read", "true").lower() == "true"
        if mark_read and items:
            ids = [r["id"] for r in items]
            placeholders = ",".join("?" * len(ids))
            db.execute(
                f"UPDATE notifications SET read=1 WHERE id IN ({placeholders})",
                ids
            )
            db.commit()
            uc = 0  # just marked them all as read
        
        return json_success({"items": items, "unread_count": uc})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/my-updates", methods=["GET"])
def api_my_updates():
    """Combined feed: comments mentioning the user + unread notifications.
    
    Returns:
        items: list of update objects, newest first
        unread_count: total unread (unread notifications only)
    """
    db = get_dict_db()
    try:
        u = getattr(request, 'current_user', None) or session.get("user", {})
        uname = u.get("username", "") if isinstance(u, dict) else getattr(u, "username", "")
        limit = request.args.get("limit", "50")
        offset = request.args.get("offset", "0")
        try:
            limit = max(1, min(int(limit), 200))
            offset = max(0, int(offset))
        except:
            limit, offset = 50, 0

        # 1. Comments where user is mentioned
        # mentions is stored as JSON array: '["Sami","Norbert"]'
        mentions_sql = """
            SELECT c.id, c.entity_type, c.entity_id, c.author, c.body, c.mentions, c.created, c.is_edited,
                   c.author_id,
                   (SELECT COUNT(*) FROM comments r WHERE r.parent_id = c.id AND r.is_deleted = 0) AS reply_count
            FROM comments c
            WHERE c.is_deleted = 0
              AND c.mentions LIKE ?
            ORDER BY c.created DESC
        """
        # LIKE pattern to find the username anywhere in the JSON array
        like_pattern = f'%"{uname}"%'
        mentioned_items = db.execute(mentions_sql, (like_pattern,)).fetchall()

        # 2. Comments where user is author AND someone replied
        replied_sql = """
            SELECT DISTINCT c.id, c.entity_type, c.entity_id, c.author, c.body, c.mentions, c.created, c.is_edited,
                   c.author_id,
                   (SELECT COUNT(*) FROM comments r WHERE r.parent_id = c.id AND r.is_deleted = 0) AS reply_count
            FROM comments c
            INNER JOIN comments r ON r.parent_id = c.id
            WHERE c.is_deleted = 0
              AND c.author = ?
              AND r.author != ?
              AND r.is_deleted = 0
            ORDER BY c.created DESC
        """
        replied_items = db.execute(replied_sql, (uname, uname)).fetchall()

        # 3. Unread notifications
        notif_sql = """
            SELECT n.id, n.message, n.link, n.created
            FROM notifications n
            WHERE n.username = ? AND n.read = 0
            ORDER BY n.created DESC
            LIMIT ?
        """
        notif_items = db.execute(notif_sql, (uname, limit)).fetchall()

        # 4. Comments on the same entity the user authored (updates to threads they're in)
        thread_sql = """
            SELECT c.id, c.entity_type, c.entity_id, c.author, c.body, c.mentions, c.created, c.is_edited,
                   c.author_id,
                   (SELECT COUNT(*) FROM comments r WHERE r.parent_id = c.id AND r.is_deleted = 0) AS reply_count
            FROM comments c
            WHERE c.is_deleted = 0
              AND c.entity_id IN (
                  SELECT DISTINCT cc.entity_id FROM comments cc
                  WHERE cc.author = ? AND cc.is_deleted = 0
              )
              AND c.author != ?
              AND c.id NOT IN (SELECT id FROM comments WHERE mentions LIKE ?)
            ORDER BY c.created DESC
            LIMIT 20
        """
        thread_items = db.execute(thread_sql, (uname, uname, like_pattern)).fetchall()

        # Combine and sort by created DESC
        combined = []
        seen_ids = set()

        for row in mentioned_items:
            item = dict(row)
            item['type'] = 'mention'
            key = f"comment_{item['id']}"
            if key not in seen_ids:
                seen_ids.add(key)
                combined.append(item)

        for row in replied_items:
            item = dict(row)
            item['type'] = 'reply'
            key = f"comment_{item['id']}"
            if key not in seen_ids:
                seen_ids.add(key)
                combined.append(item)

        for row in thread_items:
            item = dict(row)
            item['type'] = 'thread_update'
            key = f"comment_{item['id']}"
            if key not in seen_ids:
                seen_ids.add(key)
                combined.append(item)

        for row in notif_items:
            item = dict(row)
            item['type'] = 'notification'
            item['author'] = 'system'
            item['body'] = item['message']
            item['entity_type'] = 'notification'
            item['entity_id'] = item['id']
            item['reply_count'] = 0
            item['mentions'] = '[]'
            item['is_edited'] = 0
            item['author_id'] = 'system'
            key = f"notif_{item['id']}"
            if key not in seen_ids:
                seen_ids.add(key)
                combined.append(item)

        # Sort by created DESC, newest first
        combined.sort(key=lambda x: x.get('created', ''), reverse=True)

        # Paginate
        total = len(combined)
        page_items = combined[offset:offset + limit]

        # Unread count from notifications only
        uc = db.execute(
            "SELECT COUNT(*) AS c FROM notifications WHERE username=? AND read=0",
            (uname,)
        ).fetchone()["c"]

        # Get entity display names for context
        entity_names = {}
        for item in page_items:
            if item['entity_type'] in ('notification',):
                continue
            et = item['entity_type']
            eid = item['entity_id']
            try:
                if et == 'maintenance_job' or et == 'maintenance':
                    r = db.execute("SELECT reference, title FROM maintenance_jobs WHERE id=?", (eid,)).fetchone()
                    if r:
                        entity_names[f"{et}_{eid}"] = r['reference'] or r['title'] or f"Job #{eid}"
                    else:
                        entity_names[f"{et}_{eid}"] = f"Job #{eid}"
                elif et == 'property':
                    r = db.execute("SELECT name, address_line_1 FROM properties WHERE id=?", (eid,)).fetchone()
                    entity_names[f"{et}_{eid}"] = (r['name'] or r['address_line_1'] or f"Property #{eid}") if r else f"Property #{eid}"
                elif et == 'property_owner':
                    r = db.execute("SELECT name FROM property_owners WHERE id=?", (eid,)).fetchone()
                    entity_names[f"{et}_{eid}"] = r['name'] if r else f"Landlord #{eid}"
                elif et == 'unit':
                    r = db.execute("SELECT ref, name FROM units WHERE id=?", (eid,)).fetchone()
                    entity_names[f"{et}_{eid}"] = (r['ref'] or r['name'] or f"Unit #{eid}") if r else f"Unit #{eid}"
                else:
                    entity_names[f"{et}_{eid}"] = f"{et} #{eid}"
            except:
                entity_names[f"{et}_{eid}"] = f"{et} #{eid}"

        return json_success({
            "items": page_items,
            "entity_names": entity_names,
            "unread_count": uc,
            "total": total,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/notifications/<int:notification_id>/read", methods=["POST"])
def api_mark_notification_read(notification_id):
    """Mark a single notification as read."""
    db = get_dict_db()
    try:
        u = getattr(request, 'current_user', None) or session.get("user", {})
        uname = u.get("username", "") if isinstance(u, dict) else getattr(u, "username", "")
        db.execute(
            "UPDATE notifications SET read=1 WHERE id=? AND username=?",
            (notification_id, uname)
        )
        db.commit()
        return json_success({"ok": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/notifications/read-all", methods=["POST"])
def api_mark_all_read():
    """Mark all of the current user's notifications as read."""
    db = get_dict_db()
    try:
        u = getattr(request, 'current_user', None) or session.get("user", {})
        uname = u.get("username", "") if isinstance(u, dict) else getattr(u, "username", "")
        db.execute(
            "UPDATE notifications SET read=1 WHERE username=? AND read=0",
            (uname,)
        )
        db.commit()
        return json_success({"ok": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/notifications/mark-read", methods=["POST"])
def api_mark_read():
    """Legacy: mark single notification by id in JSON body, or all if no id given."""
    db = get_dict_db()
    try:
        u = getattr(request, 'current_user', None) or session.get("user", {})
        uname = u.get("username", "") if isinstance(u, dict) else getattr(u, "username", "")
        data = request.get_json() or {}
        nid = data.get("id")
        if nid:
            db.execute("UPDATE notifications SET read=1 WHERE id=? AND username=?", (nid, uname))
        else:
            db.execute("UPDATE notifications SET read=1 WHERE username=? AND read=0", (uname,))
        db.commit()
        return json_success({"ok": True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


def create_notification(username, message, link=None):
    """Standalone helper: insert a notification for a user.
    
    Args:
        username: str — the recipient's username
        message: str — notification text
        link: str or None — optional link path
    
    Returns:
        int — the new notification id, or None on failure
    """
    try:
        # get_dict_db() hands back the caller's own thread-local connection, so closing
        # it here closed the database out from under whoever called us: the tenancy
        # conversion notified four admins, then died on its own db.commit() with
        # "Cannot operate on a closed database" — after the tenancy had been written.
        # The connection is per-thread and reused by design; leave it open.
        db = get_dict_db()
        now = datetime.now(timezone.utc).isoformat()
        cur = db.execute(
            "INSERT INTO notifications (username, message, link, read, created) VALUES (?, ?, ?, 0, ?)",
            (username, message, link or "", now)
        )
        db.commit()
        return cur.lastrowid
    except Exception:
        return None


@banksia_os_bp.route("/users", methods=["GET"])
def api_users():
    import json as jm
    uf = os.path.join(os.path.dirname(__file__),"users.json")
    users_list = []
    if os.path.exists(uf):
        with open(uf) as f:
            users = jm.load(f)
            for username, info in users.items():
                if not isinstance(info, dict):
                    info = {}
                users_list.append({
                    "username": username,
                    "role": info.get("role", "user"),
                    "email": info.get("email", ""),
                    "phone": info.get("phone", ""),
                    "date_of_birth": info.get("date_of_birth", ""),
                    "department": info.get("department", ""),
                    "position": info.get("position", ""),
                    "biography": info.get("biography", ""),
                })
    return json_success(users_list)


@banksia_os_bp.route("/users", methods=["POST"])
def api_add_user():
    user = session.get("user", {})
    role = user.get("role", "")
    # Creating login accounts is super-admin only (was admin+super — a
    # privilege-escalation path).
    if role != "super_admin":
        return json_error("Forbidden — only super admins can create users", 403)
    data = request.get_json()
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    new_role = data.get("role", "viewer").strip()
    if not username or not password:
        return json_error("username and password required", 400)
    import re as _email_re
    email_val = data.get("email", "").strip()
    if email_val and not _email_re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email_val):
        return json_error("Invalid email format — enter a valid email address", 400)
    msg = _validate_password_strength(password)
    if msg:
        return json_error(msg, 400)
    if username in _load_users():
        return json_error("A user with that username already exists", 409)
    if new_role not in VALID_ROLES:
        new_role = "viewer"
    users = _load_users()
    users[username] = {"password": _hash_password(password), "role": new_role,
                       "email": data.get("email", "").strip()}
    _save_users(users)
    return json_success({"user": {"username": username, "role": new_role}})


@banksia_os_bp.route("/users/<username>", methods=["PATCH"])
def api_update_user(username):
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data", 400)
    users = _load_users()
    if username not in users:
        return json_error("User not found", 404)
    current_user = session.get("user", {})
    current_role = current_user.get("role", "")
    is_super = current_role == "super_admin"
    is_admin = current_role in ("super_admin", "admin")
    is_self = current_user.get("username") == username
    # Super admin can edit anyone. Admin can edit themselves or non-super_admin users.
    target = users[username] if isinstance(users[username], dict) else {}
    target_role = target.get("role", "admin") if isinstance(target, dict) else "admin"
    if is_super:
        pass  # can edit anyone
    elif is_admin and is_self:
        pass  # can edit self
    elif is_admin and target_role != "super_admin":
        pass  # admin can edit non-super_admin users
    else:
        return json_error("Forbidden", 403)
    allowed_fields = ["email", "phone", "date_of_birth", "biography", "department", "position"]
    for f in allowed_fields:
        if f in data:
            users[username][f] = data[f]
    # Only super admin can change role, and only to a known role.
    if is_super and "role" in data:
        _nr = str(data["role"]).strip()
        if _nr in VALID_ROLES:
            users[username]["role"] = _nr
    # Password update / reset.
    # A user may reset their own password; only super_admin may reset ANOTHER user's password.
    if "password" in data and data["password"]:
        if not is_self and not is_super:
            return json_error("Only a super admin can reset another user's password", 403)
        msg = _validate_password_strength(data["password"])
        if msg:
            return json_error(msg, 400)
        users[username]["password"] = _hash_password(data["password"])
    _save_users(users)
    return json_success({"user": {"username": username, "role": users[username].get("role")}})


@banksia_os_bp.route("/users/<username>", methods=["DELETE"])
def api_delete_user(username):
    current = session.get("user", {})
    current_role = current.get("role", "")
    # Deleting login accounts is super-admin only.
    if current_role != "super_admin":
        return json_error("Forbidden — only super admins can delete users", 403)
    if username == "Sami":
        return json_error("Cannot delete super admin", 400)
    users = _load_users()
    if username in users:
        del users[username]
        _save_users(users)
    return json_success({"deleted": True})


@banksia_os_bp.route("/users/autocomplete", methods=["GET"])
def api_users_autocomplete():
    import json as jm
    uf = os.path.join(os.path.dirname(__file__),"users.json")
    if os.path.exists(uf):
        with open(uf) as f: users = jm.load(f)
        names = list(users.keys())
    else:
        names = []
    return json_success(names)


# ═══════════════════════════════════════════════
# 6. FINANCE — Rent Schedule & Tenancy Summary
# ═══════════════════════════════════════════════


@banksia_os_bp.route("/finance/move-in/<int:tenancy_id>", methods=["GET"])
def api_get_move_in(tenancy_id):
    """Return the move-in breakdown (Holding Deposit, Deposit, Pro-Rata, Move-in Amount) for a tenancy."""
    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM tenancy_move_in WHERE tenancy_id = ?", (tenancy_id,)).fetchone()
        ten = db.execute("SELECT rent_amount, start_date FROM tenancies WHERE id = ?", (tenancy_id,)).fetchone()
        rent = ten.get("rent_amount") if ten else None
        start = ten.get("start_date") if ten else None
        if not row:
            return json_success({"tenancy_id": tenancy_id, "holding_deposit": 0, "deposit": 0,
                "pro_rata": 0, "first_month": 0, "move_in_amount": None, "needs_review": 1,
                "note": "No move-in data on file", "rent_amount": rent, "start_date": start, "exists": False})
        d = dict(row)
        d["rent_amount"] = rent
        d["start_date"] = start
        d["exists"] = True
        return json_success(d)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/move-in/<int:tenancy_id>", methods=["PATCH"])
def api_update_move_in(tenancy_id):
    """Edit the move-in breakdown. Recomputes Move-in Amount from components unless one is explicitly supplied."""
    data = request.get_json() or {}
    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM tenancy_move_in WHERE tenancy_id = ?", (tenancy_id,)).fetchone()
        cur = dict(row) if row else {"holding_deposit": 0, "deposit": 0, "pro_rata": 0,
                                     "first_month": 0, "move_in_amount": None, "needs_review": 1, "note": ""}
        for key in ("holding_deposit", "deposit", "pro_rata", "first_month"):
            if key in data and data[key] is not None:
                cur[key] = float(data[key])
        if "move_in_amount" in data and data["move_in_amount"] is not None:
            cur["move_in_amount"] = float(data["move_in_amount"])
        else:
            initial = cur.get("pro_rata") or 0
            if not initial:
                initial = cur.get("first_month") or 0
            cur["move_in_amount"] = round(float(cur.get("deposit") or 0)
                                          - float(cur.get("holding_deposit") or 0)
                                          + float(initial), 2)
        if "needs_review" in data:
            cur["needs_review"] = 1 if data["needs_review"] else 0
        if "note" in data:
            cur["note"] = data["note"]
        actor = getattr(request, "current_user", {}).get("username", "system") if hasattr(request, "current_user") else "system"
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "INSERT INTO tenancy_move_in (tenancy_id,holding_deposit,deposit,pro_rata,first_month,move_in_amount,needs_review,note,source,edited_by,updated) "
            "VALUES (?,?,?,?,?,?,?,?,'manual',?,?) "
            "ON CONFLICT(tenancy_id) DO UPDATE SET holding_deposit=excluded.holding_deposit, deposit=excluded.deposit, "
            "pro_rata=excluded.pro_rata, first_month=excluded.first_month, move_in_amount=excluded.move_in_amount, "
            "needs_review=excluded.needs_review, note=excluded.note, edited_by=excluded.edited_by, updated=excluded.updated",
            (tenancy_id, cur.get("holding_deposit") or 0, cur.get("deposit") or 0, cur.get("pro_rata") or 0,
             cur.get("first_month") or 0, cur.get("move_in_amount"), cur.get("needs_review") or 0,
             cur.get("note") or "", actor, now))
        db.commit()
        out = dict(db.execute("SELECT * FROM tenancy_move_in WHERE tenancy_id = ?", (tenancy_id,)).fetchone())
        out["updated"] = True
        return json_success(out)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 6b. TENANCY NOTICE — tenant's notice to quit
# ═══════════════════════════════════════════════
# Legal requirement from 1 May. A tenant serving notice must do so by the day
# before rent collection (the 1st of the month), and the notice period is two
# months. So notice served at any point during July is served against the
# 1 August rent day, and two months from there puts the move-out at 1 October.
# That reduces to: the 1st of the month three months after the notice month.
# Rule and worked example from Norbert, 2026-08-03.

def _notice_move_out(notice_date):
    """Move-out date for a notice served on notice_date (YYYY-MM-DD)."""
    if not notice_date:
        return None
    try:
        d = datetime.strptime(str(notice_date)[:10], "%Y-%m-%d")
    except (ValueError, TypeError):
        return None
    m = d.month + 3
    y = d.year + (m - 1) // 12
    m = (m - 1) % 12 + 1
    return "%04d-%02d-01" % (y, m)


def _notice_effective_move_out(row):
    """The move-out date that actually applies: the agreed exception if there is
    one, otherwise the two-month rule.

    Everything downstream — the tag, the end date, the letting pipeline — reads
    this rather than recomputing, so an agreed one-month notice is not silently
    overwritten by the rule the next time somebody presses Approve.
    """
    d = dict(row) if row else {}
    override = (d.get("move_out_override") or "").strip()
    return override[:10] if override else _notice_move_out(d.get("notice_date"))


def _notice_months_between(notice_date, move_out):
    """Roughly how many months of notice a move-out date represents. Used only to
    describe the exception on screen, never to decide anything."""
    try:
        a = datetime.strptime(str(notice_date)[:10], "%Y-%m-%d")
        b = datetime.strptime(str(move_out)[:10], "%Y-%m-%d")
    except (ValueError, TypeError):
        return None
    return round((b - a).days / 30.44, 1)


def _notice_payload(row):
    """Shape a tenancy_notice row for the API, including the derived tag."""
    if not row:
        return {"notice_given": 0, "notice_date": None, "move_out_date": None,
                "status": None, "tag": None, "exists": False}
    d = dict(row)
    d["exists"] = True
    d["notice_given"] = 1 if d.get("notice_given") else 0
    # The tag is DERIVED, never stored on tenancies.tags — that column is
    # overwritten wholesale by the Arthur sync (arthur_sync.py:546), so a tag
    # written there would silently vanish on the next pull.
    d["tag"] = None
    # A withdrawn notice never carries a tag, whatever it was decided before.
    d["revoked"] = d.get("status") == "revoked"
    # Both dates go out, always. Showing only the agreed one would hide the fact
    # that an exception was made; showing only the rule would be wrong.
    d["rule_move_out_date"] = _notice_move_out(d.get("notice_date"))
    d["move_out_override"] = (d.get("move_out_override") or "") or None
    d["is_override"] = bool(d["move_out_override"])
    d["move_out_date"] = _notice_effective_move_out(row)
    d["notice_months"] = _notice_months_between(d.get("notice_date"), d.get("move_out_date"))
    if d.get("status") == "approved" and d.get("move_out_date"):
        try:
            mo = datetime.strptime(d["move_out_date"][:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            mo = d["move_out_date"]
        d["tag"] = "Notice (%s)" % mo
    return d


_NOTICE_SCHEMA_READY = False


def _ensure_notice_schema(db):
    """Idempotently create the notice table and its withdrawal columns.

    A withdrawn notice keeps its original date and decision — it is revoked,
    not erased — so it needs columns of its own rather than reusing the
    approve/decline ones."""
    global _NOTICE_SCHEMA_READY
    if _NOTICE_SCHEMA_READY:
        return
    db.execute(
        "CREATE TABLE IF NOT EXISTS tenancy_notice ("
        "  tenancy_id      INTEGER PRIMARY KEY,"
        "  notice_given    INTEGER DEFAULT 0,"
        "  notice_date     TEXT,"
        "  move_out_date   TEXT,"
        "  status          TEXT DEFAULT 'pending',"
        "  decided_by      TEXT,"
        "  decided_at      TEXT,"
        "  note            TEXT,"
        "  created_by      TEXT,"
        "  created         TEXT DEFAULT CURRENT_TIMESTAMP,"
        "  updated         TEXT DEFAULT CURRENT_TIMESTAMP"
        ")")
    db.execute("CREATE INDEX IF NOT EXISTS idx_tenancy_notice_status ON tenancy_notice(status)")
    for stmt in (
        "ALTER TABLE tenancy_notice ADD COLUMN revoked_by TEXT",
        "ALTER TABLE tenancy_notice ADD COLUMN revoked_at TEXT",
        "ALTER TABLE tenancy_notice ADD COLUMN revoke_reason TEXT",
        # An agreed move-out that differs from the two-month rule. Kept separate
        # from move_out_date so the rule's answer is never lost -- the board can
        # always show what was agreed AND what the rule said.
        "ALTER TABLE tenancy_notice ADD COLUMN move_out_override TEXT",
        "ALTER TABLE tenancy_notice ADD COLUMN override_reason TEXT",
        "ALTER TABLE tenancy_notice ADD COLUMN override_by TEXT",
        "ALTER TABLE tenancy_notice ADD COLUMN override_at TEXT",
    ):
        try:
            db.execute(stmt)
        except Exception:
            pass  # column already present
    db.commit()
    _NOTICE_SCHEMA_READY = True


def _approved_move_out(db, tenancy_id):
    """The agreed move-out date, or None when there is no approved notice.

    This is the ONLY thing that may put an end date on a tenancy (Norbert,
    2026-08-03): a tenancy runs until the tenant gives notice, so an end date
    typed before that would be a guess presented as a fact."""
    try:
        row = db.execute(
            "SELECT move_out_date FROM tenancy_notice WHERE tenancy_id = ? AND status = 'approved'",
            (tenancy_id,)).fetchone()
    except Exception:
        return None  # table not created yet
    return (dict(row).get("move_out_date") if row else None) or None


def _sync_end_date_to_notice(db, tenancy_id, move_out, previous_move_out=None):
    """Keep tenancies.end_date in step with the notice.

    Approving fills the end date in, so nobody has to type it. Withdrawing or
    declining clears it again, but only when it still matches the date the
    notice put there — an end date entered for some other reason is left alone."""
    if move_out:
        db.execute("UPDATE tenancies SET end_date = ? WHERE id = ?", (move_out, tenancy_id))
        return
    cur = db.execute("SELECT end_date FROM tenancies WHERE id = ?", (tenancy_id,)).fetchone()
    cur_end = (dict(cur).get("end_date") or "")[:10] if cur else ""
    if cur_end and previous_move_out and cur_end == str(previous_move_out)[:10]:
        db.execute("UPDATE tenancies SET end_date = '' WHERE id = ?", (tenancy_id,))


def _get_notice(db, tenancy_id):
    _ensure_notice_schema(db)
    return db.execute("SELECT * FROM tenancy_notice WHERE tenancy_id = ?", (tenancy_id,)).fetchone()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice", methods=["GET"])
def api_get_tenancy_notice(tenancy_id):
    """Return the notice record for a tenancy (empty shape when none on file)."""
    db = get_dict_db()
    try:
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice", methods=["PUT"])
def api_save_tenancy_notice(tenancy_id):
    """Record that notice was given, with its date. Move-out is always derived,
    never taken from the client, so the legal calculation cannot be edited around."""
    data = request.get_json() or {}
    db = get_dict_db()
    try:
        ten = db.execute("SELECT id, status, main_tenant_name FROM tenancies WHERE id = ?", (tenancy_id,)).fetchone()
        if not ten:
            return json_error("Tenancy not found", 404)

        given = 1 if data.get("notice_given") else 0
        notice_date = (data.get("notice_date") or "").strip()[:10] or None

        if given and not notice_date:
            return json_error("A notice date is required when notice has been given.", 400)
        if notice_date:
            try:
                datetime.strptime(notice_date, "%Y-%m-%d")
            except ValueError:
                return json_error("Notice date must be a valid date.", 400)

        rule_move_out = _notice_move_out(notice_date) if given else None
        prev = _get_notice(db, tenancy_id)
        prev_d = dict(prev) if prev else {}

        # The move-out date may be typed. The two-month rule is what the field is
        # pre-filled with, not a value the user is forbidden from changing --
        # a one-month notice is a normal thing to agree, and making it hard to
        # record just means it gets recorded somewhere we cannot see.
        typed = (data.get("move_out_date") or "").strip()[:10]
        override = None
        if given and typed:
            try:
                parsed = datetime.strptime(typed, "%Y-%m-%d")
            except ValueError:
                return json_error("Move-out date must be a valid date.", 400)
            # Before the notice was served is not a short notice period, it is a
            # typo, and it would put an end date in the past on a live tenancy.
            if notice_date and parsed < datetime.strptime(notice_date, "%Y-%m-%d"):
                return json_error("The move-out date cannot be before the date notice was given.", 400)
            if typed != rule_move_out:
                override = typed

        move_out = override or rule_move_out
        override_reason = (data.get("override_reason") or "").strip() if override else ""
        override_by = (getattr(request, "current_user", {}).get("username", "system")
                       if override else None)
        override_at = datetime.now(timezone.utc).isoformat() if override else None

        # Editing the date after a decision reopens it — the approval was for a
        # different move-out date and must not silently carry over.
        status = "pending"
        actor = getattr(request, "current_user", {}).get("username", "system")
        now = datetime.now(timezone.utc).isoformat()

        db.execute(
            "INSERT INTO tenancy_notice (tenancy_id, notice_given, notice_date, move_out_date, status, note, created_by, updated, "
            " move_out_override, override_reason, override_by, override_at) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?) "
            "ON CONFLICT(tenancy_id) DO UPDATE SET notice_given=excluded.notice_given, notice_date=excluded.notice_date, "
            "move_out_date=excluded.move_out_date, status=excluded.status, note=excluded.note, "
            "decided_by=NULL, decided_at=NULL, revoked_by=NULL, revoked_at=NULL, revoke_reason=NULL, "
            "move_out_override=excluded.move_out_override, override_reason=excluded.override_reason, "
            "override_by=excluded.override_by, override_at=excluded.override_at, "
            "updated=excluded.updated",
            (tenancy_id, given, notice_date, move_out, status, data.get("note") or "", actor, now,
             override, override_reason, override_by, override_at))

        # Any edit reopens the notice as pending, so the previously agreed end
        # date is no longer agreed.
        _sync_end_date_to_notice(db, tenancy_id, None, prev_d.get("move_out_date"))

        log_activity("tenancy", tenancy_id, "notice_updated", "notice_date",
                     prev_d.get("notice_date"), notice_date, actor)
        if move_out != prev_d.get("move_out_date"):
            log_activity("tenancy", tenancy_id, "notice_move_out_set", "move_out_date",
                         prev_d.get("move_out_date"), move_out, actor)
        db.commit()
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


def _decide_notice(tenancy_id, decision):
    """Shared approve/decline handler."""
    db = get_dict_db()
    try:
        row = _get_notice(db, tenancy_id)
        if not row or not dict(row).get("notice_given"):
            return json_error("There is no notice on file for this tenancy.", 400)
        cur = dict(row)
        if not cur.get("notice_date"):
            return json_error("This notice has no date, so it cannot be decided.", 400)
        if cur.get("status") == "revoked":
            return json_error(
                "This notice was withdrawn, so it cannot be decided. "
                "Reinstate it first if the tenant is going ahead after all.", 400)

        actor = getattr(request, "current_user", {}).get("username", "system")
        now = datetime.now(timezone.utc).isoformat()
        # Recompute from the rule so a record written before a rule change is
        # decided under the current rule -- UNLESS a different date was agreed,
        # in which case the agreement is the whole point and must survive.
        move_out = _notice_effective_move_out(cur)

        db.execute(
            "UPDATE tenancy_notice SET status = ?, move_out_date = ?, decided_by = ?, decided_at = ?, updated = ? "
            "WHERE tenancy_id = ?",
            (decision, move_out if decision == "approved" else cur.get("move_out_date"),
             actor, now, now, tenancy_id))

        # The end date follows the decision rather than being typed separately.
        if decision == "approved":
            _sync_end_date_to_notice(db, tenancy_id, move_out)
        else:
            _sync_end_date_to_notice(db, tenancy_id, None, cur.get("move_out_date"))

        log_activity("tenancy", tenancy_id, "notice_" + decision, "notice_status",
                     cur.get("status"), decision, actor)
        db.commit()
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice/move-out", methods=["PUT"])
def api_set_tenancy_notice_move_out(tenancy_id):
    """Agree a move-out date that differs from the two-month rule.

    The rule is still what the board calculates and still what applies unless
    somebody deliberately changes it here, with a reason. Sending an empty date
    puts it back on the rule.

    Deliberately a separate endpoint from saving the notice: shortening someone's
    notice period is a decision, not a field edit, and it should be hard to do by
    accident while correcting a typo in a date.
    """
    data = request.get_json() or {}
    db = get_dict_db()
    try:
        row = _get_notice(db, tenancy_id)
        if not row or not dict(row).get("notice_given"):
            return json_error("There is no notice on file for this tenancy.", 400)
        cur = dict(row)
        if not cur.get("notice_date"):
            return json_error("Record the date notice was given first.", 400)
        if cur.get("status") == "revoked":
            return json_error(
                "This notice was withdrawn, so its move-out date cannot be changed. "
                "Reinstate it first if the tenant is going ahead after all.", 400)

        new_date = (data.get("move_out_date") or "").strip()[:10]
        reason = (data.get("reason") or "").strip()
        actor = getattr(request, "current_user", {}).get("username", "system")
        now = datetime.now(timezone.utc).isoformat()
        rule_date = _notice_move_out(cur["notice_date"])
        prev_effective = _notice_effective_move_out(cur)

        if not new_date:
            # Back to the rule.
            db.execute(
                "UPDATE tenancy_notice SET move_out_override = NULL, override_reason = NULL, "
                "override_by = NULL, override_at = NULL, move_out_date = ?, updated = ? "
                "WHERE tenancy_id = ?", (rule_date, now, tenancy_id))
            if cur.get("status") == "approved":
                _sync_end_date_to_notice(db, tenancy_id, rule_date)
            log_activity("tenancy", tenancy_id, "notice_move_out_cleared", "move_out_date",
                         prev_effective, rule_date, actor)
            db.commit()
            return json_success(_notice_payload(_get_notice(db, tenancy_id)))

        try:
            parsed = datetime.strptime(new_date, "%Y-%m-%d")
        except ValueError:
            return json_error("Move-out date must be a valid date.", 400)

        # A move-out before the notice was served is not a short notice period,
        # it is a mistake.
        served = datetime.strptime(str(cur["notice_date"])[:10], "%Y-%m-%d")
        if parsed < served:
            return json_error("The move-out date cannot be before the date notice was given.", 400)

        if not reason:
            return json_error(
                "Give a reason for the agreed date. The rule is two months, so a "
                "different date needs to say who agreed it and why.", 400)

        db.execute(
            "UPDATE tenancy_notice SET move_out_override = ?, override_reason = ?, "
            "override_by = ?, override_at = ?, move_out_date = ?, updated = ? "
            "WHERE tenancy_id = ?",
            (new_date, reason, actor, now, new_date, now, tenancy_id))

        # An already-approved notice keeps its approval -- the date was agreed by
        # the same people -- so the end date follows it straight away.
        if cur.get("status") == "approved":
            _sync_end_date_to_notice(db, tenancy_id, new_date)

        log_activity("tenancy", tenancy_id, "notice_move_out_agreed", "move_out_date",
                     prev_effective, new_date, actor)
        db.commit()
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice/approve", methods=["POST"])
def api_approve_tenancy_notice(tenancy_id):
    """Approve the notice. The tenancy then carries a Notice tag with the move-out date."""
    return _decide_notice(tenancy_id, "approved")


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice/decline", methods=["POST"])
def api_decline_tenancy_notice(tenancy_id):
    """Decline the notice. No tag is applied and the tenancy carries on unchanged."""
    return _decide_notice(tenancy_id, "declined")


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice/revoke", methods=["POST"])
def api_revoke_tenancy_notice(tenancy_id):
    """Withdraw a notice the tenant has changed their mind about.

    The record is kept — the date served, the decision, and who withdrew it —
    because a served-then-withdrawn notice is part of the tenancy's history and
    may matter later. The Notice tag drops and the room comes back off the
    letting pipeline, since only an approved notice feeds it."""
    data = request.get_json(silent=True) or {}
    db = get_dict_db()
    try:
        row = _get_notice(db, tenancy_id)
        if not row or not dict(row).get("notice_given"):
            return json_error("There is no notice on file for this tenancy.", 400)
        cur = dict(row)
        if cur.get("status") == "revoked":
            return json_error("This notice has already been withdrawn.", 400)

        actor = getattr(request, "current_user", {}).get("username", "system")
        now = datetime.now(timezone.utc).isoformat()
        reason = (data.get("reason") or "").strip()[:500]

        db.execute(
            "UPDATE tenancy_notice SET status = 'revoked', revoked_by = ?, revoked_at = ?, "
            "revoke_reason = ?, updated = ? WHERE tenancy_id = ?",
            (actor, now, reason, now, tenancy_id))

        # The tenant is staying, so the tenancy has no end date again.
        _sync_end_date_to_notice(db, tenancy_id, None, cur.get("move_out_date"))

        log_activity("tenancy", tenancy_id, "notice_revoked", "notice_status",
                     cur.get("status"), "revoked", actor)
        db.commit()
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice/reinstate", methods=["POST"])
def api_reinstate_tenancy_notice(tenancy_id):
    """Put a withdrawn notice back on the table.

    It returns to awaiting decision rather than straight back to approved: the
    earlier approval was withdrawn and has to be given again."""
    db = get_dict_db()
    try:
        row = _get_notice(db, tenancy_id)
        if not row:
            return json_error("There is no notice on file for this tenancy.", 400)
        cur = dict(row)
        if cur.get("status") != "revoked":
            return json_error("This notice has not been withdrawn, so there is nothing to reinstate.", 400)

        actor = getattr(request, "current_user", {}).get("username", "system")
        now = datetime.now(timezone.utc).isoformat()
        # Recompute, so a notice reinstated after a rule change lands on the
        # current calculation rather than the one it was first saved under.
        move_out = _notice_move_out(cur.get("notice_date"))

        db.execute(
            "UPDATE tenancy_notice SET status = 'pending', move_out_date = ?, decided_by = NULL, "
            "decided_at = NULL, revoked_by = NULL, revoked_at = NULL, revoke_reason = NULL, "
            "updated = ? WHERE tenancy_id = ?",
            (move_out, now, tenancy_id))

        # Back to awaiting decision, so there is no agreed end date yet.
        _sync_end_date_to_notice(db, tenancy_id, None, cur.get("move_out_date"))

        log_activity("tenancy", tenancy_id, "notice_reinstated", "notice_status",
                     "revoked", "pending", actor)
        db.commit()
        return json_success(_notice_payload(_get_notice(db, tenancy_id)))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:tenancy_id>/notice", methods=["DELETE"])
def api_clear_tenancy_notice(tenancy_id):
    """Remove the notice entirely — for one logged against the wrong tenancy."""
    db = get_dict_db()
    try:
        row = _get_notice(db, tenancy_id)
        if not row:
            return json_success(_notice_payload(None))
        actor = getattr(request, "current_user", {}).get("username", "system")
        db.execute("DELETE FROM tenancy_notice WHERE tenancy_id = ?", (tenancy_id,))
        _sync_end_date_to_notice(db, tenancy_id, None, dict(row).get("move_out_date"))
        log_activity("tenancy", tenancy_id, "notice_removed", "notice_date",
                     dict(row).get("notice_date"), None, actor)
        db.commit()
        return json_success(_notice_payload(None))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/rent-schedule/<int:tenancy_id>")
def api_rent_schedule(tenancy_id):
    """Return projected rent payment schedule for a tenancy."""
    db = get_dict_db()
    try:
        tenancy = db.execute(
            "SELECT * FROM tenancies WHERE id = ?", (tenancy_id,)
        ).fetchone()
        if not tenancy:
            return json_error("Tenancy not found", 404)

        start_date = tenancy.get("start_date")
        end_date = tenancy.get("end_date")
        rent_amount = tenancy.get("rent_amount")
        rent_frequency = tenancy.get("rent_frequency", "monthly")

        if not start_date or not rent_amount:
            return json_error("Tenancy missing start_date or rent_amount")

        try:
            cur = datetime.strptime(start_date, "%Y-%m-%d")
        except (ValueError, TypeError):
            try:
                cur = datetime.fromisoformat(start_date)
            except (ValueError, TypeError):
                return json_error("Invalid start_date format")

        if end_date:
            try:
                end = datetime.strptime(end_date, "%Y-%m-%d")
            except (ValueError, TypeError):
                try:
                    end = datetime.fromisoformat(end_date)
                except (ValueError, TypeError):
                    end = cur.replace(year=cur.year + 1)
        else:
            end = cur.replace(year=cur.year + 1)

        freq = rent_frequency.lower() if rent_frequency else "monthly"
        schedule = []
        index = 1
        cur_cursor = cur
        while cur_cursor < end:
            payment_date = cur_cursor
            if freq in ("weekly", "week"):
                delta = timedelta(weeks=1)
            elif freq in ("fortnightly", "biweekly", "2-week"):
                delta = timedelta(weeks=2)
            elif freq in ("quarterly", "quarter", "3-month"):
                delta = None
                try:
                    month = cur_cursor.month + 3
                    year = cur_cursor.year + (month - 1) // 12
                    month = ((month - 1) % 12) + 1
                    cur_cursor = cur_cursor.replace(year=year, month=month)
                except ValueError:
                    cur_cursor = cur_cursor.replace(year=cur_cursor.year + 1)
            elif freq in ("annually", "yearly", "annual", "year"):
                delta = timedelta(days=365)
            else:
                delta = timedelta(days=30)

            schedule.append({
                "payment_no": index,
                "due_date": payment_date.strftime("%Y-%m-%d"),
                "amount": float(rent_amount),
            })
            index += 1
            if delta:
                cur_cursor = cur_cursor + delta

        return json_success({
            "tenancy_id": tenancy_id,
            "rent_amount": float(rent_amount),
            "rent_frequency": rent_frequency,
            "start_date": start_date,
            "end_date": end_date,
            "schedule": schedule,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/finance/tenancy-summary/<int:tenancy_id>")
def api_tenancy_summary(tenancy_id):
    """Return financial summary for a tenancy."""
    db = get_dict_db()
    try:
        tenancy = db.execute(
            "SELECT * FROM tenancies WHERE id = ?", (tenancy_id,)
        ).fetchone()
        if not tenancy:
            return json_error("Tenancy not found", 404)

        start_date = tenancy.get("start_date")
        end_date = tenancy.get("end_date")
        rent_amount = tenancy.get("rent_amount")
        rent_frequency = tenancy.get("rent_frequency", "monthly")

        # Total paid from transactions
        total_paid = db.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total FROM transactions "
            "WHERE tenancy_id = ? AND transaction_type = 'Payment'",
            (tenancy_id,)
        ).fetchone()["total"]

        # Total expected (projected rent up to today or end_date)
        total_expected = 0.0
        if start_date and rent_amount:
            try:
                sd = datetime.strptime(start_date, "%Y-%m-%d") if isinstance(start_date, str) else start_date
            except (ValueError, TypeError):
                try:
                    sd = datetime.fromisoformat(start_date) if isinstance(start_date, str) else start_date
                except (ValueError, TypeError):
                    sd = datetime.now()

            today = datetime.now()
            if end_date:
                try:
                    ed = datetime.strptime(end_date, "%Y-%m-%d") if isinstance(end_date, str) else end_date
                except (ValueError, TypeError):
                    try:
                        ed = datetime.fromisoformat(end_date) if isinstance(end_date, str) else end_date
                    except (ValueError, TypeError):
                        ed = today
            else:
                ed = today

            freq = rent_frequency.lower() if rent_frequency else "monthly"
            cur = sd
            while cur < min(ed, today):
                if freq in ("weekly", "week"):
                    cur += timedelta(weeks=1)
                elif freq in ("fortnightly", "biweekly", "2-week"):
                    cur += timedelta(weeks=2)
                elif freq in ("quarterly", "quarter", "3-month"):
                    try:
                        month = cur.month + 3
                        year = cur.year + (month - 1) // 12
                        month = ((month - 1) % 12) + 1
                        cur = cur.replace(year=year, month=month)
                    except ValueError:
                        cur = cur.replace(year=cur.year + 1)
                elif freq in ("annually", "yearly", "annual", "year"):
                    cur += timedelta(days=365)
                else:
                    cur += timedelta(days=30)
                total_expected += float(rent_amount)

        balance = total_expected - float(total_paid) if total_paid is not None else total_expected

        # Next payment date (first date after today in the projected schedule)
        next_payment_date = None
        if start_date and rent_amount:
            try:
                sd = datetime.strptime(start_date, "%Y-%m-%d") if isinstance(start_date, str) else start_date
            except (ValueError, TypeError):
                try:
                    sd = datetime.fromisoformat(start_date) if isinstance(start_date, str) else start_date
                except (ValueError, TypeError):
                    sd = datetime.now()
            today = datetime.now()
            freq = rent_frequency.lower() if rent_frequency else "monthly"
            cur = sd
            max_iter = 500
            while cur <= today and max_iter > 0:
                if freq in ("weekly", "week"):
                    cur += timedelta(weeks=1)
                elif freq in ("fortnightly", "biweekly", "2-week"):
                    cur += timedelta(weeks=2)
                elif freq in ("quarterly", "quarter", "3-month"):
                    try:
                        month = cur.month + 3
                        year = cur.year + (month - 1) // 12
                        month = ((month - 1) % 12) + 1
                        cur = cur.replace(year=year, month=month)
                    except ValueError:
                        cur = cur.replace(year=cur.year + 1)
                elif freq in ("annually", "yearly", "annual", "year"):
                    cur += timedelta(days=365)
                else:
                    cur += timedelta(days=30)
                max_iter -= 1
            if cur > today:
                next_payment_date = cur.strftime("%Y-%m-%d")

        # Arrears from outstanding transactions
        arrears = db.execute(
            "SELECT COALESCE(SUM(amount_outstanding), 0) AS total FROM transactions "
            "WHERE tenancy_id = ? AND is_outstanding = 1",
            (tenancy_id,)
        ).fetchone()["total"]

        return json_success({
            "tenancy_id": tenancy_id,
            "total_paid": float(total_paid) if total_paid else 0.0,
            "total_expected": total_expected,
            "balance": balance,
            "next_payment_date": next_payment_date,
            "arrears": float(arrears) if arrears else 0.0,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 7. ACCESS MANAGEMENT
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/access", methods=["GET"])
def api_access_list():
    """List access records with pagination + property_id/unit_id filters."""
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    property_id = request.args.get("property_id")
    unit_id = request.args.get("unit_id")

    where_parts = []
    params = []

    if property_id:
        where_parts.append("property_id = ?")
        params.append(property_id)
    if unit_id:
        where_parts.append("unit_id = ?")
        params.append(unit_id)

    where = " AND ".join(where_parts) if where_parts else "1=1"

    rows, total = paginate(
        f"SELECT * FROM access_records WHERE {where} ORDER BY created_at DESC",
        f"SELECT COUNT(*) AS cnt FROM access_records WHERE {where}",
        params, page, per_page
    )

    return json_success(rows, total, page, per_page)


@banksia_os_bp.route("/access/<int:access_id>", methods=["GET"])
def api_access_get(access_id):
    """Get a single access record."""
    db = get_dict_db()
    try:
        record = db.execute(
            "SELECT * FROM access_records WHERE id = ?", (access_id,)
        ).fetchone()
        if not record:
            return json_error("Access record not found", 404)
        return json_success(record)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/access", methods=["POST"])
def api_access_create():
    """Create a new access record."""
    data = request.get_json(silent=True) or {}
    property_id = data.get("property_id")
    unit_id = data.get("unit_id")
    rec_type = data.get("type")  # key, fob, code
    label = data.get("label")
    identifier = data.get("identifier") or data.get("value")  # accept both
    notes = data.get("notes")
    assigned_to = data.get("assigned_to")
    issued_date = data.get("issued_date")

    if not property_id:
        return json_error("property_id is required")
    if not rec_type:
        return json_error("type is required (key, fob, code)")
    if rec_type not in ("key", "fob", "code"):
        return json_error("type must be one of: key, fob, code")

    db = get_dict_db()
    try:
        # Check property exists
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (property_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        db.execute(
            "INSERT INTO access_records (property_id, unit_id, type, label, identifier, notes, assigned_to, issued_date) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (property_id, unit_id, rec_type, label, identifier, notes, assigned_to, issued_date)
        )
        db.commit()
        new_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
        record = db.execute("SELECT * FROM access_records WHERE id = ?", (new_id,)).fetchone()
        return json_success(record)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/access/<int:access_id>", methods=["PUT"])
def api_access_update(access_id):
    """Update an existing access record."""
    data = request.get_json(silent=True) or {}

    db = get_dict_db()
    try:
        record = db.execute(
            "SELECT * FROM access_records WHERE id = ?", (access_id,)
        ).fetchone()
        if not record:
            return json_error("Access record not found", 404)

        # Build SET clause from provided fields
        allowed_fields = [
            "property_id", "unit_id", "type", "label",
            "identifier", "notes", "assigned_to",
            "issued_date", "returned_date",
        ]
        set_parts = []
        params = []
        for field in allowed_fields:
            if field == "identifier":
                val = data.get("identifier") or data.get("value")
            else:
                val = data.get(field)
            if val is not None:
                set_parts.append(f"{field} = ?")
                params.append(val)

        if not set_parts:
            return json_error("No fields to update")

        set_parts.append("updated_at = datetime('now')")
        params.append(access_id)

        db.execute(
            f"UPDATE access_records SET {', '.join(set_parts)} WHERE id = ?",
            params
        )
        db.commit()

        updated = db.execute(
            "SELECT * FROM access_records WHERE id = ?", (access_id,)
        ).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/access/available", methods=["GET"])
def api_access_available():
    """List available keys/codes for a property (unassigned records)."""
    property_id = request.args.get("property_id")
    if not property_id:
        return json_error("property_id query parameter is required")

    db = get_dict_db()
    try:
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (property_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        records = db.execute(
            "SELECT * FROM access_records "
            "WHERE property_id = ? AND (assigned_to IS NULL OR assigned_to = '') "
            "ORDER BY type ASC, label ASC",
            (property_id,)
        ).fetchall()
        return json_success(records)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 8. PROPERTY MEDIA UPLOAD
# ═══════════════════════════════════════════════

MEDIA_ROOT = os.path.join(os.path.dirname(__file__), "media")


@banksia_os_bp.route("/properties/<int:prop_id>/images", methods=["POST"])
def api_property_image_upload(prop_id):
    """Upload an image for a property. Expects multipart/form-data with 'file' field."""
    db = get_dict_db()
    try:
        prop = db.execute("SELECT id FROM properties WHERE id = ?", (prop_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)

        if "file" not in request.files:
            return json_error("No image file provided (use field 'file')")

        file = request.files["file"]
        if file.filename == "":
            return json_error("Empty filename")

        unit_id = request.form.get("unit_id", type=int) or None
        caption = request.form.get("caption", "").strip()

        # Ensure upload directory exists
        prop_dir = os.path.join(MEDIA_ROOT, "properties", str(prop_id))
        os.makedirs(prop_dir, exist_ok=True)

        # Sanitize filename — preserve extension
        orig_name = file.filename
        ext = os.path.splitext(orig_name)[1] or ""
        safe_name = f"{int(datetime.now().timestamp())}_{abs(hash(orig_name)) % 100000}{ext}"
        filepath = os.path.join(prop_dir, safe_name)
        file.save(filepath)

        # File metadata
        file_size = os.path.getsize(filepath)
        mime_type = file.content_type or "image/jpeg"
        uploaded_by = (request.current_user or {}).get("username", "") if hasattr(request, "current_user") else ""

        # Record in property_images table
        image_url = f"/api/banksia-os/media/properties/{prop_id}/{safe_name}"
        db.execute(
            "INSERT INTO property_images (property_id, unit_id, image_url, caption, category, sort_order, created_at) "
            "VALUES (?, ?, ?, ?, '', 0, datetime('now'))",
            (prop_id, unit_id, image_url, caption or "Uploaded image")
        )
        db.commit()
        image_id = db.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]

        # Return full record
        img = db.execute(
            "SELECT id, property_id, unit_id, image_url, caption, category, created_at FROM property_images WHERE id = ?",
            (image_id,)
        ).fetchone()

        return json_success(dict(img))
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/media/properties/<int:prop_id>/<filename>")
def api_serve_property_image(prop_id, filename):
    """Serve an uploaded property image."""
    prop_dir = os.path.join(MEDIA_ROOT, "properties", str(prop_id))
    filepath = os.path.join(prop_dir, filename)

    # Prevent directory traversal
    real_path = os.path.realpath(filepath)
    real_base = os.path.realpath(prop_dir)
    if not real_path.startswith(real_base):
        return json_error("Invalid path", 403)

    if not os.path.exists(filepath):
        return json_error("Image not found", 404)

    from flask import send_file
    return send_file(filepath)


@banksia_os_bp.route("/properties/images/<int:img_id>", methods=["PATCH"])
def api_property_image_patch(img_id):
    """PATCH /api/banksia-os/properties/images/:img_id — update image metadata (unit_id, caption)."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    db = get_dict_db()
    try:
        img = db.execute("SELECT * FROM property_images WHERE id = ?", (img_id,)).fetchone()
        if not img:
            return json_error("Image not found", 404)

        allowed = {"unit_id", "caption"}
        updates = {}
        for k in allowed:
            if k in data:
                updates[k] = data[k]

        if not updates:
            return json_success({"msg": "No changes"})

        set_clause = ", ".join(f"{k} = ?" for k in updates)
        vals = list(updates.values()) + [img_id]

        db.execute(
            f"UPDATE property_images SET {set_clause} WHERE id = ?",
            vals
        )
        db.commit()

        # Return updated record with unit_ref
        updated = db.execute(
            "SELECT pi.id, pi.property_id, pi.unit_id, pi.image_url, pi.caption, "
            "pi.category, pi.sort_order, pi.created_at, "
            "u.unit_ref "
            "FROM property_images pi "
            "LEFT JOIN units u ON pi.unit_id = u.id "
            "WHERE pi.id = ?",
            (img_id,)
        ).fetchone()

        d = dict(updated)
        d["url"] = d.pop("image_url")
        d["thumbnail_url"] = d["url"]

        return json_success(d)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/properties/images/<int:img_id>", methods=["DELETE"])
def api_property_image_delete(img_id):
    """DELETE /api/banksia-os/properties/images/:img_id — delete image and file."""
    db = get_dict_db()
    try:
        img = db.execute("SELECT * FROM property_images WHERE id = ?", (img_id,)).fetchone()
        if not img:
            return json_error("Image not found", 404)

        # Delete physical file
        image_url = img.get("image_url", "")
        if image_url and image_url.startswith("/api/banksia-os/media/properties/"):
            parts = image_url.split("/")
            if len(parts) >= 7:
                prop_id_str = parts[-2]
                filename = parts[-1]
                try:
                    prop_id_val = int(prop_id_str)
                    filepath = os.path.join(MEDIA_ROOT, "properties", str(prop_id_val), filename)
                    if os.path.exists(filepath):
                        os.remove(filepath)
                except (ValueError, OSError):
                    pass

        db.execute("DELETE FROM property_images WHERE id = ?", (img_id,))
        db.commit()

        return json_success({"deleted": img_id})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 11. TENANCY LIFECYCLE — Renewals, Rent Reviews, Section 21
# ═══════════════════════════════════════════════


@banksia_os_bp.route("/tenancies/<int:ten_id>/renew", methods=["POST"])
def api_renew_tenancy(ten_id):
    """Renew a tenancy — set new end date, optionally new rent."""
    data = request.get_json(silent=True) or {}
    new_end = data.get("end_date")
    new_rent = data.get("rent_amount")
    if not new_end:
        return json_error("new_end_date is required")
    db = get_dict_db()
    try:
        ten = db.execute("SELECT * FROM tenancies WHERE id = ?", (ten_id,)).fetchone()
        if not ten:
            return json_error("Tenancy not found", 404)
        now_iso = datetime.now(timezone.utc).isoformat()
        updates = {
            "renewal_start": ten.get("end_date"),
            "renewal_end": new_end,
            "end_date": new_end,
            "is_renewed": 1,
            "modified": now_iso,
        }
        if new_rent:
            updates["rent_amount"] = new_rent
        set_clause = ", ".join([f"{k} = ?" for k in updates])
        vals = list(updates.values()) + [ten_id]
        db.execute(f"UPDATE tenancies SET {set_clause} WHERE id = ?", vals)
        db.commit()
        return json_success({"renewed": True, "new_end_date": new_end})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:ten_id>/rent-review", methods=["POST"])
def api_rent_review(ten_id):
    """Record a rent review for a tenancy."""
    data = request.get_json(silent=True) or {}
    new_rent = data.get("new_rent_amount")
    review_date = data.get("review_date", datetime.now(timezone.utc).strftime("%Y-%m-%d"))
    if not new_rent:
        return json_error("new_rent_amount is required")
    db = get_dict_db()
    try:
        db.execute(
            "UPDATE tenancies SET rent_amount = ?, rent_review_date = ?, modified = ? WHERE id = ?",
            (new_rent, review_date, datetime.now(timezone.utc).isoformat(), ten_id)
        )
        db.commit()
        return json_success({"rent_reviewed": True, "new_rent": new_rent, "review_date": review_date})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/tenancies/<int:ten_id>/section-21", methods=["POST"])
def api_section_21(ten_id):
    """Record that a Section 21 notice has been served."""
    data = request.get_json(silent=True) or {}
    served_date = data.get("served_date", datetime.now(timezone.utc).strftime("%Y-%m-%d"))
    db = get_dict_db()
    try:
        db.execute(
            "UPDATE tenancies SET section_21_served = 1, modified = ? WHERE id = ?",
            (datetime.now(timezone.utc).isoformat(), ten_id)
        )
        db.commit()
        return json_success({"section_21_served": True, "served_date": served_date})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 12. TAGS SYSTEM
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/tags")
def api_tags():
    db = get_dict_db()
    try:
        tags = db.execute("SELECT * FROM tags ORDER BY name").fetchall()
        return json_success(tags)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/tags", methods=["POST"])
def api_create_tag():
    data = request.get_json()
    if not data or not data.get("name"):
        return json_error("Tag name required")
    db = get_dict_db()
    try:
        db.execute("INSERT INTO tags (name, color, category) VALUES (?,?,?)",
                   (data["name"], data.get("color","#80d8ff"), data.get("category","general")))
        db.commit()
        return json_success({"message":"Tag created"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/tags/<int:tag_id>", methods=["PATCH","DELETE"])
def api_tag(tag_id):
    if request.method == "DELETE":
        db = get_dict_db()
        try:
            db.execute("DELETE FROM tags WHERE id=?", (tag_id,))
            db.commit()
            return json_success({"deleted":True})
        except Exception as e:
            return json_error(safe_error(e), 500)
        finally:
            db.close()
    return api_update_resource("tags", tag_id)


# ═══════════════════════════════════════════════
# 13. PROPERTY OWNERS
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/property-owners")
def api_property_owners():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    search = request.args.get("search","").strip()
    order_clause = build_order_by({
        "name": "name", "company_name": "company_name",
        "main_contact_name": "main_contact_name", "email": "email",
    }, "name ASC")
    db = get_dict_db()
    try:
        if search:
            where = "WHERE name LIKE ? OR company_name LIKE ? OR main_contact_name LIKE ?"
            like = f"%{search}%"
            total = db.execute(f"SELECT COUNT(*) AS cnt FROM property_owners {where}", (like,like,like)).fetchone()["cnt"]
            rows = db.execute(f"SELECT * FROM property_owners {where} ORDER BY {order_clause} LIMIT ? OFFSET ?",
                              (like,like,like,per_page,(page-1)*per_page)).fetchall()
        else:
            total = db.execute("SELECT COUNT(*) AS cnt FROM property_owners").fetchone()["cnt"]
            rows = db.execute(f"SELECT * FROM property_owners ORDER BY {order_clause} LIMIT ? OFFSET ?",
                              (per_page,(page-1)*per_page)).fetchall()
        return json_success(rows, total, page, per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/contractors", methods=["GET"])
def api_contractors():
    """Return aggregated contractor data from maintenance jobs.
    Groups by contractor name, returns:
    - name, job_count, total_labour_cost, total_materials_cost, total_cost
    - latest_job_date, top_job_type, top_priority
    - status_breakdown: {status: count}
    """
    db = get_dict_db()
    try:
        rows = db.execute("""
            SELECT
                contractor AS name,
                COUNT(*) AS job_count,
                COALESCE(SUM(labour_cost), 0) AS total_labour_cost,
                COALESCE(SUM(materials_cost), 0) AS total_materials_cost,
                COALESCE(SUM(COALESCE(labour_cost,0) + COALESCE(materials_cost,0)), 0) AS total_cost,
                MAX(created) AS latest_job_date
            FROM maintenance_jobs
            WHERE contractor IS NOT NULL AND contractor != ''
            GROUP BY contractor
            ORDER BY total_cost DESC
        """).fetchall()

        # Add status breakdown per contractor
        status_rows = db.execute("""
            SELECT contractor, status, COUNT(*) AS cnt
            FROM maintenance_jobs
            WHERE contractor IS NOT NULL AND contractor != ''
            GROUP BY contractor, status
            ORDER BY contractor, cnt DESC
        """).fetchall()

        breakdown = {}
        for s in status_rows:
            c = s["contractor"]
            if c not in breakdown:
                breakdown[c] = {}
            breakdown[c][s["status"]] = s["cnt"]

        # Add top job type per contractor
        type_rows = db.execute("""
            SELECT contractor, type, COUNT(*) AS cnt
            FROM maintenance_jobs
            WHERE contractor IS NOT NULL AND contractor != '' AND type IS NOT NULL AND type != ''
            GROUP BY contractor, type
            ORDER BY contractor, cnt DESC
        """).fetchall()

        top_types = {}
        for t in type_rows:
            c = t["contractor"]
            if c not in top_types:
                top_types[c] = t["type"]

        result = []
        for r in rows:
            name = r["name"]
            result.append({
                "name": name,
                "job_count": r["job_count"],
                "total_labour_cost": r["total_labour_cost"],
                "total_materials_cost": r["total_materials_cost"],
                "total_cost": r["total_cost"],
                "latest_job_date": r["latest_job_date"],
                "top_job_type": top_types.get(name, ""),
                "status_breakdown": breakdown.get(name, {}),
            })

        totals = {
            "total_contractors": len(result),
            "total_jobs": sum(r["job_count"] for r in result),
            "total_labour_cost": sum(r["total_labour_cost"] for r in result),
            "total_materials_cost": sum(r["total_materials_cost"] for r in result),
            "total_cost": sum(r["total_cost"] for r in result),
        }

        return json_success({"items": result, "totals": totals})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/contractors/<contractor_name>/jobs", methods=["GET"])
def api_contractor_jobs(contractor_name):
    """Return maintenance jobs for a specific contractor."""
    from urllib.parse import unquote
    name = unquote(contractor_name)
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)

    rows, total = paginate(
        f"SELECT * FROM maintenance_jobs WHERE contractor = ? ORDER BY created DESC",
        f"SELECT COUNT(*) AS cnt FROM maintenance_jobs WHERE contractor = ?",
        [name], page, per_page
    )

    return json_success({"items": rows}, total, page, per_page)


@banksia_os_bp.route("/property-owners", methods=["POST"])
def api_create_property_owner():
    """Create a landlord and optionally auto-create a linked property."""
    data = request.get_json()
    if not data or not data.get("name"):
        return json_error("Owner name required")
    db = get_dict_db()
    try:
        cols = ["name","company_name","office_no","main_contact_name","contact_phone",
                "contact_email","address_line_1","city","postcode","status","tags","notes"]
        ins = {k:data.get(k,"") for k in cols}
        ins["modified"] = datetime.now(timezone.utc).isoformat()
        placeholders = ",".join(["?"]*len(ins))
        cursor = db.execute(f"INSERT INTO property_owners ({','.join(ins.keys())}) VALUES ({placeholders})",
                            list(ins.values()))
        owner_id = cursor.lastrowid
        db.commit()

        # Auto-create a property for this owner if requested
        created_property = None
        if data.get("create_property") and data.get("property_name"):
            prop_name = data["property_name"]
            prop_ref = data.get("property_ref", "")
            prop_addr = data.get("property_address", "")
            prop_type = data.get("property_type", "single")
            prop_city = data.get("property_city", "")
            prop_postcode = data.get("property_postcode", "")
            prop_units = int(data.get("property_units", 1))

            ins_prop = {
                "ref": prop_ref or f"OWN-{owner_id}",
                "name": prop_name,
                "address_line_1": prop_addr,
                "city": prop_city,
                "postcode": prop_postcode or "",
                "property_type": prop_type,
                "total_units": prop_units,
                "bedrooms": prop_units,
                "property_owner_id": str(owner_id),
                "property_owner_name": data["name"],
                "owner_company": data.get("company_name", ""),
                "notes": f"Auto-created from owner: {data['name']}"
            }
            ins_parts = ",".join(["?"]*len(ins_prop))
            prop_cursor = db.execute(f"INSERT INTO properties ({','.join(ins_prop.keys())}) VALUES ({ins_parts})",
                       list(ins_prop.values()))
            db.commit()
            created_property = {"id": prop_cursor.lastrowid, "name": prop_name}

        return json_success({"id": owner_id, "message": "Owner created", "created_property": created_property}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

# ── Landlord → property link ─────────────────────────────────────────────────
# One definition used by the cards endpoint, the detail endpoint and delete, so
# the count you see on a card is the same count that blocks a delete.
# property_owner_id is authoritative; the name is only consulted when a property
# carries no owner id at all (legacy Arthur rows).
OWNER_LINK_SQL = ("(property_owner_id = ? OR "
                  "(COALESCE(property_owner_id,'') = '' AND property_owner_name = ?))")


def owner_link_params(owner_id, owner_name):
    return (str(owner_id), owner_name or "")


@banksia_os_bp.route("/property-owners/cards")
def api_property_owners_cards():
    """Every landlord with their portfolio totals — powers the landlord card grid.

    Returns one row per landlord: property/unit/let counts, plus the flags the
    cards and filter chips need (management, active, inactive, archived-only).
    Counts are computed unfiltered so a search in the box never shrinks them.
    """
    db = get_dict_db()
    try:
        owners = db.execute("SELECT * FROM property_owners ORDER BY name COLLATE NOCASE").fetchall()
        props = db.execute(
            "SELECT id, ref, name, address_line_1, property_owner_id, property_owner_name, "
            "management_type, status, COALESCE(is_active,1) AS is_active FROM properties"
        ).fetchall()

        # unit totals per property (archived units excluded)
        unit_rows = db.execute(
            "SELECT property_id, COUNT(*) AS units, "
            "SUM(CASE WHEN LOWER(COALESCE(unit_status,'')) = 'occupied' THEN 1 ELSE 0 END) AS let_units "
            "FROM units WHERE archived_at IS NULL GROUP BY property_id"
        ).fetchall()
        units_by_prop = {str(u["property_id"]): (u["units"] or 0, u["let_units"] or 0) for u in unit_rows}

        by_id = {}
        by_name = {}
        for p in props:
            pid = str(p["property_owner_id"] or "")
            if pid:
                by_id.setdefault(pid, []).append(p)
            else:
                by_name.setdefault(p["property_owner_name"] or "", []).append(p)

        items = []
        for o in owners:
            linked = list(by_id.get(str(o["id"]), [])) + list(by_name.get(o["name"] or "", []))
            live = [p for p in linked if (p["status"] or "") != "archived"]
            archived = [p for p in linked if (p["status"] or "") == "archived"]
            units = lets = 0
            for p in live:
                u, l = units_by_prop.get(str(p["id"]), (0, 0))
                units += u
                lets += l
            is_mgmt = any((p["management_type"] or "") == "Management Fee" for p in live)
            active_props = [p for p in live if p["is_active"] == 1]
            row = dict(o)
            cancelled = (o.get("status") or "").lower() == "cancelled"
            row.update({
                "is_cancelled": cancelled,
                "property_count": len(live),
                "archived_property_count": len(archived),
                "unit_count": units,
                "let_count": lets,
                "is_management": is_mgmt,
                "is_active_landlord": len(active_props) > 0,
                "is_inactive_landlord": len(active_props) == 0,
                "archived_only": len(live) == 0 and len(archived) > 0,
                "property_refs": [p["address_line_1"] or p["ref"] or p["name"] for p in live],
            })
            items.append(row)

        items.sort(key=lambda r: (-r["property_count"], (r["name"] or "").lower()))
        # A cancelled landlord is off the books: it stays on record and can be
        # restored, but it does not count towards the working views.
        live = [r for r in items if not r["is_cancelled"]]
        totals = {
            "all": len(live),
            "active": sum(1 for r in live if r["is_active_landlord"]),
            "management": sum(1 for r in live if r["is_management"]),
            "inactive": sum(1 for r in live if r["is_inactive_landlord"]),
            "cancelled": sum(1 for r in items if r["is_cancelled"]),
            "properties": sum(r["property_count"] for r in live),
            "units": sum(r["unit_count"] for r in live),
        }
        return json_success({"items": items, "totals": totals})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/property-owners/all")
def api_property_owners_all():
    """Lightweight list for dropdowns — returns id, name, company_name only."""
    db = get_dict_db()
    try:
        rows = db.execute("SELECT id, name, company_name FROM property_owners ORDER BY name").fetchall()
        return json_success(rows)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/property-owners/<int:owner_id>", methods=["GET","PATCH","DELETE"])
def api_property_owner(owner_id):
    if request.method == "PATCH":
        return api_update_resource("property_owners", owner_id)
    if request.method == "DELETE":
        return api_delete_property_owner(owner_id)
    db = get_dict_db()
    try:
        owner = db.execute("SELECT * FROM property_owners WHERE id=?", (owner_id,)).fetchone()
        if not owner: return json_error("Not found", 404)
        # Count + list linked properties
        props = db.execute(
            "SELECT id, ref, name, address_line_1, city, property_type FROM properties "
            f"WHERE {OWNER_LINK_SQL} ORDER BY name",
            owner_link_params(owner_id, owner.get("name", ""))
        ).fetchall()
        owner["property_count"] = len(props)
        owner["properties"] = [dict(p) for p in props]
        return json_success(owner)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

def api_delete_property_owner(owner_id):
    db = get_dict_db()
    try:
        owner = db.execute("SELECT * FROM property_owners WHERE id=?", (owner_id,)).fetchone()
        if not owner:
            return json_error("Not found", 404)
        # Check linked properties — same rule the landlord cards count with, so
        # a card showing "no properties" can always be deleted.
        # Archived stock is excluded, exactly as the landlord cards count it —
        # otherwise a landlord reading "no properties on record" could never be
        # deleted and the refusal would name a property nobody can see.
        linked = db.execute(
            "SELECT id, ref, name, address_line_1 FROM properties "
            f"WHERE {OWNER_LINK_SQL} AND COALESCE(status,'') != 'archived' ORDER BY name",
            owner_link_params(owner_id, owner.get("name", ""))
        ).fetchall()
        if linked:
            names = [(p["address_line_1"] or p["ref"] or p["name"] or f"#{p['id']}") for p in linked]
            shown = ", ".join(names[:3]) + (f" and {len(names) - 3} more" if len(names) > 3 else "")
            return json_error(
                f"{owner.get('name') or 'This landlord'} still has "
                f"{len(linked)} propert{'y' if len(linked) == 1 else 'ies'} linked ({shown}). "
                "Reassign or unlink them first, then delete.", 409)
        db.execute("DELETE FROM property_owners WHERE id=?", (owner_id,))
        db.commit()
        _log_activity("property_owner", owner_id, "deleted",
                      notes=f"Landlord deleted: {owner.get('name') or ''}", db=db)
        db.commit()
        return json_success({"deleted": True, "name": owner.get("name", "")})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 14. MESSAGING SYSTEM (Threaded)
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/threads")
def api_threads():
    status = request.args.get("status","").strip()
    db = get_dict_db()
    try:
        where = "1=1"
        params = []
        if status:
            where = "status=?"
            params.append(status)
        threads = db.execute(
            f"SELECT * FROM message_threads WHERE {where} ORDER BY modified DESC LIMIT 50", params
        ).fetchall()
        # Get last message for each thread
        for t in threads:
            last = db.execute("SELECT author, body, created FROM messages WHERE thread_id=? ORDER BY id DESC LIMIT 1",
                              (t["id"],)).fetchone()
            t["last_message"] = last
            msg_count = db.execute("SELECT COUNT(*) AS cnt FROM messages WHERE thread_id=?", (t["id"],)).fetchone()["cnt"]
            t["message_count"] = msg_count
        return json_success(threads)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/threads", methods=["POST"])
def api_create_thread():
    data = request.get_json()
    if not data:
        return json_error("No data")
    db = get_dict_db()
    try:
        cols = ["title","entity_type","entity_id","tenancy_id","property_id",
                "status","priority","task_type","raised_by","assigned_to","participants"]
        ins = {k:data.get(k,"") for k in cols}
        ins["modified"] = datetime.now(timezone.utc).isoformat()
        pl = ",".join(["?"]*len(ins))
        cursor = db.execute(f"INSERT INTO message_threads ({','.join(ins.keys())}) VALUES ({pl})", list(ins.values()))
        db.commit()
        tid = cursor.lastrowid
        # If there's a body, create first message
        if data.get("body"):
            db.execute("INSERT INTO messages (thread_id, author, body) VALUES (?,?,?)",
                       (tid, data.get("author","System"), data["body"]))
            db.commit()
        return json_success({"id": tid, "message":"Thread created"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/threads/<int:thread_id>")
def api_thread(thread_id):
    db = get_dict_db()
    try:
        thread = db.execute("SELECT * FROM message_threads WHERE id=?", (thread_id,)).fetchone()
        if not thread: return json_error("Not found", 404)
        messages = db.execute(
            "SELECT * FROM messages WHERE thread_id=? ORDER BY id ASC", (thread_id,)
        ).fetchall()
        thread["messages"] = messages
        return json_success(thread)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/threads/<int:thread_id>/status", methods=["PATCH"])
def api_update_thread_status(thread_id):
    data = request.get_json()
    if not data or not data.get("status"):
        return json_error("Status required")
    db = get_dict_db()
    try:
        db.execute("UPDATE message_threads SET status=?, modified=? WHERE id=?",
                   (data["status"], datetime.now(timezone.utc).isoformat(), thread_id))
        db.commit()
        return json_success({"updated":True})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/threads/<int:thread_id>/attachments", methods=["POST"])
def api_upload_thread_attachment(thread_id):
    """Upload a file attachment to a message thread."""
    if "file" not in request.files:
        return json_error("No file provided")
    file = request.files["file"]
    if file.filename == "":
        return json_error("Empty filename")
    docs_dir = os.path.join(os.path.dirname(__file__), "documents")
    os.makedirs(docs_dir, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_name = f"thread_{thread_id}_{ts}_{file.filename}"
    save_path = os.path.join(docs_dir, safe_name)
    file.save(save_path)
    author = request.form.get("author", session.get("user", {}).get("username", "User"))
    db = get_dict_db()
    try:
        t = db.execute("SELECT id FROM message_threads WHERE id=?", (thread_id,)).fetchone()
        if not t:
            return json_error("Thread not found", 404)
        attachment_url = f"/api/banksia-os/threads/{thread_id}/attachments/{safe_name}"
        body = f"[File attached: {file.filename}]({attachment_url})"
        db.execute("INSERT INTO messages (thread_id, author, author_role, body) VALUES (?,?,?,?)",
                   (thread_id, author, "team", body))
        db.execute("UPDATE message_threads SET modified=? WHERE id=?",
                   (datetime.now(timezone.utc).isoformat(), thread_id))
        db.commit()
        return json_success({"filename": file.filename, "path": save_path, "url": attachment_url, "message_id": db.lastrowid}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/threads/<int:thread_id>/attachments/<path:filename>")
def api_serve_thread_attachment(thread_id, filename):
    """Serve a file attachment from the documents folder."""
    from flask import send_from_directory
    docs_dir = os.path.join(os.path.dirname(__file__), "documents")
    return send_from_directory(docs_dir, f"thread_{thread_id}_{filename}", as_attachment=True)

@banksia_os_bp.route("/messages", methods=["POST"])
def api_post_message():
    data = request.get_json()
    if not data or not data.get("thread_id") or not data.get("body"):
        return json_error("thread_id and body required")
    db = get_dict_db()
    try:
        db.execute("INSERT INTO messages (thread_id, author, author_role, body) VALUES (?,?,?,?)",
                   (data["thread_id"], data.get("author","User"), data.get("author_role","team"), data["body"]))
        db.execute("UPDATE message_threads SET modified=? WHERE id=?",
                   (datetime.now(timezone.utc).isoformat(), data["thread_id"]))
        db.commit()
        return json_success({"message":"Sent"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/messages/<int:msg_id>")
def api_get_message(msg_id):
    db = get_dict_db()
    try:
        msg = db.execute("SELECT * FROM messages WHERE id=? AND (is_deleted IS NULL OR is_deleted=0)", (msg_id,)).fetchone()
        if not msg: return json_error("Not found", 404)
        return json_success(msg)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/messages/<int:msg_id>", methods=["PATCH"])
def api_edit_message(msg_id):
    data = request.get_json()
    if not data or not data.get("body"):
        return json_error("body required")
    db = get_dict_db()
    try:
        msg = db.execute("SELECT * FROM messages WHERE id=? AND (is_deleted IS NULL OR is_deleted=0)", (msg_id,)).fetchone()
        if not msg: return json_error("Not found", 404)
        db.execute("UPDATE messages SET body=?, edited=1, edited_at=? WHERE id=?",
                   (data["body"], datetime.now(timezone.utc).isoformat(), msg_id))
        db.execute("UPDATE message_threads SET modified=? WHERE id=?",
                   (datetime.now(timezone.utc).isoformat(), msg["thread_id"]))
        db.commit()
        return json_success({"message":"Updated"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/messages/<int:msg_id>", methods=["DELETE"])
def api_delete_message(msg_id):
    db = get_dict_db()
    try:
        msg = db.execute("SELECT * FROM messages WHERE id=? AND (is_deleted IS NULL OR is_deleted=0)", (msg_id,)).fetchone()
        if not msg: return json_error("Not found", 404)
        db.execute("UPDATE messages SET body='[deleted]', is_deleted=1, edited=0 WHERE id=?", (msg_id,))
        db.commit()
        return json_success({"message":"Deleted"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 15. INVOICES
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/invoices")
def api_invoices():
    status = request.args.get("status","").strip()
    db = get_dict_db()
    try:
        where = "1=1"; params=[]
        if status:
            where = "status=?"; params=[status]
        invoices = db.execute(f"SELECT * FROM invoices WHERE {where} ORDER BY due_date DESC", params).fetchall()
        return json_success(invoices)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/invoices/summary")
def api_invoice_summary():
    db = get_dict_db()
    try:
        unpaid = db.execute("SELECT COALESCE(SUM(amount-amount_paid),0) AS total FROM invoices WHERE status!='paid'").fetchone()
        overdue = db.execute("SELECT COALESCE(SUM(amount-amount_paid),0) AS total FROM invoices WHERE due_date<date('now') AND status!='paid'").fetchone()
        due_today = db.execute("SELECT COALESCE(SUM(amount-amount_paid),0) AS total FROM invoices WHERE due_date=date('now') AND status!='paid'").fetchone()
        return json_success({
            "unpaid_total": round(unpaid["total"], 2),
            "overdue_total": round(overdue["total"], 2),
            "due_today": round(due_today["total"], 2)
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/invoices", methods=["POST"])
def api_create_invoice():
    data = request.get_json()
    if not data:
        return json_error("No data")
    db = get_dict_db()
    try:
        db.execute("INSERT INTO invoices (tenancy_id, tenant_id, invoice_ref, description, amount, due_date, status, type) VALUES (?,?,?,?,?,?,?,?)",
                   (data.get("tenancy_id"), data.get("tenant_id"), data.get("invoice_ref"),
                    data.get("description"), data.get("amount",0), data.get("due_date"),
                    data.get("status","pending"), data.get("type","rent")))
        db.commit()
        return json_success({"message":"Invoice created"}), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 16. COMPANY SETTINGS
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/company-settings")
def api_company_settings():
    db = get_dict_db()
    try:
        rows = db.execute("SELECT key, value FROM company_settings").fetchall()
        return json_success({r["key"]: r["value"] for r in rows})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

@banksia_os_bp.route("/company-settings", methods=["POST"])
def api_update_company_settings():
    data = request.get_json()
    if not data:
        return json_error("No data")
    db = get_dict_db()
    try:
        for key, value in data.items():
            db.execute("INSERT INTO company_settings (key, value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                       (key, value))
        db.commit()
        return json_success({"message":"Settings saved"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 17. ENHANCED PROPERTIES — filtered/tagged
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/properties/enhanced")
def api_properties_enhanced():
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    search = request.args.get("search","").strip()
    condition = request.args.get("condition","").strip()  # Comma-separated filter codes
    tag_filter = request.args.get("tag","").strip()

    where_parts = ["1=1"]
    params = []
    if search:
        like = f"%{search}%"
        where_parts.append("(ref LIKE ? OR address_line_1 LIKE ? OR city LIKE ? OR postcode LIKE ?)")
        params.extend([like]*4)
    if condition:
        # Map filter codes to LIKE patterns
        # HMOGuaranteed → "HMO / Guaranteed", HMOManagement → "HMO / Management"
        # SingleGuaranteed → "Single / Guaranteed", SingleManagement → "Single / Management"
        code_map = {
            "HMOGuaranteed": "HMO / Guaranteed",
            "HMOManagement": "HMO / Management",
            "SingleGuaranteed": "Single / Guaranteed",
            "SingleManagement": "Single / Management",
        }
        codes = [c.strip() for c in condition.split(",") if c.strip()]
        like_clauses = []
        for code in codes:
            mapped = code_map.get(code)
            if mapped:
                like_clauses.append("property_type=?")
                params.append(mapped)
        if like_clauses:
            where_parts.append("(" + " OR ".join(like_clauses) + ")")
    if tag_filter:
        where_parts.append("tags LIKE ?")
        params.append(f"%{tag_filter}%")

    where = " AND ".join(where_parts)
    db = get_dict_db()
    try:
        total = db.execute(f"SELECT COUNT(*) AS cnt FROM properties WHERE {where}", params).fetchone()["cnt"]
        props = db.execute(
            f"SELECT * FROM properties WHERE {where} ORDER BY name ASC LIMIT ? OFFSET ?",
            params + [per_page, (page-1)*per_page]
        ).fetchall()
        # Enrich with unit counts and owner info
        for p in props:
            total_u = db.execute("SELECT COUNT(*) AS cnt FROM units WHERE property_id=?", (p["id"],)).fetchone()["cnt"]
            avail_u = db.execute("SELECT COUNT(*) AS cnt FROM units WHERE property_id=? AND unit_vacant=1", (p["id"],)).fetchone()["cnt"]
            occupied_u = db.execute("SELECT COUNT(*) AS cnt FROM units WHERE property_id=? AND unit_vacant=0", (p["id"],)).fetchone()["cnt"]
            p["total_unit_count"] = total_u
            p["available_units"] = avail_u
            p["occupied_units"] = occupied_u
            # Resolve owner display from property_owners table
            owner_id = p.get("property_owner_id", "")
            if owner_id:
                try:
                    oid = int(float(owner_id))
                    owner_info = db.execute("SELECT id, name, company_name FROM property_owners WHERE id=?", (oid,)).fetchone()
                    if owner_info:
                        p["owner_display_name"] = owner_info["name"]
                        p["owner_display_id"] = owner_info["id"]
                    else:
                        p["owner_display_name"] = p.get("property_owner_name", "")
                        p["owner_display_id"] = None
                except (ValueError, TypeError):
                    p["owner_display_name"] = p.get("property_owner_name", "")
                    p["owner_display_id"] = None
            else:
                p["owner_display_name"] = p.get("property_owner_name", "")
                p["owner_display_id"] = None
            # Compute monthly rent from active tenancies
            rent = db.execute(
                "SELECT COALESCE(SUM(rent_amount), 0) AS total FROM tenancies WHERE property_id=? AND status IN ('Current','current','Periodic','periodic','Active','active')",
                (p["id"],)
            ).fetchone()["total"]
            p["monthly_rent"] = rent
            # Property status
            p["property_status"] = "Active" if occupied_u > 0 else ("Vacant" if total_u > 0 else "No Units")
            # Parse tags from JSON string
            if p.get("tags"):
                try:
                    import json as jmod
                    parsed = jmod.loads(p["tags"])
                    p["tags_list"] = parsed if isinstance(parsed, list) else [str(parsed)]
                except (json.JSONDecodeError, TypeError):
                    p["tags_list"] = [t.strip() for t in p["tags"].split(",") if t.strip()]
            else:
                p["tags_list"] = []
        return json_success(props, total, page, per_page)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════
# 18. INVOICE DETAIL / PAY / CANCEL
# ═══════════════════════════════════════════════

@banksia_os_bp.route("/invoices/<int:invoice_id>", methods=["GET"])
def api_invoice_detail(invoice_id):
    db = get_dict_db()
    try:
        inv = db.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,)).fetchone()
        if not inv:
            return json_error("Not found", 404)
        if inv.get("tenancy_id"):
            tn = db.execute("SELECT * FROM tenancies WHERE id=?", (inv["tenancy_id"],)).fetchone()
            if tn:
                inv["tenant_name"] = tn.get("main_tenant_name") or tn.get("tenant_name")
                prop = db.execute("SELECT * FROM properties WHERE id=?", (tn.get("property_id"),)).fetchone()
                if prop:
                    inv["property_name"] = prop.get("name") or prop.get("ref") or prop.get("address_line_1")
        return json_success(inv)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/invoices/<int:invoice_id>", methods=["PATCH"])
def api_update_invoice(invoice_id):
    """PATCH /api/banksia-os/invoices/{id} — update invoice fields."""
    data = request.get_json(silent=True)
    if not data:
        return json_error("No data provided", 400)

    db = get_dict_db()
    try:
        inv = db.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,)).fetchone()
        if not inv:
            return json_error("Not found", 404)

        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(invoices)").fetchall()}
        protected_keys = {"id", "created", "modified"}

        set_parts = []
        params = []
        for key, val in data.items():
            if key in protected_keys or key not in real_cols:
                continue
            set_parts.append(f"{key} = ?")
            params.append(val)

        if not set_parts:
            return json_error("No valid fields to update", 400)

        now = datetime.now(timezone.utc).isoformat()
        set_parts.append("modified = ?")
        params.append(now)
        params.append(invoice_id)

        db.execute(f"UPDATE invoices SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()

        updated = db.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,)).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/invoices/<int:invoice_id>/pay", methods=["POST"])
def api_pay_invoice(invoice_id):
    db = get_dict_db()
    try:
        inv = db.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,)).fetchone()
        if not inv:
            return json_error("Not found", 404)
        from datetime import datetime, timezone
        db.execute("UPDATE invoices SET status='paid', paid_date=? WHERE id=?",
                   (datetime.now(timezone.utc).isoformat(), invoice_id))
        db.commit()
        return json_success({"message": "Invoice marked as paid"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/invoices/<int:invoice_id>", methods=["DELETE"])
def api_cancel_invoice(invoice_id):
    db = get_dict_db()
    try:
        inv = db.execute("SELECT * FROM invoices WHERE id=?", (invoice_id,)).fetchone()
        if not inv:
            return json_error("Not found", 404)
        db.execute("UPDATE invoices SET status='cancelled' WHERE id=?", (invoice_id,))
        db.commit()
        return json_success({"message": "Invoice cancelled"})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9. APPLICANT-TO-TENANCY WORKFLOW
# ═══════════════════════════════════════════════════════════════
#
# Endpoints:
#   Applicant CRUD + status transitions
#   Referencing lifecycle
#   Guarantor CRUD
#   Single-transaction applicant-to-tenancy conversion
#   Unit occupancy check
# ═══════════════════════════════════════════════════════════════

# ── Applicant status machine ──
APPLICANT_VALID_TRANSITIONS = {
    "new":              ["form_sent", "withdrawn"],
    "form_sent":        ["submitted", "withdrawn"],
    "submitted":        ["under_review", "more_info", "withdrawn"],
    "under_review":     ["approved", "declined", "more_info", "withdrawn"],
    "more_info":        ["submitted", "withdrawn"],
    "approved":         ["tenancy_created", "withdrawn"],
    "declined":         ["withdrawn"],
    "tenancy_created":  ["withdrawn"],
    "withdrawn":        [],
}

# ── Referencing status machine ──
REFERENCING_VALID_TRANSITIONS = {
    "new":              ["form_sent", "withdrawn"],
    "form_sent":        ["submitted", "withdrawn"],
    "submitted":        ["under_review", "more_info", "withdrawn"],
    "under_review":     ["approved", "declined", "more_info", "withdrawn"],
    "more_info":        ["submitted", "withdrawn"],
    "approved":         ["withdrawn"],
    "declined":         ["withdrawn"],
    "tenancy_created":  ["withdrawn"],
    "withdrawn":        [],
}


# ═══════════════════════════════════════════════════════════════
# 9A. APPLICANT ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@banksia_os_bp.route("/applicants", methods=["GET"])
def api_applicants_list():
    """List applicants with search, pagination, status filter."""
    page = int_param(request.args.get("page"))
    per_page = int_param(request.args.get("per_page"), 20, max_val=MAX_PAGE_SIZE)
    status_filter = request.args.get("status", "").strip()
    search = request.args.get("search", "").strip()

    where_parts = ["1=1"]
    params = []

    if status_filter and status_filter.lower() == "all":
        pass  # explicit 'show me everything', cancelled included
    elif status_filter:
        where_parts.append("status = ?")
        params.append(status_filter)
    else:
        # Cancelled applicants are off the main list (Norbert, 2026-08-03): cancelling
        # is how the team clears the pipeline, so leaving them in defeats the point.
        # They are not deleted — filter Pipeline Stage = Cancelled (or status=all) and
        # they are all still there, with their audit trail.
        where_parts.append("LOWER(COALESCE(status,'')) != 'cancelled'")

    if search:
        search_clause, search_params = build_search_clause(
            ["first_name", "last_name", "email", "mobile", "phone"], search
        )
        where_parts.append(search_clause)
        params.extend(search_params)

    where = " AND ".join(where_parts)

    order_clause = build_order_by({
        "created": "created", "status": "status",
        "last_name": "last_name", "first_name": "first_name",
        "email": "email",
    }, "created DESC")

    rows, total = paginate(
        f"SELECT * FROM applicants WHERE {where} ORDER BY {order_clause}",
        f"SELECT COUNT(*) AS cnt FROM applicants WHERE {where}",
        params, page, per_page
    )

    for r in rows:
        bool_fields(r, "has_guarantor")

    return json_success(rows, total, page, per_page)


@banksia_os_bp.route("/applicants/<int:app_id>", methods=["GET"])
def api_applicant_detail(app_id):
    """Get full applicant detail with linked referencing(s) and guarantor(s)."""
    db = get_dict_db()
    try:
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)

        bool_fields(app, "has_guarantor")

        # Load referencing records (new-style referencing_forms)
        refs = db.execute(
            "SELECT * FROM referencing_forms WHERE applicant_id = ? ORDER BY created DESC",
            (app_id,)
        ).fetchall()
        app["referencing_forms"] = refs

        # Load new-style referencing_checks
        ref_checks = db.execute(
            "SELECT rc.* FROM referencing_checks rc "
            "JOIN referencing_forms rf ON rf.id = rc.form_id "
            "WHERE rf.applicant_id = ? ORDER BY rc.created DESC",
            (app_id,)
        ).fetchall()
        app["referencing_checks"] = ref_checks

        # Load guarantor(s)
        gs = db.execute(
            "SELECT * FROM guarantors WHERE applicant_id = ?", (app_id,)
        ).fetchall()
        app["guarantors"] = gs

        # Load converted tenancy (if applicant moved through to tenancy)
        converted_tenancy = None
        for ref in refs:
            if ref.get("status") == "tenancy_created":
                # Find tenancy created from this referencing form
                tn = db.execute(
                    "SELECT id, ref, status, start_date, end_date, rent_amount, full_address "
                    "FROM tenancies WHERE sync_origin='banksia' AND notes LIKE ? ORDER BY id DESC LIMIT 1",
                    (f"%form_id={ref['id']}%",)
                ).fetchone()
                if tn:
                    converted_tenancy = dict(tn)
                    # Also find linked tenant
                    tenant_info = db.execute(
                        "SELECT id, first_name, last_name, email, mobile FROM tenants WHERE tenancy_id=? LIMIT 1",
                        (tn["id"],)
                    ).fetchone()
                    if tenant_info:
                        converted_tenancy["tenant"] = dict(tenant_info)
                    break
        app["converted_tenancy"] = converted_tenancy

        return json_success(app)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants", methods=["POST"])
def api_create_applicant():
    """Create a new applicant."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    first_name = (data.get("first_name") or "").strip()
    last_name = (data.get("last_name") or "").strip()
    if not first_name or not last_name:
        return json_error("first_name and last_name are required")

    email = (data.get("email") or "").strip()
    phone = (data.get("phone") or "").strip()
    property_id = data.get("property_id")
    unit_id = data.get("unit_id")
    proposed_rent = data.get("proposed_rent")
    proposed_deposit = data.get("proposed_deposit")
    desired_move_in = data.get("desired_move_in")
    assigned_to = (data.get("assigned_to") or "").strip()
    notes = data.get("notes", "").strip()
    has_guarantor = 1 if data.get("has_guarantor") else 0

    db = get_dict_db()
    try:
        now = datetime.now(timezone.utc).isoformat()

        # Validate property exists if property_id provided
        if property_id:
            prop = db.execute("SELECT id FROM properties WHERE id = ?", (property_id,)).fetchone()
            if not prop:
                return json_error(f"Property {property_id} not found", 404)

        # Validate unit belongs to property if both provided
        if property_id and unit_id:
            unit = db.execute(
                "SELECT id FROM units WHERE id = ? AND property_id = ?",
                (unit_id, property_id)
            ).fetchone()
            if not unit:
                return json_error(f"Unit {unit_id} does not belong to property {property_id}", 400)

        branch_id = getattr(request, "current_user", {}).get("branch_id", "")

        cur = db.execute(
            "INSERT INTO applicants (first_name, last_name, email, phone, property_id, unit_id, "
            "proposed_rent, proposed_deposit, desired_move_in, assigned_to, applicant_note, "
            "has_guarantor, status, created, modified) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'new', ?, ?)",
            [first_name, last_name, email, phone, property_id, unit_id,
             proposed_rent, proposed_deposit, desired_move_in, assigned_to, notes,
             has_guarantor, now, now]
        )
        db.commit()
        new_id = cur.lastrowid

        _log_activity("applicant", new_id, "created",
                       notes=f"Applicant {first_name} {last_name} created",
                       db=db)

        app = db.execute("SELECT * FROM applicants WHERE id = ?", (new_id,)).fetchone()
        bool_fields(app, "has_guarantor")
        return json_success(app), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:app_id>", methods=["PATCH"])
def api_update_applicant(app_id):
    """Update applicant fields (title-safe fields only)."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    db = get_dict_db()
    try:
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)

        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(applicants)").fetchall()}
        protected_keys = {"id", "arthur_id", "sync_dirty", "local_modified", "sync_origin", "pushed_at"}

        set_parts = []
        params = []
        changed_fields = []

        for key, val in data.items():
            if key in protected_keys or key not in real_cols:
                continue
            old_val = app.get(key)
            set_parts.append(f"{key} = ?")
            params.append(val)
            changed_fields.append((key, old_val, val))

        if not set_parts:
            return json_error("No valid fields to update")

        now = datetime.now(timezone.utc).isoformat()
        set_parts.append("modified = ?")
        params.append(now)
        params.append(app_id)

        db.execute(f"UPDATE applicants SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()

        for field, old_val, new_val in changed_fields:
            _log_activity("applicant", app_id, "update",
                           field_changed=field, old_value=str(old_val) if old_val else None,
                           new_value=str(new_val) if new_val else None, db=db)

        updated = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        bool_fields(updated, "has_guarantor")
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:app_id>/status", methods=["POST"])
def api_transition_applicant_status(app_id):
    """Transition applicant status with validation."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    new_status = (data.get("status") or "").strip().lower()
    if not new_status:
        return json_error("status is required")

    db = get_dict_db()
    try:
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)

        current = (app.get("status") or "new").strip().lower()
        allowed = APPLICANT_VALID_TRANSITIONS.get(current, [])

        if new_status not in allowed:
            # Also allow remaining transitions from tenancy_created state
            if current == "approved" and new_status == "tenancy_created":
                pass  # explicit allow
            else:
                return json_error(
                    f"Cannot transition from '{current}' to '{new_status}'. "
                    f"Allowed: {allowed or '(none)'}",
                    400
                )

        now = datetime.now(timezone.utc).isoformat()
        db.execute("UPDATE applicants SET status = ?, modified = ? WHERE id = ?",
                   (new_status, now, app_id))
        db.commit()

        _log_activity("applicant", app_id, "status_change",
                       field_changed="status",
                       old_value=current, new_value=new_status,
                       notes=f"Status changed from {current} to {new_status}",
                       db=db)

        updated = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        bool_fields(updated, "has_guarantor")
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


APPLICANT_CANCELLED_STATUS = "Cancelled"
BULK_CANCEL_MAX = 200


@banksia_os_bp.route("/applicants/bulk-cancel", methods=["POST"])
def api_applicants_bulk_cancel():
    """Cancel several applicants in one go (Norbert, 2026-08-03).

    Cancelling is a status change, never a delete: the applicant, their referencing
    form and their documents all stay, so a cancellation made by mistake is undone
    by setting the status back. Each row is logged to activity_log individually so
    the audit trail reads the same as if they had been cancelled one at a time.

    Already-cancelled and unknown ids are reported back rather than failing the
    whole batch — a stale tab must not be able to turn one bad id into a no-op for
    the other 40 rows the user picked.
    """
    data = request.get_json(silent=True) or {}
    raw_ids = data.get("ids")
    if not isinstance(raw_ids, list) or not raw_ids:
        return json_error("ids must be a non-empty list of applicant ids")

    ids = []
    for value in raw_ids:
        try:
            ids.append(int(value))
        except (TypeError, ValueError):
            return json_error("ids must all be numeric applicant ids")
    ids = list(dict.fromkeys(ids))

    if len(ids) > BULK_CANCEL_MAX:
        return json_error("Too many applicants at once (limit %d)" % BULK_CANCEL_MAX, 413)

    reason = (data.get("reason") or "").strip()[:200]

    db = get_dict_db()
    try:
        placeholders = ",".join("?" * len(ids))
        rows = db.execute(
            "SELECT id, first_name, last_name, status FROM applicants WHERE id IN (%s)" % placeholders,
            ids
        ).fetchall()
        found = {r["id"]: r for r in rows}

        not_found = [i for i in ids if i not in found]
        already = [i for i in ids
                   if i in found
                   and (found[i].get("status") or "").strip().lower() == APPLICANT_CANCELLED_STATUS.lower()]
        to_cancel = [i for i in ids if i in found and i not in already]

        if not to_cancel:
            return json_success({
                "cancelled": 0, "cancelled_ids": [],
                "already_cancelled": already, "not_found": not_found,
            })

        now = datetime.now(timezone.utc).isoformat()
        marks = ",".join("?" * len(to_cancel))
        db.execute(
            "UPDATE applicants SET status = ?, modified = ? WHERE id IN (%s)" % marks,
            [APPLICANT_CANCELLED_STATUS, now] + to_cancel
        )

        # Log BEFORE the commit, not after. _log_activity does not commit when it is
        # handed a connection, and this connection is thread-local and closed at the
        # end of the request — so an audit row written after db.commit() sits in a
        # fresh implicit transaction that close() throws away. Logging first puts the
        # status change and its audit trail in one transaction: both land, or neither.
        for app_id in to_cancel:
            row = found[app_id]
            name = ("%s %s" % (row.get("first_name") or "", row.get("last_name") or "")).strip() or "applicant"
            note = "Cancelled in bulk (%d selected)" % len(to_cancel)
            if reason:
                note += " — %s" % reason
            _log_activity("applicant", app_id, "status_change",
                          field_changed="status",
                          old_value=row.get("status") or "",
                          new_value=APPLICANT_CANCELLED_STATUS,
                          notes="%s: %s" % (name, note), db=db)

        db.commit()

        return json_success({
            "cancelled": len(to_cancel), "cancelled_ids": to_cancel,
            "already_cancelled": already, "not_found": not_found,
        })
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9B. REFERENCING ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@banksia_os_bp.route("/referencing/create", methods=["POST"])
def api_create_referencing_standalone():
    """Create a referencing record manually (the manual version of the applicant portal).

    Mirrors the portal signup: creates an applicant record (so property/unit linkage and
    the referencing detail view work) plus a referencing form linked to it, in one
    atomic transaction. Property and unit are optional.
    """
    data = request.get_json(silent=True) or {}
    first_name = (data.get("first_name") or "").strip()
    last_name = (data.get("last_name") or "").strip()
    email = (data.get("email") or "").strip()

    if not first_name or not last_name or not email:
        return json_error("first_name, last_name, and email are required", 400)

    import secrets
    db = get_dict_db()
    try:
        now = datetime.now(timezone.utc).isoformat()
        phone = (data.get("phone") or "").strip()
        property_id = data.get("property_id") or None
        unit_ref = data.get("unit_ref")

        # Resolve property + unit (both optional). unit_ref may be a unit_ref string or a unit id.
        unit_id = None
        if property_id:
            prop = db.execute("SELECT id FROM properties WHERE id = ?", (property_id,)).fetchone()
            if not prop:
                return json_error(f"Property {property_id} not found", 404)
            if unit_ref:
                unit = db.execute(
                    "SELECT id FROM units WHERE property_id = ? AND (unit_ref = ? OR CAST(id AS TEXT) = ?)",
                    (property_id, str(unit_ref), str(unit_ref))
                ).fetchone()
                if unit:
                    unit_id = unit["id"]

        # Create the applicant (manual entry) so property/unit link the same way the portal does
        acur = db.execute(
            "INSERT INTO applicants (first_name, last_name, email, phone, property_id, unit_id, "
            "status, created, modified) VALUES (?, ?, ?, ?, ?, ?, 'new', ?, ?)",
            [first_name, last_name, email, phone, property_id, unit_id, now, now]
        )
        applicant_id = acur.lastrowid

        # Create the referencing form linked to the new applicant
        form_token = secrets.token_urlsafe(32)
        cur = db.execute(
            "INSERT INTO referencing_forms (applicant_id, form_token, status, first_name, last_name, "
            "email, mobile_phone, created, modified) "
            "VALUES (?, ?, 'new', ?, ?, ?, ?, ?, ?)",
            [applicant_id, form_token, first_name, last_name, email, phone, now, now]
        )
        form_id = cur.lastrowid

        _log_activity("applicant", applicant_id, "created",
                       notes=f"Manual applicant {first_name} {last_name} created via New Referencing", db=db)
        _log_activity("referencing_form", form_id, "created",
                       notes=f"Standalone referencing created for {first_name} {last_name}", db=db)

        form = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (form_id,)).fetchone()
        db.commit()
        return json_success(form), 201
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:app_id>/referencing", methods=["POST"])
def api_create_referencing(app_id):
    """Create a referencing for an applicant (links to applicant_id)."""
    db = get_dict_db()
    try:
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)

        now = datetime.now(timezone.utc).isoformat()

        # Insert into referencing_forms (new-style)
        import secrets
        form_token = secrets.token_urlsafe(32)

        cur = db.execute(
            "INSERT INTO referencing_forms (applicant_id, form_token, status, first_name, last_name, "
            "email, created, modified) "
            "VALUES (?, ?, 'new', ?, ?, ?, ?, ?)",
            [app_id, form_token, app.get("first_name", ""), app.get("last_name", ""),
             app.get("email", ""), now, now]
        )
        form_id = cur.lastrowid

        _log_activity("referencing_form", form_id, "created",
                       notes=f"Referencing created for applicant #{app_id}", db=db)

        ref = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (form_id,)).fetchone()
        return json_success(ref), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/applicants/<int:app_id>/referencing", methods=["GET"])
def api_applicant_referencing(app_id):
    """Latest referencing form (full submitted data) for the Applicant detail Referencing tab."""
    db = get_dict_db()
    try:
        app = db.execute("SELECT id FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)
        form = db.execute(
            "SELECT * FROM referencing_forms WHERE applicant_id = ? ORDER BY created DESC LIMIT 1",
            (app_id,)
        ).fetchone()
        if not form:
            return json_success(None)
        if form.get("status") == "draft":
            form["status"] = "new"
        form["checks"] = db.execute(
            "SELECT * FROM referencing_checks WHERE form_id = ? ORDER BY created",
            (form["id"],)
        ).fetchall()
        form["documents"] = db.execute(
            "SELECT * FROM referencing_documents WHERE form_id = ? ORDER BY uploaded_at",
            (form["id"],)
        ).fetchall()
        return json_success(form)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/referencing/<int:ref_id>", methods=["GET"])
def api_get_referencing(ref_id):
    """Get referencing detail with forms, documents, history."""
    db = get_dict_db()
    try:
        form = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        if not form:
            return json_error("Referencing not found", 404)

        # Map DB draft → frontend new
        if form.get("status") == "draft":
            form["status"] = "new"

        # Attach property info via applicant → unit → property chain
        property_name = ""
        unit_ref = ""
        property_id = None
        unit_id = None
        if form.get("applicant_id"):
            app_info = db.execute(
                "SELECT a.property_id AS apid, a.unit_id AS auid, "
                "COALESCE(p.name, p2.name, p2.address_line_1) AS pname, u.unit_ref AS uref "
                "FROM applicants a "
                "LEFT JOIN units u ON a.unit_id = u.id "
                "LEFT JOIN properties p ON u.property_id = p.id "
                "LEFT JOIN properties p2 ON a.property_id = p2.id "
                "WHERE a.id = ?",
                (form["applicant_id"],)
            ).fetchone()
            if app_info:
                property_name = app_info["pname"] or ""
                unit_ref = app_info["uref"] or ""
                property_id = app_info["apid"]
                unit_id = app_info["auid"]
        form["property_name"] = property_name
        form["unit_ref"] = unit_ref
        form["property_id"] = property_id
        form["unit_id"] = unit_id

        # Attach referencing_checks
        checks = db.execute(
            "SELECT * FROM referencing_checks WHERE form_id = ? ORDER BY created",
            (ref_id,)
        ).fetchall()
        form["checks"] = checks

        # Attach referencing_documents
        docs = db.execute(
            "SELECT * FROM referencing_documents WHERE form_id = ? ORDER BY uploaded_at",
            (ref_id,)
        ).fetchall()
        form["documents"] = docs

        # Attach applicant info
        if form.get("applicant_id"):
            app = db.execute("SELECT id, first_name, last_name, email, phone, status FROM applicants WHERE id = ?",
                             (form["applicant_id"],)).fetchone()
            form["applicant"] = app

        return json_success(form)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/referencing/<int:ref_id>", methods=["PATCH"])
def api_update_referencing(ref_id):
    """Update referencing fields."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    db = get_dict_db()
    try:
        form = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        if not form:
            return json_error("Referencing not found", 404)

        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(referencing_forms)").fetchall()}
        protected_keys = {"id", "form_token", "submitted_at", "created"}

        set_parts = []
        params = []

        for key, val in data.items():
            if key in protected_keys or key not in real_cols:
                continue
            set_parts.append(f"{key} = ?")
            params.append(val)

        if not set_parts:
            return json_error("No valid fields to update")

        now = datetime.now(timezone.utc).isoformat()
        set_parts.append("modified = ?")
        params.append(now)
        params.append(ref_id)

        db.execute(f"UPDATE referencing_forms SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()

        updated = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/referencing/<int:ref_id>", methods=["DELETE"])
def api_delete_referencing(ref_id):
    """Delete a referencing record."""
    db = get_dict_db()
    try:
        form = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        if not form:
            return json_error("Referencing not found", 404)
        applicant_id = form.get("applicant_id")
        who = ((str(form.get("first_name") or "") + " " + str(form.get("last_name") or "")).strip()
               or form.get("email") or f"#{ref_id}")
        status = form.get("status") or ""
        db.execute("DELETE FROM referencing_forms WHERE id = ?", (ref_id,))
        db.commit()
        _log_activity("referencing", ref_id, "deleted",
                      notes=f"Referencing application for {who}"
                            + (f" (status: {status})" if status else "")
                            + " permanently deleted",
                      db=db)
        return json_success({"deleted": ref_id})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/referencing/<int:ref_id>/status", methods=["POST"])
def api_transition_referencing_status(ref_id):
    """Transition referencing status. If approved, auto-sets applicant status."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    new_status = (data.get("status") or "").strip().lower()
    if not new_status:
        return json_error("status is required")

    db = get_dict_db()
    try:
        form = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        if not form:
            return json_error("Referencing not found", 404)

        current = (form.get("status") or "new").strip().lower()
        allowed = REFERENCING_VALID_TRANSITIONS.get(current, [])

        if new_status not in allowed:
            return json_error(
                f"Cannot transition referencing from '{current}' to '{new_status}'. "
                f"Allowed: {allowed or '(none)'}",
                400
            )

        now = datetime.now(timezone.utc).isoformat()

        # If approved, set the submitted_at/reviewed_at
        extra_updates = []
        if new_status == "approved":
            extra_updates.append("reviewed_at = ?")
            extra_updates.append("reviewed_by = ?")
            user_name = getattr(request, "current_user", {}).get("username", "system")
            params = [now, user_name, new_status, now, ref_id]
        else:
            params = [new_status, now, ref_id]

        if extra_updates:
            db.execute(
                f"UPDATE referencing_forms SET status = ?, modified = ?, {', '.join(extra_updates)} WHERE id = ?",
                params
            )
        else:
            db.execute(
                "UPDATE referencing_forms SET status = ?, modified = ? WHERE id = ?",
                params
            )
        db.commit()

        _log_activity("referencing_form", ref_id, "status_change",
                       field_changed="status",
                       old_value=current, new_value=new_status, db=db)

        # If approved, auto-set applicant status to 'approved'
        if new_status == "approved" and form.get("applicant_id"):
            app = db.execute("SELECT status FROM applicants WHERE id = ?",
                             (form["applicant_id"],)).fetchone()
            if app and app.get("status", "").strip().lower() != "approved":
                app_current = app["status"]
                db.execute("UPDATE applicants SET status = 'approved', modified = ? WHERE id = ?",
                           (now, form["applicant_id"]))
                db.commit()
                _log_activity("applicant", form["applicant_id"], "status_change",
                               field_changed="status",
                               old_value=app_current, new_value="approved",
                               notes="Auto-approved via referencing approval", db=db)

        updated = db.execute("SELECT * FROM referencing_forms WHERE id = ?", (ref_id,)).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9C. GUARANTOR ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@banksia_os_bp.route("/applicants/<int:app_id>/guarantor", methods=["POST"])
def api_create_guarantor(app_id):
    """Create a guarantor for an applicant."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    first_name = (data.get("first_name") or "").strip()
    last_name = (data.get("last_name") or "").strip()
    if not first_name or not last_name:
        return json_error("guarantor first_name and last_name are required")

    db = get_dict_db()
    try:
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            return json_error("Applicant not found", 404)

        now = datetime.now(timezone.utc).isoformat()

        cur = db.execute(
            "INSERT INTO guarantors (applicant_id, first_name, last_name, email, phone, "
            "address, city, postcode, country, employment, income, relation, created, modified) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            [app_id, first_name, last_name,
             (data.get("email") or "").strip(),
             (data.get("phone") or "").strip(),
             (data.get("address") or "").strip(),
             (data.get("city") or "").strip(),
             (data.get("postcode") or "").strip(),
             (data.get("country") or "").strip(),
             (data.get("employment") or "").strip(),
             data.get("income"),
             (data.get("relation") or "").strip(),
             now, now]
        )
        db.commit()
        new_id = cur.lastrowid

        # Update has_guarantor on applicant
        db.execute("UPDATE applicants SET has_guarantor = 1, modified = ? WHERE id = ?",
                   (now, app_id))
        db.commit()

        _log_activity("guarantor", new_id, "created",
                       notes=f"Guarantor {first_name} {last_name} created for applicant #{app_id}",
                       db=db)

        g = db.execute("SELECT * FROM guarantors WHERE id = ?", (new_id,)).fetchone()
        return json_success(g), 201
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/guarantors/<int:g_id>", methods=["GET"])
def api_get_guarantor(g_id):
    """Get guarantor detail."""
    db = get_dict_db()
    try:
        # The guarantors list is built from tenant rows, so the ids it links to
        # are tenant ids. Tenant ids start well above the handful of rows in the
        # guarantors table, so there is no ambiguity between the two.
        t = db.execute(
            f"SELECT * FROM tenants t WHERE t.id = ? AND {GUARANTOR_BASE_WHERE}", (g_id,)
        ).fetchone()
        if t:
            status = db.execute(
                f"SELECT {GUARANTOR_STATUS_SQL} AS s FROM tenants t WHERE t.id=?", (g_id,)
            ).fetchone()["s"]
            return json_success({
                "id": t["id"],
                "first_name": t.get("guarantor_first_name") or "",
                "last_name": t.get("guarantor_last_name") or "",
                "email": t.get("guarantor_email") or "",
                "phone": t.get("guarantor_phone") or "",
                "mobile": t.get("guarantor_mobile") or "",
                "date_of_birth": t.get("guarantor_date_of_birth") or "",
                "relationship": t.get("guarantor_relation") or "",
                "address": t.get("guarantor_address") or "",
                "city": t.get("guarantor_city") or "",
                "postcode": t.get("guarantor_postcode") or "",
                "country": t.get("guarantor_country") or "",
                "employment": t.get("guarantor_profession") or "",
                "status": status,
                "linked_tenant_id": t["id"],
                "linked_tenant_name": f"{t.get('first_name','')} {t.get('last_name','')}".strip(),
                "tenancy_id": t.get("tenancy_id"),
                "documents": [],
                "source": "tenant",
            })

        g = db.execute("SELECT * FROM guarantors WHERE id = ?", (g_id,)).fetchone()
        if not g:
            return json_error("Guarantor not found", 404)

        # Also load the linked applicant info
        if g.get("applicant_id"):
            app = db.execute("SELECT id, first_name, last_name, email FROM applicants WHERE id = ?",
                             (g["applicant_id"],)).fetchone()
            g["applicant"] = app

        return json_success(g)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/guarantors/<int:g_id>", methods=["PATCH"])
def api_update_guarantor(g_id):
    """Update guarantor fields."""
    data = request.get_json()
    if not data:
        return json_error("No data provided")

    db = get_dict_db()
    try:
        # Same resolution as the GET: an id off the list is a tenant id, and the
        # editable fields map onto the tenant's guarantor_* columns. Without
        # this the inline edits on the list wrote to whatever row in the
        # guarantors table happened to share the number.
        t = db.execute(
            f"SELECT * FROM tenants t WHERE t.id = ? AND {GUARANTOR_BASE_WHERE}", (g_id,)
        ).fetchone()
        if t:
            field_map = {
                "first_name": "guarantor_first_name", "last_name": "guarantor_last_name",
                "email": "guarantor_email", "phone": "guarantor_phone",
                "mobile": "guarantor_mobile", "relationship": "guarantor_relation",
                "address": "guarantor_address", "city": "guarantor_city",
                "postcode": "guarantor_postcode", "country": "guarantor_country",
                "date_of_birth": "guarantor_date_of_birth",
            }
            sets, vals = [], []
            for key, val in data.items():
                col = field_map.get(key)
                if not col:
                    continue
                sets.append(f"{col} = ?")
                vals.append(val)
            if not sets:
                return json_error("No valid fields to update")
            now = datetime.now(timezone.utc).isoformat()
            sets += ["modified = ?", "sync_dirty = 1", "local_modified = ?", "sync_origin = 'banksia_os'"]
            vals += [now, now, g_id]
            db.execute(f"UPDATE tenants SET {', '.join(sets)} WHERE id = ?", vals)
            db.commit()
            for key, val in data.items():
                if key in field_map:
                    _log_activity("guarantor", g_id, "update", field_changed=key,
                                  old_value=t.get(field_map[key]), new_value=val, db=db)
            db.commit()
            return json_success({"id": g_id, "updated": True})

        g = db.execute("SELECT * FROM guarantors WHERE id = ?", (g_id,)).fetchone()
        if not g:
            return json_error("Guarantor not found", 404)

        real_cols = {r["name"] for r in db.execute("PRAGMA table_info(guarantors)").fetchall()}
        protected_keys = {"id", "applicant_id", "created"}

        set_parts = []
        params = []

        for key, val in data.items():
            if key in protected_keys or key not in real_cols:
                continue
            set_parts.append(f"{key} = ?")
            params.append(val)

        if not set_parts:
            return json_error("No valid fields to update")

        now = datetime.now(timezone.utc).isoformat()
        set_parts.append("modified = ?")
        params.append(now)
        params.append(g_id)

        db.execute(f"UPDATE guarantors SET {', '.join(set_parts)} WHERE id = ?", params)
        db.commit()

        updated = db.execute("SELECT * FROM guarantors WHERE id = ?", (g_id,)).fetchone()
        return json_success(updated)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9D. SINGLE-TRANSACTION TENANCY CREATION
# ═══════════════════════════════════════════════════════════════

@banksia_os_bp.route("/applicants/<int:app_id>/create-tenancy", methods=["POST"])
def api_create_tenancy_from_applicant(app_id):
    """Convert an approved applicant to a tenancy in one atomic transaction.

    Creates tenancy + tenant + deposit, updates unit occupancy,
    logs activity, and sends notifications.
    """
    data = request.get_json() or {}
    override_occupancy = data.get("override_occupancy", False)

    db = get_dict_db()
    try:
        # Begin transaction explicitly
        db.execute("BEGIN IMMEDIATE")

        # 1. Validate applicant
        app = db.execute("SELECT * FROM applicants WHERE id = ?", (app_id,)).fetchone()
        if not app:
            db.execute("ROLLBACK")
            return json_error("Applicant not found", 404)

        app_status = (app.get("status") or "").strip().lower()
        if app_status not in ("approved", "tenancy_created"):
            db.execute("ROLLBACK")
            return json_error(
                f"Applicant status must be 'approved' to create tenancy, got '{app_status}'",
                400
            )

        property_id = app.get("property_id")
        unit_id = app.get("unit_id")
        if not property_id or not unit_id:
            db.execute("ROLLBACK")
            return json_error("Applicant must have property_id and unit_id set", 400)

        # Validate property exists
        prop = db.execute("SELECT id, name FROM properties WHERE id = ?", (property_id,)).fetchone()
        if not prop:
            db.execute("ROLLBACK")
            return json_error(f"Property {property_id} not found", 404)

        # Validate unit exists and belongs to property
        unit = db.execute(
            "SELECT * FROM units WHERE id = ? AND property_id = ?",
            (unit_id, property_id)
        ).fetchone()
        if not unit:
            db.execute("ROLLBACK")
            return json_error(f"Unit {unit_id} not found under property {property_id}", 404)

        # 3. Check unit availability
        if not override_occupancy:
            active_tenancy = db.execute(
                "SELECT id, status, start_date, end_date FROM tenancies "
                "WHERE unit_id = ? AND status IN ('Active', 'active', 'Periodic', 'periodic') "
                "LIMIT 1",
                (unit_id,)
            ).fetchone()
            if active_tenancy:
                db.execute("ROLLBACK")
                return json_error(
                    f"Unit {unit_id} already has an active tenancy (#{active_tenancy['id']}). "
                    "Use override_occupancy=true to force.",
                    409
                )

        now = datetime.now(timezone.utc)
        now_iso = now.isoformat()
        user_name = getattr(request, "current_user", {}).get("username", "system")

        first_name = app.get("first_name", "")
        last_name = app.get("last_name", "")
        email = app.get("email", "")
        phone = app.get("phone", "") or app.get("mobile", "")
        main_tenant_name = f"{first_name} {last_name}".strip()

        # Determine start_date
        start_date = app.get("desired_move_in") or now_iso[:10]
        # end_date = start_date + 6 months
        try:
            from dateutil.relativedelta import relativedelta
            start_dt = datetime.fromisoformat(start_date) if "T" in start_date else datetime.strptime(start_date, "%Y-%m-%d")
            end_dt = start_dt + relativedelta(months=6)
        except ImportError:
            # Fallback: add ~182 days
            from datetime import timedelta
            try:
                start_dt = datetime.fromisoformat(start_date) if "T" in start_date else datetime.strptime(start_date, "%Y-%m-%d")
            except (ValueError, TypeError):
                start_dt = now
            end_dt = start_dt + timedelta(days=182)
        end_date = end_dt.strftime("%Y-%m-%d")
        rent_amount = app.get("proposed_rent")
        deposit_amount = app.get("proposed_deposit")

        # 4. Create the tenancy — or promote the prospective one.
        # Signing the tenant application already created a Prospective tenancy + tenant
        # for this applicant. Converting must move those same records forward, not open
        # a second tenancy on the same room: two tenancies for one let is how a unit ends
        # up double-counted in occupancy, rent and arrears.
        prospective = db.execute(
            "SELECT ta.tenancy_id AS tenancy_id, ta.tenant_id AS tenant_id "
            "FROM tenant_applications ta JOIN tenancies t ON t.id = ta.tenancy_id "
            "WHERE ta.applicant_id = ? AND LOWER(COALESCE(t.status,'')) = 'prospective' "
            "ORDER BY ta.id DESC LIMIT 1",
            (app_id,)
        ).fetchone()

        if prospective and prospective.get("tenancy_id"):
            tenancy_id = prospective["tenancy_id"]
            db.execute(
                "UPDATE tenancies SET property_id = ?, unit_id = ?, main_tenant_name = ?, "
                "status = 'active', start_date = ?, end_date = ?, rent_amount = ?, "
                "rent_frequency = 'pcm', modified = ? WHERE id = ?",
                [property_id, unit_id, main_tenant_name, start_date, end_date,
                 rent_amount, now_iso, tenancy_id]
            )
            _log_activity("tenancy", tenancy_id, "status_change",
                           field_changed="status", old_value="Prospective", new_value="active",
                           notes=f"Prospective tenancy confirmed from applicant #{app_id} ({main_tenant_name})",
                           db=db)
        else:
            tenancy_cur = db.execute(
                "INSERT INTO tenancies (property_id, unit_id, main_tenant_name, status, "
                "start_date, end_date, rent_amount, rent_frequency, created, modified) "
                "VALUES (?, ?, ?, 'active', ?, ?, ?, 'pcm', ?, ?)",
                [property_id, unit_id, main_tenant_name, start_date, end_date,
                 rent_amount, now_iso, now_iso]
            )
            tenancy_id = tenancy_cur.lastrowid

            _log_activity("tenancy", tenancy_id, "created",
                           notes=f"Tenancy created from applicant #{app_id} ({main_tenant_name})",
                           db=db)

        # 5. Create or promote the tenant record
        if prospective and prospective.get("tenant_id"):
            tenant_id = prospective["tenant_id"]
            db.execute(
                "UPDATE tenants SET first_name = ?, last_name = ?, email = ?, phone_home = ?, "
                "mobile = ?, property_id = ?, unit_id = ?, tenancy_id = ?, main_tenant = 1, "
                "status = 'active', modified = ? WHERE id = ?",
                [first_name, last_name, email, phone, phone,
                 property_id, unit_id, tenancy_id, now_iso, tenant_id]
            )
            _log_activity("tenant", tenant_id, "status_change",
                           field_changed="status", old_value="Prospective", new_value="active",
                           notes=f"Prospective tenant {main_tenant_name} confirmed from applicant #{app_id}",
                           db=db)
        else:
            tenant_cur = db.execute(
                "INSERT INTO tenants (first_name, last_name, email, phone_home, mobile, "
                "property_id, unit_id, tenancy_id, main_tenant, status, created, modified) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, 'active', ?, ?)",
                [first_name, last_name, email, phone, phone,
                 property_id, unit_id, tenancy_id, now_iso, now_iso]
            )
            tenant_id = tenant_cur.lastrowid

            _log_activity("tenant", tenant_id, "created",
                           notes=f"Tenant {main_tenant_name} created from applicant #{app_id}",
                           db=db)

        # 6. Deposit record — through the shared step, which is idempotent per
        # tenancy. Signing the application already opened this tenancy and its
        # deposit; converting the applicant must not open a second one.
        _dep = ensure_deposit_for_tenancy(
            db, tenancy_id, amount=deposit_amount, origin="applicant #%s" % app_id)
        deposit_id = _dep["id"] if _dep else None

        # 7. Update applicant status to 'tenancy_created'
        old_app_status = app.get("status", "")
        db.execute("UPDATE applicants SET status = 'tenancy_created', modified = ? WHERE id = ?",
                   (now_iso, app_id))

        _log_activity("applicant", app_id, "status_change",
                       field_changed="status",
                       old_value=old_app_status, new_value="tenancy_created",
                       notes="Status updated to tenancy_created via tenancy creation", db=db)

        # 8. Update referencing status to 'tenancy_created' if referencing exists
        refs = db.execute(
            "SELECT id, status FROM referencing_forms WHERE applicant_id = ? AND status NOT IN ('tenancy_created', 'withdrawn')",
            (app_id,)
        ).fetchall()
        for ref in refs:
            old_ref_status = ref["status"]
            db.execute("UPDATE referencing_forms SET status = 'tenancy_created', modified = ? WHERE id = ?",
                       (now_iso, ref["id"]))
            _log_activity("referencing_form", ref["id"], "status_change",
                           field_changed="status",
                           old_value=old_ref_status, new_value="tenancy_created", db=db)

        # 9. Update unit occupancy
        is_occupied = False
        if start_date and start_date <= now_iso[:10]:
            unit_status_map = {
                "available": "Occupied",
                "Available": "Occupied",
                "Available To Let": "Occupied",
                "Let": "Occupied",
            }
            new_unit_status = unit_status_map.get(unit.get("unit_status", ""), "Occupied")
            db.execute(
                "UPDATE units SET unit_status = ?, unit_vacant = 0, status = ?, modified = ? WHERE id = ?",
                (new_unit_status, "occupied", now_iso, unit_id)
            )
            is_occupied = True

        _log_activity("unit", unit_id, "updated",
                       field_changed="unit_status",
                       old_value=unit.get("unit_status", ""),
                       new_value="Occupied" if is_occupied else unit.get("unit_status", ""),
                       notes="Unit occupied via tenancy creation" if is_occupied else "Unit linked to new tenancy",
                       db=db)

        # 10. Create notifications
        notify_message = (
            f"Tenancy created for {main_tenant_name} — "
            f"tenancy #{tenancy_id}, unit #{unit_id}"
        )
        notify_link = f"/banksia-os?entity=tenancies&id={tenancy_id}"

        # Notify assigned_to
        # dict.get's default only fires when the key is ABSENT; applicants.assigned_to
        # exists and is NULL for anyone unassigned, so the default never applied and
        # .strip() blew up — 500ing every conversion for an applicant with no assignee.
        assigned_to = (app.get("assigned_to") or "").strip()
        notified = set()
        if assigned_to:
            create_notification(assigned_to, notify_message, notify_link)
            notified.add(assigned_to)

        # Notify super_admins
        super_admins = ["Sami", "Roo", "Norbert", "Sadman"]
        for sa in super_admins:
            if sa not in notified:
                create_notification(sa, notify_message, notify_link)
                notified.add(sa)

        db.commit()

        # 11. Fetch and return results using a fresh connection
        db2 = get_dict_db()
        try:
            tenancy = db2.execute("SELECT * FROM tenancies WHERE id = ?", (tenancy_id,)).fetchone()
            tenant = db2.execute("SELECT * FROM tenants WHERE id = ?", (tenant_id,)).fetchone()
            deposit = db2.execute("SELECT * FROM deposits WHERE id = ?", (deposit_id,)).fetchone()
        finally:
            db2.close()

        return json_success({
            "tenancy_id": tenancy_id,
            "tenant_id": tenant_id,
            "deposit_id": deposit_id,
            "tenancy": tenancy,
            "tenant": tenant,
            "deposit": deposit,
        })
    except Exception as e:
        db.execute("ROLLBACK")
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9E. UNIT OCCUPANCY CHECK
# ═══════════════════════════════════════════════════════════════

@banksia_os_bp.route("/units/<int:unit_id>/occupancy", methods=["GET"])
def api_unit_occupancy(unit_id):
    """Check if a unit is available for a new tenancy.

    Returns:
        - is_available: boolean
        - current_tenancy: active/perodic tenancy on this unit, or null
        - next_tenancy: upcoming tenancy (future start_date), or null
        - future_tenancies: all future/recent tenancies
    """
    db = get_dict_db()
    try:
        unit = db.execute("SELECT * FROM units WHERE id = ?", (unit_id,)).fetchone()
        if not unit:
            return json_error("Unit not found", 404)

        now_iso = datetime.now(timezone.utc).isoformat()

        # Current active/periodic tenancy
        current = db.execute(
            "SELECT * FROM tenancies WHERE unit_id = ? AND status IN ('Active', 'active', 'Periodic', 'periodic') "
            "ORDER BY start_date DESC LIMIT 1",
            (unit_id,)
        ).fetchone()

        # Future tenancies (future start_date)
        future = db.execute(
            "SELECT * FROM tenancies WHERE unit_id = ? AND start_date > ? AND status NOT IN ('ended', 'Ended', 'cancelled', 'Cancelled') "
            "ORDER BY start_date ASC",
            (unit_id, now_iso[:10])
        ).fetchall()

        # Next tenancy (closest future)
        next_tenancy = future[0] if future else None

        # Also check tenant count
        tenant_count = db.execute(
            "SELECT COUNT(*) AS cnt FROM tenants WHERE unit_id = ? AND status = 'active'",
            (unit_id,)
        ).fetchone()["cnt"]

        is_available = current is None and tenant_count == 0

        return json_success({
            "is_available": is_available,
            "unit": {
                "id": unit["id"],
                "unit_ref": unit.get("unit_ref"),
                "unit_status": unit.get("unit_status"),
                "unit_vacant": bool(unit.get("unit_vacant")),
            },
            "current_tenancy": current,
            "next_tenancy": next_tenancy,
            "future_tenancies": future,
        })
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ═══════════════════════════════════════════════════════════════
# 9F. GLOBAL SEARCH
# ═══════════════════════════════════════════════════════════════


@banksia_os_bp.route("/search/global", methods=["GET"])
def global_search():
    """
    Global search across all entity types.

    GET /api/banksia-os/search?q=4+Studd&limit=10

    Searches properties, units, tenants, applicants, tenancies,
    guarantors, property_owners, maintenance_jobs, documents, and invoices.
    Returns grouped results with type/id/label/url/match/parent for each hit.
    """
    q_raw = request.args.get("q", "").strip()
    limit_per_type = request.args.get("limit", 5, type=int)
    if limit_per_type < 1:
        limit_per_type = 5
    if limit_per_type > 50:
        limit_per_type = 50

    if not q_raw:
        return jsonify({
            "success": True,
            "data": {
                "query": "",
                "total": 0,
                "results": [],
                "grouped": {},
            }
        })

    like = f"%{q_raw}%"
    db = get_dict_db()

    results = []
    grouped = {}

    def add_result(etype, eid, label, url, match, parent=None):
        entry = {
            "type": etype,
            "id": eid,
            "label": label,
            "url": url,
            "match": match,
        }
        if parent:
            entry["parent"] = parent
        results.append(entry)
        grouped.setdefault(etype, []).append(entry)

    try:
        # ── Properties ──
        rows = db.execute(
            """SELECT id, name, address_line_1, address_line_2, city, postcode, property_ref
               FROM properties
               WHERE name LIKE ? OR address_line_1 LIKE ? OR address_line_2 LIKE ?
                     OR city LIKE ? OR postcode LIKE ? OR property_ref LIKE ?
               LIMIT ?""",
            (like, like, like, like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = r["name"] or r["address_line_1"] or r["property_ref"] or f"Property #{r['id']}"
            if r["address_line_1"] and r["city"]:
                label = f"{label}, {r['city']}"
            match_field = "name"
            if not (r["name"] and q_raw.lower() in r["name"].lower()):
                if r["address_line_1"] and q_raw.lower() in r["address_line_1"].lower():
                    match_field = "address_line_1"
                elif r["address_line_2"] and q_raw.lower() in r["address_line_2"].lower():
                    match_field = "address_line_2"
                elif r["city"] and q_raw.lower() in r["city"].lower():
                    match_field = "city"
                elif r["postcode"] and q_raw.lower() in r["postcode"].lower():
                    match_field = "postcode"
                else:
                    match_field = "property_ref"
            add_result("property", r["id"], label, f"/properties/{r['id']}", match_field)

        # ── Units ──
        rows = db.execute(
            """SELECT u.id, u.unit_ref, u.property_id, p.name AS pname, p.address_line_1 AS paddr
               FROM units u LEFT JOIN properties p ON u.property_id = p.id
               WHERE u.unit_ref LIKE ?
               LIMIT ?""",
            (like, limit_per_type)
        ).fetchall()
        for r in rows:
            parent_label = r["pname"] or r["paddr"] or f"Property #{r['property_id']}" if r["property_id"] else None
            parent = None
            if r["property_id"]:
                parent = {"type": "property", "id": r["property_id"], "label": parent_label}
            add_result("unit", r["id"], r["unit_ref"] or f"Unit #{r['id']}",
                       f"/units/{r['id']}", "unit_ref", parent)

        # ── Tenants ──
        rows = db.execute(
            """SELECT t.id, t.first_name, t.last_name, t.email, t.phone_home, t.mobile,
                      t.property_id, p.name AS pname, p.address_line_1 AS paddr
               FROM tenants t LEFT JOIN properties p ON t.property_id = p.id
               WHERE t.first_name LIKE ? OR t.last_name LIKE ? OR t.email LIKE ?
                     OR t.phone_home LIKE ? OR t.mobile LIKE ?
               LIMIT ?""",
            (like, like, like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = f"{r['first_name'] or ''} {r['last_name'] or ''}".strip() or f"Tenant #{r['id']}"
            match_field = "first_name"
            if not (r["first_name"] and q_raw.lower() in r["first_name"].lower()):
                if r["last_name"] and q_raw.lower() in r["last_name"].lower():
                    match_field = "last_name"
                elif r["email"] and q_raw.lower() in r["email"].lower():
                    match_field = "email"
                elif r["phone_home"] and q_raw.lower() in r["phone_home"].lower():
                    match_field = "phone_home"
                else:
                    match_field = "mobile"
            parent_label = r["pname"] or r["paddr"] or f"Property #{r['property_id']}" if r["property_id"] else None
            parent = None
            if r["property_id"]:
                parent = {"type": "property", "id": r["property_id"], "label": parent_label}
            add_result("tenant", r["id"], label, f"/tenants/{r['id']}", match_field, parent)

        # ── Applicants ──
        rows = db.execute(
            """SELECT id, first_name, last_name, email
               FROM applicants
               WHERE first_name LIKE ? OR last_name LIKE ? OR email LIKE ?
               LIMIT ?""",
            (like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = f"{r['first_name'] or ''} {r['last_name'] or ''}".strip() or f"Applicant #{r['id']}"
            match_field = "first_name"
            if not (r["first_name"] and q_raw.lower() in r["first_name"].lower()):
                if r["last_name"] and q_raw.lower() in r["last_name"].lower():
                    match_field = "last_name"
                else:
                    match_field = "email"
            add_result("applicant", r["id"], label, f"/applicants/{r['id']}", match_field)

        # ── Tenancies ──
        rows = db.execute(
            """SELECT tn.id, tn.ref, tn.main_tenant_name, tn.property_id,
                      p.name AS pname, p.address_line_1 AS paddr
               FROM tenancies tn LEFT JOIN properties p ON tn.property_id = p.id
               WHERE tn.main_tenant_name LIKE ? OR tn.ref LIKE ?
               LIMIT ?""",
            (like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = r["main_tenant_name"] or r["ref"] or f"Tenancy #{r['id']}"
            match_field = "main_tenant_name"
            if not (r["main_tenant_name"] and q_raw.lower() in r["main_tenant_name"].lower()):
                match_field = "ref"
            parent_label = r["pname"] or r["paddr"] or f"Property #{r['property_id']}" if r["property_id"] else None
            parent = None
            if r["property_id"]:
                parent = {"type": "property", "id": r["property_id"], "label": parent_label}
            add_result("tenancy", r["id"], label, f"/tenancies/{r['id']}", match_field, parent)

        # ── Guarantors ──
        rows = db.execute(
            """SELECT g.id, g.first_name, g.last_name, g.email, g.applicant_id,
                      a.first_name AS afn, a.last_name AS aln
               FROM guarantors g LEFT JOIN applicants a ON g.applicant_id = a.id
               WHERE g.first_name LIKE ? OR g.last_name LIKE ? OR g.email LIKE ?
               LIMIT ?""",
            (like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = f"{r['first_name'] or ''} {r['last_name'] or ''}".strip() or f"Guarantor #{r['id']}"
            match_field = "first_name"
            if not (r["first_name"] and q_raw.lower() in r["first_name"].lower()):
                if r["last_name"] and q_raw.lower() in r["last_name"].lower():
                    match_field = "last_name"
                else:
                    match_field = "email"
            parent_label = f"{r['afn'] or ''} {r['aln'] or ''}".strip() if r["applicant_id"] else None
            parent = None
            if r["applicant_id"]:
                parent = {"type": "applicant", "id": r["applicant_id"], "label": parent_label or f"Applicant #{r['applicant_id']}"}
            add_result("guarantor", r["id"], label, f"/guarantors/{r['id']}", match_field, parent)

        # ── Landlords ──
        rows = db.execute(
            """SELECT id, name, company_name, contact_email
               FROM property_owners
               WHERE name LIKE ? OR company_name LIKE ? OR contact_email LIKE ?
               LIMIT ?""",
            (like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = r["name"] or r["company_name"] or f"Owner #{r['id']}"
            match_field = "name"
            if not (r["name"] and q_raw.lower() in r["name"].lower()):
                if r["company_name"] and q_raw.lower() in r["company_name"].lower():
                    match_field = "company_name"
                else:
                    match_field = "contact_email"
            add_result("property_owner", r["id"], label, f"/property-owners/{r['id']}", match_field)

        # ── Maintenance Jobs ──
        rows = db.execute(
            """SELECT mj.id, mj.title, mj.reference, mj.address, mj.property_id,
                      p.name AS pname, p.address_line_1 AS paddr
               FROM maintenance_jobs mj LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.title LIKE ? OR mj.reference LIKE ? OR mj.address LIKE ?
               LIMIT ?""",
            (like, like, like, limit_per_type)
        ).fetchall()
        for r in rows:
            label = r["title"] or r["reference"] or f"Maintenance #{r['id']}"
            match_field = "title"
            if not (r["title"] and q_raw.lower() in r["title"].lower()):
                if r["reference"] and q_raw.lower() in r["reference"].lower():
                    match_field = "reference"
                else:
                    match_field = "address"
            parent_label = r["pname"] or r["paddr"] or f"Property #{r['property_id']}" if r["property_id"] else None
            parent = None
            if r["property_id"]:
                parent = {"type": "property", "id": r["property_id"], "label": parent_label}
            add_result("maintenance", r["id"], label, f"/maintenance/{r['id']}", match_field, parent)

        # ── Documents ──
        try:
            rows = db.execute(
                """SELECT id, filename, category AS name, related_to, related_id
                   FROM documents
                   WHERE (filename IS NOT NULL AND filename LIKE ?)
                      OR (category IS NOT NULL AND category LIKE ?)
                   LIMIT ?""",
                (like, like, limit_per_type)
            ).fetchall()
            for r in rows:
                label = r["name"] or r["filename"] or f"Document #{r['id']}"
                match_field = "name" if (r["name"] and q_raw.lower() in r["name"].lower()) else "filename"
                parent = None
                if r["related_to"] and r["related_id"]:
                    parent = {"type": r["related_to"], "id": r["related_id"], "label": f"{r['related_to'].title()} #{r['related_id']}"}
                add_result("document", r["id"], label, f"/documents/{r['id']}", match_field, parent)
        except Exception:
            pass  # documents table might not have these columns

        # ── Invoices ──
        try:
            rows = db.execute(
                """SELECT id, invoice_ref, description, tenancy_id
                   FROM invoices
                   WHERE invoice_ref LIKE ? OR description LIKE ?
                   LIMIT ?""",
                (like, like, limit_per_type)
            ).fetchall()
            for r in rows:
                label = r["invoice_ref"] or r["description"] or f"Invoice #{r['id']}"
                match_field = "invoice_ref" if (r["invoice_ref"] and q_raw.lower() in r["invoice_ref"].lower()) else "description"
                parent = None
                if r["tenancy_id"]:
                    parent = {"type": "tenancy", "id": r["tenancy_id"], "label": f"Tenancy #{r['tenancy_id']}"}
                add_result("invoice", r["id"], label, f"/invoices/{r['id']}", match_field, parent)
        except Exception:
            pass  # invoices table might not exist

        total = len(results)

        return jsonify({
            "success": True,
            "data": {
                "query": q_raw,
                "total": total,
                "results": results,
                "grouped": grouped,
            }
        })

    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

# ═══════════════════════════════════════════════
# 2i. UNIVERSAL TIMELINE — activity for entity + all related sub-entities
# ═══════════════════════════════════════════════

TIMELINE_TYPE_MAP = {
    ("property", "update"): ("property_updated", "edit"),
    ("property", "created"): ("property_created", "plus"),
    ("property", "archived"): ("property_archived", "archive"),
    ("property", "restored"): ("property_restored", "rotate-ccw"),
    ("property", "deleted"): ("property_deleted", "trash-2"),
    ("unit", "created"): ("unit_added", "door-open"),
    ("unit", "linked"): ("unit_added", "door-open"),
    ("unit", "update"): ("unit_updated", "edit"),
    ("unit", "updated"): ("unit_updated", "edit"),
    ("unit", "archived"): ("unit_archived", "archive"),
    ("unit", "deleted"): ("unit_deleted", "trash-2"),
    ("tenancy", "created"): ("tenancy_created", "file-text"),
    ("tenancy", "update"): ("tenancy_updated", "edit"),
    ("tenant", "created"): ("tenant_created", "user-plus"),
    ("tenant", "update"): ("tenant_updated", "edit"),
    ("applicant", "created"): ("applicant_created", "user-plus"),
    ("applicant", "update"): ("applicant_updated", "edit"),
    ("applicant", "status_change"): ("applicant_status", "arrow-right"),
    ("deposit", "created"): ("deposit_received", "shield"),
    ("deposit", "update"): ("deposit_updated", "edit"),
    ("referencing_form", "created"): ("referencing_created", "clipboard"),
    ("referencing_form", "status_change"): ("referencing_updated", "clipboard"),
    ("guarantor", "created"): ("guarantor_added", "user-plus"),
    ("maintenance_job", "created"): ("maintenance_raised", "wrench"),
    ("maintenance_job", "update"): ("maintenance_updated", "wrench"),
}


def _get_entity_label(db, entity_type, entity_id):
    """Look up a human-readable label for an entity."""
    try:
        if entity_type == "property":
            row = db.execute(
                "SELECT COALESCE(NULLIF(name,''), NULLIF(ref,''), 'Property #'||CAST(id AS TEXT)) AS label FROM properties WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Property #{entity_id}"
        elif entity_type == "unit":
            row = db.execute(
                "SELECT COALESCE(NULLIF(unit_ref,''), 'Unit #'||CAST(id AS TEXT)) AS label FROM units WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Unit #{entity_id}"
        elif entity_type == "tenancy":
            row = db.execute(
                "SELECT COALESCE(NULLIF(ref,''), 'Tenancy #'||CAST(id AS TEXT)) AS label FROM tenancies WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Tenancy #{entity_id}"
        elif entity_type == "tenant":
            row = db.execute(
                "SELECT COALESCE(NULLIF(first_name||' '||last_name, ' '), 'Tenant #'||CAST(id AS TEXT)) AS label FROM tenants WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Tenant #{entity_id}"
        elif entity_type == "applicant":
            row = db.execute(
                "SELECT COALESCE(NULLIF(first_name||' '||last_name, ' '), 'Applicant #'||CAST(id AS TEXT)) AS label FROM applicants WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Applicant #{entity_id}"
        elif entity_type == "maintenance_job":
            row = db.execute(
                "SELECT COALESCE(NULLIF(title,''), NULLIF(reference,''), 'Job #'||CAST(id AS TEXT)) AS label FROM maintenance_jobs WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Job #{entity_id}"
        elif entity_type == "deposit":
            row = db.execute(
                "SELECT 'Deposit #'||CAST(id AS TEXT) AS label FROM deposits WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Deposit #{entity_id}"
        elif entity_type == "referencing_form":
            row = db.execute(
                "SELECT 'Referencing #'||CAST(id AS TEXT) AS label FROM referencing_forms WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Referencing #{entity_id}"
        elif entity_type == "guarantor":
            row = db.execute(
                "SELECT COALESCE(NULLIF(first_name||' '||last_name, ' '), 'Guarantor #'||CAST(id AS TEXT)) AS label FROM guarantors WHERE id = ?",
                (entity_id,)
            ).fetchone()
            return row["label"] if row else f"Guarantor #{entity_id}"
    except Exception as _e:
        current_app.logger.error(f"Error in line ~9613: {_e}")
        pass
    return f"{entity_type.title()} #{entity_id}"


def _derive_timeline_type(action, entity_type, field_changed):
    """Derive type and icon from action + entity_type + field_changed."""
    key = (entity_type, action)
    if key in TIMELINE_TYPE_MAP:
        return TIMELINE_TYPE_MAP[key]
    return (f"{entity_type}_{action}", "circle")


def _redact_sensitive_fields(item_dict):
    """Replace sensitive field values with '[REDACTED]' in-place.
    
    Uses the SENSITIVE_FIELDS set defined near _log_activity.
    """
    for field in SENSITIVE_FIELDS:
        if field in item_dict and item_dict[field] is not None:
            item_dict[field] = "[REDACTED]"
    return item_dict


def _enhance_timeline_item(item):
    """Add derived type, icon, and entity_label to an activity log row."""
    row = dict(item)
    ttype, icon = _derive_timeline_type(
        row.get("action", ""),
        row.get("entity_type", ""),
        row.get("field_changed", "")
    )
    row["type"] = ttype
    row["icon"] = icon
    _redact_sensitive_fields(row)
    return row


@banksia_os_bp.route("/timeline", methods=["GET"])
def api_universal_timeline():
    """Return paginated activity log for an entity + all related sub-entities."""
    entity_type = request.args.get("entity_type", "").strip().lower()
    entity_id_str = request.args.get("entity_id", "").strip()
    page = int_param(request.args.get("page"), default=1)
    per_page = int_param(request.args.get("per_page"), default=20, max_val=MAX_PAGE_SIZE)

    if not entity_type or not entity_id_str:
        return json_error("entity_type and entity_id are required", 400)

    try:
        entity_id = int(entity_id_str)
    except ValueError:
        return json_error("entity_id must be an integer", 400)

    db = get_dict_db()
    try:
        union_parts = []
        params_list = []

        # 1. Direct entity activity
        union_parts.append(
            "SELECT *, 0 AS sort_order FROM activity_log WHERE entity_type = ? AND entity_id = ?"
        )
        params_list.append([entity_type, entity_id])

        # 2. For property, expand to sub-entities
        if entity_type == "property":
            union_parts.append(
                "SELECT al.*, 1 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'unit' AND al.entity_id IN (SELECT id FROM units WHERE property_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 2 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'tenancy' AND al.entity_id IN (SELECT id FROM tenancies WHERE property_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 3 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'tenant' AND al.entity_id IN (SELECT id FROM tenants WHERE property_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 4 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'applicant' AND al.entity_id IN (SELECT id FROM applicants WHERE property_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 5 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'maintenance_job' AND al.entity_id IN (SELECT id FROM maintenance_jobs WHERE property_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 6 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'deposit' AND al.entity_id IN (SELECT id FROM deposits WHERE property_id = ?)"
            )
            params_list.append([entity_id])

        # 3. For tenancy, include deposits and tenants
        elif entity_type == "tenancy":
            union_parts.append(
                "SELECT al.*, 1 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'deposit' AND al.entity_id IN (SELECT id FROM deposits WHERE tenancy_id = ?)"
            )
            params_list.append([entity_id])

            union_parts.append(
                "SELECT al.*, 2 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'tenant' AND al.entity_id IN (SELECT id FROM tenants WHERE tenancy_id = ?)"
            )
            params_list.append([entity_id])

        # 4. For tenant, include related tenancy
        elif entity_type == "tenant":
            union_parts.append(
                "SELECT al.*, 1 AS sort_order FROM activity_log al "
                "WHERE al.entity_type = 'tenancy' AND al.entity_id IN (SELECT tenancy_id FROM tenants WHERE id = ?) AND tenancy_id IS NOT NULL"
            )
            params_list.append([entity_id])

        full_query_parts = []
        all_params = []
        for sql, params in zip(union_parts, params_list):
            full_query_parts.append(sql)
            all_params.extend(params)

        combined_sql = " UNION ALL ".join(full_query_parts)

        count_sql = f"SELECT COUNT(*) AS cnt FROM ({combined_sql})"
        total = db.execute(count_sql, all_params).fetchone()["cnt"]

        offset = (page - 1) * per_page
        data_sql = f"""
            SELECT * FROM ({combined_sql}) AS combined
            ORDER BY created DESC, sort_order ASC
            LIMIT ? OFFSET ?
        """
        params_with_pagination = all_params + [per_page, offset]
        rows = db.execute(data_sql, params_with_pagination).fetchall()

        items = []
        for row in rows:
            item = _enhance_timeline_item(row)
            item["entity_label"] = _get_entity_label(db, row["entity_type"], row["entity_id"])
            items.append(item)

        return jsonify({
            "success": True,
            "data": {
                "items": items,
                "total": total,
                "page": page,
                "per_page": per_page,
            }
        })

    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ── Monday.com Comment Import & 2-Way Sync ──

def _ensure_monday_update_id_column():
    """Add monday_update_id column to comments table if it doesn't exist."""
    db = get_dict_db()
    try:
        rows = db.execute("PRAGMA table_info(comments)").fetchall()
        cols = [r["name"] for r in rows]
        if "monday_update_id" not in cols:
            db.execute("ALTER TABLE comments ADD COLUMN monday_update_id TEXT")
            db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/import-monday-comments", methods=["POST"])
def api_import_monday_comments():
    """Import ALL Monday.com Updates as comments for all jobs with a monday_id."""
    _ensure_monday_update_id_column()

    mtok = None
    try:
        mtok = open("/root/.hermes/secrets/monday_token.txt").read().strip()
    except Exception:
        pass
    if not mtok:
        return json_error("Monday token not found")

    db = get_dict_db()
    try:
        jobs = db.execute(
            "SELECT id, monday_id FROM maintenance_jobs WHERE monday_id IS NOT NULL AND monday_id != ''"
        ).fetchall()

        if not jobs:
            return json_success({"imported": 0, "total_jobs": 0, "message": "No jobs with monday_id found"})

        users_data = _load_comment_users()
        imported = 0
        total_updates = 0

        for job in jobs:
            job_id = job["id"]
            monday_id = job["monday_id"]

            q = ('{ items(ids: [%s]) { updates(limit: 200) {'
                 ' id body text_body created_at updated_at'
                 ' creator { id name }'
                 ' assets { id name url file_extension }'
                 ' } } }') % monday_id

            try:
                data = _monday_graphql(mtok, q)
            except Exception:
                continue

            items = data.get("data", {}).get("items", [])
            if not items:
                continue

            updates = items[0].get("updates", []) if items else []
            total_updates += len(updates)

            for upd in updates:
                update_id = str(upd["id"])
                existing = db.execute(
                    "SELECT id FROM comments WHERE monday_update_id = ?",
                    [update_id]
                ).fetchone()
                if existing:
                    continue

                creator = upd.get("creator", {})
                monday_author_name = creator.get("name", "Monday User")
                author_id = monday_author_name
                for uname, uinfo in users_data.items():
                    display = uinfo.get("display_name", "") or uname
                    if display == monday_author_name or uname == monday_author_name:
                        author_id = uname
                        break

                text_body = upd.get("text_body") or upd.get("body") or ""
                created_at = upd.get("created_at", "")
                updated_at = upd.get("updated_at", "")
                is_edited = 1 if (updated_at and created_at and updated_at != created_at) else 0

                assets = upd.get("assets", [])
                media_paths = []
                if assets:
                    for asset in assets:
                        asset_url = asset.get("url", "")
                        if asset_url:
                            media_paths.append(asset_url)

                db.execute(
                    "INSERT INTO comments "
                    "(entity_type, entity_id, author, author_id, body, media_paths, created, modified, is_edited, monday_update_id) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    ("maintenance_jobs", job_id, monday_author_name, author_id,
                     text_body, json.dumps(media_paths), created_at,
                     updated_at or created_at, is_edited, update_id)
                )
                imported += 1

        db.commit()
        return json_success({
            "imported": imported,
            "total_jobs": len(jobs),
            "total_updates_found": total_updates,
            "message": f"Imported {imported} new comments from Monday",
        })

    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


@banksia_os_bp.route("/maintenance/monday-sync", methods=["POST"])
def api_monday_sync():
    """Full 2-way sync between Monday.com and Banksia OS maintenance jobs."""
    mtok = None
    try:
        mtok = open("/root/.hermes/secrets/monday_token.txt").read().strip()
    except Exception:
        pass
    if not mtok:
        return json_error("Monday token not found")

    _ensure_monday_update_id_column()

    db = get_dict_db()
    try:
        users_data = _load_comment_users()
        results = {
            "monday_to_local": {"jobs_synced": 0, "comments_imported": 0},
            "local_to_monday": {"jobs_pushed": 0, "comments_pushed": 0, "errors": []},
        }

        # ── Direction A: Monday to Banksia OS ──
        all_items = []
        cursor = None
        while True:
            page_ql = f"items_page(limit:200" + (f',cursor:"{cursor}"' if cursor else "") + ")"
            q = ("{ boards(ids: [18401159622]) { id name "
                 + page_ql
                 + """ { cursor items {
                        id name column_values { id text value }
                    } } } }""")
            data = _monday_graphql(mtok, q)
            page_data = (
                data.get("data", {})
                .get("boards", [{}])[0]
                .get("items_page", {})
            )
            items = page_data.get("items", [])
            cursor = page_data.get("cursor")
            all_items.extend(items)
            if not cursor or len(items) < 200:
                break

        jobs_synced = 0
        for item in all_items:
            cols = _parse_monday_cols(item.get("column_values", []))
            monday_id = item["id"]
            title = item.get("name", "")

            status = _safe_status(cols.get("status", "PENDING"))
            priority = _safe_priority(cols.get("color_mm0p8qna", "Medium"))
            maint_type = cols.get("color_mm0vfxmq", "")
            address = (cols.get("short_text041ydfbp", "")
                       or cols.get("long_text_mm50g0j6", "")
                       or cols.get("board_relation_mm0p7cv6", ""))
            contractor = cols.get("color_mm0p4947", "")
            location = cols.get("dropdown_mm0p6nzm", "")
            labour_raw = cols.get("numeric_mm0pndmj", "") or "0"
            materials_raw = cols.get("numeric_mm0p7jdn", "") or "0"
            try:
                labour_cost = float(labour_raw.replace("\u00a3", "").replace(",", "").strip())
            except (ValueError, AttributeError):
                labour_cost = 0.0
            try:
                materials_cost = float(materials_raw.replace("\u00a3", "").replace(",", "").strip())
            except (ValueError, AttributeError):
                materials_cost = 0.0
            bill_ll = 1 if cols.get("boolean_mm0phkaq", "") == "checked" else 0
            emergency = 1 if cols.get("boolean2hbqq7ey", "") == "checked" else 0
            reporter_name = cols.get("short_textcvckh2h3", "")
            reporter_email = cols.get("emailzit7svgb", "")
            photo_paths = _parse_photo_paths(cols)
            invoice_paths = _parse_invoice_paths(cols)

            existing = db.execute(
                "SELECT id, status, priority, type, address, contractor, "
                "labour_cost, materials_cost, bill_ll, emergency, "
                "reporter_name, reporter_email, photo_paths, invoice_paths, "
                "location, description, team_notes "
                "FROM maintenance_jobs WHERE monday_id = ?",
                [monday_id],
            ).fetchone()

            if existing:
                changed = False
                updates = {}
                compare_map = {
                    "title": title,
                    "status": status,
                    "priority": priority,
                    "type": maint_type,
                    "address": address,
                    "contractor": contractor,
                    "location": location,
                    "labour_cost": labour_cost,
                    "materials_cost": materials_cost,
                    "bill_ll": bill_ll,
                    "emergency": emergency,
                    "reporter_name": reporter_name,
                    "reporter_email": reporter_email,
                    "photo_paths": photo_paths,
                    "invoice_paths": invoice_paths,
                }
                for field, new_val in compare_map.items():
                    old_val = existing[field]
                    if old_val is None:
                        old_val = ""
                    if isinstance(old_val, float) or isinstance(new_val, float):
                        if abs(float(old_val or 0) - float(new_val or 0)) > 0.001:
                            updates[field] = new_val
                            changed = True
                    elif str(old_val).strip() != str(new_val).strip():
                        updates[field] = new_val
                        changed = True

                if changed:
                    updates["modified"] = "datetime('now')"
                    set_clause = ", ".join(f"{k} = ?" for k in updates)
                    values = list(updates.values())
                    values.append(existing["id"])
                    db.execute(
                        f"UPDATE maintenance_jobs SET {set_clause} WHERE id = ?",
                        values,
                    )
                    jobs_synced += 1
            else:
                db.execute(
                    "INSERT INTO maintenance_jobs "
                    "(monday_id, title, status, priority, type, address, contractor, location, "
                    "labour_cost, materials_cost, bill_ll, emergency, "
                    "reporter_name, reporter_email, photo_paths, invoice_paths) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (monday_id, title, status, priority, maint_type,
                     address, contractor, location, labour_cost,
                     materials_cost, bill_ll, emergency,
                     reporter_name, reporter_email, photo_paths, invoice_paths),
                )
                jobs_synced += 1

        results["monday_to_local"]["jobs_synced"] = jobs_synced

        # Import new Monday Updates as comments
        jobs_with_monday = db.execute(
            "SELECT id, monday_id FROM maintenance_jobs WHERE monday_id IS NOT NULL AND monday_id != ''"
        ).fetchall()

        comments_imported = 0
        for job in jobs_with_monday:
            job_id = job["id"]
            monday_id = job["monday_id"]

            q = ('{ items(ids: [%s]) { updates(limit: 200) {'
                 ' id body text_body created_at updated_at'
                 ' creator { id name }'
                 ' assets { id name url file_extension }'
                 ' } } }') % monday_id

            try:
                data = _monday_graphql(mtok, q)
            except Exception:
                continue

            items = data.get("data", {}).get("items", [])
            if not items:
                continue
            updates = items[0].get("updates", []) if items else []

            for upd in updates:
                update_id = str(upd["id"])
                existing = db.execute(
                    "SELECT id FROM comments WHERE monday_update_id = ?",
                    [update_id]
                ).fetchone()
                if existing:
                    continue

                creator = upd.get("creator", {})
                monday_author_name = creator.get("name", "Monday User")
                author_id = monday_author_name
                for uname, uinfo in users_data.items():
                    display = uinfo.get("display_name", "") or uname
                    if display == monday_author_name or uname == monday_author_name:
                        author_id = uname
                        break

                text_body = upd.get("text_body") or upd.get("body") or ""
                created_at = upd.get("created_at", "")
                updated_at = upd.get("updated_at", "")
                is_edited = 1 if (updated_at and created_at and updated_at != created_at) else 0

                assets = upd.get("assets", [])
                media_paths = []
                if assets:
                    for asset in assets:
                        asset_url = asset.get("url", "")
                        if asset_url:
                            media_paths.append(asset_url)

                db.execute(
                    "INSERT INTO comments "
                    "(entity_type, entity_id, author, author_id, body, media_paths, created, modified, is_edited, monday_update_id) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    ("maintenance_jobs", job_id, monday_author_name, author_id,
                     text_body, json.dumps(media_paths), created_at,
                     updated_at or created_at, is_edited, update_id)
                )
                comments_imported += 1

        results["monday_to_local"]["comments_imported"] = comments_imported
        db.commit()

        # ── Direction B: Banksia OS to Monday ──
        # Push pending field changes
        try:
            from monday_push import push_all_pending
            push_result = push_all_pending(db)
            results["local_to_monday"]["jobs_pushed"] = push_result.get("pushed", 0)
            if push_result.get("errors"):
                results["local_to_monday"]["errors"].extend(push_result["errors"])
        except Exception as e:
            results["local_to_monday"]["errors"].append(f"push_all_pending error: {e}")

        # Push local comments without monday_update_id as new Monday Updates
        local_comments = db.execute(
            "SELECT c.id, c.entity_id, c.author, c.author_id, c.body, c.media_paths "
            "FROM comments c "
            "JOIN maintenance_jobs j ON j.id = c.entity_id "
            "WHERE c.entity_type = 'maintenance_jobs' "
            "AND (c.monday_update_id IS NULL OR c.monday_update_id = '') "
            "AND j.monday_id IS NOT NULL AND j.monday_id != ''"
        ).fetchall()

        comments_pushed = 0
        for comment in local_comments:
            job = db.execute(
                "SELECT monday_id FROM maintenance_jobs WHERE id = ?",
                [comment["entity_id"]]
            ).fetchone()
            if not job:
                continue

            monday_item_id = job["monday_id"]
            body_text = (comment.get("body") or "").strip()
            if not body_text:
                continue

            escaped_body = body_text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n")

            mutation = ('mutation { create_update('
                        f'item_id: "{monday_item_id}", '
                        f'body: "{escaped_body}"'
                        ') { id } }')

            try:
                resp = _monday_graphql(mtok, mutation)
                if resp.get("data") and resp["data"].get("create_update"):
                    monday_update_id = str(resp["data"]["create_update"]["id"])
                    db.execute(
                        "UPDATE comments SET monday_update_id = ? WHERE id = ?",
                        [monday_update_id, comment["id"]]
                    )
                    comments_pushed += 1
                else:
                    errors = resp.get("errors", [])
                    err_msg = str(errors) if errors else "unknown error"
                    results["local_to_monday"]["errors"].append(
                        f"comment {comment['id']}: {err_msg}"
                    )
            except Exception as e:
                results["local_to_monday"]["errors"].append(
                    f"comment {comment['id']}: {e}"
                )

        results["local_to_monday"]["comments_pushed"] = comments_pushed
        db.commit()

        return json_success(results)

    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ══════════════════════════════════════════════════════════════════════════════
# ZOLT WORK ORDERS — super-admin-only board
#
# Mirrors the Monday "Zolt" workspace (9234995), board "Workorder List"
# (9250109545). Each work order carries a cost price and a bill-to-landlord
# flag; the cost/revenue lines underneath are what make it commercially
# sensitive, which is why the whole board is restricted to super_admin rather
# than hidden in the UI only. Every route below re-checks the role, so the API
# refuses even if the nav is bypassed.
# ══════════════════════════════════════════════════════════════════════════════

ZOLT_STATUSES = [
    "Pending Job", "Job Completed", "Invoice Raised",
    "Contractor Paid", "Client Paid", "On Hold",
]


def _zolt_ensure_tables():
    """Created on demand so the board works on a database that predates it."""
    db = get_dict_db()
    db.execute(
        "CREATE TABLE IF NOT EXISTS zolt_workorders ("
        "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
        "  ref TEXT DEFAULT '',"
        "  title TEXT NOT NULL,"
        "  period TEXT DEFAULT '',"
        "  period_order INTEGER DEFAULT 0,"
        "  position INTEGER DEFAULT 0,"
        "  issue_raised TEXT DEFAULT '',"
        "  invoice_date TEXT DEFAULT '',"
        "  status TEXT DEFAULT 'Pending Job',"
        "  bill_ll INTEGER DEFAULT 0,"
        "  contractor TEXT DEFAULT '',"
        "  cost_price REAL DEFAULT 0,"
        "  invoice_url TEXT DEFAULT '',"
        "  invoice_name TEXT DEFAULT '',"
        "  notes TEXT DEFAULT '',"
        "  monday_id TEXT DEFAULT '',"
        "  created TEXT DEFAULT (datetime('now')),"
        "  updated TEXT DEFAULT (datetime('now'))"
        ")"
    )
    db.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_zolt_monday ON zolt_workorders(monday_id)"
        " WHERE monday_id != ''"
    )
    db.execute(
        "CREATE TABLE IF NOT EXISTS zolt_workorder_lines ("
        "  id INTEGER PRIMARY KEY AUTOINCREMENT,"
        "  workorder_id INTEGER NOT NULL,"
        "  name TEXT DEFAULT '',"
        "  cost REAL DEFAULT 0,"
        "  revenue REAL DEFAULT 0,"
        "  position INTEGER DEFAULT 0,"
        "  monday_id TEXT DEFAULT '',"
        "  created TEXT DEFAULT (datetime('now'))"
        ")"
    )
    db.execute(
        "CREATE INDEX IF NOT EXISTS idx_zolt_lines_wo ON zolt_workorder_lines(workorder_id)"
    )
    db.commit()


def _zolt_guard():
    """Refuse anyone who is not a super admin. Returns a response, or None to continue.

    The board exposes per-job margin, so this is a hard gate rather than a UI
    nicety -- an admin hitting the URL directly still gets 403.
    """
    user = getattr(request, "current_user", None) or session.get("user") or {}
    if (user.get("role") or "").lower() != "super_admin":
        return json_error("This board is restricted to super admins", 403)
    return None


def _zolt_actor():
    user = getattr(request, "current_user", None) or session.get("user") or {}
    return user.get("username") or user.get("name") or "system"


def _zolt_log(entity_id, action, field=None, old=None, new=None):
    """Field-level history for this board.

    Deliberately not services.activity_service.log_activity: that helper inserts
    into activity_log.created_at, a column this database does not have, so every
    call it makes fails silently. Raised separately -- it is app-wide, not ours
    to change from one board.
    """
    try:
        db = get_dict_db()
        db.execute(
            "INSERT INTO activity_log (entity_type, entity_id, action, field_changed,"
            " old_value, new_value, user_name) VALUES (?, ?, ?, ?, ?, ?, ?)",
            ("zolt_workorder", entity_id, action, field,
             None if old is None else str(old)[:200],
             None if new is None else str(new)[:200], _zolt_actor()),
        )
        db.commit()
    except Exception:
        pass


def _zolt_lines_for(db, ids):
    """All lines for the given work orders, grouped by work order id."""
    if not ids:
        return {}
    marks = ",".join("?" * len(ids))
    rows = db.execute(
        "SELECT id, workorder_id, name, cost, revenue FROM zolt_workorder_lines"
        f" WHERE workorder_id IN ({marks}) ORDER BY position, id",
        list(ids),
    ).fetchall()
    out = {}
    for r in rows:
        out.setdefault(r["workorder_id"], []).append(dict(r))
    return out


def _zolt_shape(row, lines):
    """One board row: the work order plus its money, with profit derived.

    Revenue only exists on the cost/revenue lines. A job with no lines has not
    been priced yet, so revenue and profit are left null rather than reported as
    zero -- charging nothing and not having said what we charge are different
    things, and averaging them together would understate every total.
    """
    cost_lines = sum(float(l.get("cost") or 0) for l in lines)
    revenue = sum(float(l.get("revenue") or 0) for l in lines)
    cost_price = float(row.get("cost_price") or 0)
    item = dict(row)
    item["bill_ll"] = bool(row.get("bill_ll"))
    item["lines"] = lines
    item["cost_price"] = cost_price
    item["line_cost"] = cost_lines
    item["cost"] = cost_lines or cost_price
    item["priced"] = bool(lines)
    item["revenue"] = round(revenue, 2) if lines else None
    item["profit"] = round(revenue - cost_lines, 2) if lines else None
    item["has_invoice"] = bool(row.get("invoice_url"))
    return item


@banksia_os_bp.route("/zolt/workorders", methods=["GET"])
def api_zolt_workorders():
    """The board, grouped exactly as Monday grouped it (newest month first)."""
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    db = get_dict_db()
    try:
        rows = [dict(r) for r in db.execute(
            "SELECT * FROM zolt_workorders ORDER BY period_order, position, id"
        ).fetchall()]
        lines = _zolt_lines_for(db, [r["id"] for r in rows])
    finally:
        db.close()

    groups, index = [], {}
    for r in rows:
        item = _zolt_shape(r, lines.get(r["id"], []))
        key = r.get("period") or "Ungrouped"
        if key not in index:
            index[key] = {
                "period": key,
                "period_order": r.get("period_order") or 0,
                "items": [],
                "cost": 0.0, "revenue": 0.0, "profit": 0.0,
                "priced": 0, "unpriced": 0,
            }
            groups.append(index[key])
        g_ = index[key]
        g_["items"].append(item)
        g_["cost"] += item["cost"]
        if item["priced"]:
            g_["revenue"] += item["revenue"]
            g_["profit"] += item["profit"]
            g_["priced"] += 1
        else:
            g_["unpriced"] += 1

    for g_ in groups:
        for k in ("cost", "revenue", "profit"):
            g_[k] = round(g_[k], 2)

    status_counts = {}
    for r in rows:
        s = r.get("status") or "Pending Job"
        status_counts[s] = status_counts.get(s, 0) + 1

    totals = {
        "count": len(rows),
        "cost": round(sum(g_["cost"] for g_ in groups), 2),
        "revenue": round(sum(g_["revenue"] for g_ in groups), 2),
        "profit": round(sum(g_["profit"] for g_ in groups), 2),
        # Revenue and profit above cover the priced jobs only; `unpriced` is how
        # many jobs carry a cost with no revenue lines yet, so the margin is
        # never read as if it covered the whole board.
        "priced": sum(g_["priced"] for g_ in groups),
        "unpriced": sum(g_["unpriced"] for g_ in groups),
        "priced_cost": round(sum(i["cost"] for g_ in groups for i in g_["items"] if i["priced"]), 2),
        "billable": sum(1 for r in rows if r.get("bill_ll")),
        "with_invoice": sum(1 for r in rows if r.get("invoice_url")),
    }
    return json_success({
        "groups": groups,
        "totals": totals,
        "status_counts": status_counts,
        "statuses": ZOLT_STATUSES,
        "contractors": sorted({(r.get("contractor") or "").strip()
                               for r in rows if (r.get("contractor") or "").strip()}),
        "periods": [g_["period"] for g_ in groups],
    })


def _zolt_clean(data, existing=None):
    """Validate an incoming work order. Returns (fields, error)."""
    out = {}
    title = " ".join(str(data.get("title", (existing or {}).get("title", ""))).split())
    if not title:
        return None, "Give the job a title"
    if len(title) > 300:
        return None, "That title is too long (300 characters max)"
    out["title"] = title

    status = str(data.get("status", (existing or {}).get("status") or "Pending Job")).strip()
    if status not in ZOLT_STATUSES:
        return None, "Unknown status: " + status
    out["status"] = status

    for f in ("ref", "period", "issue_raised", "invoice_date", "contractor",
              "invoice_url", "invoice_name", "notes"):
        if f in data:
            out[f] = str(data.get(f) or "").strip()
        elif existing is not None:
            out[f] = existing.get(f) or ""
        else:
            out[f] = ""

    for f in ("issue_raised", "invoice_date"):
        v = out.get(f) or ""
        if v and not re.match(r"^\d{4}-\d{2}-\d{2}$", v):
            return None, "Dates must look like 2026-08-06"

    try:
        out["cost_price"] = float(data.get("cost_price", (existing or {}).get("cost_price") or 0) or 0)
    except (TypeError, ValueError):
        return None, "Cost price must be a number"
    if out["cost_price"] < 0:
        return None, "Cost price cannot be negative"

    raw_bill = data.get("bill_ll", (existing or {}).get("bill_ll") or 0)
    out["bill_ll"] = 1 if raw_bill in (1, True, "1", "true", "True", "yes", "v") else 0
    return out, None


@banksia_os_bp.route("/zolt/workorders", methods=["POST"])
def api_zolt_workorder_create():
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    fields, err = _zolt_clean(request.get_json(silent=True) or {})
    if err:
        return json_error(err, 422)
    db = get_dict_db()
    try:
        order_row = db.execute(
            "SELECT COALESCE(MIN(period_order), 0) AS po FROM zolt_workorders WHERE period = ?",
            (fields["period"],),
        ).fetchone()
        cols = list(fields.keys()) + ["period_order", "position"]
        vals = list(fields.values()) + [order_row["po"] if order_row else 0, 0]
        cur = db.execute(
            f"INSERT INTO zolt_workorders ({','.join(cols)}) VALUES ({','.join('?' * len(cols))})",
            vals,
        )
        new_id = cur.lastrowid
        db.commit()
        _zolt_log(new_id, "created", None, None, fields["title"])
        row = dict(db.execute("SELECT * FROM zolt_workorders WHERE id = ?", (new_id,)).fetchone())
    finally:
        db.close()
    return json_success(_zolt_shape(row, []))


@banksia_os_bp.route("/zolt/workorders/<int:wid>", methods=["PUT", "PATCH"])
def api_zolt_workorder_update(wid):
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    db = get_dict_db()
    try:
        existing = db.execute("SELECT * FROM zolt_workorders WHERE id = ?", (wid,)).fetchone()
        if not existing:
            return json_error("That work order no longer exists", 404)
        existing = dict(existing)
        fields, err = _zolt_clean(request.get_json(silent=True) or {}, existing)
        if err:
            return json_error(err, 422)
        sets = ", ".join(f"{k} = ?" for k in fields)
        db.execute(
            f"UPDATE zolt_workorders SET {sets}, updated = datetime('now') WHERE id = ?",
            list(fields.values()) + [wid],
        )
        db.commit()
        for k, v in fields.items():
            if str(existing.get(k) or "") != str(v or ""):
                _zolt_log(wid, "updated", k,
                           str(existing.get(k) or ""), str(v or ""))
        row = dict(db.execute("SELECT * FROM zolt_workorders WHERE id = ?", (wid,)).fetchone())
        lines = _zolt_lines_for(db, [wid]).get(wid, [])
    finally:
        db.close()
    return json_success(_zolt_shape(row, lines))


@banksia_os_bp.route("/zolt/workorders/<int:wid>", methods=["DELETE"])
def api_zolt_workorder_delete(wid):
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    db = get_dict_db()
    try:
        row = db.execute("SELECT title FROM zolt_workorders WHERE id = ?", (wid,)).fetchone()
        if not row:
            return json_error("That work order no longer exists", 404)
        db.execute("DELETE FROM zolt_workorder_lines WHERE workorder_id = ?", (wid,))
        db.execute("DELETE FROM zolt_workorders WHERE id = ?", (wid,))
        db.commit()
        _zolt_log(wid, "deleted", None, row["title"], None)
    finally:
        db.close()
    return json_success({"deleted": wid})


@banksia_os_bp.route("/zolt/workorders/<int:wid>/lines", methods=["POST"])
def api_zolt_line_create(wid):
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    data = request.get_json(silent=True) or {}
    name = " ".join(str(data.get("name", "")).split()) or "Line"
    try:
        cost = float(data.get("cost") or 0)
        revenue = float(data.get("revenue") or 0)
    except (TypeError, ValueError):
        return json_error("Cost and revenue must be numbers", 422)
    db = get_dict_db()
    try:
        if not db.execute("SELECT id FROM zolt_workorders WHERE id = ?", (wid,)).fetchone():
            return json_error("That work order no longer exists", 404)
        pos = db.execute(
            "SELECT COALESCE(MAX(position), 0) + 1 AS p FROM zolt_workorder_lines WHERE workorder_id = ?",
            (wid,),
        ).fetchone()["p"]
        cur = db.execute(
            "INSERT INTO zolt_workorder_lines (workorder_id, name, cost, revenue, position)"
            " VALUES (?, ?, ?, ?, ?)",
            (wid, name, cost, revenue, pos),
        )
        db.commit()
        _zolt_log(wid, "updated", "line added", None, name)
        row = dict(db.execute(
            "SELECT id, workorder_id, name, cost, revenue FROM zolt_workorder_lines WHERE id = ?",
            (cur.lastrowid,),
        ).fetchone())
    finally:
        db.close()
    return json_success(row)


@banksia_os_bp.route("/zolt/lines/<int:lid>", methods=["PUT", "PATCH"])
def api_zolt_line_update(lid):
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    data = request.get_json(silent=True) or {}
    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM zolt_workorder_lines WHERE id = ?", (lid,)).fetchone()
        if not row:
            return json_error("That line no longer exists", 404)
        row = dict(row)
        name = " ".join(str(data.get("name", row["name"])).split()) or "Line"
        try:
            cost = float(data.get("cost", row["cost"]) or 0)
            revenue = float(data.get("revenue", row["revenue"]) or 0)
        except (TypeError, ValueError):
            return json_error("Cost and revenue must be numbers", 422)
        db.execute(
            "UPDATE zolt_workorder_lines SET name = ?, cost = ?, revenue = ? WHERE id = ?",
            (name, cost, revenue, lid),
        )
        db.commit()
        _zolt_log(row["workorder_id"], "updated", "line " + name,
                  f"{row['cost']}/{row['revenue']}", f"{cost}/{revenue}")
        out = dict(db.execute(
            "SELECT id, workorder_id, name, cost, revenue FROM zolt_workorder_lines WHERE id = ?",
            (lid,),
        ).fetchone())
    finally:
        db.close()
    return json_success(out)


@banksia_os_bp.route("/zolt/lines/<int:lid>", methods=["DELETE"])
def api_zolt_line_delete(lid):
    denied = _zolt_guard()
    if denied:
        return denied
    _zolt_ensure_tables()
    db = get_dict_db()
    try:
        row = db.execute("SELECT * FROM zolt_workorder_lines WHERE id = ?", (lid,)).fetchone()
        if not row:
            return json_error("That line no longer exists", 404)
        db.execute("DELETE FROM zolt_workorder_lines WHERE id = ?", (lid,))
        db.commit()
        _zolt_log(row["workorder_id"], "updated", "line removed",
                  row["name"], None)
    finally:
        db.close()
    return json_success({"deleted": lid})


# ── Automations panel on the compliance board ─────────────────────────────────
# Norbert, 2026-08-06: "tell me if the 2 automations are on hold or active, and
# give me a section where I can see them."
#
# Everything below is READ-ONLY and computed live. Nothing here is a stored
# status flag that somebody has to remember to update -- a panel that says
# "Active" because a human once typed Active is worse than no panel at all. So:
#   * the send counts come from compliance_emails, the real send log
#   * the workload counts come from the compliance rows themselves, using the
#     same 15-day window and the same guards the board uses
#   * the contractor chase reads automation_runs, which is empty until the job
#     actually runs, so "never run" is a fact rather than an assumption.

COMPLIANCE_DUE_WINDOW_DAYS = 15  # matches the board's "Due for Renew" section

# Certificate key -> the date column behind it. floor-plan is deliberately absent:
# it is scored on the document being held, not on an expiry date, so it has no
# renewal to chase.
COMPLIANCE_CERT_DATE_COLUMN = {
    "gas": "gas_date",
    "electric": "electrical_date",
    "epc": "epc_date",
    "fra": "fra_date",
    "emergency-lighting": "emergency_lighting_date",
    "fire-alarm": "fire_alarm_date",
    "fire-doors": "fire_doors_date",
    "fire-blanket": "fire_blanket_date",
    "co2-alarm": "co2_alarm_date",
}


def _automation_ensure_tables():
    """A run log any scheduled compliance job writes to. Created on demand so the
    board works on a database that predates it."""
    db = get_dict_db()
    try:
        db.execute("""
            CREATE TABLE IF NOT EXISTS automation_runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                automation TEXT NOT NULL,
                mode TEXT DEFAULT '',
                outcome TEXT DEFAULT '',
                detail TEXT DEFAULT '',
                items INTEGER DEFAULT 0,
                ran_at TEXT DEFAULT (datetime('now'))
            )
        """)
        db.execute("CREATE INDEX IF NOT EXISTS idx_automation_runs_name "
                   "ON automation_runs (automation, ran_at)")
        db.commit()
    finally:
        db.close()


def _compliance_due_counts():
    """How much work each certificate automation is actually looking at right now.

    Applies the board's own guards in the same order the chase script does, so
    the number here and the number that would be messaged cannot drift apart:
    Property Returned is exempt, a certificate hidden on a property is not a job,
    and a property with no boiler has no gas certificate to renew.
    """
    db = get_dict_db()
    try:
        rows = db.execute("SELECT * FROM compliance").fetchall()
    finally:
        db.close()

    today = datetime.now().date()
    horizon = today + timedelta(days=COMPLIANCE_DUE_WINDOW_DAYS)
    due = expired = 0

    for r in rows:
        row = dict(r)
        if int(row.get("automation_exempt") or 0) == 1:
            continue
        hidden = {c.strip() for c in (row.get("hidden_certs") or "").split(",") if c.strip()}
        no_gas = int(row.get("no_gas") or 0) == 1

        for cert, col in COMPLIANCE_CERT_DATE_COLUMN.items():
            if cert in hidden:
                continue
            if cert == "gas" and no_gas:
                continue
            raw = (row.get(col) or "").strip()
            if not raw:
                continue
            try:
                d = datetime.strptime(raw[:10], "%Y-%m-%d").date()
            except ValueError:
                continue
            if d < today:
                expired += 1
            elif d <= horizon:
                due += 1

    return {"due": due, "expired": expired}


def _automation_email_stats(kind):
    db = get_dict_db()
    try:
        row = db.execute(
            "SELECT COUNT(*) AS total, MAX(sent_at) AS last_at FROM compliance_emails WHERE kind = ?",
            (kind,)
        ).fetchone()
        last = db.execute(
            "SELECT to_email, sent_by, sent_at FROM compliance_emails WHERE kind = ? "
            "ORDER BY sent_at DESC LIMIT 1", (kind,)
        ).fetchone()
    finally:
        db.close()
    d = dict(row or {})
    l = dict(last) if last else {}
    return {
        "sent_total": int(d.get("total") or 0),
        "last_sent_at": d.get("last_at") or None,
        "last_sent_to": l.get("to_email") or None,
        "last_sent_by": l.get("sent_by") or None,
    }


def _automation_last_run(name):
    db = get_dict_db()
    try:
        row = db.execute(
            "SELECT mode, outcome, items, ran_at FROM automation_runs "
            "WHERE automation = ? ORDER BY ran_at DESC LIMIT 1", (name,)
        ).fetchone()
        cnt = db.execute(
            "SELECT COUNT(*) AS n FROM automation_runs WHERE automation = ?", (name,)
        ).fetchone()
    finally:
        db.close()
    return {
        "runs_total": int(dict(cnt or {}).get("n") or 0),
        "last_run": dict(row) if row else None,
    }


@banksia_os_bp.route("/compliance/automations", methods=["GET"])
def api_compliance_automations():
    """Live status of every automation that can touch the compliance board.

    `state` is one of:
      manual    -- built and working, but a person presses Send every time
      scheduled -- runs on its own
      off       -- built, not running, nothing scheduled
    """
    _automation_ensure_tables()

    counts = _compliance_due_counts()
    renewal = _automation_email_stats("renewal")
    certificate = _automation_email_stats("certificate")
    nudge = _automation_last_run("contractor_nudge")

    automations = [
        {
            "key": "renewal_email",
            "name": "Renewal reminder to the landlord",
            "channel": "Email",
            "state": "manual",
            "trigger": f"Certificate within {COMPLIANCE_DUE_WINDOW_DAYS} days of expiry",
            "sends_itself": False,
            "summary": (
                "The board works out which certificates are due and drafts the email. "
                "Nothing leaves until somebody opens the row and presses Send."
            ),
            "guards": [
                "Refuses to send while [QUOTE] is still in the body",
                "Refuses to send if the body mentions VAT",
                "Property Returned is exempt",
                "Expired certificates are not chased automatically",
            ],
            "waiting": counts["due"],
            "waiting_label": "certificates in Due for Renew right now",
            "sent_total": renewal["sent_total"],
            "last_sent_at": renewal["last_sent_at"],
            "last_sent_to": renewal["last_sent_to"],
            "last_sent_by": renewal["last_sent_by"],
        },
        {
            "key": "certificate_email",
            "name": "Certificate to the landlord",
            "channel": "Email",
            "state": "manual",
            "trigger": "A certificate is uploaded on a property",
            "sends_itself": False,
            "summary": (
                "Uploading a certificate opens the email ready to go, with the landlord, "
                "the address, the certificate type and the PDF attached. You press Send."
            ),
            "guards": [
                "Refuses to send if no certificate is on file",
                "Refuses to send if the body mentions VAT",
                "Property Returned is exempt",
                "Asks you to confirm before sending the same certificate twice",
            ],
            "waiting": None,
            "waiting_label": "",
            "sent_total": certificate["sent_total"],
            "last_sent_at": certificate["last_sent_at"],
            "last_sent_to": certificate["last_sent_to"],
            "last_sent_by": certificate["last_sent_by"],
        },
        {
            "key": "contractor_nudge",
            "name": "Quote chase to the contractor",
            "channel": "WhatsApp",
            "state": "off",
            "trigger": f"Certificate within {COMPLIANCE_DUE_WINDOW_DAYS} days of expiry",
            "sends_itself": True,
            "summary": (
                "Built and ready, but nothing is scheduled, so it has never run. It is the "
                "only one of the three that would message somebody without a person in the "
                "loop, which is why it is waiting on a decision rather than switched on."
            ),
            "guards": [
                "Maximum 8 properties per contractor per run",
                "One batched message per contractor, not one per property",
                "Remembers what it chased, keyed to the expiry date",
                "Property Returned, hidden certificates and no-boiler properties skipped",
                "Expired certificates reported, never messaged",
            ],
            "waiting": counts["due"],
            "waiting_label": "certificates it would look at today",
            "sent_total": 0,
            "last_sent_at": None,
            "last_sent_to": None,
            "last_sent_by": None,
            "runs_total": nudge["runs_total"],
            "last_run": nudge["last_run"],
        },
    ]

    return json_success({
        "automations": automations,
        "due_window_days": COMPLIANCE_DUE_WINDOW_DAYS,
        "certificates_due": counts["due"],
        "certificates_expired": counts["expired"],
    })


# ── Standard quotes per certificate ───────────────────────────────────────────
# Norbert, 2026-08-06: "standard quotes for some of the certificates, as later on
# we will automate this part as well until we don't know all the correct quotes
# from the contractors."
#
# These are what we CHARGE THE LANDLORD, not what the contractor charges us. They
# live in the database rather than in the code so the price can be corrected by the
# people who negotiate it, without a deploy — and they are read-only as far as the
# renewal email is concerned: the email still makes a human type the figure, because
# a standing price that is only right for zones 1-3 must not auto-fill an email
# about a property in zone 5.

# Not quoted (Norbert, 2026-08-06): fire doors and fire blankets are not jobs we
# price for the landlord, so they are off the quote list entirely rather than
# sitting there permanently unpriced. They stay on the certificate boards.
COMPLIANCE_QUOTE_EXCLUDED = {"fire-doors", "fire-blanket"}

COMPLIANCE_QUOTE_SEED = [
    ("gas", "60", "East London, North London — zones 1, 2 and 3"),
    ("electric", "120", "Zones 1, 2 and 3. Any other zone needs confirming."),
    ("epc", "100", ""),
    ("fire-alarm", "120", ""),
    ("emergency-lighting", "120", ""),
    ("floor-plan", "80", ""),
    ("co2-alarm", "25", ""),
    ("fra", "", "Depends on the location."),
]


def _ensure_compliance_quote_table():
    """Created on demand and seeded once with Norbert's figures. The seed only
    fills an empty table, so a corrected price is never overwritten on restart."""
    db = get_dict_db()
    db.execute("""
        CREATE TABLE IF NOT EXISTS compliance_quotes (
            cert_key TEXT PRIMARY KEY,
            amount TEXT DEFAULT '',
            coverage TEXT DEFAULT '',
            updated_by TEXT DEFAULT '',
            updated_at TEXT DEFAULT (datetime('now'))
        )
    """)
    existing = db.execute("SELECT COUNT(*) AS n FROM compliance_quotes").fetchone()
    if not int(dict(existing or {}).get("n") or 0):
        for cert, amount, coverage in COMPLIANCE_QUOTE_SEED:
            db.execute(
                "INSERT OR IGNORE INTO compliance_quotes (cert_key, amount, coverage, updated_by) "
                "VALUES (?, ?, ?, 'Norbert')", (cert, amount, coverage)
            )
    db.commit()


@banksia_os_bp.route("/compliance/quotes", methods=["GET"])
def api_compliance_quotes():
    """Every certificate gets a row, priced or not.

    Certificates with no agreed price are returned with an empty amount rather than
    left out — "we have not agreed a price for FRA yet" is the useful thing to see,
    and a list that silently omits them looks complete when it is not.
    """
    _ensure_compliance_quote_table()
    db = get_dict_db()
    try:
        rows = {r["cert_key"]: dict(r) for r in
                db.execute("SELECT * FROM compliance_quotes").fetchall()}
    finally:
        db.close()

    out = []
    for cert in sorted(COMPLIANCE_CERT_KEYS - COMPLIANCE_QUOTE_EXCLUDED):
        r = rows.get(cert, {})
        out.append({
            "cert_key": cert,
            "amount": r.get("amount") or "",
            "coverage": r.get("coverage") or "",
            "updated_by": r.get("updated_by") or "",
            "updated_at": r.get("updated_at") or "",
            "priced": bool((r.get("amount") or "").strip()),
        })
    return json_success({"quotes": out})


@banksia_os_bp.route("/compliance/quotes/<cert>", methods=["PUT"])
def api_compliance_quote_save(cert):
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    if cert in COMPLIANCE_QUOTE_EXCLUDED:
        return json_error("That certificate is not quoted.", 400)
    data = request.get_json(silent=True) or {}
    amount = str(data.get("amount", "")).strip()
    coverage = str(data.get("coverage", "")).strip()

    # Accept "60", "£60", "60.00" and store the number. A price that reads back
    # differently from what was typed is the kind of thing nobody notices until it
    # is in a landlord's inbox.
    if amount:
        cleaned = amount.replace("£", "").replace(",", "").strip()
        try:
            value = float(cleaned)
            if value < 0:
                return json_error("A quote cannot be negative.", 400)
            amount = ("%g" % value)
        except ValueError:
            return json_error("That does not look like an amount — use a number such as 120.", 400)

    _ensure_compliance_quote_table()
    db = get_dict_db()
    try:
        who = ""
        try:
            who = (getattr(request, "current_user", None) or {}).get("name") or ""
        except Exception:
            who = ""
        db.execute(
            "INSERT INTO compliance_quotes (cert_key, amount, coverage, updated_by, updated_at) "
            "VALUES (?, ?, ?, ?, datetime('now')) "
            "ON CONFLICT(cert_key) DO UPDATE SET amount = excluded.amount, "
            "coverage = excluded.coverage, updated_by = excluded.updated_by, "
            "updated_at = datetime('now')",
            (cert, amount, coverage, who)
        )
        db.commit()
        row = db.execute("SELECT * FROM compliance_quotes WHERE cert_key = ?", (cert,)).fetchone()
    finally:
        db.close()
    d = dict(row or {})
    d["priced"] = bool((d.get("amount") or "").strip())
    return json_success(d)


# ─── Contractor quote rounds ─────────────────────────────────────────────────
# Norbert, 2026-08-07: "you ask every gas engineer and who is cheaper and has an
# earlier availability will get the job."
#
# So a certificate is no longer routed to one trade. Every contractor who covers
# it is asked, their answers sit side by side, and the job is awarded from the
# comparison. Cheapest wins, the date breaks a tie on price, and the first-call
# star breaks a tie after that (Norbert: "cheaper first then who is faster"), so
# the star decides nothing while the prices differ.
#
# The round is deliberately its own table rather than columns on `compliance`.
# Three contractors answering about one boiler is three facts, and the losing two
# are worth keeping -- next time this certificate comes round, what Zakir quoted
# in August is the reason someone rings him first.

QUOTE_ROUND_LABELS = {
    "gas": "gas certificate",
    "electric": "electrical certificate (EICR)",
    "epc": "EPC",
    "fire-alarm": "fire alarm certificate",
    "emergency-lighting": "emergency lighting certificate",
    "fra": "fire risk assessment",
    "floor-plan": "floor plan",
    "fire-doors": "fire door inspection",
    "fire-blanket": "fire blanket check",
    "co2-alarm": "CO2 alarm check",
}

QUOTE_ROUND_DATE_FIELD = {
    "gas": "gas_date",
    "electric": "electrical_date",
    "epc": "epc_date",
    "fire-alarm": "fire_alarm_date",
    "emergency-lighting": "emergency_lighting_date",
    "fra": "fra_date",
    "fire-doors": "fire_doors_date",
    "fire-blanket": "fire_blanket_date",
}


def _ensure_quote_round(db):
    db.execute("""
        CREATE TABLE IF NOT EXISTS compliance_quote_round (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            compliance_id INTEGER NOT NULL,
            cert_key TEXT NOT NULL,
            contractor_id INTEGER NOT NULL,
            contractor_name TEXT DEFAULT '',
            expiry_date TEXT DEFAULT '',
            -- asked -> quoted | declined, then won / lost once it is awarded
            status TEXT NOT NULL DEFAULT 'asked',
            quote REAL,
            earliest_date TEXT DEFAULT '',
            note TEXT DEFAULT '',
            reference TEXT DEFAULT '',
            maintenance_job_id INTEGER,
            asked_at TEXT DEFAULT (datetime('now')),
            quoted_at TEXT,
            decided_at TEXT,
            recorded_by TEXT DEFAULT '',
            modified TEXT DEFAULT (datetime('now'))
        )
    """)
    # One row per contractor per certificate per property: asking the same trade
    # twice for the same job is a mistake, not a second quote.
    db.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_quote_round_one "
               "ON compliance_quote_round (compliance_id, cert_key, contractor_id)")
    db.commit()


def _quote_money(value):
    """Accept "60", "£60", "60.00". Same rule as the standard quotes page."""
    raw = str(value if value is not None else "").replace("£", "").replace(",", "").strip()
    if not raw:
        return None, None
    try:
        amount = float(raw)
    except ValueError:
        return None, "That does not look like an amount — use a number such as 120."
    if amount < 0:
        return None, "A quote cannot be negative."
    return amount, None


def _quote_round_actor():
    try:
        return (getattr(request, "current_user", None) or {}).get("name") or ""
    except Exception:
        return ""


def _quote_recommendation(entries, expiry=""):
    """Who gets the job.

    Norbert, 2026-08-07: "prioritise who is cheaper first then who is faster."
    So price decides it. The date is the tie-break when two come back at the same
    money, and the first-call star is the tie-break after that. The round still
    names anyone who could come sooner for less delay, because that is worth
    knowing before you press book -- but it is information, not the decision.

    A date after the certificate expires is not a cheaper option, it is a gap in
    the property's compliance. Those are set aside before price is compared at
    all, and only considered if nobody can come in time -- in which case the round
    says so rather than quietly recommending a late visit.
    """
    priced = [e for e in entries if e.get("quote") is not None and e["status"] in ("quoted", "won")]
    if not priced:
        return {"winner_id": None, "reason": "No quotes back yet.",
                "cheapest_id": None, "earliest_id": None}

    late = [e for e in priced if e.get("after_expiry")]
    on_time = [e for e in priced if not e.get("after_expiry")]
    late_note = ""
    if late and on_time:
        priced = on_time
        late_note = " %s ruled out — cannot come before it expires." % (
            ", ".join(e["contractor_name"] for e in late))
    elif late and not on_time:
        late_note = " Every date offered is after the certificate expires on %s." % expiry

    def by_price(e):
        # Price, then who can come first, then the first-call star. Name last so
        # the same set of quotes always produces the same answer.
        return (e["quote"], e.get("earliest_date") or "9999-12-31",
                0 if e.get("first_call") else 1, e["contractor_name"])

    winner = sorted(priced, key=by_price)[0]
    dated = [e for e in priced if (e.get("earliest_date") or "").strip()]
    earliest = sorted(dated, key=lambda e: (e["earliest_date"], e["quote"]))[0] if dated else None

    when = winner.get("earliest_date") or ""
    if when:
        reason = "%s wins on price at £%s, coming %s." % (
            winner["contractor_name"], ("%g" % winner["quote"]), _uk_date(when))
    else:
        reason = "%s wins on price at £%s — no date from him yet, so ask when he can go." % (
            winner["contractor_name"], ("%g" % winner["quote"]))

    # Said plainly rather than hidden: paying more to be seen sooner is a decision
    # somebody may want to make, and they can only make it if they are told.
    if earliest and earliest["contractor_id"] != winner["contractor_id"]:
        reason += " %s could come sooner (%s) but is £%s." % (
            earliest["contractor_name"], _uk_date(earliest["earliest_date"]),
            ("%g" % earliest["quote"]))

    return {"winner_id": winner["contractor_id"],
            "cheapest_id": winner["contractor_id"],
            "earliest_id": earliest["contractor_id"] if earliest else None,
            "reason": reason + late_note}


def _quote_round_view(row_id, cert):
    """Every contractor who covers the certificate, with whatever they have said.

    Contractors with no row yet are returned as "not asked" rather than left out.
    A comparison that only lists the two who replied looks complete when a third
    was never contacted, which is exactly the mistake this is meant to prevent.

    Takes no connection on purpose. `_ensure_compliance_contractor_table()` closes
    the shared thread-local connection, so anything holding one across this call
    would be operating on a closed database -- re-acquire after calling this.
    """
    _ensure_compliance_contractor_table()
    db = get_dict_db()
    _ensure_quote_round(db)
    contractors = db.execute(
        "SELECT id, name, group_id, certs, preferred_certs, trades FROM compliance_contractors "
        "ORDER BY name COLLATE NOCASE"
    ).fetchall()
    rows = db.execute(
        "SELECT * FROM compliance_quote_round WHERE compliance_id = ? AND cert_key = ?",
        (row_id, cert)
    ).fetchall()
    by_contractor = {int(dict(r)["contractor_id"]): dict(r) for r in rows}

    # Read the expiry from the property rather than the stored round: a certificate
    # renewed while the round was open moves the deadline, and the comparison
    # should judge the dates against the date that is true now.
    expiry = ""
    field = QUOTE_ROUND_DATE_FIELD.get(cert)
    if field:
        prop = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        expiry = str(dict(prop or {}).get(field) or "").strip()

    entries = []
    for c in contractors:
        c = dict(c)
        covers = [k.strip() for k in str(c.get("certs") or "").split(",") if k.strip()]
        first = [k.strip() for k in str(c.get("preferred_certs") or "").split(",") if k.strip()]
        if cert not in covers:
            continue
        r = by_contractor.get(int(c["id"])) or {}
        when = str(r.get("earliest_date") or "").strip()
        entries.append({
            "after_expiry": bool(expiry and when and when > expiry),
            "contractor_id": int(c["id"]),
            "contractor_name": c["name"],
            "group_id": c.get("group_id") or "",
            "first_call": cert in first,
            "status": r.get("status") or "not_asked",
            "quote": r.get("quote"),
            "earliest_date": r.get("earliest_date") or "",
            "note": r.get("note") or "",
            "reference": r.get("reference") or "",
            "maintenance_job_id": r.get("maintenance_job_id"),
            "asked_at": r.get("asked_at") or "",
            "quoted_at": r.get("quoted_at") or "",
            "recorded_by": r.get("recorded_by") or "",
        })

    awarded = next((e for e in entries if e["status"] == "won"), None)
    return {
        "cert_key": cert,
        "label": QUOTE_ROUND_LABELS.get(cert, cert),
        "expiry_date": expiry,
        "entries": entries,
        "awarded": awarded,
        "recommendation": _quote_recommendation(entries, expiry),
    }


@banksia_os_bp.route("/compliance/<int:row_id>/quote-round/<cert>", methods=["GET"])
def api_quote_round(row_id, cert):
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    db = get_dict_db()
    prop = db.execute("SELECT id, property_name FROM compliance WHERE id = ?", (row_id,)).fetchone()
    if not prop:
        return json_error("Property not found", 404)
    name = dict(prop).get("property_name") or ""
    view = _quote_round_view(row_id, cert)
    view["property_name"] = name
    return json_success(view)


@banksia_os_bp.route("/compliance/<int:row_id>/quote-round/<cert>/ask", methods=["POST"])
def api_quote_round_ask(row_id, cert):
    """Open the round: mark every contractor who covers this certificate as asked.

    Sending is the chase script's job, not this route's -- this records that the
    question went out so the comparison has a full field to fill in. Contractors
    who have already answered are left exactly as they are.
    """
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    db = get_dict_db()
    prop = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
    if not prop:
        return json_error("Property not found", 404)
    prop = dict(prop)
    expiry = str(prop.get(QUOTE_ROUND_DATE_FIELD.get(cert, ""), "") or "")

    view = _quote_round_view(row_id, cert)
    if not view["entries"]:
        return json_error("Nobody on the Contractors page covers that certificate yet.", 422)
    db = get_dict_db()          # the view re-opens the connection; do not reuse the old one
    who = _quote_round_actor()
    added = 0
    for e in view["entries"]:
        if e["status"] != "not_asked":
            continue
        db.execute(
            "INSERT OR IGNORE INTO compliance_quote_round "
            "(compliance_id, cert_key, contractor_id, contractor_name, expiry_date, "
            " status, recorded_by) VALUES (?, ?, ?, ?, ?, 'asked', ?)",
            (row_id, cert, e["contractor_id"], e["contractor_name"], expiry, who)
        )
        added += 1
    db.commit()
    view = _quote_round_view(row_id, cert)
    view["property_name"] = prop.get("property_name") or ""
    view["added"] = added
    return json_success(view)


@banksia_os_bp.route("/compliance/<int:row_id>/quote-round/<cert>/<int:contractor_id>",
                     methods=["PUT"])
def api_quote_round_save(row_id, cert, contractor_id):
    """Record what one contractor came back with."""
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    data = request.get_json(silent=True) or {}
    status = str(data.get("status") or "quoted").strip()
    if status not in ("asked", "quoted", "declined"):
        return json_error("Unknown status", 400)

    quote, err = _quote_money(data.get("quote"))
    if err:
        return json_error(err, 422)
    # "Quoted" with no figure is the state that would quietly drop a contractor out
    # of the comparison while looking answered.
    if status == "quoted" and quote is None:
        return json_error("Give the price they quoted, or mark them as declined.", 422)
    earliest = str(data.get("earliest_date") or "").strip()
    if earliest and not re.match(r"^\d{4}-\d{2}-\d{2}$", earliest):
        return json_error("Use a date like 2026-08-20.", 422)
    note = str(data.get("note") or "").strip()[:500]

    _ensure_compliance_contractor_table()
    db = get_dict_db()
    _ensure_quote_round(db)
    prop = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
    if not prop:
        return json_error("Property not found", 404)
    c = _contractor_row(db, contractor_id)
    if not c:
        return json_error("Contractor not found", 404)
    c = dict(c)
    covers = [k.strip() for k in str(c.get("certs") or "").split(",") if k.strip()]
    if cert not in covers:
        return json_error("%s is not marked as doing that certificate." % c["name"], 422)

    existing = db.execute(
        "SELECT * FROM compliance_quote_round WHERE compliance_id = ? AND cert_key = ? "
        "AND contractor_id = ?", (row_id, cert, contractor_id)
    ).fetchone()
    # A booked job is not a quote any more. Changing the winner's price after
    # the work order exists would put two different numbers on one job.
    if existing and dict(existing).get("status") == "won":
        return json_error("That job is already booked with %s — reopen it first." % c["name"], 409)

    expiry = str(dict(prop).get(QUOTE_ROUND_DATE_FIELD.get(cert, ""), "") or "")
    who = _quote_round_actor()
    quoted_at = "datetime('now')" if status == "quoted" else "NULL"
    db.execute(
        "INSERT INTO compliance_quote_round "
        "(compliance_id, cert_key, contractor_id, contractor_name, expiry_date, status, "
        " quote, earliest_date, note, recorded_by, quoted_at, modified) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, %s, datetime('now')) "
        "ON CONFLICT(compliance_id, cert_key, contractor_id) DO UPDATE SET "
        "  status = excluded.status, quote = excluded.quote, "
        "  earliest_date = excluded.earliest_date, note = excluded.note, "
        "  recorded_by = excluded.recorded_by, quoted_at = %s, "
        "  modified = datetime('now')" % (quoted_at, quoted_at),
        (row_id, cert, contractor_id, c["name"], expiry, status,
         quote, earliest, note, who)
    )
    db.commit()
    prop_name = dict(prop).get("property_name") or ""
    view = _quote_round_view(row_id, cert)
    view["property_name"] = prop_name
    _log_activity("compliance", row_id, "update", "quote", "",
                  "%s: %s" % (c["name"], ("£%g" % quote) if quote is not None else status),
                  notes="%s quote recorded" % cert)
    return json_success(view)


def _uk_date(value):
    """25 Aug 2026, not 2026-08-25. This one goes to a contractor."""
    raw = str(value or "").strip()
    try:
        return datetime.strptime(raw[:10], "%Y-%m-%d").strftime("%-d %b %Y")
    except (ValueError, TypeError):
        return raw


def _booking_message(contractor_name, cert_label, property_name, quote, when, reference):
    first = (str(contractor_name or "").strip().split() or ["there"])[0]
    lines = ["Hello %s," % first, ""]
    lines.append("Thanks — that is booked in.")
    lines.append("")
    lines.append("%s at %s" % (cert_label[:1].upper() + cert_label[1:], property_name))
    if when:
        lines.append("Date: %s" % _uk_date(when))
    if quote is not None:
        lines.append("Agreed price: £%g" % quote)
    lines.append("Our reference: %s" % reference)
    lines.append("")
    lines.append("Please quote that reference on your invoice. Thanks")
    return "\n".join(lines)


@banksia_os_bp.route("/compliance/<int:row_id>/quote-round/<cert>/award", methods=["POST"])
def api_quote_round_award(row_id, cert):
    """Give the job to one contractor: work order first, then the reference.

    The reference the contractor is given IS the work order reference. Minting a
    separate booking number would mean an invoice quoting one number and a board
    holding another, which is a reconciliation problem nobody asked for.

    The message is drafted and returned, not sent. Booking commits us to a price
    and a date with an outside company; every other outbound in this system waits
    for a person, and this is the one that should wait most.
    """
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    data = request.get_json(silent=True) or {}
    try:
        contractor_id = int(data.get("contractor_id") or 0)
    except (TypeError, ValueError):
        contractor_id = 0
    if not contractor_id:
        return json_error("Which contractor?", 422)

    db = get_dict_db()
    try:
        _ensure_quote_round(db)
        _ensure_maintenance_cost_ll(db)
        _ensure_compliance_quote_table()
        prop = db.execute("SELECT * FROM compliance WHERE id = ?", (row_id,)).fetchone()
        if not prop:
            return json_error("Property not found", 404)
        prop = dict(prop)

        already = db.execute(
            "SELECT * FROM compliance_quote_round WHERE compliance_id = ? AND cert_key = ? "
            "AND status = 'won'", (row_id, cert)
        ).fetchone()
        if already:
            already = dict(already)
            return json_error(
                "Already booked with %s (%s)." % (already.get("contractor_name"),
                                                  already.get("reference") or "no reference"), 409)

        win = db.execute(
            "SELECT * FROM compliance_quote_round WHERE compliance_id = ? AND cert_key = ? "
            "AND contractor_id = ?", (row_id, cert, contractor_id)
        ).fetchone()
        if not win:
            return json_error("That contractor has not been asked for this job yet.", 404)
        win = dict(win)
        if win.get("quote") is None:
            return json_error("No price recorded for %s yet — a job cannot be booked on no price."
                              % win.get("contractor_name"), 422)

        # What we charge the landlord: the standard quote from the Quotes panel.
        # Left empty rather than guessed when the certificate has no agreed price,
        # so an unpriced FRA reaches a human instead of billing a landlord £0.
        cost_ll = 0.0
        try:
            std = db.execute("SELECT amount FROM compliance_quotes WHERE cert_key = ?",
                             (cert,)).fetchone()
            cost_ll = float(str(dict(std or {}).get("amount") or "").strip() or 0)
        except (ValueError, TypeError):
            cost_ll = 0.0

        label = QUOTE_ROUND_LABELS.get(cert, cert)
        _ensure_maintenance_cert_key(db)
        reference = _next_maintenance_reference(db)
        address = prop.get("property_name") or ""
        cur = db.execute(
            """INSERT INTO maintenance_jobs
               (reference, title, description, type, cert_key, priority, status, address,
                contractor, labour_cost, materials_cost, bill_ll, emergency,
                team_notes, source, cost_ll)
               VALUES (?, ?, ?, 'Certificate', ?, 'Medium', 'PENDING', ?, ?, ?, 0, 1, 0, ?, 'compliance', ?)""",
            [reference,
             "%s — %s" % (label[:1].upper() + label[1:], address),
             "Booked from the compliance board after a quote round.",
             cert,
             address,
             win.get("contractor_name") or "",
             float(win.get("quote") or 0),
             "Attending %s. Quote round: %s." % (win.get("earliest_date") or "date to confirm",
                                                 win.get("note") or "no notes"),
             cost_ll]
        )
        job_id = cur.lastrowid

        db.execute(
            "UPDATE compliance_quote_round SET status = 'won', reference = ?, "
            "maintenance_job_id = ?, decided_at = datetime('now'), modified = datetime('now') "
            "WHERE id = ?", (reference, job_id, win["id"])
        )
        db.execute(
            "UPDATE compliance_quote_round SET status = 'lost', decided_at = datetime('now'), "
            "modified = datetime('now') WHERE compliance_id = ? AND cert_key = ? AND id != ? "
            "AND status IN ('asked', 'quoted')", (row_id, cert, win["id"])
        )
        db.commit()

        view = _quote_round_view(row_id, cert)
        view["property_name"] = address
        view["reference"] = reference
        view["maintenance_job_id"] = job_id
        view["cost_ll"] = cost_ll
        view["booking_message"] = _booking_message(
            win.get("contractor_name"), label, address, win.get("quote"),
            win.get("earliest_date") or "", reference)
        view["booking_target"] = next(
            (e["group_id"] for e in view["entries"] if e["contractor_id"] == contractor_id), "")
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()
    _log_activity("compliance", row_id, "create", "booking", "",
                  "%s — %s" % (win.get("contractor_name"), reference),
                  notes="%s booked, work order %s" % (cert, reference))
    return json_success(view)


@banksia_os_bp.route("/compliance/<int:row_id>/quote-round/<cert>/reopen", methods=["POST"])
def api_quote_round_reopen(row_id, cert):
    """Undo a booking that was made in error.

    The work order is deliberately NOT deleted -- somebody may already have that
    reference. It is left for a human to cancel on the maintenance board, and the
    round says so rather than quietly tidying up behind them.
    """
    if cert not in COMPLIANCE_CERT_KEYS:
        return json_error("Unknown certificate", 400)
    db = get_dict_db()
    _ensure_quote_round(db)
    won = db.execute(
        "SELECT * FROM compliance_quote_round WHERE compliance_id = ? AND cert_key = ? "
        "AND status = 'won'", (row_id, cert)
    ).fetchone()
    if not won:
        return json_error("That job is not booked.", 404)
    won = dict(won)
    db.execute(
        "UPDATE compliance_quote_round SET status = CASE WHEN quote IS NULL THEN 'asked' "
        "ELSE 'quoted' END, decided_at = NULL, reference = '', maintenance_job_id = NULL, "
        "modified = datetime('now') WHERE compliance_id = ? AND cert_key = ? "
        # Only the decision is undone. A contractor who said no is still a
        # contractor who said no -- resetting him to "asked" would put a question
        # back on the board that has already been answered.
        "AND status IN ('won', 'lost')",
        (row_id, cert)
    )
    db.commit()
    view = _quote_round_view(row_id, cert)
    view["orphan_reference"] = won.get("reference") or ""
    return json_success(view)


# ─── Maintenance board: groups, Cost LL, evidence ────────────────────────────
# Norbert, 2026-08-07. Four groups, one line per job, and the landlord's price
# derived from the labour rather than typed twice.

# Cancelled is a group rather than a delete (Norbert, 2026-08-08). A job called
# off still has to be answerable for -- who raised it, against which property and
# what it would have cost. Deleting it loses that, so cancelling moves it aside.
# New Report first: it is the intake, not a stage of the work. A tenant's
# report arrives there and is triaged out of it (Norbert, 2026-08-08).
MAINT_BOARD_GROUPS = ["NEW REPORT", "URGENT", "TO BE ARRANGED", "LIVE",
                      "COMPLETED", "CANCELLED"]

# What we charge the landlord on top of the LABOUR only. Materials are passed on
# at cost -- marking them up as well was never the agreement, and it is the kind
# of thing that is invisible until a landlord adds the invoice up himself.
MAINT_LL_MARKUP = 0.15

MAINT_EVIDENCE_DIR = os.path.join(os.path.dirname(__file__), "media", "maintenance")
os.makedirs(MAINT_EVIDENCE_DIR, exist_ok=True)

_EVIDENCE_ALLOWED_EXT = {
    ".jpg", ".jpeg", ".png", ".webp", ".heic", ".gif",
    ".mp4", ".mov", ".m4v", ".webm", ".pdf",
}
# Video is the reason this is not the 25MB the certificates get: a minute of a
# contractor walking round a flat on a phone is comfortably past that.
_EVIDENCE_MAX_BYTES = 60 * 1024 * 1024


def _ensure_cost_ll_override(db):
    """`cost_ll_override` marks a Cost LL somebody typed by hand.

    Without it there is no way to tell a figure the rule produced from a figure a
    person chose, so the next edit to the labour would silently overwrite their
    number.
    """
    try:
        db.execute("ALTER TABLE maintenance_jobs ADD COLUMN cost_ll_override INTEGER DEFAULT 0")
        db.commit()
    except Exception:
        pass  # already present


def _evidence_name(raw):
    base = os.path.basename(str(raw or "")).strip().replace("\\", "_").replace("/", "_")
    base = re.sub(r"[^A-Za-z0-9._-]", "_", base)
    return base[:120] or "file"


def _sync_cost_ll(db, job_id, data):
    """Cost LL follows the labour at +15% until somebody types their own figure.

    An explicit `cost_ll` in the payload IS that decision, so it latches and the
    rule stops touching it. Sending `cost_ll_override: 0` hands it back to the
    rule. With Bill LL unticked there is nothing to charge, so it reads zero.
    """
    _ensure_cost_ll_override(db)
    _ensure_maintenance_cert_key(db)
    row = db.execute(
        "SELECT bill_ll, labour_cost, cost_ll, cost_ll_override, cert_key "
        "FROM maintenance_jobs WHERE id = ?",
        [job_id]
    ).fetchone()
    if not row:
        return
    row = dict(row)

    if "cost_ll" in data:
        db.execute("UPDATE maintenance_jobs SET cost_ll_override = 1 WHERE id = ?", [job_id])
        db.commit()
        return
    if "cost_ll_override" in data and not data.get("cost_ll_override"):
        row["cost_ll_override"] = 0
    if int(row.get("cost_ll_override") or 0):
        return

    if not int(row.get("bill_ll") or 0):
        target = 0.0
    else:
        # A certificate is sold at the agreed price, not at cost plus a margin
        # (Norbert, 2026-08-08). The landlord is quoted from the Quotes panel, so
        # marking up whatever the contractor happened to charge would contradict
        # the number he was given.
        quoted = _standard_quote(db, row.get("cert_key"))
        target = quoted if quoted is not None else round(
            float(row.get("labour_cost") or 0) * (1 + MAINT_LL_MARKUP), 2)
    if float(row.get("cost_ll") or 0) != target:
        db.execute("UPDATE maintenance_jobs SET cost_ll = ? WHERE id = ?", [target, job_id])
        db.commit()


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>/evidence", methods=["POST"])
def api_maintenance_evidence_upload(job_id):
    """Attach photos or video of the work to a job.

    Files are appended, never replaced: two contractors sending a picture each
    should end up with two pictures on the job, not the second overwriting the
    first.
    """
    import secrets
    db = get_dict_db()
    combined = ""
    try:
        job = db.execute("SELECT id, photo_paths FROM maintenance_jobs WHERE id = ?", [job_id]).fetchone()
        if not job:
            return json_error("Job not found", 404)
        job = dict(job)

        files = request.files.getlist("file")
        if not files:
            return json_error("No file provided (use field 'file')", 400)

        added = []
        for f in files:
            if not f or not f.filename:
                continue
            name = _evidence_name(f.filename)
            ext = os.path.splitext(name)[1].lower()
            if ext not in _EVIDENCE_ALLOWED_EXT:
                return json_error("Cannot attach %s — allowed: %s"
                                  % (ext or "that file", ", ".join(sorted(_EVIDENCE_ALLOWED_EXT))), 415)
            payload = f.read()
            if not payload:
                return json_error("%s is empty" % name, 400)
            if len(payload) > _EVIDENCE_MAX_BYTES:
                return json_error("%s is larger than 60MB" % name, 413)
            stored = "%s_%s_%s" % (job_id, secrets.token_hex(4), name)
            with open(os.path.join(MAINT_EVIDENCE_DIR, stored), "wb") as fh:
                fh.write(payload)
            added.append("/api/banksia-os/maintenance/evidence/%s" % stored)

        if not added:
            return json_error("Nothing to attach", 400)

        existing = [p.strip() for p in str(job.get("photo_paths") or "").split(",") if p.strip()]
        combined = ",".join(existing + added)
        db.execute(
            "UPDATE maintenance_jobs SET photo_paths = ?, modified = datetime('now') WHERE id = ?",
            [combined, job_id]
        )
        db.commit()
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()
    return json_success({"photo_paths": combined, "added": added})


@banksia_os_bp.route("/maintenance/evidence/<path:stored>", methods=["GET"])
def api_maintenance_evidence_file(stored):
    name = _evidence_name(stored)
    path = os.path.join(MAINT_EVIDENCE_DIR, name)
    if not os.path.exists(path):
        return json_error("Not found", 404)
    from flask import send_file
    return send_file(path, as_attachment=False)


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>/evidence", methods=["DELETE"])
def api_maintenance_evidence_delete(job_id):
    """Take one file off a job.

    The file itself is left on disk. Somebody may have the link, and a wrongly
    removed photo of a completed repair is worth more than the megabyte.
    """
    target = (request.args.get("path") or "").strip()
    if not target:
        return json_error("Which file?", 400)
    db = get_dict_db()
    try:
        job = db.execute("SELECT photo_paths FROM maintenance_jobs WHERE id = ?", [job_id]).fetchone()
        if not job:
            return json_error("Job not found", 404)
        kept = [p.strip() for p in str(dict(job).get("photo_paths") or "").split(",")
                if p.strip() and p.strip() != target]
        db.execute(
            "UPDATE maintenance_jobs SET photo_paths = ?, modified = datetime('now') WHERE id = ?",
            [",".join(kept), job_id]
        )
        db.commit()
    finally:
        db.close()
    return json_success({"photo_paths": ",".join(kept)})


# ─── Acting on several jobs at once ──────────────────────────────────────────


@banksia_os_bp.route("/maintenance/jobs/bulk-status", methods=["POST"])
def api_maintenance_bulk_status():
    """Move a set of jobs into one group -- cancelling a handful being the point
    of it (Norbert, 2026-08-08).

    Every job is judged on its own. One that cannot go where it is being sent is
    refused *by name* and the rest still move, because the alternative is a
    selection of twenty failing on account of one and no way to see which.
    """
    data = request.get_json(force=True, silent=True) or {}
    raw_ids = data.get("ids") or []
    ids = []
    for i in raw_ids:
        try:
            ids.append(int(i))
        except (TypeError, ValueError):
            continue
    status = str(data.get("status") or "").strip().upper()

    if not ids:
        return json_error("No jobs selected", 400)
    if status not in MAINT_BOARD_GROUPS:
        return json_error("%s is not a group on this board" % (status or "That"), 400)

    db = get_dict_db()
    moved, refused = [], []
    try:
        _ensure_maintenance_cost_ll(db)
        for job_id in ids:
            row = db.execute(
                "SELECT id, reference, status, contractor, labour_cost, materials_cost, "
                "photo_paths, type FROM maintenance_jobs WHERE id = ?", [job_id]
            ).fetchone()
            if not row:
                refused.append({"id": job_id, "reference": "#%s" % job_id,
                                "reason": "no longer on the board"})
                continue
            row = dict(row)
            ref = row.get("reference") or "#%s" % job_id

            if str(row.get("status") or "").strip().upper() == status:
                # Already where it is being sent. Counted as moved rather than
                # skipped: somebody who cancels five and is told "three moved"
                # reasonably wonders what happened to the other two.
                moved.append({"id": job_id, "reference": ref})
                continue

            if status == "COMPLETED":
                missing = _completed_blockers(row)
                if missing:
                    refused.append({"id": job_id, "reference": ref,
                                    "reason": "still needs %s" % _join_words(missing)})
                    continue
            if status == "LIVE":
                missing = _live_blockers(row)
                if missing:
                    refused.append({"id": job_id, "reference": ref,
                                    "reason": "still needs %s" % _join_words(missing)})
                    continue

            db.execute(
                "UPDATE maintenance_jobs SET status = ?, modified = datetime('now') WHERE id = ?",
                [status, job_id]
            )
            if status == "LIVE":
                db.execute(
                    "UPDATE maintenance_jobs SET start_date = ? WHERE id = ? "
                    "AND (start_date IS NULL OR TRIM(start_date) = '')",
                    [_london_today(), job_id]
                )
            if status == "COMPLETED":
                db.execute(
                    "UPDATE maintenance_jobs SET completed_date = datetime('now') "
                    "WHERE id = ? AND completed_date IS NULL", [job_id]
                )
            db.commit()
            _sync_cost_ll(db, job_id, {})
            try:
                db.execute("UPDATE maintenance_jobs SET sync_pending = 1 WHERE id = ?", [job_id])
                db.commit()
            except Exception:
                db.rollback()
            moved.append({"id": job_id, "reference": ref})
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()

    return json_success({"status": status, "moved": moved, "refused": refused,
                         "moved_count": len(moved), "refused_count": len(refused)})


# ─── LL invoice ──────────────────────────────────────────────────────────────
# A printable invoice per completed job (Norbert, 2026-08-08), modelled on the
# Zolt invoice he sent. Generated on demand rather than stored: the job is the
# record, and a stored PDF would go stale the moment a figure is corrected.

INVOICE_BILL_TO = [
    "Verv Rooms",
    "35a Highbury Corner",
    "N5 1RA",
    "London",
]

# Banksia green, so the document does not look like it came from somewhere else.
INVOICE_ACCENT = (0.32, 0.73, 0.19)
INVOICE_INK = (0.11, 0.13, 0.16)
INVOICE_MUTED = (0.45, 0.49, 0.55)


def _money(n):
    return "£%s" % format(float(n or 0), ",.2f")


def _invoice_date(job):
    """The day the work was signed off.

    Falls back to the day the job was raised, never to "today": an invoice that
    carries a different date each time it is printed is not an invoice. Older
    jobs completed before the board stamped a completion date rely on this.
    """
    from datetime import datetime
    for field in ("completed_date", "created"):
        raw = str(job.get(field) or "").strip()
        for fmt, width in (("%Y-%m-%d %H:%M:%S", 19), ("%Y-%m-%d", 10)):
            try:
                return datetime.strptime(raw[:width], fmt)
            except ValueError:
                continue
    try:
        from zoneinfo import ZoneInfo
        return datetime.now(ZoneInfo("Europe/London"))
    except Exception:
        return datetime.now()


def _invoice_pdf_bytes(job):
    """Draw the invoice. Kept to one page: it is one job, and an invoice that
    runs over is an invoice somebody has to check twice."""
    import io as _io
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas

    buf = _io.BytesIO()
    W, H = A4
    c = rl_canvas.Canvas(buf, pagesize=A4)
    c.setTitle("Invoice %s" % (job.get("reference") or ""))

    L, R = 48, W - 48

    def ink(col=INVOICE_INK):
        c.setFillColorRGB(*col)

    # Accent rule across the head.
    c.setFillColorRGB(*INVOICE_ACCENT)
    c.rect(0, H - 10, W, 10, stroke=0, fill=1)

    y = H - 62
    ink()
    c.setFont("Helvetica-Bold", 26)
    c.drawString(L, y, "INVOICE")

    number = "INV-%s" % (job.get("reference") or job.get("id"))
    when = _invoice_date(job)
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(R, y + 8, number)
    ink(INVOICE_MUTED)
    c.setFont("Helvetica", 10)
    c.drawRightString(R, y - 6, when.strftime("%-d %b %Y") if hasattr(when, "strftime") else "")

    y -= 24
    c.setStrokeColorRGB(0.85, 0.87, 0.89)
    c.setLineWidth(1)
    c.line(L, y, R, y)

    # ── Bill to (left) and From (right) ──
    y -= 30
    top = y
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(L, y, "BILL TO")
    ink()
    c.setFont("Helvetica-Bold", 11.5)
    y -= 16
    c.drawString(L, y, INVOICE_BILL_TO[0])
    c.setFont("Helvetica", 10.5)
    for line in INVOICE_BILL_TO[1:]:
        y -= 14
        c.drawString(L, y, line)

    ry = top
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawRightString(R, ry, "FROM")
    ink()
    c.setFont("Helvetica-Bold", 11.5)
    ry -= 16
    c.drawRightString(R, ry, str(job.get("contractor") or "Contractor not recorded"))
    ink(INVOICE_MUTED)
    c.setFont("Helvetica", 10.5)
    ry -= 14
    c.drawRightString(R, ry, "Reference %s" % (job.get("reference") or ""))

    # ── What the invoice is for ──
    y = min(y, ry) - 34
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(L, y, "WORK CARRIED OUT")
    y -= 17
    ink()
    c.setFont("Helvetica-Bold", 12)
    title = str(job.get("title") or "Maintenance")
    c.drawString(L, y, title[:78])

    where = str(job.get("property_name") or "").strip()
    unit = str(job.get("unit") or "").strip()
    if unit:
        where = ("%s, %s" % (where, unit)).strip(", ")
    if where:
        y -= 15
        ink(INVOICE_MUTED)
        c.setFont("Helvetica", 10.5)
        c.drawString(L, y, where[:88])

    done = str(job.get("completed_date") or "")[:10]
    if done:
        y -= 14
        c.setFont("Helvetica", 9.5)
        c.drawString(L, y, "Completed %s" % done)

    # ── Items ──
    y -= 34
    c.setFillColorRGB(0.96, 0.97, 0.98)
    c.rect(L, y - 6, R - L, 24, stroke=0, fill=1)
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(L + 12, y + 3, "ITEMS")
    c.drawRightString(R - 12, y + 3, "AMOUNT")

    labour = float(job.get("labour_cost") or 0)
    materials = float(job.get("materials_cost") or 0)
    total = round(labour + materials, 2)

    y -= 14
    for label, amount in (("Labour", labour), ("Materials", materials)):
        y -= 22
        ink()
        c.setFont("Helvetica", 11)
        c.drawString(L + 12, y, label)
        c.drawRightString(R - 12, y, _money(amount))
        c.setStrokeColorRGB(0.91, 0.93, 0.94)
        c.setLineWidth(0.6)
        c.line(L + 12, y - 8, R - 12, y - 8)

    y -= 30
    c.setStrokeColorRGB(*INVOICE_ACCENT)
    c.setLineWidth(1.6)
    c.line(L + 12, y + 14, R - 12, y + 14)
    ink()
    c.setFont("Helvetica-Bold", 13)
    c.drawString(L + 12, y - 4, "Total Due")
    c.drawRightString(R - 12, y - 4, _money(total))

    # ── Payment, directly under the total where it is read ──
    y -= 44
    c.setFillColorRGB(0.96, 0.97, 0.98)
    c.rect(L, y - 34, R - L, 52, stroke=0, fill=1)
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(L + 12, y + 6, "PAYMENT")
    ink()
    c.setFont("Helvetica", 10)
    c.drawString(L + 12, y - 10, "Please quote %s on payment." % number)
    ink(INVOICE_MUTED)
    c.setFont("Helvetica", 9)
    c.drawString(L + 12, y - 25, "Bank details to be confirmed.")

    # ── Foot ──
    c.setStrokeColorRGB(0.85, 0.87, 0.89)
    c.setLineWidth(1)
    c.line(L, 74, R, 74)
    ink(INVOICE_MUTED)
    c.setFont("Helvetica-Oblique", 9.5)
    c.drawString(L, 56, "Thank you for your business.")
    c.setFont("Helvetica", 8.5)
    c.drawRightString(R, 56, "%s · %s" % (number, INVOICE_BILL_TO[0]))

    c.showPage()
    c.save()
    return buf.getvalue()


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>/ll-invoice", methods=["GET"])
def api_maintenance_ll_invoice(job_id):
    """The invoice for a completed job, as a PDF.

    Completed only: an invoice for work still being argued about is a document
    somebody will send by accident.
    """
    denied = _require_super_admin()
    if denied:
        return denied

    db = get_dict_db()
    try:
        job = db.execute(
            """SELECT mj.*,
                      COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN ''
                                           ELSE p.name END, ''),
                               p.address_line_1, p.ref, p.name) AS property_name
               FROM maintenance_jobs mj
               LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.id = ?""", [job_id]
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)
        job = dict(job)
        if str(job.get("status") or "").strip().upper() != "COMPLETED":
            return json_error("Only a completed job has an invoice.", 422)
        pdf = _invoice_pdf_bytes(job)
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()

    from flask import Response
    return Response(pdf, mimetype="application/pdf", headers={
        "Content-Disposition": 'inline; filename="Invoice-%s.pdf"' % (job.get("reference") or job_id),
    })


# ─── Landlord Report ─────────────────────────────────────────────────────────


def _require_super_admin():
    """Refuse anyone who is not super_admin. Returns a response, or None to carry on.

    Hiding the tab in the browser is not a control -- the endpoint has to say no
    on its own (Norbert, 2026-08-08). This report puts what we charge a landlord
    next to what the work cost us, which is the whole reason it is restricted.
    """
    _, role = _get_current_user()
    if str(role or "").strip().lower() != "super_admin":
        return json_error("This report is restricted to super admins.", 403)
    return None


@banksia_os_bp.route("/maintenance/landlord-report", methods=["GET"])
def api_maintenance_landlord_report():
    """Maintenance the landlord is being charged for, by property.

    Only jobs with Bill LL ticked (Norbert, 2026-08-08). A job we are carrying is
    not part of what a landlord owes, and listing it under their property invites
    somebody to invoice for it.

    Two categories, because the agreement decides who pays: on a management fee
    the landlord carries the repair, on guaranteed rent we do. A billed job on a
    guaranteed-rent property is therefore an exception -- a recharge, usually --
    and worth being able to see on its own rather than mixed in.
    """
    denied = _require_super_admin()
    if denied:
        return denied

    db = get_dict_db()
    try:
        rows = db.execute(
            """SELECT mj.id, mj.reference, mj.title, mj.status, mj.type, mj.contractor,
                      mj.start_date, mj.completed_date, mj.labour_cost, mj.materials_cost,
                      mj.cost_ll, mj.unit, mj.property_id,
                      COALESCE(NULLIF(CASE WHEN LOWER(p.name) IN ('multi','single') THEN ''
                                           ELSE p.name END, ''),
                               p.address_line_1, p.ref, p.name) AS property_name,
                      p.property_owner_name, p.owner_company, p.management_type
               FROM maintenance_jobs mj
               LEFT JOIN properties p ON mj.property_id = p.id
               WHERE mj.bill_ll = 1
                 AND UPPER(COALESCE(mj.status, '')) != 'CANCELLED'
               ORDER BY mj.created DESC"""
        ).fetchall()

        buckets = {
            "management": {"key": "management", "label": "Management Fee",
                           "blurb": "The landlord carries the cost of repairs.",
                           "props": {}},
            "guaranteed": {"key": "guaranteed", "label": "Guaranteed Rent",
                           "blurb": "We carry the cost, so anything billed here is a recharge.",
                           "props": {}},
        }
        unassigned = []

        for r in rows:
            r = dict(r)
            ours = round(float(r.get("labour_cost") or 0) + float(r.get("materials_cost") or 0), 2)
            charged = round(float(r.get("cost_ll") or 0), 2)
            job = {
                "id": r["id"], "reference": r.get("reference"), "title": r.get("title"),
                "status": r.get("status"), "type": r.get("type"),
                "contractor": r.get("contractor"), "unit": r.get("unit"),
                "start_date": r.get("start_date"), "completed_date": r.get("completed_date"),
                "labour_cost": r.get("labour_cost"), "materials_cost": r.get("materials_cost"),
                "our_cost": ours, "charged": charged, "margin": round(charged - ours, 2),
            }

            # A billed job with no property cannot be filed under one. Kept aside and
            # shown as such rather than dropped, or the totals would not add up.
            if not r.get("property_id"):
                unassigned.append(job)
                continue

            mt = str(r.get("management_type") or "").strip()
            key = "management" if "management fee" in mt.lower() else "guaranteed"
            props = buckets[key]["props"]
            pid = r["property_id"]
            entry = props.setdefault(pid, {
                "property_id": pid,
                "property_name": r.get("property_name") or "Property %s" % pid,
                "landlord": (str(r.get("property_owner_name") or "").strip()
                             or str(r.get("owner_company") or "").strip()
                             or "No landlord on the property"),
                "management_type": mt or "Not set",
                "jobs": [], "our_cost": 0.0, "charged": 0.0,
            })
            entry["jobs"].append(job)
            entry["our_cost"] = round(entry["our_cost"] + ours, 2)
            entry["charged"] = round(entry["charged"] + charged, 2)

        categories = []
        for key in ("management", "guaranteed"):
            b = buckets[key]
            props = []
            for e in b["props"].values():
                e["job_count"] = len(e["jobs"])
                e["margin"] = round(e["charged"] - e["our_cost"], 2)
                props.append(e)
            props.sort(key=lambda e: (-e["charged"], e["property_name"]))
            categories.append({
                "key": b["key"], "label": b["label"], "blurb": b["blurb"],
                "properties": props,
                "property_count": len(props),
                "job_count": sum(p["job_count"] for p in props),
                "our_cost": round(sum(p["our_cost"] for p in props), 2),
                "charged": round(sum(p["charged"] for p in props), 2),
                "margin": round(sum(p["margin"] for p in props), 2),
            })

        totals = {
            "properties": sum(c["property_count"] for c in categories),
            "jobs": sum(c["job_count"] for c in categories) + len(unassigned),
            "our_cost": round(sum(c["our_cost"] for c in categories)
                              + sum(j["our_cost"] for j in unassigned), 2),
            "charged": round(sum(c["charged"] for c in categories)
                             + sum(j["charged"] for j in unassigned), 2),
        }
        totals["margin"] = round(totals["charged"] - totals["our_cost"], 2)

        return json_success({"categories": categories, "unassigned": unassigned,
                             "totals": totals})
    except Exception as e:
        return json_error(safe_error(e), 500)
    finally:
        db.close()


# ─── Contractor invoices ─────────────────────────────────────────────────────
# Separate from the evidence photos because the two answer different questions
# and are read by different people: evidence settles whether the work happened,
# the invoice settles what we were actually charged for it (Norbert, 2026-08-08).

MAINT_INVOICE_DIR = os.path.join(MAINT_EVIDENCE_DIR, "invoices")
os.makedirs(MAINT_INVOICE_DIR, exist_ok=True)

# Invoices arrive as a PDF or a photograph of a paper one. No video: there is no
# such thing as a video invoice, and allowing it only invites 60MB of nothing.
_INVOICE_ALLOWED_EXT = {".pdf", ".jpg", ".jpeg", ".png", ".webp", ".heic"}
_INVOICE_MAX_BYTES = 25 * 1024 * 1024


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>/invoice", methods=["POST"])
def api_maintenance_invoice_upload(job_id):
    """Attach the contractor's invoice to a job. Appended, never replaced -- a
    job split across two invoices is ordinary."""
    import secrets
    db = get_dict_db()
    combined = ""
    try:
        job = db.execute(
            "SELECT id, invoice_paths FROM maintenance_jobs WHERE id = ?", [job_id]
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)
        job = dict(job)

        files = request.files.getlist("file")
        if not files:
            return json_error("No file provided (use field 'file')", 400)

        added = []
        for f in files:
            if not f or not f.filename:
                continue
            name = _evidence_name(f.filename)
            ext = os.path.splitext(name)[1].lower()
            if ext not in _INVOICE_ALLOWED_EXT:
                return json_error(
                    "Cannot attach %s as an invoice — allowed: %s"
                    % (ext or "that file", ", ".join(sorted(_INVOICE_ALLOWED_EXT))), 415)
            payload = f.read()
            if not payload:
                return json_error("%s is empty" % name, 400)
            if len(payload) > _INVOICE_MAX_BYTES:
                return json_error("%s is larger than 25MB" % name, 413)
            stored = "%s_%s_%s" % (job_id, secrets.token_hex(4), name)
            with open(os.path.join(MAINT_INVOICE_DIR, stored), "wb") as fh:
                fh.write(payload)
            added.append("/api/banksia-os/maintenance/invoice/%s" % stored)

        if not added:
            return json_error("Nothing to attach", 400)

        existing = [p.strip() for p in str(job.get("invoice_paths") or "").split(",") if p.strip()]
        combined = ",".join(existing + added)
        db.execute(
            "UPDATE maintenance_jobs SET invoice_paths = ?, modified = datetime('now') WHERE id = ?",
            [combined, job_id]
        )
        db.commit()
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()
    return json_success({"invoice_paths": combined, "added": added})


@banksia_os_bp.route("/maintenance/invoice/<path:stored>", methods=["GET"])
def api_maintenance_invoice_file(stored):
    name = _evidence_name(stored)
    path = os.path.join(MAINT_INVOICE_DIR, name)
    if not os.path.exists(path):
        return json_error("Not found", 404)
    from flask import send_file
    return send_file(path, as_attachment=False)


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>/invoice", methods=["DELETE"])
def api_maintenance_invoice_delete(job_id):
    """Take one invoice off a job. The file stays on disk -- an invoice removed
    by a mis-click is a document we are supposed to be able to produce."""
    target = (request.args.get("path") or "").strip()
    if not target:
        return json_error("Which file?", 400)
    db = get_dict_db()
    try:
        job = db.execute(
            "SELECT invoice_paths FROM maintenance_jobs WHERE id = ?", [job_id]
        ).fetchone()
        if not job:
            return json_error("Job not found", 404)
        kept = [p.strip() for p in str(dict(job).get("invoice_paths") or "").split(",")
                if p.strip() and p.strip() != target]
        db.execute(
            "UPDATE maintenance_jobs SET invoice_paths = ?, modified = datetime('now') WHERE id = ?",
            [",".join(kept), job_id]
        )
        db.commit()
    finally:
        db.close()
    return json_success({"invoice_paths": ",".join(kept)})


@banksia_os_bp.route("/maintenance/jobs/<int:job_id>", methods=["DELETE"])
def api_maintenance_job_delete(job_id):
    """Remove a job from the board.

    The evidence files are left on disk. They cost almost nothing to keep and a
    photo of a finished repair deleted by a mis-click is not recoverable.
    """
    db = get_dict_db()
    try:
        job = db.execute("SELECT reference, title FROM maintenance_jobs WHERE id = ?", [job_id]).fetchone()
        if not job:
            return json_error("Job not found", 404)
        job = dict(job)
        db.execute("DELETE FROM maintenance_orders WHERE job_id = ?", [job_id])
        db.execute("DELETE FROM ll_communications WHERE job_id = ?", [job_id])
        db.execute("DELETE FROM maintenance_jobs WHERE id = ?", [job_id])
        db.commit()
    except Exception as e:
        db.rollback()
        return json_error(safe_error(e), 500)
    finally:
        db.close()
    _log_activity("maintenance", job_id, "delete", "job", job.get("reference") or "", "",
                  notes="job %s deleted" % (job.get("reference") or job.get("title") or job_id))
    return json_success({"deleted": job_id, "reference": job.get("reference") or ""})


def _join_words(items):
    """"a, b and c" -- the message is read by a person, not parsed."""
    items = [str(i) for i in items if i]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return "%s and %s" % (", ".join(items[:-1]), items[-1])


def _known_contractor(name):
    """Return the contractor's stored name, or None if we do not have them.

    Matched case-insensitively and returned in the list's own spelling, so a job
    can never hold a name that differs from the Contractors page by a capital.
    """
    wanted = " ".join(str(name or "").split()).lower()
    if not wanted:
        return ""
    db = get_dict_db()
    try:
        row = db.execute(
            "SELECT name FROM compliance_contractors WHERE LOWER(name) = ?", (wanted,)
        ).fetchone()
    except Exception:
        # No contractors table yet means nobody is on the page yet.
        return None
    return dict(row)["name"] if row else None
