#!/usr/bin/env python3
"""
Banksia OS — Document Generation Engine.
Mail-merge style Word document generation from .docx templates.
Uses python-docx to find and replace merge fields in paragraphs and tables.
"""
import os, json, shutil, uuid, io
from datetime import datetime, date
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from banksia_os_db import get_dict_db

DOCUMENTS_DIR = os.path.join(os.path.dirname(__file__), "documents")
TEMPLATES_DIR = os.path.join(DOCUMENTS_DIR, "templates")
GENERATED_DIR = os.path.join(DOCUMENTS_DIR, "generated")
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

# ── Merge data field map ──
def get_merge_data(tenancy_id):
    """Fetch tenancy data and build a flat dict of merge fields."""
    db = get_dict_db()
    try:
        tenancy = db.execute("SELECT * FROM tenancies WHERE id = ?", (tenancy_id,)).fetchone()
        if not tenancy:
            return None, "Tenancy not found"

        # Get property
        property_data = {}
        if tenancy.get("property_id"):
            prop = db.execute("SELECT * FROM properties WHERE id = ?", (tenancy["property_id"],)).fetchone()
            if prop:
                property_data = dict(prop)

        # Get unit
        unit_data = {}
        if tenancy.get("unit_id"):
            unit = db.execute("SELECT * FROM units WHERE id = ?", (tenancy["unit_id"],)).fetchone()
            if unit:
                unit_data = dict(unit)

        # Get tenants on this tenancy
        tenants = db.execute(
            "SELECT * FROM tenants WHERE tenancy_id = ? OR property_id = ? LIMIT 10",
            (tenancy.get("arthur_id", ""), tenancy.get("property_id"))
        ).fetchall()

        # Get property owner / landlord info
        landlord_name = property_data.get("property_owner_name", "Landlord")
        landlord_address = f"{property_data.get('address_line_1', '')}, {property_data.get('city', '')}, {property_data.get('postcode', '')}"

        # Build merged fields
        main_tenant = tenants[0] if tenants else {}
        tenant_names = ", ".join([f"{t.get('first_name','')} {t.get('last_name','')}".strip() for t in tenants if t.get('first_name') or t.get('last_name')]) or tenancy.get("main_tenant_name", "Tenant")

        fields = {
            # Tenancy
            "TenancyRef": tenancy.get("ref", ""),
            "TenancyStatus": tenancy.get("status", ""),
            "TenancyType": tenancy.get("tenancy_type", ""),
            "ContractType": tenancy.get("contract_type", ""),

            # Tenant(s)
            "TenantName": tenant_names,
            "MainTenantName": main_tenant.get("first_name", "") + " " + main_tenant.get("last_name", "") or tenancy.get("main_tenant_name", ""),
            "TenantFirstName": main_tenant.get("first_name", ""),
            "TenantLastName": main_tenant.get("last_name", ""),
            "TenantEmail": main_tenant.get("email", ""),
            "TenantPhone": main_tenant.get("mobile", "") or main_tenant.get("phone_home", ""),
            "TenantDOB": str(main_tenant.get("date_of_birth", "") or ""),
            "TenantEmployer": main_tenant.get("employment_company", ""),
            "TenantNI": main_tenant.get("ni_number", ""),
            "TenantPassport": main_tenant.get("passport_number", ""),
            "GuarantorName": f"{main_tenant.get('guarantor_first_name','')} {main_tenant.get('guarantor_last_name','')}".strip(),
            "GuarantorEmail": main_tenant.get("guarantor_email", ""),

            # Property
            "PropertyName": property_data.get("name", ""),
            "PropertyRef": property_data.get("ref", ""),
            "PropertyAddress": f"{property_data.get('address_line_1','')}, {property_data.get('address_line_2','')}, {property_data.get('city','')}, {property_data.get('postcode','')}".strip(", "),
            "PropertyAddressLine1": property_data.get("address_line_1", ""),
            "PropertyCity": property_data.get("city", ""),
            "PropertyPostcode": property_data.get("postcode", ""),
            "CouncilTaxBand": property_data.get("council_tax_band", ""),

            # Unit
            "UnitRef": unit_data.get("unit_ref", ""),
            "UnitType": unit_data.get("unit_type", ""),
            "UnitAddress": unit_data.get("full_address", ""),
            "UnitBedrooms": str(unit_data.get("bedrooms", "") or ""),
            "UnitMaxOccupancy": str(unit_data.get("max_occupancy", "") or ""),

            # Rent & Deposit
            "RentAmount": f"£{tenancy.get('rent_amount', 0):,.2f}",
            "RentAmountNumeric": str(tenancy.get("rent_amount", 0)),
            "RentFrequency": tenancy.get("rent_frequency", "pcm"),
            "RentReviewDate": str(tenancy.get("rent_review_date", "") or ""),
            "DepositAmount": f"£{tenancy.get('deposit_registered_amount', 0):,.2f}",
            "DepositScheme": tenancy.get("deposit_scheme", ""),
            "DepositHeldBy": tenancy.get("deposit_held_by", ""),

            # Dates
            "TenancyStartDate": str(tenancy.get("start_date", "") or ""),
            "TenancyEndDate": str(tenancy.get("end_date", "") or ""),
            "RenewalStart": str(tenancy.get("renewal_start", "") or ""),
            "RenewalEnd": str(tenancy.get("renewal_end", "") or ""),
            "BreakClauseDate": str(tenancy.get("break_clause_date", "") or ""),
            "MoveInDate": str(tenancy.get("move_in_date", "") or ""),
            "MoveOutDate": str(tenancy.get("move_out_date", "") or ""),
            "NoticePeriod": tenancy.get("notice_period", ""),

            # Landlord / Agent
            "LandlordName": landlord_name,
            "LandlordAddress": landlord_address,
            "AgentName": "Banksia London",
            "AgentAddress": "Banksia London, London, UK",
            "AgentEmail": "hello@banksialondon.com",
            "AgentPhone": "020 1234 5678",

            # Today
            "Date": date.today().strftime("%d/%m/%Y"),
            "DateLong": date.today().strftime("%d %B %Y"),
            "Year": str(date.today().year),
            "Month": date.today().strftime("%B"),
            "Day": date.today().strftime("%d"),
        }
        return fields, None
    finally:
        db.close()


def replace_in_paragraph(paragraph, fields):
    """Replace {{FieldName}} placeholders in a paragraph's runs."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return False

    modified = False
    for key, value in fields.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
            modified = True

    if modified:
        # Clear existing runs and set the full text in the first run
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            # If no runs, add one
            try:
                from docx.oxml.ns import qn
                new_run = paragraph.add_run(full_text)
            except:
                pass
    return modified


def replace_in_table(table, fields):
    """Replace placeholders in all cells of a table."""
    modified = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, fields):
                    modified = True
    return modified


def replace_in_header_footer(section, fields):
    """Replace placeholders in header/footer."""
    modified = False
    for header in [section.header, section.footer]:
        if header is None:
            continue
        if header.paragraphs:
            for p in header.paragraphs:
                if replace_in_paragraph(p, fields):
                    modified = True
        # Also check tables in header/footer
        if header.tables:
            for table in header.tables:
                if replace_in_table(table, fields):
                    modified = True
    return modified


def generate_document(template_path, tenancy_id, output_path=None):
    """
    Generate a document from a .docx template using tenancy data.
    
    Args:
        template_path: Path to the .docx template file
        tenancy_id: ID of the tenancy to pull data from
        output_path: Path for the output file (auto-generated if None)
    
    Returns:
        (output_path, error_message)
    """
    fields, error = get_merge_data(tenancy_id)
    if error:
        return None, error

    if not os.path.exists(template_path):
        return None, "Template file not found"

    try:
        doc = Document(template_path)
    except Exception as e:
        return None, f"Failed to open template: {str(e)}"

    try:
        # Replace in main body paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, fields)

        # Replace in tables
        for table in doc.tables:
            replace_in_table(table, fields)

        # Replace in headers/footers
        for section in doc.sections:
            replace_in_header_footer(section, fields)

        # Generate output path
        if not output_path:
            tenancy = get_merge_data(tenancy_id)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            tenant_name = fields.get("MainTenantName", "Tenant").replace(" ", "_")
            output_path = os.path.join(GENERATED_DIR, f"document_{tenancy_id}_{tenant_name}_{ts}.docx")

        doc.save(output_path)
        return output_path, None
    except Exception as e:
        return None, f"Document generation failed: {str(e)}"


def get_template_info(template_id):
    """Get template metadata from the JSON index."""
    idx_path = os.path.join(TEMPLATES_DIR, "index.json")
    if not os.path.exists(idx_path):
        return None
    idx = json.load(open(idx_path))
    return idx.get(template_id)


def list_templates():
    """List all uploaded templates."""
    idx_path = os.path.join(TEMPLATES_DIR, "index.json")
    if not os.path.exists(idx_path):
        return []
    return list(json.load(open(idx_path)).values())


def save_template(file_obj, name, description=""):
    """
    Save an uploaded template file.
    Returns (template_id, error)
    """
    template_id = str(uuid.uuid4())[:8]
    filename = f"{template_id}_{name.replace(' ', '_')}.docx"
    dest_path = os.path.join(TEMPLATES_DIR, filename)

    try:
        file_obj.save(dest_path)
    except Exception as e:
        return None, f"Failed to save template: {str(e)}"

    # Update index
    idx_path = os.path.join(TEMPLATES_DIR, "index.json")
    idx = {}
    if os.path.exists(idx_path):
        idx = json.load(open(idx_path))

    idx[template_id] = {
        "id": template_id,
        "name": name,
        "description": description,
        "filename": filename,
        "created": datetime.now().isoformat(),
    }
    json.dump(idx, open(idx_path, "w"), indent=2)

    return template_id, None


def delete_template(template_id):
    """Delete a template by ID."""
    idx_path = os.path.join(TEMPLATES_DIR, "index.json")
    if not os.path.exists(idx_path):
        return False
    idx = json.load(open(idx_path))
    info = idx.pop(template_id, None)
    if not info:
        return False

    # Delete file
    filepath = os.path.join(TEMPLATES_DIR, info["filename"])
    if os.path.exists(filepath):
        os.remove(filepath)

    json.dump(idx, open(idx_path, "w"), indent=2)
    return True


def list_generated_documents():
    """List all generated documents."""
    idx_path = os.path.join(GENERATED_DIR, "index.json")
    if not os.path.exists(idx_path):
        return []
    return list(json.load(open(idx_path)).values())


def record_generated_document(output_path, template_name, tenancy_id, tenant_name):
    """Record a generated document in the index."""
    idx_path = os.path.join(GENERATED_DIR, "index.json")
    idx = {}
    if os.path.exists(idx_path):
        idx = json.load(open(idx_path))

    doc_id = str(uuid.uuid4())[:8]
    idx[doc_id] = {
        "id": doc_id,
        "filename": os.path.basename(output_path),
        "template": template_name,
        "tenancy_id": tenancy_id,
        "tenant_name": tenant_name,
        "generated": datetime.now().isoformat(),
    }
    json.dump(idx, open(idx_path, "w"), indent=2)
    return doc_id